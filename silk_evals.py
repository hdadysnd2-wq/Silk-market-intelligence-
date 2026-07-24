"""أداة تقييم جودة البحث العميق — Silk deep-research quality eval harness
(الموجة ٥، V5).

قاعدة القبول ذاتها المطبَّقة على المنصة بأكملها: **لا اختلاق**. المحور
الأول (استشهاد الأرقام) برمجي حتمي — يفحص هل كل رقم ورد في نص التقرير
المكتوب موجود فعلاً في نتائج البعثات الخام، بلا أي نداء كلود؛ فشله = صفر
فوري لهذا المحور (لا تقدير جزئي — رقم مختلَق واحد يكفي لإسقاطه). الأربعة
الباقية (اكتمال الأقسام، إعلان الفجوات، توصية بلا ادّعاءات غير مسندة،
جودة التقاطعات الخمسة) حَكَمٌ كلود (`_FAST_MODEL`) — تحتاج مفتاح.

**قرار تصميم موثَّق**: نتيجة `run_llm_agent` (`silk_llm_runtime.py`)
تحوّل كل بند مُستشهَد إلى AgentReport.findings حيث `.value` هو **نص
الادّعاء** الذي كتبه كلود (وليس الرقم الخام من الأداة) — سجل نقاط
البيانات الخام (dp1, dp2...) لجولة تشغيل واحدة لا يُخزَّن مستقلاً. لذا
"الأرقام المعروفة" هنا تُستخرَج من نص ادّعاءات/ملاحظات البعثات المخزَّنة
(`finding.value` + `finding.note`) — وهي عملياً نفس الأرقام التي استشهد
بها الوكيل من الأداة (التعليمات تُلزمه بذكرها كما وردت). تحسين لاحق طبيعي:
حفظ السجل الخام نفسه لكل بعثة لفحصٍ أدق — غير مطبَّق هنا (نطاق الموجة ٥).

الحالات الذهبية (`evals/golden_cases.json`) — WS10 أغلق الفجوة النظامية:
الملف يحمل حالةً قياسية **بنيوية** (قطر × HS 200811) تعمل هرمتياً في CI
بلا مفتاح (أقسام/نظافة متن/سلامة مراجع/سقف فجوات)، **بلا `expected`
مُختلَق**. حقل `expected` (رقم محقَّق يدوياً بمصدرٍ حي) يبقى اختيارياً
يُملأ في بيئةٍ بمفتاح — وبقيّة هذا الشرح للجزء الحيّ:
حالةٌ رقمية حقيقية تتطلب أرقاماً مُتحقَّقة يدوياً من مصدر رسمي حيّ
(مثال: Comtrade مباشرة) — هذه البيئة بلا مفتاح Anthropic وبلا وصول شبكي
لمصادر البيانات (Comtrade/WorldBank/GDELT/WITS)، فإضافة حالة الآن تعني
إما اختلاق أرقام أو ترك الحقول فارغة زوراً بمظهر التحقق. البنية والمخطط
(`golden_cases.schema.json`) جاهزان بالكامل ومُختبَران؛ أول حالة حقيقية
مؤجَّلة صراحةً لبيئة بمفتاح حيّ (راجع docs/DEEP_RESEARCH_DECISIONS.md).
"""
from __future__ import annotations

import argparse
import datetime
import json
import logging
import os
import re
import sys

log = logging.getLogger(__name__)

_HERE = os.path.dirname(os.path.abspath(__file__))
_GOLDEN_PATH = os.path.join(_HERE, "evals", "golden_cases.json")
_SCORES_PATH = os.path.join(_HERE, "evals", "scores.json")

# أوزان المحاور الخمسة (تجمع ١٫٠) — المحور الأول وحده برمجي، البقية حَكَم كلود.
AXIS_WEIGHTS = {
    "citation_correctness": 0.35,
    "section_completeness": 0.15,
    "gaps_declared": 0.15,
    "recommendation_grounded": 0.20,
    "intersections_quality": 0.15,
}
_REGRESSION_DROP_THRESHOLD = 10  # نقطة — انخفاض أكبر منها = فشل معلن

_NUM_RE = re.compile(r"-?\d[\d,]*\.?\d*")


def _extract_numbers(text: str) -> list[float]:
    """أرقام مرشَّحة من نص — parses "950,000"/"5.2"/"12" into floats.

    يستثني علامات ترقيم العناوين (أسطر تبدأ بـ"## ") — رقم القسم ليس
    ادّعاءً يحتاج استشهاداً. Header-numbering lines are excluded first.
    """
    body = "\n".join(ln for ln in (text or "").splitlines()
                     if not ln.strip().startswith("## "))
    out: list[float] = []
    for m in _NUM_RE.finditer(body):
        raw = m.group().replace(",", "")
        try:
            out.append(float(raw))
        except ValueError:
            continue
    return out


def _report_fields(rep: object) -> dict:
    """AgentReport حيّة أو dict مُعاد تحميله — نفس نمط silk_render._report_fields."""
    if isinstance(rep, dict):
        return {"findings": rep.get("findings") or []}
    return {"findings": getattr(rep, "findings", None) or []}


def _finding_text(f: object) -> str:
    if isinstance(f, dict):
        return f"{f.get('value')} {f.get('note') or ''}"
    return f"{getattr(f, 'value', '')} {getattr(f, 'note', '') or ''}"


def _known_numbers(mission_reports: dict) -> set[float]:
    """الأرقام المعروفة عبر كل البعثات — من نص الادّعاء/الملاحظة (راجع تعليق
    التصميم أعلى الملف)."""
    known: set[float] = set()
    for rep in (mission_reports or {}).values():
        for f in _report_fields(rep)["findings"]:
            known.update(_extract_numbers(_finding_text(f)))
    return known



# ── محور حسابي (TAM/SAM/SOM وأشباهها) — formula-aware validation ────────────
# رقم مشتق (شريحة سوق، TAM/SAM/SOM، مدى إيراد) ليس اختلاقاً إن كانت مدخلاته
# أرقاماً مسندة فعلاً وكانت المعادلة صحيحة حسابياً — القرار الموثَّق أعلاه
# (`debt 1`، تدقيق ما بعد الموجة ١١): الفحص الحرفي وحده كان يُسقط كل رقم
# TAM/SAM/SOM لأنه ناتج ضرب لا يرد حرفياً في أي بعثة، رغم أن مدخلاته حقيقية.
_EQ_RE = re.compile(
    r"(?P<a>-?\d[\d,]*\.?\d*)(?P<apct>\s*%)?"
    r"[^\n=×xX*÷/+\-−\d]{0,20}"
    r"(?P<op>[×xX*÷/+\-−])\s*"
    r"(?P<b>-?\d[\d,]*\.?\d*)(?P<bpct>\s*%)?"
    r"[^\n=\d]{0,60}"
    r"=\s*"
    r"(?P<c>-?\d[\d,]*\.?\d*)(?P<cpct>\s*%)?"
)
_ASSUMPTION_MARKERS = ("افتراض", "assumption", "assumed", "نفترض", "بافتراض")


def _num_literal(raw: str) -> float:
    return float(raw.replace(",", ""))


def _num_value(raw: str, pct: str | None) -> float:
    v = _num_literal(raw)
    return v / 100.0 if pct else v


def _apply_op(a: float, op: str, b: float) -> float | None:
    if op in ("×", "x", "X", "*"):
        return a * b
    if op in ("÷", "/"):
        return (a / b) if b else None
    if op == "+":
        return a + b
    if op in ("-", "−"):
        return a - b
    return None


def _num_close(actual: float, expected: float | None,
              rel: float = 0.03, abs_tol: float = 1.0) -> bool:
    """تسامح تقريب — كلود يقرّب للمنازل/الآلاف أحياناً عند عرض الناتج."""
    if expected is None:
        return False
    return abs(actual - expected) <= max(abs_tol, abs(expected) * rel)


def _nearby_has_assumption(text: str, start: int, end: int, window: int = 80) -> bool:
    seg = text[max(0, start - window): end + window]
    return any(marker in seg for marker in _ASSUMPTION_MARKERS)


def formula_grounded_numbers(report_text: str, known: set) -> set:
    """أرقام مشتقة صحّت معادلتها ومدخلاتها — derived numbers whose equation
    checks out arithmetically AND at least one operand traces back to a real
    cited number (مباشرة أو عبر سلسلة معادلات سابقة) — لا يُقبل ناتج معادلة
    كلا طرفيها افتراض غير مسند، فهذا يعادل اختلاق سلسلة كاملة من لا شيء.
    """
    grounded: set = set()
    rooted: set = set()
    for m in _EQ_RE.finditer(report_text or ""):
        try:
            a_lit, b_lit, c_lit = (_num_literal(m.group(g)) for g in ("a", "b", "c"))
            expected = _apply_op(_num_value(m.group("a"), m.group("apct")),
                                 m.group("op"),
                                 _num_value(m.group("b"), m.group("bpct")))
            c_val = _num_value(m.group("c"), m.group("cpct"))
        except (ValueError, ZeroDivisionError):
            continue
        if not _num_close(c_val, expected):
            continue
        a_known = a_lit in known or a_lit in rooted
        b_known = b_lit in known or b_lit in rooted
        a_ok = a_known or a_lit in grounded or _nearby_has_assumption(
            report_text, m.start("a"), m.end("a"))
        b_ok = b_known or b_lit in grounded or _nearby_has_assumption(
            report_text, m.start("b"), m.end("b"))
        if a_ok and b_ok and (a_known or b_known):
            grounded.add(c_lit)
            rooted.add(c_lit)
            # حصة/افتراض مُعلَن صراحة بجوار المعادلة الصحيحة — ليس ادّعاءً
            # يحتاج استشهاداً خارجياً؛ استثنِ رقمه الحرفي أيضاً (مثال: "15%"
            # في "TAM × 15% (افتراض) = SAM").
            if _nearby_has_assumption(report_text, m.start("a"), m.end("a")) and not a_known:
                grounded.add(a_lit)
            if _nearby_has_assumption(report_text, m.start("b"), m.end("b")) and not b_known:
                grounded.add(b_lit)
    return grounded


def citation_correctness_score(report_text: str,
                               mission_reports: dict) -> dict:
    """المحور البرمجي — كل رقم في التقرير يجب أن يرد حرفياً في بعثة ما، أو أن
    يكون ناتج معادلة صحيحة مدخلاتها مسندة (TAM/SAM/SOM وأشباهها).

    يعيد {"score": 0|100, "violations": [أرقام غير مسندة]} — لا نداء كلود،
    لا تقدير جزئي (رقم واحد مختلَق = صفر الفور، نفس مبدأ لا اختلاق).
    """
    known = _known_numbers(mission_reports)
    grounded = formula_grounded_numbers(report_text, known)
    report_numbers = _extract_numbers(report_text)
    effective = known | grounded
    violations = sorted({n for n in report_numbers if n not in effective})
    return {"score": 0 if violations else 100, "violations": violations,
           "checked": len(report_numbers), "known_pool": len(known),
           "formula_grounded": sorted(grounded)}


def _judge_prompt(result: dict) -> str:
    from silk_ai_judge import _isolate
    dr = result.get("deep_research") or {}
    report_text = (dr.get("report") or {}).get("report") or ""
    missions = dr.get("missions") or {}
    mission_summaries = "\n".join(
        f"- [{k}] فشل={_report_fields(v).get('failed', False)}"
        for k, v in missions.items())
    return (
        "قيّم تقرير بحث عميق على أربعة محاور (٠-١٠٠ لكل محور)، بعيداً عن "
        "استشهاد الأرقام (يُفحص برمجياً بمعزل عنك):\n"
        "1. section_completeness: هل الأقسام الخمسة عشر المطلوبة حاضرة "
        "بمحتوى فعلي لا عناوين فارغة؟\n"
        "2. gaps_declared: هل الفجوات (بيانات غائبة/بعثات فاشلة) مُعلَنة "
        "صراحة بدل تجاهلها؟\n"
        "3. recommendation_grounded: هل التوصية النهائية خالية من ادّعاءات "
        "غير مسندة لحقائق البعثات؟\n"
        "4. intersections_quality: هل التقاطعات الخمسة (الطلب/تكلفة الدخول/"
        "التنافسية السعرية/أبواب الدخول/SWOT) مبنية منطقياً من الأدلة؟\n\n"
        f"ملخص البعثات:\n{_isolate(mission_summaries)}\n\n"
        f"نص التقرير:\n{_isolate(report_text[:8000])}\n\n"
        'أعد JSON فقط: {"section_completeness":N,"gaps_declared":N,'
        '"recommendation_grounded":N,"intersections_quality":N,'
        '"reasoning":"سبب موجز"}')


def evaluate_report(result: dict) -> dict | None:
    """قيّم تقريراً كاملاً — المحور البرمجي أولاً (لا كلود)، ثم أربعة محاور
    حَكَم كلود إن توفّر مفتاح. None فقط إن تعذّر النداء ولم يوجد نص تقرير
    أصلاً؛ استشهاد الأرقام يُحسب دوماً بلا حاجة لمفتاح."""
    dr = result.get("deep_research") or {}
    report_text = (dr.get("report") or {}).get("report") or ""
    mission_reports = dr.get("missions") or {}
    citation = citation_correctness_score(report_text, mission_reports)

    axes = {"citation_correctness": citation["score"]}
    from silk_ai_judge import _call, _FAST_MODEL, _PRINCIPLE, available
    llm_axes: dict | None = None
    if report_text and available():
        raw = _call(_PRINCIPLE, _judge_prompt(result), max_tokens=700,
                    model=_FAST_MODEL, timeout=30)
        if raw:
            try:
                start, end = raw.find("{"), raw.rfind("}")
                llm_axes = json.loads(raw[start:end + 1]) if start >= 0 else None
            except Exception:  # noqa: BLE001 — رد غير-JSON = محاور كلود غائبة
                llm_axes = None
    for axis in ("section_completeness", "gaps_declared",
                "recommendation_grounded", "intersections_quality"):
        val = (llm_axes or {}).get(axis)
        try:
            axes[axis] = max(0.0, min(100.0, float(val)))
        except (TypeError, ValueError):
            axes[axis] = None  # محور غير محسوب (لا مفتاح/فشل) — فجوة معلنة

    scored = {k: v for k, v in axes.items() if v is not None}
    total_weight = sum(AXIS_WEIGHTS[k] for k in scored)
    overall = (round(sum(axes[k] * AXIS_WEIGHTS[k] for k in scored)
                     / total_weight, 1) if total_weight else None)
    return {
        "overall": overall, "axes": axes,
        "citation_violations": citation["violations"],
        "reasoning": (llm_axes or {}).get("reasoning", ""),
        "grounded_axes": sorted(scored),
        "note": ("محاور كلود غير محسوبة — بلا مفتاح ANTHROPIC_API_KEY"
                 if llm_axes is None else ""),
    }


# ── الحالات الذهبية · golden cases ───────────────────────────────────────────

# الحقول الإلزامية — `expected` (أرقام مُتحقَّقة يدوياً) صار **اختيارياً**: حالةٌ
# ذهبية بنيوية (أقسام/نظافة متن/سلامة مراجع/سقف فجوات) لا تتطلّب رقماً محقَّقاً
# حياً (لا يمكن التحقّق منه في بيئةٍ بلا شبكة — لا اختلاق). كل حالة تحمل إمّا
# `expected` (أرقام) أو `structural` (بوّابة بنيوية) أو كليهما.
_REQUIRED_CASE_FIELDS = ("key", "product", "market", "hs_code",
                         "verified_at", "verified_by")

_STRUCTURAL_KEYS = ("required_sections", "clean_body", "references_integrity",
                    "gap_rate_max",
                    # الهوتفكس (بلاغ قطر ٢٠٢٦-٠٧-٢٣):
                    "no_truncation_artifacts",   # HF2: لا خليةٌ مبتورةٌ داخل رقم/
                                                 # قوسٌ غير متوازن/مجموعةٌ فارغة
                    "plausibility_reconciled")   # HF3: لا مقدارٌ متعارضٌ بلا تحفّظ


def validate_case(case: dict) -> list[str]:
    """تحقّق خفيف من مخطط الحالة — لا مكتبة jsonschema (stdlib-first)؛
    يعيد قائمة الأخطاء (فارغة = صالحة)."""
    errors = []
    for field in _REQUIRED_CASE_FIELDS:
        if field not in case:
            errors.append(f"missing field: {field}")
    hs = case.get("hs_code", "")
    if hs and not re.fullmatch(r"\d{6}", str(hs)):
        errors.append(f"hs_code must be 6 digits: {hs!r}")
    expected = case.get("expected")
    if isinstance(expected, dict):
        for k, v in expected.items():
            if not isinstance(v, dict) or "value" not in v or "source_url" not in v:
                errors.append(f"expected.{k} missing value/source_url")
    elif "expected" in case:
        errors.append("expected must be an object")
    struct = case.get("structural")
    if struct is not None:
        if not isinstance(struct, dict):
            errors.append("structural must be an object")
        else:
            rs = struct.get("required_sections")
            if rs is not None and not (isinstance(rs, list)
                                       and all(isinstance(s, str) for s in rs)):
                errors.append("structural.required_sections must be a list of strings")
            gm = struct.get("gap_rate_max")
            if gm is not None and not (isinstance(gm, (int, float))
                                       and 0.0 <= float(gm) <= 1.0):
                errors.append("structural.gap_rate_max must be a number in [0,1]")
            for b in ("clean_body", "references_integrity",
                      "no_truncation_artifacts", "plausibility_reconciled"):
                if b in struct and not isinstance(struct[b], bool):
                    errors.append(f"structural.{b} must be a boolean")
    # حالةٌ بلا `expected` **ولا** `structural` = بلا معيار قبول (مرفوضة).
    if "expected" not in case and struct is None:
        errors.append("case must declare 'expected' and/or 'structural'")
    return errors


def load_golden_cases(path: str = _GOLDEN_PATH) -> list[dict]:
    """حمّل الحالات الذهبية — صفوف صالحة فقط؛ صفّ فاسد يُسجَّل ويُستبعد لا يُسقط التحميل."""
    try:
        with open(path, encoding="utf-8") as f:
            raw = json.load(f)
    except Exception as e:  # noqa: BLE001 — ملف غائب/فاسد = قائمة فارغة، لا اختلاق
        log.warning("golden cases unavailable (%s): %s", path, e)
        return []
    if not isinstance(raw, list):
        return []
    out = []
    for case in raw:
        errs = validate_case(case) if isinstance(case, dict) else ["not an object"]
        if errs:
            log.warning("golden case rejected (%s): %s", case.get("key", "?"), errs)
            continue
        out.append(case)
    return out


def load_scores(path: str = _SCORES_PATH) -> dict:
    try:
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    except Exception:  # noqa: BLE001 — لا سجل سابق = قاموس فارغ
        return {}


def save_scores(scores: dict, path: str = _SCORES_PATH) -> None:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(scores, f, ensure_ascii=False, indent=2, sort_keys=True)


def compare_to_last_score(case_key: str, new_score: float | None,
                          scores: dict | None = None,
                          path: str = _SCORES_PATH) -> dict:
    """قارن بالنتيجة المحفوظة — انخفاض > 10 نقطة = فشل معلن (لا يُسقط شيئاً،
    يُعلَن). أول تشغيلة لحالة (لا سجل سابق) = نجاح دوماً (لا أساس للمقارنة)."""
    scores = scores if scores is not None else load_scores(path)
    prev = scores.get(case_key)
    if prev is None or new_score is None:
        return {"regression": False, "previous": prev, "new": new_score,
               "note": "لا نتيجة سابقة للمقارنة" if prev is None
                       else "لا نتيجة جديدة محسوبة"}
    drop = prev - new_score
    return {"regression": drop > _REGRESSION_DROP_THRESHOLD,
           "previous": prev, "new": new_score, "drop": round(drop, 1)}


def run_case(case: dict) -> dict:
    """أعد تشغيل حالة ذهبية حياً — ١٢ بعثة + محلل + توليف + كاتب، يتطلب
    شبكة ومفتاح Anthropic فعليين (لا يعمل في هذه البيئة/CI). غير مُختبَر
    هيرمتياً بتصميم — الاختبارات تموّه هذه الدالة عند فحص منطق الـCLI."""
    from silk_market_resolver import resolve_market
    from silk_missions import run_all_missions
    from silk_market_analyst import analyze_market, to_synthesis_input
    from silk_synthesis import synthesize
    from silk_ai_judge import write_reviewed_report

    ref, suggestions = resolve_market(case["market"])
    if ref is None:
        raise ValueError(f"cannot resolve market {case['market']!r}: "
                         f"suggestions={suggestions}")
    mission_reports = run_all_missions(ref, product=case["product"],
                                       hs_code=case["hs_code"])
    analyst_out = analyze_market(ref, case["product"], mission_reports,
                                 hs_code=case["hs_code"])
    analyst_input = to_synthesis_input(analyst_out)
    verdict = synthesize(list(mission_reports.values()), product=case["product"],
                         market=ref.name_en, with_ai=True,
                         analyst_assessment=analyst_input)
    report_out = write_reviewed_report(mission_reports, analyst_input["summary"],
                                       verdict, case["product"], ref.name_en)
    return {"product": case["product"], "hs_code": case["hs_code"],
           "market": {"iso3": ref.iso3, "m49": ref.m49, "name_en": ref.name_en,
                     "name_ar": ref.name_ar},
           "markets": [],
           "deep_research": {"missions": mission_reports, "analyst": analyst_out,
                            "verdict": verdict, "report": report_out}}


# ── البوّابة البنيوية (WS10) — هرمتية بالكامل، صفر مفتاح، صفر تكلفة ───────────
# تُشغَّل على **نتيجةٍ مُولَّدة** (حيّة في بيئةٍ بمفتاح، أو مموّهة في CI): حضور
# الأقسام، نظافة المتن (لا شارات/أعمدة إسناد)، سلامة المراجع (WS9)، وسقف نسبة
# الفجوات. لا تستدعي أيّ نموذج — تفحص بنية التقرير المُصيَّر وحقائق البعثات.

_FORBIDDEN_BODY_TOKENS = ("قوة الدليل", "مستوى التوثيق", "✓ موثّق",
                          "◐ ثانوي", "○ غير", "✓", "◐", "○")


def _rendered_client_body(result: dict):
    """(نصّ متن تقرير العميل، ترويسات الجداول) — hermetic، عبر render_client_docx.

    يُصيَّر المستند الفعلي ويُستخرَج نصّه (فقرات + خلايا) كي تفحص البوّابة ما
    يراه العميل حقيقةً لا مسوّدة الكاتب الخام. يتطلّب python-docx (مثبّت في
    الاختبارات والبيئة المفتاحية)."""
    import os
    import tempfile
    import silk_render
    import silk_reports
    from docx import Document
    view = silk_render.build_view(result)
    path = os.path.join(tempfile.mkdtemp(), "eval_client.docx")
    silk_reports.render_client_docx(view, path)
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    headers: list[str] = []
    cells: list[str] = []
    for t in doc.tables:
        if t.rows:
            headers += [c.text for c in t.rows[0].cells]
        for row in t.rows:
            row_cells = [c.text for c in row.cells]
            cells += row_cells
            parts += row_cells
    # HF2: خلايا الجداول وحداتٌ **ذرّية** (حدُّها حقيقيّ) — تُفحَص فيها نهايةُ
    # الرقم/توازنُ الأقواس. الفقراتُ النثريةُ يُقسِّمها المُصيِّر للعرض فلا يصحّ
    # فحصُ توازنِها الموضعيّ (تعارضٌ مشروعٌ عبر حدود الفقرة) — تُفحَص فيها
    # المجموعةُ الفارغة فقط (لا تُشرَّع أبداً)، وذلك على المتن الكامل أدناه.
    return "\n".join(parts), headers, cells


def gap_rate(result: dict) -> tuple:
    """(نسبة الفجوات، عدد الفجوات، الإجمالي) — نتيجةٌ بلا قيمة (None) أو
    `status=fetch_failed` = فجوة معلنة؛ النسبة = الفجوات ÷ إجمالي نتائج البعثات."""
    missions = (result.get("deep_research") or {}).get("missions") or {}
    total = gaps = 0
    for m in missions.values():
        if not isinstance(m, dict):
            continue
        for f in (m.get("findings") or []):
            if not isinstance(f, dict):
                continue
            total += 1
            if f.get("value") is None or f.get("status") == "fetch_failed":
                gaps += 1
    return ((gaps / total) if total else 0.0), gaps, total


def structural_checks(result: dict, case: dict) -> dict:
    """بوّابة بنيوية هرمتية — {passed, checks, failures}. لا نموذج، لا مفتاح."""
    struct = case.get("structural") or {}
    text, headers, cells = _rendered_client_body(result)
    checks: dict = {}
    failures: list[str] = []

    req = struct.get("required_sections") or []
    missing = [s for s in req if s not in text]
    checks["required_sections"] = {"missing": missing, "passed": not missing}
    if missing:
        failures.append(f"أقسام مطلوبة مفقودة: {missing}")

    if struct.get("clean_body"):
        hits = sorted({tok for tok in _FORBIDDEN_BODY_TOKENS if tok in text})
        if "المصدر" in headers:
            hits.append("عمود «المصدر» لكل صف")
        if "قوة الدليل" in headers:
            hits.append("عمود «قوة الدليل»")
        checks["clean_body"] = {"hits": hits, "passed": not hits}
        if hits:
            failures.append(f"شارات/أعمدة إسناد في المتن: {hits}")

    if struct.get("references_integrity"):
        ok, detail = _references_integrity(result, text)
        checks["references_integrity"] = {"passed": ok, "detail": detail}
        if not ok:
            failures.append(f"سلامة المراجع (WS9): {detail}")

    # HF2: لا أثرَ بترٍ — خليةُ جدولٍ تنتهي داخل رقم أو بقوسٍ غير متوازن،
    # أو مجموعةُ استشهادٍ فارغة «()»/«(/)»/«(///)» في أيّ موضعٍ من المتن.
    if struct.get("no_truncation_artifacts"):
        hits = _truncation_hits(cells)
        if _EVAL_EMPTY_GROUP_RE.search(text):
            hits.append("مجموعةُ استشهادٍ فارغةٌ في المتن")
        checks["no_truncation_artifacts"] = {"hits": hits[:12],
                                             "passed": not hits}
        if hits:
            failures.append(f"آثارُ بترٍ في المتن (HF2): {hits[:6]}")

    # HF3: كلُّ مقدارٍ متعارضٍ مع مرتكزات التشغيلة إمّا مُسقَطٌ أو مُتحفَّظٌ عليه
    # صراحةً — لا رقمٌ غيرُ مُصالَحٍ يصل العميلَ صامتاً.
    if struct.get("plausibility_reconciled"):
        ok, detail = _plausibility_reconciled(result, text)
        checks["plausibility_reconciled"] = {"passed": ok, "detail": detail}
        if not ok:
            failures.append(f"مقدارٌ غيرُ مُصالَح (HF3): {detail}")

    if "gap_rate_max" in struct:
        rate, g, tot = gap_rate(result)
        passed = rate <= float(struct["gap_rate_max"])
        checks["gap_rate"] = {"rate": round(rate, 3), "gaps": g, "total": tot,
                              "max": struct["gap_rate_max"], "passed": passed}
        if not passed:
            failures.append(
                f"نسبة الفجوات {rate:.1%} تتجاوز السقف "
                f"{float(struct['gap_rate_max']):.0%}")

    return {"passed": not failures, "checks": checks, "failures": failures}


def _references_integrity(result: dict, body_text: str) -> tuple:
    """WS9 (+HF1) — قسم المراجع حاضرٌ حين توجد نتائجُ قابلةٌ للاستشهاد، بلا
    مصدرٍ نائبٍ عام، **بلا معرّفٍ مركّب**، واتّحادُ المصادر مُتحقَّقٌ في الاتجاهين:
    كلُّ مصدرٍ مُستعمَلٍ يظهر (used-but-missing) وكلُّ مصدرٍ مُدرَجٍ مُستعمَلٌ فعلاً
    (listed-but-unused — يلتقط «OpenAlex مُدرَجٌ ولم يُستعمَل»)."""
    # HF1: اقرأ حقائقَ **النموذج المُصيَّر** (view["deep_research"]) لا النتيجةَ
    # الخام — فممرّ المصالحة (WP-3) في build_view يكتب وسومَ عرضٍ (شارة الدليل)
    # على بنود العرض، فتطابق مجموعةُ المراجع المُعاد حسابُها هنا ما يُصيَّر فعلاً.
    import silk_render
    missions = ((silk_render.build_view(result) or {}).get("deep_research")
                or {}).get("missions") or {}
    has_citable = any(
        isinstance(f, dict) and f.get("value") is not None
        for m in missions.values() if isinstance(m, dict)
        for f in (m.get("findings") or []))
    if "المراجع" not in body_text:
        if has_citable:
            return False, "قسم «المراجع» غائب رغم وجود نتائج قابلة للاستشهاد"
        return True, "لا نتائج قابلة للاستشهاد ⇒ لا قسم مراجع (سليم)"
    for placeholder in ("Web Search", "مرجع سلك", "Silk L1 reference",
                        "Silk requirements"):
        if placeholder in body_text:
            return False, f"مصدرٌ نائبٌ عامٌّ في المراجع: {placeholder!r}"

    # HF1: أعِدْ حسابَ اتحاد المصادر الذرّية المُستعمَلة (نفسُ منطق بناء المراجع)
    # فنقارنه بالمُدرَج فعلاً في المتن — بلا معرّفٍ مركّب في أيّ اتجاه.
    import silk_reports as R
    from silk_data_layer import atomic_source_ids, is_atomic_source_id
    from silk_narrative import RECONCILED_OUT_TAG, evidence_badge_for
    expected: dict = {}   # name.casefold -> name
    for m in missions.values():
        if not isinstance(m, dict):
            continue
        for f in (m.get("findings") or []):
            if not isinstance(f, dict):
                continue
            badge = evidence_badge_for(f)
            if badge.startswith("○") or badge.startswith(RECONCILED_OUT_TAG):
                continue
            for sid in atomic_source_ids(f.get("source"), f.get("source_ids")):
                if not is_atomic_source_id(sid):
                    return False, f"معرّفُ مصدرٍ **مركّب** (غير ذرّي): {sid!r}"
                row = R._reference_row_from_finding(
                    sid, f.get("value"), f.get("note"))
                if row:
                    expected[row[0].strip().casefold()] = row[0]
    # (أ) used-but-missing: كلُّ مصدرٍ مُستعمَلٍ يحلّ لرابطٍ يجب أن يظهر في المتن.
    for name in expected.values():
        if name not in body_text:
            return False, f"مصدرٌ مُستعمَلٌ غائبٌ عن المراجع: {name!r}"
    # (ب) listed-but-unused: كلُّ سطرِ مرجعٍ مُدرَجٍ («name — http…») لا بدّ أن
    # يكون مصدراً مُستعمَلاً فعلاً (WS9 union) — يلتقط مرجعاً مُدرَجاً بلا استعمال.
    listed = _listed_reference_names(body_text)
    for name in listed:
        if name.strip().casefold() not in expected:
            return False, f"مرجعٌ مُدرَجٌ لم يُستعمَل في أيّ نتيجة: {name!r}"
    # (ج) لا فاصلَ دمجٍ عربيّ في اسم مرجعٍ مُدرَج (شبكةُ أمانٍ نصّية).
    for name in listed:
        if not is_atomic_source_id(name):
            return False, f"اسمُ مرجعٍ مُدرَجٍ مركّب: {name!r}"
    return True, (f"المراجع سليمة: {len(expected)} مصدرٌ ذرّيٌّ مُستعمَلٌ "
                  "بروابطَ صحيحة، بلا مركّبٍ ولا مُدرَجٍ زائد")


def _listed_reference_names(body_text: str) -> list:
    """أسماءُ المصادر المُدرَجةُ فعلاً في قسم «المراجع» من المتن المُصيَّر — أسطرُ
    «الاسم — https://…» بعد ترويسة «المراجع». (لعكس اتّحاد WS9.)"""
    lines = str(body_text or "").split("\n")
    try:
        start = next(i for i, ln in enumerate(lines) if ln.strip() == "المراجع")
    except StopIteration:
        return []
    out: list = []
    for ln in lines[start + 1:]:
        m = re.match(r"\s*(.+?)\s+—\s+https?://", ln)
        if m:
            out.append(m.group(1).strip())
    return out


# HF2: أنماطُ آثار البتر — فاصلٌ عشريٌّ متدلٍّ في نهاية خلية، ومجموعةُ استشهادٍ
# فارغة. القوسُ غيرُ المتوازن يُحسَب عدّاً.
_EVAL_DANGLING_DECIMAL_RE = re.compile(r"[0-9٠-٩][.,،٬٫]\s*$")
_EVAL_EMPTY_GROUP_RE = re.compile(r"[\(（]\s*[/／،,;\s]*[\)）]")


def _truncation_hits(units: list) -> list:
    """آثارُ البتر في وحدات المتن (فقرات + خلايا) — HF2. لكلٍّ: لا تنتهي بفاصلٍ
    عشريٍّ متدلٍّ («…7.»)، لا مجموعةَ استشهادٍ فارغة («()»)، والأقواسُ متوازنة."""
    hits: list = []
    for u in units:
        t = str(u or "").strip()
        if not t:
            continue
        if _EVAL_DANGLING_DECIMAL_RE.search(t):
            hits.append(f"نهايةٌ داخل رقم: …{t[-20:]!r}")
        if _EVAL_EMPTY_GROUP_RE.search(t):
            hits.append(f"مجموعةٌ فارغة: {t[:40]!r}")
        if (t.count("(") + t.count("（")) != (t.count(")") + t.count("）")):
            hits.append(f"قوسٌ غير متوازن: {t[:40]!r}")
    return hits


def _plausibility_reconciled(result: dict, body_text: str) -> tuple:
    """HF3 — كلُّ مقدارٍ متعارضٍ مع مرتكزات التشغيلة إمّا مُسقَطٌ (action=drop)
    أو مُتحفَّظٌ عليه صراحةً في المتن (تحفّظُ نطاق). لا رقمٌ غيرُ مُصالَحٍ صامت."""
    import silk_plausibility as P
    flags = P.check_magnitudes(result)
    if not flags:
        return True, "لا مقادير متعارضة مع المرتكزات"
    if P.action() == "drop":
        return True, f"{len(flags)} مقدارٌ مُسقَطٌ بسببٍ مُسجَّلٍ في المانيفست"
    if "تنبيه تحقّق" in body_text or "يتعذّر التوفيق" in body_text:
        return True, f"{len(flags)} مقدارٌ مُتحفَّظٌ عليه صراحةً في المتن"
    return False, f"{len(flags)} مقدارٌ متعارضٌ بلا تحفّظٍ ظاهرٍ للعميل"


def _has_key() -> bool:
    """هل يتوفّر مفتاح Anthropic لتشغيل الجزء الحيّ؟"""
    return bool(os.environ.get("ANTHROPIC_API_KEY", "").strip())


def aggregate_gap_rates(results_by_case: dict, cases_by_key: dict) -> dict:
    """نسبة الفجوات **لكل حالة** (مع سقفها الخاصّ ونجاحها) **والمجموع الكلّي**.

    قرار مقصود: يُعرَض الاثنان معاً — المجموع وحده يُخفي سوقاً منهارة (سوقٌ
    نظيفةٌ كبيرة تُميّع نسبةَ سوقٍ صغيرةٍ انهارت تغطيتها). كل حالةٍ تُقيَّم مقابل
    `gap_rate_max` الخاصّ بها (per-case لا سقفٌ عالميّ — السوق الصعبة تحتمل
    أرضيةً أعلى). hermetic بالكامل — لا نموذج، لا مفتاح."""
    per_case: dict = {}
    tot_gaps = tot_total = 0
    for key, result in results_by_case.items():
        rate, g, t = gap_rate(result)
        cap = ((cases_by_key.get(key) or {}).get("structural") or {}).get(
            "gap_rate_max")
        per_case[key] = {"rate": round(rate, 3), "gaps": g, "total": t,
                         "max": cap,
                         "passed": (cap is None or rate <= float(cap))}
        tot_gaps += g
        tot_total += t
    agg = (tot_gaps / tot_total) if tot_total else 0.0
    return {
        "per_case": per_case,
        "aggregate": {"rate": round(agg, 3), "gaps": tot_gaps,
                      "total": tot_total},
        "any_case_over_ceiling": any(not v["passed"]
                                     for v in per_case.values()),
    }


def run_suite(cases: list) -> dict:
    """شغّل كل الحالات حياً (يتطلب مفتاحاً وشبكة) — {key: (case, result,
    evaluation, structural)}. غير مُختبَر هرمتياً (يموّه المُستدعي run_case)."""
    out: dict = {}
    for case in cases:
        result = run_case(case)
        out[case["key"]] = {
            "case": case, "result": result,
            "evaluation": evaluate_report(result),
            "structural": structural_checks(result, case)}
    return out


def main(argv: list[str] | None = None) -> int:
    logging.basicConfig(level=logging.INFO)
    ap = argparse.ArgumentParser(
        description="أعد تشغيل حالة/حالات ذهبية: بوّابة بنيوية هرمتية + تقييمُ "
                    "حَكَمٍ حيّ. الجزء الحيّ يتطلب شبكة ومفتاح Anthropic — يتخطّى "
                    "بسببٍ معلن (لا يفشل) حين لا مفتاح، فيبقى مخرَج CI نظيفاً.")
    grp = ap.add_mutually_exclusive_group(required=True)
    grp.add_argument("--case", help="مفتاح حالة واحدة (key)")
    grp.add_argument("--all", action="store_true",
                     help="كل الحالات + تقرير نسبة فجوات لكل حالة والمجموع")
    args = ap.parse_args(argv)

    cases = {c["key"]: c for c in load_golden_cases()}
    if args.case is not None and args.case not in cases:
        log.error("unknown golden case %r (available: %s)", args.case,
                  sorted(cases) or "none — evals/golden_cases.json فارغ حالياً")
        return 1
    selected = (list(cases.values()) if args.all
                else [cases[args.case]])
    keys = [c["key"] for c in selected]

    # تخطٍّ صريحٌ بسبب (لا فشل) حين لا مفتاح — الجزء الحيّ يحتاج شبكةً ومفتاحاً؛
    # فبيئةُ CI/الصندوق تبقى إشارتُها نظيفة (نفس عقد #166).
    if not _has_key():
        print(json.dumps({
            "cases": keys, "skipped": True,
            "reason": "ANTHROPIC_API_KEY غير مضبوط — الجزء الحيّ (run_case + "
                      "تقييم الحَكَم) يتطلّب مفتاحاً وشبكة. البوّابة البنيوية "
                      "الهرمتية (بما فيها نسبة الفجوات لكل حالة والمجموع) "
                      "تُشغَّل في CI عبر tests/test_ws10_golden_case.py.",
            "how_to_run_live": ("ANTHROPIC_API_KEY=<key> python3 silk_evals.py "
                                + ("--all" if args.all
                                   else f"--case {args.case}")),
        }, ensure_ascii=False, indent=2))
        return 0

    suite = run_suite(selected)
    scores = load_scores()
    results_by_case = {k: v["result"] for k, v in suite.items()}
    gap_report = aggregate_gap_rates(results_by_case, cases)

    regressed = False
    structural_failed = False
    per_case_out: dict = {}
    for key, s in suite.items():
        overall = (s["evaluation"] or {}).get("overall")
        cmp = compare_to_last_score(key, overall, scores)
        per_case_out[key] = {"evaluation": s["evaluation"],
                             "structural": s["structural"], "comparison": cmp}
        if overall is not None:
            scores[key] = overall
        if cmp["regression"]:
            regressed = True
            log.error("quality regression on %s: %s -> %s (drop %.1f)",
                     key, cmp["previous"], cmp["new"], cmp["drop"])
        if not s["structural"]["passed"]:
            structural_failed = True
            log.error("structural gate failed on %s: %s",
                     key, s["structural"]["failures"])

    save_scores(scores)
    print(json.dumps({"cases": per_case_out, "gap_rates": gap_report},
                     ensure_ascii=False, indent=2))
    return 1 if (regressed or structural_failed) else 0


if __name__ == "__main__":
    sys.exit(main())
