"""اختبارات جداول التقرير الاحترافية — real docx tables + market-entry strategy.

مراجعة المشروع (بلاغ المالك: "التقرير غير مقنعة، شوف كيف تكون احترافية"):
بالمقارنة بمنصات مرجعية (Country Commercial Guides، ITC Trade Map،
Euromonitor) — العنصر المشترك: جداول حقيقية لا نقاط سردية، وفصل مستقل
لاستراتيجية دخول السوق بدل سطر مدفون. يقفل هذا الملف: (١) قرار الدخول/
حجم السوق/التسعير/الاشتراطات تُعرض كجداول Word فعلية، (٢) SWOT شبكة ٢×٢
حقيقية لا أربع عناوين متتالية، (٣) فصل "استراتيجية دخول السوق" الجديد
يركّب توصية من أرقامٍ مرصودة فعلاً — لا رقم جديد ولا اختلاق.
Run:  python3 -m pytest tests/test_report_professional_tables.py -q
"""
import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from conftest import block_network, docx_all_text  # noqa: E402

import silk_store  # noqa: E402


def _seed_store():
    silk_store.migrate()
    silk_store.upsert_trade_flows([
        {"hs6": "080410", "reporter_iso3": "CHN", "partner_iso3": "WLD",
         "year": 2021, "flow": "M", "value_usd": 4.0e7},
        {"hs6": "080410", "reporter_iso3": "CHN", "partner_iso3": "WLD",
         "year": 2022, "flow": "M", "value_usd": 5.0e7},
        {"hs6": "080410", "reporter_iso3": "CHN", "partner_iso3": "WLD",
         "year": 2023, "flow": "M", "value_usd": 6.0e7, "qty_kg": 2.0e7},
        {"hs6": "080410", "reporter_iso3": "CHN", "partner_iso3": "IRN",
         "year": 2023, "flow": "M", "value_usd": 3.0e7},
        {"hs6": "080410", "reporter_iso3": "CHN", "partner_iso3": "SAU",
         "year": 2023, "flow": "M", "value_usd": 1.8e7, "qty_kg": 8.0e6},
        {"hs6": "080410", "reporter_iso3": "CHN", "partner_iso3": "TUN",
         "year": 2023, "flow": "M", "value_usd": 1.2e7}])


def _view():
    import silk_engine
    from silk_render import build_view
    with block_network():
        result = silk_engine.analyze(
            "تمور", countries=[{"iso3": "CHN", "m49": "156"}], year=2023,
            with_research=True, with_requirements=True, with_risk=True)
    return build_view(result)


def _render(view):
    import pytest
    pytest.importorskip("docx")
    from silk_reports import render_docx
    return render_docx(view, os.path.join(tempfile.mkdtemp(), "r.docx"))


def test_pillars_market_size_and_regulatory_are_real_tables():
    _seed_store()
    from docx import Document
    path = _render(_view())
    doc = Document(path)
    assert doc.tables, "render_docx must emit at least one real table"
    header_rows = ["".join(c.text for c in t.rows[0].cells) for t in doc.tables]
    joined_headers = "\n".join(header_rows)
    # قرار الدخول: جدول الأعمدة الأربعة (يوازي render_markdown أخيراً).
    assert any("العمود" in h and "الأساس" in h for h in header_rows), joined_headers
    # حجم السوق TAM/SAM/SOM: جدول لا نقاط.
    assert any("المؤشر" in h and "النوع" in h for h in header_rows), joined_headers


def test_swot_is_a_real_2x2_grid_not_four_headings():
    _seed_store()
    from docx import Document
    path = _render(_view())
    doc = Document(path)
    swot_tables = [t for t in doc.tables
                  if len(t.rows) == 2 and len(t.rows[0].cells) == 2
                  and "القوة" in t.rows[0].cells[0].text]
    assert swot_tables, "expected a 2x2 SWOT grid table"
    grid = swot_tables[0]
    assert "الضعف" in grid.rows[0].cells[1].text
    assert "الفرص" in grid.rows[1].cells[0].text
    assert "التهديدات" in grid.rows[1].cells[1].text


def test_market_entry_strategy_section_synthesizes_from_existing_numbers():
    # الفصل الجديد: يركّب من HHI + بوابة الأهلية + عدد المرشّحين — لا رقم جديد.
    _seed_store()
    view = _view()
    m = view["markets"][0]
    comp = ((m.get("research") or {}).get("agents") or {}).get("competitor") or {}
    hhi = next((f["value"] for f in comp.get("findings", [])
               if f.get("metric") == "hhi"), None)
    path = _render(view)
    texts = docx_all_text(path)
    assert "استراتيجية دخول السوق" in texts
    idx = texts.find("استراتيجية دخول السوق")
    section = texts[idx:idx + 800]
    assert "النموذج الموصى به" in section
    if hhi is not None:
        assert str(hhi) in section        # نفس رقم HHI المرصود أصلاً، لا رقم جديد


def test_docx_tables_do_not_break_existing_gap_declarations():
    # بلا with_research: كل الأقسام الجدولية تتدهور لفقرة غياب معلنة — لا جدول فارغ.
    import pytest
    pytest.importorskip("docx")
    from docx import Document
    from silk_render import build_view
    from silk_reports import render_docx
    view = build_view({"product": "x", "hs_code": "1", "year": 2023,
                       "classified": True,
                       "markets": [{"country": "c", "iso3": "ARE",
                                    "total_score": 0.1, "confidence": 0.1,
                                    "components": {}}]})
    path = render_docx(view, os.path.join(tempfile.mkdtemp(), "r3.docx"))
    doc = Document(path)
    # لا حزمة بحث => أقسام الجداول (قرار الدخول/حجم السوق/التسعير/الاشتراطات)
    # تتدهور لفقرة غياب قبل بناء أي جدول؛ SWOT وحدها تبقى شبكة ٢×٢ فارغة —
    # نفس سلوكها القديم (كانت تطبع ٤ عناوين "لا بند مرصوداً" بلا حزمة بحث
    # أيضاً)، فقط بشكل شبكة بدل عناوين متتالية.
    assert len(doc.tables) == 1
    grid = doc.tables[0]
    assert all("لا بند مرصوداً" in c.text for row in grid.rows for c in row.cells)
    joined = docx_all_text(path)
    assert "استراتيجية دخول السوق" in joined
    assert "بيانات غير كافية لتركيب استراتيجية دخول" in joined
