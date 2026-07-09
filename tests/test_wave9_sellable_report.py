"""اختبارات الموجة ٩ (V5) — تقرير قابل للبيع + تحليل كامل العمق.

بلاغ المالك على تشغيلة هولندا #٣: المحتوى حقيقي (١١/١٢ بعثة مؤسَّسة) لكن
المستند "غير مقنع، سيّئ التنسيق، مليء بدرجات تبدو بلا سند"، والبيانات
الموجودة لا تُحلَّل بعمق كافٍ. يغطي: P0-1 (تصنيف تقاطعات المحلل)،
P0-2 (تعليمات تعميق البعثات)، P0-3 (أداة OpenAlex)، P1-4 (تراجُع GDELT)،
P1-5 (مسار بطاقة المنتج)، P0-A (تنسيق docx احترافي)، P0-B (شارات الثقة)،
P0-C (برومبت الكاتب كحجّة لا معلومات).
Run:  python3 -m pytest tests/test_wave9_sellable_report.py -q
"""
import json
import os
import sys
import tempfile
from unittest.mock import patch

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from conftest import docx_all_text  # noqa: E402


# ── P0-1: تصنيف تقاطعات المحلل — مطابقة حالة الأحرف/المسافات ─────────────

def test_category_matching_is_case_and_whitespace_tolerant():
    # بلاغ حي: خمس تقاطعات ظهرت "دليل غير كافٍ" رغم أدلة حقيقية — المطابقة
    # الحرفية الصارمة كانت تُسقط أي بند بحالة أحرف مختلفة صمتاً.
    from silk_agents import AgentReport
    from silk_data_layer import DataPoint
    from silk_market_analyst import analyze_market
    from silk_market_resolver import resolve_market

    def fake_call_tools(system, messages, tools=None, max_tokens=1600,
                        model=None, timeout=None):
        return {"stop_reason": "end_turn", "content": [
            {"type": "text", "text": json.dumps({
                "findings": [
                    {"claim": "طلب مقدَّر بحساب واضح",
                     "datapoint_ids": ["dp1"], "confidence": 0.6,
                     "category": "Demand"},          # حرف كبير
                    {"claim": "تكلفة دخول محسوبة",
                     "datapoint_ids": ["dp1"], "confidence": 0.6,
                     "category": " entry_cost "},    # مسافات زائدة
                ],
                "gaps": [], "summary": "تحليل"})}]}

    ref, _ = resolve_market("Netherlands")
    reports = {"trade_flow": AgentReport(
        "LLMAgent:trade_flow", [DataPoint(88000000, "UN Comtrade", 0.9, "n")],
        False, "ok")}
    with patch("silk_llm_runtime._call_tools", side_effect=fake_call_tools):
        out = analyze_market(ref, "تمور", reports, hs_code="080410")
    assert out["by_category"]["demand"], "Demand (مكبَّرة) لم تُصنَّف"
    assert out["by_category"]["entry_cost"], "entry_cost بمسافات لم يُصنَّف"
    assert "demand" not in out["missing_categories"]
    assert "entry_cost" not in out["missing_categories"]


def test_zero_tool_mission_does_not_get_a_premature_finalize_nudge():
    # بلاغ حي: إصلاح الموجة ٨ (جولة إنهاء قسرية) كان يرسل توجيه "لا أدوات
    # متاحة" في الجولة صفر لبعثات بلا أدوات أصلاً (المحلل الشامل) — قبل أي
    # فرصة تحليل فعلية. الآن: التوجيه محصور ببعثات امتلكت أدوات فعلاً.
    import silk_llm_runtime as rt
    from silk_market_resolver import resolve_market

    seen_messages = []

    def fake_call_tools(system, messages, tools=None, max_tokens=1600,
                        model=None, timeout=None):
        seen_messages.append([dict(m) for m in messages])
        return {"stop_reason": "end_turn", "content": [
            {"type": "text", "text": json.dumps(
                {"findings": [], "gaps": [], "summary": "ok"})}]}

    mission = {"key": "market_analyst", "name": "المحلل الشامل",
              "allowed_tools": [], "instructions": "test"}
    ref, _ = resolve_market("Netherlands")
    with patch("silk_llm_runtime._call_tools", side_effect=fake_call_tools):
        rt.run_llm_agent(mission, ref, product="تمور")
    # أول رسالة مُرسَلة لكلود يجب ألا تحوي نص "الإنهاء القسري"
    first_round_messages = seen_messages[0]
    assert not any("لا مزيد من نداءات الأدوات" in str(m.get("content"))
                  for m in first_round_messages)


def test_analyst_prompt_requires_explicit_arithmetic_over_insufficient():
    from silk_market_analyst import _ANALYST_MISSION
    txt = _ANALYST_MISSION["instructions"]
    assert "يُمنَع كتابة" in txt
    assert "المعادلة" in txt or "الحساب الحسابي" in txt


def test_deep_research_missions_get_a_higher_tool_budget():
    from silk_missions import _budget_for, _DEEP_RESEARCH_MISSIONS, _MISSION_BUDGET
    for key in _DEEP_RESEARCH_MISSIONS:
        b = _budget_for(key)
        assert b["tool_calls"] > _MISSION_BUDGET["tool_calls"]
    assert _budget_for("trade_flow") == _MISSION_BUDGET


def test_demand_trends_instructions_require_multi_term_multi_timeframe():
    from silk_missions import MISSIONS
    txt = MISSIONS["demand_trends"]["instructions"]
    assert "5-y" in txt
    assert "12-m" in txt
    assert "رمضان" in txt


def test_research_missions_require_minimum_four_search_angles():
    from silk_missions import MISSIONS
    for key in ("pricing_scout", "consumer_culture", "channels_importers",
               "competitors", "risk_news"):
        assert "أربعة استعلامات" in MISSIONS[key]["instructions"], key


def test_risk_news_has_web_search_fallback_and_openalex():
    from silk_missions import MISSIONS
    tools = MISSIONS["risk_news"]["allowed_tools"]
    assert "web_search" in tools
    assert "openalex_search" in tools
    assert "بديل موثَّق" in MISSIONS["risk_news"]["instructions"]


# ── P0-3: أداة OpenAlex — أدبيات مجانية بديلة لـScopus ────────────────────

def test_openalex_empty_query_returns_none_no_network():
    from silk_openalex_agent import openalex_search
    from conftest import block_network
    with block_network():
        out = openalex_search("")
    assert len(out) == 1 and out[0].value is None and out[0].confidence == 0.0


def test_openalex_network_cut_degrades_to_tagged_none():
    from silk_openalex_agent import openalex_search
    from conftest import block_network
    with block_network():
        out = openalex_search("halal consumer market")
    assert len(out) == 1
    assert out[0].value is None
    assert out[0].confidence == 0.0
    assert out[0].source == "OpenAlex"


def test_openalex_parses_real_shaped_response():
    from silk_openalex_agent import openalex_search

    class _Resp:
        status_code = 200
        def raise_for_status(self):
            pass
        def json(self):
            return {"results": [{
                "title": "Halal food consumption in Western Europe",
                "publication_year": 2021,
                "primary_location": {"source": {"display_name": "Food Policy"}},
                "doi": "https://doi.org/10.1/x",
                "abstract_inverted_index": {"Halal": [0], "market": [1],
                                            "grows": [2]}}]}

    with patch("requests.get", return_value=_Resp()):
        out = openalex_search("halal food Netherlands")
    assert len(out) == 1
    v = out[0].value
    assert v["title"] == "Halal food consumption in Western Europe"
    assert v["year"] == 2021
    assert v["venue"] == "Food Policy"
    assert v["abstract_snippet"] == "Halal market grows"
    assert v["doi"]


def test_openalex_no_results_declares_gap_not_crash():
    from silk_openalex_agent import openalex_search

    class _Resp:
        status_code = 200
        def raise_for_status(self):
            pass
        def json(self):
            return {"results": []}

    with patch("requests.get", return_value=_Resp()):
        out = openalex_search("a very obscure query with zero hits")
    assert len(out) == 1 and out[0].value is None


def test_openalex_wired_as_an_llm_tool():
    from silk_llm_runtime import TOOLS
    assert "openalex_search" in TOOLS
    assert TOOLS["openalex_search"]["spec"]["name"] == "openalex_search"


def test_openalex_allowed_for_the_four_named_missions():
    from silk_missions import MISSIONS
    for key in ("consumer_culture", "demand_trends", "risk_news",
               "opportunity_gaps"):
        assert "openalex_search" in MISSIONS[key]["allowed_tools"], key


# ── P1-4: تراجُع GDELT عند 429 — محاولة واحدة مهذّبة، لا إلحاح ─────────────

def test_gdelt_429_then_success_recovers_on_retry():
    import silk_gdelt_agent as gd

    class _Resp429:
        status_code = 429
        headers = {"Retry-After": "0"}

    class _Resp200:
        status_code = 200
        headers = {"content-type": "application/json"}
        def raise_for_status(self):
            pass
        def json(self):
            return {"articles": [{"title": "خبر", "url": "u", "seendate": "d",
                                  "domain": "x.com"}]}

    calls = {"n": 0}

    def fake_get(url, params=None, headers=None, timeout=None):
        calls["n"] += 1
        return _Resp429() if calls["n"] == 1 else _Resp200()

    with patch("requests.get", side_effect=fake_get), \
         patch("time.sleep"):
        out = gd.gdelt_news("تمور", "Netherlands")
    assert calls["n"] == 2
    assert out[0].value["title"] == "خبر"


def test_gdelt_429_twice_declares_a_distinct_gap_no_retry_spam():
    import silk_gdelt_agent as gd

    class _Resp429:
        status_code = 429
        headers = {}

    calls = {"n": 0}

    def fake_get(url, params=None, headers=None, timeout=None):
        calls["n"] += 1
        return _Resp429()

    with patch("requests.get", side_effect=fake_get), \
         patch("time.sleep") as sleep_mock:
        out = gd.gdelt_news("تمور", "Netherlands")
    assert calls["n"] == 2  # محاولة أصلية + محاولة واحدة فقط، لا أكثر
    assert sleep_mock.call_count == 1
    assert out[0].value is None
    assert "429" in out[0].note
    assert "web_search" in out[0].note


# ── P1-5: مسار بطاقة المنتج — كانت تُجمَع ولا تصل أي بعثة/محلل ────────────

def test_product_card_context_builds_readable_narrative():
    from silk_missions import _product_card_context
    ctx = _product_card_context({
        "cost_per_unit": 2.1, "own_price": 4.5, "tier": "premium",
        "monthly_capacity": 500, "certifications": ["HALAL", "ISO22000"]})
    assert "2.1" in ctx and "4.5" in ctx and "premium" in ctx
    assert "HALAL" in ctx and "ISO22000" in ctx


def test_product_card_context_empty_when_no_card():
    from silk_missions import _product_card_context
    assert _product_card_context(None) == ""
    assert _product_card_context({}) == ""


def test_run_all_missions_threads_product_card_into_mission_context():
    import json as _json
    import silk_missions as sm
    from silk_market_resolver import resolve_market

    seen = []

    def fake_call_tools(system, messages, tools=None, max_tokens=1600,
                        model=None, timeout=None):
        seen.append(str(messages))
        return {"stop_reason": "end_turn", "content": [
            {"type": "text", "text": _json.dumps(
                {"findings": [], "gaps": [], "summary": "ok"})}]}

    ref, _ = resolve_market("Netherlands")
    with patch("silk_llm_runtime._call_tools", side_effect=fake_call_tools):
        sm.run_all_missions(ref, product="تمور", hs_code="080410",
                            product_card={"cost_per_unit": 2.1,
                                         "own_price": 4.5})
    assert any("2.1" in m and "4.5" in m for m in seen)


def test_analyze_market_receives_and_forwards_product_card():
    import json as _json
    from silk_agents import AgentReport
    from silk_market_analyst import analyze_market
    from silk_market_resolver import resolve_market

    seen = []

    def fake_call_tools(system, messages, tools=None, max_tokens=1600,
                        model=None, timeout=None):
        seen.append(str(messages))
        return {"stop_reason": "end_turn", "content": [
            {"type": "text", "text": _json.dumps(
                {"findings": [], "gaps": [], "summary": "ok"})}]}

    ref, _ = resolve_market("Netherlands")
    with patch("silk_llm_runtime._call_tools", side_effect=fake_call_tools):
        analyze_market(ref, "تمور", {}, hs_code="080410",
                       product_card={"cost_per_unit": 2.1, "own_price": 4.5})
    assert any("2.1" in m and "4.5" in m for m in seen)


def test_research_endpoint_forwards_product_card_and_own_price(monkeypatch):
    # إعادة إنتاج البلاغ حرفياً: كان product_card/own_price يُقبَلان في
    # النموذج ولا يصلان deep_research/analyze_market إطلاقاً.
    from fastapi.testclient import TestClient
    import api

    captured = {}
    real_analyze_market = None

    def fake_deep_research(market_ref, product, hs_code=None,
                           product_card=None, **kw):
        captured["deep_research_card"] = product_card
        from silk_agents import AgentReport
        return {"reports": {"trade_flow": AgentReport(
            "LLMAgent:trade_flow", [], True, "no key")},
               "trace_id": "t1"}

    def fake_analyze_market(market_ref, product, reports, hs_code=None,
                            product_card=None, **kw):
        captured["analyst_card"] = product_card
        from silk_agents import AgentReport
        return {"report": AgentReport("A", [], True, ""),
               "by_category": {}, "missing_categories": []}

    with patch("silk_missions.deep_research", side_effect=fake_deep_research), \
         patch("silk_market_analyst.analyze_market",
              side_effect=fake_analyze_market), \
         patch("requests.get", side_effect=OSError("no net")):
        r = TestClient(api.app).post("/research", json={
            "product": "تمور", "market": "Netherlands", "hs_code": "080410",
            "persist": False, "allow_degraded": True,
            "product_card": {"cost_per_unit": 2.1, "tier": "premium"},
            "own_price": 4.5})
    assert r.status_code == 200
    assert captured["deep_research_card"]["cost_per_unit"] == 2.1
    assert captured["deep_research_card"]["own_price"] == 4.5
    assert captured["analyst_card"]["own_price"] == 4.5


def test_product_card_has_certifications_field():
    import inspect
    import api
    src = inspect.getsource(api.create_app)
    idx = src.find("class ProductCard")
    assert "certifications" in src[idx:idx + 400]


# ── P0-A: تنسيق docx احترافي ───────────────────────────────────────────────

def _report_with_table_and_markdown():
    from silk_agents import AgentReport
    from silk_data_layer import DataPoint
    from silk_market_resolver import resolve_market

    ref, _ = resolve_market("Netherlands")
    return {
        "product": "تمور", "hs_code": "080410", "year": None,
        "market": {"iso3": ref.iso3, "m49": ref.m49, "iso2": ref.iso2,
                  "name_en": ref.name_en, "name_ar": ref.name_ar},
        "markets": [],
        "deep_research": {
            "missions": {"trade_flow": AgentReport(
                "LLMAgent:trade_flow",
                [DataPoint("قيمة استيراد مرصودة", "UN Comtrade", 0.9, "n")],
                False, "ok")},
            "analyst": {"report": AgentReport("A", [], True, ""),
                       "by_category": {c: [] for c in (
                           "demand", "entry_cost", "price_competitiveness",
                           "entry_door", "swot")},
                       "missing_categories": []},
            "verdict": {"verdict": "PRELIMINARY GO",
                       "ai": {"verdict": "WATCH — مراقبة قبل الدخول",
                             "confidence": 0.55, "reasoning": "سبب"}},
            "report": {"report": (
                "## 3. الواردات وتدفقات التجارة\n"
                "استيراد هولندا للتمور **نما بثبات** خلال ثلاث سنوات "
                "(المصدر: `UN Comtrade`).\n"
                "| السنة | القيمة (USD) | النمو % |\n"
                "|---|---|---|\n"
                "| 2021 | 80000000 | — |\n"
                "| 2022 | 84000000 | 5.0% |\n"
                "| 2023 | 88000000 | 4.8% |\n"
                "فقرة بعد الجدول تتابع السرد بلا انقطاع."),
                      "review_cycles": 1, "unresolved_notes": []},
        },
    }


def test_markdown_table_becomes_a_real_docx_table(monkeypatch):
    import pytest
    pytest.importorskip("docx")
    from docx import Document
    from silk_render import build_view
    from silk_reports import render_docx
    monkeypatch.setenv("SILK_HERMETIC", "1")
    view = build_view(_report_with_table_and_markdown())
    path = render_docx(view, os.path.join(tempfile.mkdtemp(), "tbl.docx"))
    doc = Document(path)
    header_rows = ["".join(c.text for c in t.rows[0].cells) for t in doc.tables]
    assert any("السنة" in h and "القيمة" in h and "النمو" in h
              for h in header_rows)
    matching = [t for t in doc.tables
               if "السنة" in "".join(c.text for c in t.rows[0].cells)]
    assert matching
    body_rows = ["".join(c.text for c in r.cells) for r in matching[0].rows[1:]]
    assert any("2021" in r and "80000000" in r for r in body_rows)
    text = docx_all_text(path)
    assert "|---|---|---|" not in text
    assert "| 2021 | 80000000 |" not in text  # لا جدول Markdown خام كنص


def test_inline_markdown_stripped_from_paragraphs(monkeypatch):
    import pytest
    pytest.importorskip("docx")
    from silk_render import build_view
    from silk_reports import render_docx
    monkeypatch.setenv("SILK_HERMETIC", "1")
    view = build_view(_report_with_table_and_markdown())
    path = render_docx(view, os.path.join(tempfile.mkdtemp(), "md.docx"))
    text = docx_all_text(path)
    assert "**" not in text
    assert "`UN Comtrade`" not in text
    assert "نما بثبات" in text  # المحتوى نفسه بقي، فقط التنسيق أُزيل


def test_truncate_at_word_never_cuts_mid_word():
    from silk_reports import _truncate_at_word
    long_text = "هذه فقرة طويلة تحوي كلمات متعددة لاختبار القصّ الصحيح دون تشويه أي كلمة منتصفها إطلاقاً"
    cut = _truncate_at_word(long_text, 40)
    assert cut.endswith("…")
    body = cut[:-1].rstrip()
    assert long_text.startswith(body)
    # الحرف التالي مباشرة بعد القصّ يجب أن يكون بداية كلمة (مسافة أو نهاية)
    next_idx = len(body)
    assert next_idx == len(long_text) or long_text[next_idx] == " "


def test_llm_runtime_citation_note_truncation_is_word_safe():
    from silk_llm_runtime import _truncate_at_word
    text = "مبني على: " + ("كلمة " * 200)
    cut = _truncate_at_word(text, 100)
    assert cut.endswith("…")
    assert not cut[:-1].rstrip().endswith("كلم")  # لا قطع منتصف الكلمة


def test_cover_page_has_verdict_badge_and_branding(monkeypatch):
    import pytest
    pytest.importorskip("docx")
    from silk_render import build_view
    from silk_reports import render_docx
    monkeypatch.setenv("SILK_HERMETIC", "1")
    view = build_view(_report_with_table_and_markdown())
    path = render_docx(view, os.path.join(tempfile.mkdtemp(), "cover.docx"))
    text = docx_all_text(path)
    assert "أُعد بواسطة منصة سِلك لذكاء الأسواق" in text
    assert "WATCH — مراقبة" in text
    idx = text.find("سِلك — تقرير بحث عميق")
    assert idx != -1
    # الشارة قريبة من الغلاف (أول ٥٠٠ حرف) لا مدفونة في آخر المستند.
    assert "WATCH" in text[idx:idx + 500]


def test_verdict_tone_matches_dashboard_logic():
    from silk_reports import _verdict_tone
    assert _verdict_tone("NO-GO — غير موصى به") == "nogo"
    assert _verdict_tone("WATCH — مراقبة") == "watch"
    assert _verdict_tone("PRELIMINARY GO") == "go"
    assert _verdict_tone("") == "unknown"


# ── P0-B: شارات الأدلة بدل أرقام الثقة الخام ──────────────────────────────

def test_evidence_badge_thresholds():
    from silk_reports import _evidence_badge
    assert _evidence_badge(0.9) == "✓ موثّق"
    assert _evidence_badge(0.8) == "✓ موثّق"
    assert _evidence_badge(0.79) == "◐ ثانوي"
    assert _evidence_badge(0.5) == "◐ ثانوي"
    assert _evidence_badge(0.49) == "○ غير متحقق"
    assert _evidence_badge(0.0) == "○ غير متحقق"
    assert _evidence_badge(None) == "○ غير متحقق"


def _deep_research_result_for_badges():
    from silk_agents import AgentReport
    from silk_data_layer import DataPoint
    from silk_market_resolver import resolve_market

    ref, _ = resolve_market("Netherlands")
    demand_finding = DataPoint("قطاع مسلم 1.27M × استيراد 88M", "x", 0.85, "n")
    analyst_report = AgentReport("LLMAgent:market_analyst",
                                 [demand_finding], False, "تحليل")
    return {
        "product": "تمور", "hs_code": "080410", "year": None,
        "market": {"iso3": ref.iso3, "m49": ref.m49, "iso2": ref.iso2,
                  "name_en": ref.name_en, "name_ar": ref.name_ar},
        "markets": [],
        "deep_research": {
            "missions": {"trade_flow": AgentReport(
                "LLMAgent:trade_flow",
                [DataPoint("رقم مستشهَد", "UN Comtrade", 0.6, "n")],
                False, "ok")},
            "analyst": {"report": analyst_report,
                       "by_category": {"demand": [demand_finding],
                                      "entry_cost": [], "price_competitiveness": [],
                                      "entry_door": [], "swot": []},
                       "missing_categories": []},
            "verdict": {"verdict": "PRELIMINARY GO",
                       "ai": {"verdict": "WATCH", "confidence": 0.5,
                             "reasoning": "سبب"}},
            "report": {"report": (
                "## 13. الحكم والتوصية\n"
                "الحكم WATCH مبني على أدلة (المصدر: UN Comtrade).\n"
                "**ماذا يعني هذا لقرارك:** راقب قبل الدخول الكامل."),
                      "review_cycles": 1, "unresolved_notes": []},
        },
    }


def test_no_raw_confidence_number_in_rendered_docx_body(monkeypatch):
    import pytest
    pytest.importorskip("docx")
    from silk_render import build_view
    from silk_reports import render_docx
    monkeypatch.setenv("SILK_HERMETIC", "1")
    view = build_view(_deep_research_result_for_badges())
    path = render_docx(view, os.path.join(tempfile.mkdtemp(), "badge.docx"))
    text = docx_all_text(path)
    assert "ثقة 0" not in text
    assert "✓ موثّق" in text  # المحلل: 0.85 >= 0.8


def test_technical_appendix_carries_full_confidence_numbers(monkeypatch):
    import pytest
    pytest.importorskip("docx")
    from docx import Document
    from silk_render import build_view
    from silk_reports import render_docx
    monkeypatch.setenv("SILK_HERMETIC", "1")
    view = build_view(_deep_research_result_for_badges())
    path = render_docx(view, os.path.join(tempfile.mkdtemp(), "appx.docx"))
    doc = Document(path)
    header_rows = ["".join(c.text for c in t.rows[0].cells) for t in doc.tables]
    assert any("الادّعاء" in h and "الثقة" in h for h in header_rows)
    matching = [t for t in doc.tables
               if "الادّعاء" in "".join(c.text for c in t.rows[0].cells)]
    body = ["".join(c.text for c in r.cells) for r in matching[0].rows[1:]]
    assert any("0.6" in r for r in body)


def test_takeaway_line_renders_bold_and_detected_without_asterisks(
        monkeypatch):
    import pytest
    pytest.importorskip("docx")
    from docx import Document
    from silk_render import build_view
    from silk_reports import render_docx
    monkeypatch.setenv("SILK_HERMETIC", "1")
    view = build_view(_deep_research_result_for_badges())
    path = render_docx(view, os.path.join(tempfile.mkdtemp(), "take.docx"))
    doc = Document(path)
    hit = [p for p in doc.paragraphs
          if p.text.startswith("ماذا يعني هذا لقرارك:")]
    assert hit, "takeaway paragraph not found"
    assert "**" not in hit[0].text
    assert hit[0].runs and hit[0].runs[0].bold


def test_writer_prompt_forbids_raw_confidence_in_prose():
    import silk_ai_judge
    import inspect
    src = inspect.getsource(silk_ai_judge.deep_report)
    assert "لا تكتب رقم ثقة خاماً" in src
    assert "شارة أدلة" in src


def test_writer_prompt_requires_thesis_and_roadmap_and_takeaway():
    import silk_ai_judge
    import inspect
    src = inspect.getsource(silk_ai_judge.deep_report)
    assert "أطروحة" in src
    assert "خارطة طريق الدخول" in src
    assert "ماذا يعني هذا لقرارك" in src
    # الموجة ١٠: الخارطة أصبحت فرعاً داخل "التوصيات الاستراتيجية" (البنية
    # الدولية بأحد عشر قسماً) لا قسماً مستقلاً — راجع test_wave10_*.
    assert "التوصيات الاستراتيجية" in silk_ai_judge._REPORT_SECTIONS
    assert len(silk_ai_judge._REPORT_SECTIONS) == 11


# P0-C (تكملة): _extract_json الآمن للسياج + قاعدة المراجع الموسّعة.
# بلاغ حي (الموجة ٨ ثم ٩): سياج ```json + تعليق ختامي بعده كان يُفسد
# rfind('}') الساذج في خمسة مواضع من silk_ai_judge.py فيُسقط الرد بأكمله.


def test_extract_json_handles_fenced_reply_with_trailing_prose():
    import silk_ai_judge
    raw = ('مرحباً، إليك الناتج:\n```json\n{"issues": ["a"], "approved": '
          'false}\n```\nملاحظة: هذا تحليل أولي وقد يحتاج مراجعة} إضافية.')
    obj = silk_ai_judge._extract_json(raw)
    assert obj == {"issues": ["a"], "approved": False}


def test_extract_json_falls_back_to_whole_text_without_fence():
    import silk_ai_judge
    raw = '{"insights": [{"point": "x", "evidence": [1]}], "note": "n"}'
    obj = silk_ai_judge._extract_json(raw)
    assert obj["insights"][0]["point"] == "x"


def test_extract_json_returns_none_never_fabricates_empty_object():
    import silk_ai_judge
    assert silk_ai_judge._extract_json("") is None
    assert silk_ai_judge._extract_json(None) is None
    assert silk_ai_judge._extract_json("ليس JSON إطلاقاً، لا أقواس هنا.") is None


def test_review_report_uses_extract_json(monkeypatch):
    import silk_ai_judge
    monkeypatch.setattr(silk_ai_judge, "available", lambda: True)
    raw = ('```json\n{"issues": [], "approved": true}\n```\n'
          'تعليق ختامي يحوي قوساً } زائداً.')
    monkeypatch.setattr(silk_ai_judge, "_call", lambda *a, **k: raw)
    # مسوّدة كاملة الأقسام حتى لا يتدخّل الفحص البنيوي الحتمي (الموجة ١٠) —
    # هذا الاختبار عن تفسير JSON المسيَّج تحديداً، لا اكتمال البنية.
    draft = "\n".join(f"## {i}. {s}\nنص." for i, s in
                      enumerate(silk_ai_judge._REPORT_SECTIONS, 1))
    result = silk_ai_judge.review_report(draft, {})
    assert result == {"issues": [], "approved": True}


def test_reviewer_rubric_checks_thesis_roadmap_takeaway_and_arithmetic():
    import silk_ai_judge
    import inspect
    src = inspect.getsource(silk_ai_judge.review_report)
    assert "أطروحة" in src
    assert "خارطة طريق الدخول" in src
    assert "ماذا يعني هذا لقرارك" in src
    assert "حساب حسابي صريح" in src


# قاعدة ١٠.٦ (بوابة تسليم الموجة ٩): نموذج DOCX فعلي محفوظ بالمستودع، وليس
# مجرد وصف — يُراجَع من الملف نفسه لا عبر قناة إرفاق. مولَّد عبر
# tools/gen_research_sample.py (نتيجة هولندا×تمور مموّهة بكامل التنسيق).


def test_research_sample_docx_meets_wave9_delivery_gate():
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    path = os.path.join(root, "samples", "research_report_latest.docx")
    assert os.path.exists(path), "شغّل tools/gen_research_sample.py"
    text = docx_all_text(path)
    assert "GO" in text  # شارة حكم على الغلاف
    assert "خارطة طريق الدخول" in text  # قسم ٩٠ يوماً
    assert "ماذا يعني هذا لقرارك" in text  # سطور الخلاصة
    assert "##" not in text and "**" not in text and "```" not in text
    assert "ثقة 0" not in text  # لا رقم ثقة خام مسرَّب للسرد
    assert "ملحق تقني" in text  # الأرقام الكاملة انتقلت للملحق
    assert "منصة سِلك لذكاء الأسواق" in text  # سطر الهوية على الغلاف
