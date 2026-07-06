"""اختبارات الموجة ٩ — إثراء القالب الموحّد + تقرير الورد الكامل.

يقفل:
1. build_view يحمل الأسعار/المنافسين/الموردين/الثقافة (من الطبقات المرصودة)،
   ويعلن [] عند الغياب (لا اختلاق).
2. render_docx يضيف أقسام: أسعار السوق، المنافسون، الموردون، الاتجاه، ثقافة
   المستهلك — ويعرض «غير مرصود» صراحةً عند غياب البيانات (تقرير لا «ناقص»).
Run:  python3 -m pytest tests/ -q
"""
import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from silk_data_layer import DataPoint
from silk_render import build_view


def _dp(v):
    return DataPoint(v, "src", 0.9 if v is not None else 0.0, "n")


def _result_with_enrichment():
    return {
        "product": "تمور", "hs_code": "080410", "hs_confidence": 0.9,
        "year": 2023, "classified": True,
        "websearch": [_dp({"title": "التمور رائجة في رمضان", "link": "http://x"})],
        "markets": [{
            "country": "الإمارات", "iso3": "ARE", "total_score": 0.7,
            "confidence": 0.7,
            "components": {"market_size": _dp(4.0), "saudi_position": _dp(34.0),
                           "demand_capacity": _dp(78000.0), "competition": _dp(None)},
            "competitors": [{"partner": "إيران", "share": 31, "value_usd": 1.2e8}],
            "competitors_named": [_dp({"title": "Bateel"})],
            "localprice": [_dp({"title": "Bateel Dates", "price": 120,
                                "currency": "AED", "store": "Noon"})],
            "maps": [_dp({"name": "Al Foah"})],
        }],
    }


def test_build_view_carries_enrichment():
    view = build_view(_result_with_enrichment())
    m0 = view["markets"][0]
    assert m0["prices"] and m0["prices"][0]["price"] == 120
    assert m0["named_competitors"] == ["Bateel"]
    assert m0["supplier_countries"][0]["partner"] == "إيران"
    assert any(s["name"] == "Al Foah" for s in m0["suppliers"])
    assert view["culture"] and "رمضان" in view["culture"][0]["title"]


def test_build_view_declares_empty_enrichment_when_absent():
    # بلا طبقات إثراء => قوائم فارغة، لا خطأ ولا اختلاق.
    view = build_view({"product": "x", "hs_code": "1", "year": 2023,
                       "classified": True,
                       "markets": [{"country": "c", "iso3": "ARE",
                                    "total_score": 0.1, "confidence": 0.1,
                                    "components": {}}]})
    m0 = view["markets"][0]
    assert m0["prices"] == [] and m0["named_competitors"] == []
    assert m0["suppliers"] == [] and view["culture"] == []


def test_render_docx_includes_new_sections():
    import pytest
    docx = pytest.importorskip("docx")  # noqa: F841
    from silk_reports import render_docx
    view = build_view(_result_with_enrichment())
    path = render_docx(view, os.path.join(tempfile.mkdtemp(), "r.docx"))
    from docx import Document
    texts = [p.text for p in Document(path).paragraphs]
    joined = "\n".join(texts)
    # الأقسام الجديدة حاضرة كعناوين
    for heading in ("أسعار المنتجات في السوق", "المنافسون",
                    "الموردون والأعمال بالاسم", "اتجاه الاستيراد متعدد السنوات",
                    "ثقافة المستهلك ونبض السوق"):
        assert heading in texts, heading
    # وبياناتها المرصودة ظهرت
    assert "Bateel" in joined and "إيران" in joined and "رمضان" in joined


def test_render_docx_declares_gaps_when_enrichment_missing():
    import pytest
    pytest.importorskip("docx")
    from silk_reports import render_docx
    from docx import Document
    view = build_view({"product": "x", "hs_code": "1", "year": 2023,
                       "classified": True,
                       "markets": [{"country": "c", "iso3": "ARE",
                                    "total_score": 0.1, "confidence": 0.1,
                                    "components": {}}]})
    path = render_docx(view, os.path.join(tempfile.mkdtemp(), "r2.docx"))
    joined = "\n".join(p.text for p in Document(path).paragraphs)
    assert "أسعار المنتجات في السوق" in joined      # القسم موجود
    # عقد 2B الجديد: دون العتبة يُعلن «بيانات غير كافية» + المصادر المُحاوَلة —
    # جملة النقص الوحيدة المسموح بها، لا نثر حشو (بديل «غير مرصود» القديم).
    assert "بيانات غير كافية" in joined
    assert "المصادر المُحاوَلة" in joined
