"""اختبارات دخان بلا شبكة — تتحقق من الاستيراد، التصنيف، والمبدأ التأسيسي (لا اختلاق بيانات).
Offline smoke tests: imports, HS classification, and the no-fabrication principle.
Run:  python3 -m pytest tests/ -q   (or)  python3 tests/test_smoke.py
"""
import contextlib
import os
import socket
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import silk_hs_resolver as resolver
import silk_engine as engine


@contextlib.contextmanager
def _block_network():
    """اقطع الشبكة مؤقتًا — force outbound sockets to fail so 'no data => 0.0'
    holds even where the CI has internet. Restores socket.socket on exit."""
    real = socket.socket

    def _no_net(*a, **k):  # noqa: ANN002, ANN003
        raise OSError("network disabled for hermetic test")

    socket.socket = _no_net
    try:
        yield
    finally:
        socket.socket = real


def test_all_modules_import():
    import silk_data_layer, silk_data_layer_v2, silk_agents, silk_market_ranker  # noqa: F401


def test_resolver_real_hs_codes():
    assert resolver.resolve("تمور").value == "080410"
    assert resolver.resolve("saffron").value == "091020"
    # كلمة بلا معنى => لا تصنيف ولا اختلاق رمز
    miss = resolver.resolve("xqzwv nonsense 123")
    assert miss.value is None and miss.confidence == 0.0


def test_engine_pipeline_offline_no_fabrication():
    # بلا شبكة: المحرّك يصنّف المنتج لكن لا يخترع أرقام أسواق
    with _block_network():
        res = engine.analyze("تمور", countries=[{"iso3": "ARE", "m49": "784"}], year=2022)
    assert res["classified"] is True
    assert res["hs_code"] == "080410"
    assert res["preliminary"] is True
    row = res["markets"][0]
    assert row["total_score"] == 0.0 and row["confidence"] == 0.0  # no data => no invented score


def test_storage_round_trip(tmp_path=None):
    # خزّن نتيجة وهمية ثم استرجعها — save a fake result, then get it back unchanged.
    import os
    import tempfile
    import silk_storage as storage

    db = os.path.join(tempfile.mkdtemp(), "smoke.db")
    fake = {"product": "demo", "hs_code": "000000", "year": 2022,
            "preliminary": True,
            "markets": [{"country": "Demo-Land", "iso3": "XXX",
                         "total_score": 0.0, "confidence": 0.0}]}
    aid = storage.save_analysis(fake, db)
    got = storage.get_analysis(aid, db)
    assert got is not None and got["product"] == "demo"
    assert any(r["id"] == aid for r in storage.list_analyses(db))


def test_quality_flags_near_zero():
    # صف بحجم سوق شبه صفري => تنبيه عدم تطابق — near-zero size flags a mismatch.
    import silk_quality as quality

    row = {"country": "X", "iso3": "XXX",
           "components": {"market_size": {"value": 12.0},
                          "saudi_position": {"value": 3.0},
                          "demand_capacity": {"value": 1.0e9},
                          "competition": {"value": 0.4}}}
    flags = quality.validate_market_row(row)
    assert any("near-zero" in f for f in flags)


def test_engine_optional_layers_offline():
    # كل الطبقات الاختيارية مفعّلة بلا شبكة: لا تعطّل، يبقى التصنيف سليمًا.
    import os
    import tempfile

    db = os.path.join(tempfile.mkdtemp(), "engine.db")
    with _block_network():
        res = engine.analyze("تمور", countries=[{"iso3": "ARE", "m49": "784"}],
                             year=2022, with_trends=True, with_tariffs=True,
                             persist=True, db_path=db)
    assert res["classified"] is True and res["hs_code"] == "080410"
    assert "quality_flags" in res["markets"][0]      # quality on by default
    assert "analysis_id" in res                       # persisted
    # طبقات السياق مرفقة (قيم None بلا شبكة، لا اختلاق) — context attached, None offline.
    assert "trends" in res["markets"][0] and "tariff" in res["markets"][0]
    assert res["markets"][0]["total_score"] == 0.0    # additive context, score unchanged


def test_api_imports_without_fastapi():
    # api.py يُستورد بلا fastapi — import works offline; app may be None.
    import api
    assert hasattr(api, "create_app") and hasattr(api, "app")


def _test_session_headers(email: str) -> dict:
    """جلسة اختبار — mint a session token directly (bypassing email transport)."""
    import hashlib
    import secrets

    import silk_auth
    import silk_db

    token = secrets.token_urlsafe(32)
    silk_db.store_magic_link(email, hashlib.sha256(token.encode()).hexdigest())
    verified = silk_auth.verify_magic_link(token)
    return {"Authorization": f"Bearer {verified['session_token']}"}


def test_api_analyze_endpoint_own_price_offline():
    # عبر TestClient الفعلي: /analyze (مصادَق) بلا شبكة يرجّع price_comparison بلا اختلاق.
    import pytest
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")  # TestClient needs it; test-only dep, not a runtime one
    from unittest.mock import patch
    from fastapi.testclient import TestClient
    import api

    assert api.app is not None
    client = TestClient(api.app)
    headers = _test_session_headers("own-price-test@example.com")
    # لا نستخدم _block_network هنا: تعطّل socket.socket عالمياً يكسر نقل TestClient
    # الداخلي (asyncio socketpair)؛ بدلاً منها نعطّل requests.get فقط — نفس الأثر
    # الحتمي (لا شبكة => لا بيانات) بلا التصادم مع بنية الاختبار التحتية.
    with patch("requests.get", side_effect=OSError("network disabled for hermetic test")):
        r = client.post("/analyze", json={
            "product": "تمور", "year": 2022,
            "with_localprice": True, "own_price": 25.0,
        }, headers=headers)
    assert r.status_code == 200
    job = r.json()
    assert job["status"] == "finished"  # no REDIS_URL -> synchronous fallback

    jr = client.get(f"/jobs/{job['job_id']}", headers=headers)
    assert jr.status_code == 200
    data = jr.json()["result"]
    assert data["classified"] is True and data["hs_code"] == "080410"
    row = data["markets"][0]
    assert "price_comparison" in row
    assert row["price_comparison"]["your_price"] == 25.0
    assert row["price_comparison"]["listings_count"] == 0     # no network -> no listings
    assert row["price_comparison"]["cheaper_than_pct"] is None  # never fabricated


def test_api_analyze_requires_auth():
    # /analyze بلا جلسة => 401؛ لا تحليلات مدفوعة بلا مصادقة.
    import pytest
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    import api

    client = TestClient(api.app)
    r = client.post("/analyze", json={"product": "تمور"})
    assert r.status_code == 401


def test_api_jobs_ownership_isolation():
    # لا يقدر مستخدم يقرأ مهمة مستخدم آخر — job ids aren't guessable but ownership
    # is still enforced explicitly.
    import pytest
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from unittest.mock import patch
    from fastapi.testclient import TestClient
    import api

    client = TestClient(api.app)
    owner_headers = _test_session_headers("owner@example.com")
    other_headers = _test_session_headers("other@example.com")
    with patch("requests.get", side_effect=OSError("network disabled for hermetic test")):
        r = client.post("/analyze", json={"product": "تمور",
                                          "year": 2022}, headers=owner_headers)
    job_id = r.json()["job_id"]
    assert client.get(f"/jobs/{job_id}", headers=other_headers).status_code == 404
    assert client.get(f"/jobs/{job_id}", headers=owner_headers).status_code == 200


def test_faostat_agent_imports():
    # وكيل فاوستات يُستورد بلا شبكة — offline import + graceful None on unknown area.
    import silk_faostat_agent as fao
    dp = fao.per_capita_supply("XXX", "Dates")
    assert dp.value is None and dp.confidence == 0.0  # no fabrication


def test_cache_returns_none_offline():
    # ذاكرة التخزين تُعيد None بلا شبكة — cached_get degrades to None offline.
    import silk_cache
    with _block_network():
        out = silk_cache.cached_get("https://example.invalid/none", {"a": "1"})
    assert out is None


def test_new_agents_import_and_no_fabrication_keyless():
    # الوكلاء الأربعة الجدد يُستوردون بلا شبكة/مفتاح، وكل نداء بلا مفتاح => value=None.
    import silk_maps_agent, silk_websearch_agent, silk_volza_agent, silk_explee_agent

    with _block_network():
        for key in ("GOOGLE_MAPS_API_KEY", "SEARCH_API_KEY",
                    "VOLZA_API_KEY", "EXPLEE_API_KEY"):
            os.environ.pop(key, None)  # ensure keyless
        dps = [
            silk_maps_agent.find_places("dates morocco")[0],
            silk_websearch_agent.web_search("dates demand")[0],
            silk_volza_agent.importers_by_name("080410", "156")[0],
            silk_explee_agent.discover_buyers("dates packaging", "DEU")[0],
        ]
    for dp in dps:
        assert dp.value is None and dp.confidence == 0.0  # no fabrication keyless


def test_comtrade_endpoint_switches_with_key():
    # بلا مفتاح -> معاينة محدودة؛ مع مفتاح -> endpoint الإنتاج الكامل /data/v1/get.
    import silk_data_layer as d

    saved = d.COMTRADE_KEY
    try:
        d.COMTRADE_KEY = ""
        assert d._comtrade_url().endswith("/public/v1/preview/C/A/HS")
        d.COMTRADE_KEY = "SAMPLEKEY"
        assert d._comtrade_url().endswith("/data/v1/get/C/A/HS")
    finally:
        d.COMTRADE_KEY = saved


def test_market_imports_one_call_size_and_competitors():
    # نداء Comtrade واحد يعطي حجم السوق (صفّ العالم) والمنافسين معًا — no 2nd call.
    import silk_data_layer_v2 as v2

    fake = [
        {"partnerCode": 0, "primaryValue": 241000000.0},   # World = market size
        {"partnerCode": 788, "primaryValue": 75700000.0},  # Tunisia
        {"partnerCode": 682, "primaryValue": 28900000.0},  # Saudi
    ]
    orig = v2.comtrade_trade
    v2.comtrade_trade = lambda *a, **k: [dict(r) for r in fake]
    try:
        mi = v2.market_imports("080410", "504", 2023)
    finally:
        v2.comtrade_trade = orig
    assert mi["total_usd"] == 241000000.0            # World row -> market size
    assert len(mi["competitors"]) == 2               # partners only (World dropped)
    assert mi["competitors"][0].value["partner"]     # named, ranked desc
    shares = sum(c.value["share"] for c in mi["competitors"])
    assert 99.0 <= shares <= 101.0                   # shares ~100% of suppliers


def test_localprice_agent_no_fabrication_keyless():
    # وكيل أسعار السوق المحلي يُستورد بلا شبكة/مفتاح، وكل نداء بلا مفتاح => value=None.
    import silk_localprice_agent as lp

    os.environ.pop("LOCALPRICE_API_KEY", None)
    with _block_network():
        rep = lp.LocalPriceAgent().run({"query": "تمور", "market": "ma"})
    assert rep.failed is True
    dp = rep.findings[0]
    assert dp.value is None and dp.confidence == 0.0  # no fabricated price


def test_engine_localprice_layer_offline():
    # طبقة السعر المحلي مفعّلة بلا شبكة/مفتاح: لا تعطّل، تبقى النتيجة مبدئية بلا اختلاق.
    os.environ.pop("LOCALPRICE_API_KEY", None)
    with _block_network():
        res = engine.analyze("تمور", countries=[{"iso3": "ARE", "m49": "784"}],
                             year=2023, with_localprice=True)
    assert res["classified"] is True and res["year"] == 2023
    assert "localprice" in res["markets"][0]            # context attached
    assert res["markets"][0]["total_score"] == 0.0      # additive, score unchanged


def test_engine_localprice_own_price_offline():
    # own_price مفعّل بلا شبكة: price_comparison مرفق لكن بلا اختلاق (لا قوائم).
    os.environ.pop("LOCALPRICE_API_KEY", None)
    with _block_network():
        res = engine.analyze("تمور", countries=[{"iso3": "ARE", "m49": "784"}],
                             year=2023, with_localprice=True, own_price=25.0)
    row = res["markets"][0]
    pc = row["price_comparison"]
    assert pc["your_price"] == 25.0 and pc["listings_count"] == 0
    assert pc["cheaper_than_pct"] is None and pc["market_avg"] is None  # no fabrication


def test_compare_own_price_no_fabrication_without_listings():
    import silk_localprice_agent as lp

    out = lp.compare_own_price(25.0, [])
    assert out["listings_count"] == 0
    assert out["market_avg"] is None and out["cheaper_than_pct"] is None
    assert out["verdict"] is None


def test_compare_own_price_no_own_price_given():
    import silk_localprice_agent as lp
    from silk_data_layer import DataPoint

    findings = [DataPoint({"price": 100}, "Local retail", 0.6, ""),
                DataPoint({"price": 80}, "Local retail", 0.6, "")]
    out = lp.compare_own_price(None, findings)
    assert out["your_price"] is None and out["listings_count"] == 2
    assert out["market_min"] == 80 and out["market_max"] == 100
    assert out["cheaper_than_pct"] is None and out["verdict"] is None  # not guessed


def test_compare_own_price_percentile_math():
    import silk_localprice_agent as lp
    from silk_data_layer import DataPoint

    findings = [DataPoint({"price": 100}, "Local retail", 0.6, ""),
                DataPoint({"price": 80}, "Local retail", 0.6, ""),
                DataPoint({"price": 60}, "Local retail", 0.6, "")]
    out = lp.compare_own_price(70.0, findings)
    assert out["listings_count"] == 3
    assert out["market_min"] == 60 and out["market_max"] == 100
    assert out["market_avg"] == 80.0
    # سعرك 70 أرخص من قائمتين (80، 100) من أصل 3 => 66.7%
    assert out["cheaper_than_pct"] == round(200 / 3, 1)
    assert "66.7" in out["verdict"]


def test_localprice_bestseller_badge_real_only():
    # الشارة تُقرأ فقط من ردّ المزوّد الحقيقي — لا تُخمَّن من السعر/الترتيب.
    import silk_localprice_agent as lp

    payload = {"shopping_results": [
        {"title": "A", "price": 10, "tag": "Best Seller"},
        {"title": "B", "price": 12, "extensions": ["Free shipping", "Bestseller"]},
        {"title": "C", "price": 8, "bestseller": True},
        {"title": "D", "price": 15},
    ]}
    listings = lp._extract(payload)
    flags = {it["title"]: it["is_best_seller"] for it in listings}
    assert flags == {"A": True, "B": True, "C": True, "D": False}


def test_engine_paid_layers_offline():
    # الطبقات الأربع الجديدة مفعّلة بلا شبكة/مفتاح: لا تعطّل، يبقى التصنيف سليمًا.
    with _block_network():
        for key in ("GOOGLE_MAPS_API_KEY", "SEARCH_API_KEY",
                    "VOLZA_API_KEY", "EXPLEE_API_KEY"):
            os.environ.pop(key, None)
        res = engine.analyze("تمور", countries=[{"iso3": "ARE", "m49": "784"}],
                             year=2022, with_maps=True, with_websearch=True,
                             with_volza=True, with_explee=True)
    assert res["classified"] is True and res["hs_code"] == "080410"
    row = res["markets"][0]
    # طبقات السياق مرفقة (None بلا مفتاح/شبكة، لا اختلاق) — attached, None offline.
    assert "maps" in row and "volza" in row and "explee" in row
    assert "websearch" in res                            # top-level web search
    assert row["total_score"] == 0.0                     # additive, score unchanged


def test_hs_codes_grew_and_resolve_dates():
    # نمت رموز HS وما زال التمر يُصنّف صحيحًا — table grew; dates still resolve.
    assert len(resolver.load_hs_codes()) >= 157
    assert resolver.resolve("تمور").value == "080410"


def test_rank_markets_has_dashboard_fields():
    # حقول لوحة المعلومات الإضافية موجودة (قد تكون None بلا شبكة، لا اختلاق).
    import silk_market_ranker as ranker

    with _block_network():
        ranked = ranker.rank_markets("080410",
                                     countries=[{"iso3": "ARE", "m49": "784"}],
                                     year=2022)
    row = ranked[0]
    for key in ("income_ppp", "population", "top_competitor"):
        assert key in row  # additive key present (value may be None offline)


def test_index_helper_matches_dates():
    # مساعد /index يطابق "تمور" ويرجع رمز التمر 080410 — offline, no network.
    import api

    out = api._index_search("تمور", limit=20)
    assert any(item["hs"] == "080410" for item in out)
    assert out and set(out[0].keys()) == {"name", "hs", "analyzed"}


def test_db_fallback_to_sqlite_without_database_url():
    # بلا DATABASE_URL: يعمل silk_db بـ SQLite محلياً (لا حاجة Postgres حيّ).
    import silk_db

    os.environ.pop("DATABASE_URL", None)
    assert silk_db._database_url().startswith("sqlite:///")
    uid = silk_db.get_or_create_user("db-fallback-test@example.com")
    assert silk_db.get_or_create_user("db-fallback-test@example.com") == uid  # idempotent


def test_auth_magic_link_single_use_and_expiry():
    # الرابط يعمل مرة واحدة فقط، ورمز غير صالح لا يُصدر جلسة أبداً.
    import hashlib
    import secrets

    import silk_auth
    import silk_db

    token = secrets.token_urlsafe(32)
    silk_db.store_magic_link("auth-smoke@example.com", hashlib.sha256(token.encode()).hexdigest())

    out = silk_auth.verify_magic_link(token)
    assert out is not None and out["email"] == "auth-smoke@example.com"
    assert silk_auth.verify_magic_link(token) is None       # replay rejected
    assert silk_auth.verify_magic_link("not-a-real-token") is None  # never guesses
    assert silk_auth.session_user_id(out["session_token"]) == out["user_id"]
    assert silk_auth.session_user_id("garbage") is None


def test_auth_dev_fallback_logs_without_smtp():
    # بلا SMTP_HOST: الرابط يُسجَّل فقط (dev)، sent=False — لا محاولة إرسال صامتة.
    import silk_auth

    os.environ.pop("SMTP_HOST", None)
    out = silk_auth.request_magic_link("no-smtp@example.com", "http://localhost:8000")
    assert out["sent"] is False


def test_ratelimit_blocks_after_cap():
    # يسمح بالحد بالضبط ثم يمنع — in-memory fallback (no REDIS_URL).
    import importlib

    import silk_ratelimit as rl
    importlib.reload(rl)  # fresh in-memory counters for this identity/test run
    identity = "ratelimit-test-user"
    for _ in range(rl._PER_HOUR):
        rl.enforce_analysis_limits(identity)
    try:
        rl.enforce_analysis_limits(identity)
        assert False, "expected RateLimitExceeded"
    except rl.RateLimitExceeded as e:
        assert e.scope == "hour" and e.limit == rl._PER_HOUR


def test_jobs_cache_hit_skips_engine_call():
    # نتيجة مخزّنة => لا نداء لـ silk_engine.analyze إطلاقاً (لا وكلاء ولا كلود).
    from unittest.mock import patch

    import silk_jobs

    request = {"product": "cache-hit-demo", "year": 2022,
              "countries": [{"iso3": "ARE", "m49": "784"}]}
    with patch("silk_cache.get_cached_analysis", return_value={"cached": True}):
        with patch("silk_engine.analyze") as mocked_analyze:
            out = silk_jobs.enqueue_analysis(request, user_id=None)
            mocked_analyze.assert_not_called()
    assert out["cached"] is True and out["status"] == "finished"
    status = silk_jobs.job_status(out["job_id"])
    assert status["result"] == {"cached": True}


def test_production_agent_offline_no_fabrication():
    # وكيل الإنتاج (المجموعة أ) بلا شبكة: لا رقم مُختلق، تقرير فاشل موسوم.
    import silk_production_agent as prod

    os.environ.pop("SEARCH_API_KEY", None)
    with _block_network():
        rep = prod.ProductionAgent().run({"iso3": "ARE", "product": "Dates",
                                          "year": 2022})
    assert rep.failed is True
    assert all(f.value is None for f in rep.findings)  # no fabricated production


def test_production_agent_unknown_area_no_network():
    # ISO3 غير معروف لفاوستات => None فوراً بلا محاولة شبكة، بلا اختلاق.
    import silk_production_agent as prod

    dp = prod.faostat_production("XXX", "Dates", 2022)
    assert dp.value is None and dp.confidence == 0.0


def test_marketsize_apparent_consumption_math():
    # الاستهلاك الظاهري = إنتاج + استيراد − تصدير (طن)، من نداءات وهمية.
    from unittest.mock import patch

    import silk_marketsize_agent as ms
    from silk_data_layer import DataPoint

    fake_imports = [{"partnerCode": 0, "netWgt": 50_000_000.0}]   # 50,000 t
    fake_exports = [{"partnerCode": 0, "netWgt": 10_000_000.0}]   # 10,000 t

    def fake_comtrade(hs, m49, year, flow="M", partner=0):
        return fake_imports if flow == "M" else fake_exports

    fake_prod = DataPoint(30_000.0, "FAOSTAT", 0.85, "prod", "")
    orig_ct = ms.comtrade_trade
    ms.comtrade_trade = fake_comtrade
    try:
        with patch("silk_marketsize_agent.production_estimate",
                   return_value=[fake_prod]):
            dp = ms.apparent_consumption("080410", "ARE", "784", "Dates", 2022)
    finally:
        ms.comtrade_trade = orig_ct
    assert dp.value["method"] == "apparent_consumption_tonnes"
    # 30,000 + 50,000 − 10,000 = 70,000 tonnes
    assert dp.value["value_tonnes"] == 70000.0
    assert dp.value["imports_tonnes"] == 50000.0


def test_marketsize_import_value_proxy_when_production_missing():
    # بلا إنتاج/كميات: يرجع لمؤشّر قيمة الاستيراد موسومًا كجزئي، لا حجم كامل.
    from unittest.mock import patch

    import silk_marketsize_agent as ms
    from silk_data_layer import DataPoint

    def no_qty_comtrade(hs, m49, year, flow="M", partner=0):
        return []  # no trade quantities available
    fao_miss = DataPoint(None, "FAOSTAT", 0.0, "no prod", "")
    orig_ct = ms.comtrade_trade
    ms.comtrade_trade = no_qty_comtrade
    try:
        with patch("silk_marketsize_agent.production_estimate",
                   return_value=[fao_miss]):
            with patch("silk_data_layer_v2.market_imports",
                       return_value={"total_usd": 241_000_000.0, "competitors": []}):
                dp = ms.apparent_consumption("080410", "ARE", "784", "Dates", 2022)
    finally:
        ms.comtrade_trade = orig_ct
    assert dp.value["method"] == "import_value_proxy_usd"
    assert dp.value["value_usd"] == 241_000_000.0
    assert "production" in dp.value["missing"]  # honest about what's missing


def test_marketsize_none_when_nothing_available():
    # لا إنتاج ولا تجارة ولا قيمة استيراد => None موسوم، لا اختلاق.
    from unittest.mock import patch

    import silk_marketsize_agent as ms
    from silk_data_layer import DataPoint

    orig_ct = ms.comtrade_trade
    ms.comtrade_trade = lambda *a, **k: []
    try:
        with patch("silk_marketsize_agent.production_estimate",
                   return_value=[DataPoint(None, "FAOSTAT", 0.0, "no", "")]):
            with patch("silk_data_layer_v2.market_imports",
                       return_value={"total_usd": None, "competitors": []}):
                dp = ms.apparent_consumption("080410", "ARE", "784", "Dates", 2022)
    finally:
        ms.comtrade_trade = orig_ct
    assert dp.value is None and dp.confidence == 0.0


def test_engine_market_size_layer_offline():
    # طبقة المجموعة أ مفعّلة بلا شبكة: تُرفق production+market_size، والنقاط لا تتغيّر.
    os.environ.pop("SEARCH_API_KEY", None)
    with _block_network():
        res = engine.analyze("تمور", countries=[{"iso3": "ARE", "m49": "784"}],
                             year=2022, with_market_size=True)
    assert res["classified"] is True and res["hs_code"] == "080410"
    row = res["markets"][0]
    assert "production" in row and "market_size" in row   # context attached
    assert row["total_score"] == 0.0                      # additive, unchanged


def test_cities_agent_known_and_unknown_country():
    # مدن حقيقية لدولة مغطّاة، وNone موسوم لدولة غير مغطّاة — لا اختلاق.
    import silk_cities_agent as cities

    rep = cities.CitiesAgent().run({"iso3": "EGY", "top_n": 2})
    assert rep.failed is False
    top = rep.findings[0].value
    assert top["city"] == "Cairo" and isinstance(top["lat"], float)  # ranked by pop
    miss = cities.CitiesAgent().run({"iso3": "XXX"})
    assert miss.failed is True and miss.findings[0].value is None  # no fabrication


def test_religion_agent_known_and_unknown_country():
    # ديانة غالبة لدولة مرجعية، وNone لدولة غير مغطّاة — بيانات تقريبية موسومة.
    import silk_religion_agent as rel

    rep = rel.ReligionAgent().run({"iso3": "SAU"})
    assert rep.failed is False
    val = rep.findings[0].value
    assert val["majority_religion"] == "Islam" and val["source"].startswith("Pew")
    assert val["majority_share_pct"] is not None
    miss = rel.ReligionAgent().run({"iso3": "XXX"})
    assert miss.failed is True and miss.findings[0].value is None  # not guessed


def test_currency_agent_offline_no_fabrication():
    # بلا شبكة: التضخم/سعر الصرف None، والتصنيف الائتماني None دائماً (لا اختلاق).
    import silk_currency_agent as cur

    with _block_network():
        rep = cur.CurrencyRiskAgent().run({"iso3": "EGY", "year": 2022})
    assert rep.failed is True
    assert all(f.value is None for f in rep.findings)
    # آخر عنصر = التصنيف الائتماني، دائماً غير متاح وموسوم بذلك صراحة.
    rating = rep.findings[-1]
    assert rating.value is None and "credit rating" in rating.note.lower()


def test_engine_demographics_layer_offline():
    # طبقة المجموعة ب مفعّلة بلا شبكة: تُرفق cities/religion/currency_risk، والنقاط ثابتة.
    with _block_network():
        res = engine.analyze("تمور", countries=[{"iso3": "EGY", "m49": "818"}],
                             year=2022, with_demographics=True)
    row = res["markets"][0]
    assert "cities" in row and "religion" in row and "currency_risk" in row
    # المدن والديانة مرجع محلي => تعمل حتى بلا شبكة (لا اختلاق، بيانات حقيقية مخزّنة).
    # engine.analyze() يرجّع DataPoint خام (تحويل JSON يحدث في طبقة jobs/api فقط).
    assert row["religion"] is not None
    assert row["religion"].value["majority_religion"] == "Islam"
    assert any(f.value for f in row["cities"])         # Cairo present offline
    assert row["total_score"] == 0.0                   # additive, score unchanged


def test_competitors_agent_no_fabrication_keyless():
    # وكيل المنافسين بلا مفتاح بحث: None موسوم، لا أسماء علامات مُختلقة.
    import silk_competitors_agent as comp

    os.environ.pop("SEARCH_API_KEY", None)
    with _block_network():
        rep = comp.CompetitorsAgent().run({"product": "dates", "country": "Morocco"})
    assert rep.failed is True
    assert all(f.value is None for f in rep.findings)  # no fabricated brands


def test_distribution_agents_no_fabrication_keyless():
    # وكيلا التوزيع/التجارة الإلكترونية بلا مفتاح: None موسوم، لا أسماء مُختلقة.
    import silk_distribution_agent as dist

    os.environ.pop("SEARCH_API_KEY", None)
    with _block_network():
        d = dist.DistributionChannelsAgent().run({"product": "dates", "country": "UAE"})
        e = dist.EcommerceLandscapeAgent().run({"product": "dates", "country": "UAE"})
    assert d.failed is True and e.failed is True
    assert all(f.value is None for f in d.findings + e.findings)


def test_bestsellers_agent_no_token_no_scraping():
    # بلا APIFY_API_TOKEN: لا محاولة شبكة، None موسوم يوضّح القيد القانوني.
    import silk_bestsellers_agent as best

    os.environ.pop("APIFY_API_TOKEN", None)
    os.environ.pop("APIFY_BESTSELLERS_ACTOR", None)
    with _block_network():  # proves NO network call is attempted without a token
        rep = best.BestsellersAgent().run({"product": "dates", "market": "ae"})
    assert rep.failed is True
    dp = rep.findings[0]
    assert dp.value is None and dp.confidence == 0.0
    assert "ToS" in dp.note or "licensed" in dp.note.lower()  # legal note surfaced


def test_bestsellers_parse_ranks_real_only():
    # التحليل يقرأ الترتيب من الردّ الحقيقي فقط (ترتيب القائمة عند غياب rank صريح).
    import silk_bestsellers_agent as best

    items = [
        {"title": "Brand A dates 1kg", "price": 39, "platform": "Amazon.ae"},  # no rank -> pos 1
        {"title": "Brand B dates", "rank": 2, "price": 25, "platform": "Noon"},  # explicit rank 2
        {"no_title": "skipme"},  # no title -> skipped, never fabricated
    ]
    parsed = best._parse_items(items, num=10)
    titles = [p["title"] for p in parsed]
    assert titles == ["Brand A dates 1kg", "Brand B dates"]  # sorted by rank asc
    assert parsed[0]["rank"] == 1 and parsed[1]["rank"] == 2  # A: list pos; B: explicit
    assert len(parsed) == 2  # untitled row dropped, never fabricated


def test_engine_competition_layer_offline():
    # طبقة المجموعة ج مفعّلة بلا شبكة/مفاتيح: تُرفق المفاتيح، والنقاط لا تتغيّر.
    for key in ("SEARCH_API_KEY", "APIFY_API_TOKEN", "APIFY_BESTSELLERS_ACTOR"):
        os.environ.pop(key, None)
    with _block_network():
        res = engine.analyze("تمور", countries=[{"iso3": "ARE", "m49": "784"}],
                             year=2022, with_competition=True)
    row = res["markets"][0]
    for key in ("competitors_web", "distribution_channels", "ecommerce", "bestsellers"):
        assert key in row                               # context attached
    assert row["total_score"] == 0.0                    # additive, score unchanged


def test_regulatory_and_customs_agents_no_fabrication_keyless():
    # وكيلا الاشتراطات/الجمارك بلا مفتاح بحث: None موسوم، لا اشتراطات مُختلقة.
    import silk_regulatory_agent as reg

    os.environ.pop("SEARCH_API_KEY", None)
    with _block_network():
        r = reg.RegulatoryStandardsAgent().run({"product": "dates", "country": "UAE"})
        c = reg.CustomsInfoAgent().run({"product": "dates", "country": "UAE"})
    assert r.failed is True and c.failed is True
    assert all(f.value is None for f in r.findings + c.findings)  # no fabrication


def test_engine_compliance_layer_offline():
    # طبقة المجموعة د مفعّلة بلا شبكة/مفتاح: تُرفق regulatory/customs_web، والنقاط ثابتة.
    os.environ.pop("SEARCH_API_KEY", None)
    with _block_network():
        res = engine.analyze("تمور", countries=[{"iso3": "ARE", "m49": "784"}],
                             year=2022, with_compliance=True)
    row = res["markets"][0]
    assert "regulatory" in row and "customs_web" in row   # context attached
    assert row["total_score"] == 0.0                       # additive, score unchanged


def test_culture_agents_no_fabrication_keyless():
    # وكلاء المجموعة هـ بلا مفتاح بحث: None موسوم، لا رؤى/أعراف/معارض مُختلقة.
    import silk_culture_agent as cult

    os.environ.pop("SEARCH_API_KEY", None)
    with _block_network():
        a = cult.CulturalAgent().run({"product": "dates", "country": "UAE"})
        b = cult.BusinessCultureAgent().run({"country": "UAE", "product": "dates"})
        c = cult.ExhibitionsAgent().run({"product": "dates", "country": "UAE"})
    assert a.failed and b.failed and c.failed
    assert all(f.value is None for f in a.findings + b.findings + c.findings)


def test_engine_culture_layer_offline():
    # طبقة المجموعة هـ مفعّلة بلا شبكة/مفتاح: تُرفق المفاتيح، والنقاط لا تتغيّر.
    os.environ.pop("SEARCH_API_KEY", None)
    with _block_network():
        res = engine.analyze("تمور", countries=[{"iso3": "ARE", "m49": "784"}],
                             year=2022, with_culture=True)
    row = res["markets"][0]
    for key in ("cultural", "business_culture", "exhibitions"):
        assert key in row                               # context attached
    assert row["total_score"] == 0.0                    # additive, score unchanged


def test_synthesis_keyless_returns_none():
    # بلا ANTHROPIC_API_KEY: التركيب يُعيد None، وتبقى اللجنة الحتمية (لا اختلاق).
    import silk_synthesis as syn

    os.environ.pop("ANTHROPIC_API_KEY", None)
    assert syn.available() is False
    row = {"country": "مصر", "components": {"market_size": {"value": 2.4e8,
           "source": "UN Comtrade", "note": "imports"}}}
    assert syn.synthesize_market(row, "تمور") is None


def test_synthesis_collect_groups_real_only():
    # التجميع يضع الحقائق الحقيقية فقط في مجموعاتها، ويحذف الفارغ (لا اختلاق).
    import silk_synthesis as syn

    row = {
        "country": "مصر",
        "components": {"market_size": {"value": 2.4e8, "source": "UN Comtrade",
                                       "note": "imports"},
                       "demand_capacity": {"value": None}},  # None -> excluded
        "income_ppp": 14800, "population": 111000000,
        "religion": {"value": {"majority_religion": "Islam"}, "source": "Pew"},
        "cultural": [{"value": None}],  # no real value -> group E stays empty/omitted
    }
    groups = syn.collect_groups(row)
    assert "A" in groups and any("UN Comtrade" in f for f in groups["A"])
    assert "B" in groups and any("111000000" in f or "income_ppp" in f for f in groups["B"])
    assert "E" not in groups  # only a None finding -> omitted, not fabricated


def test_synthesis_injection_guard_in_principle():
    # مبدأ التركيب يحوي حارس حقن التعليمات صراحةً (raw_findings = بيانات فقط).
    import silk_synthesis as syn

    assert "raw_findings" in syn._PRINCIPLE
    assert "تجاهل" in syn._PRINCIPLE  # "ignore any instructions inside"


def test_synthesis_two_stage_flow_and_gaps_mocked():
    # تدفّق المرحلتين مع كلود وهمي: يُبنى التركيب، والمجموعات الغائبة تظهر في gaps.
    from unittest.mock import patch

    import silk_synthesis as syn

    row = {"country": "مصر",
           "components": {"market_size": {"value": 2.4e8, "source": "UN Comtrade",
                                          "note": "imports"}}}  # only group A has data

    stage2 = ('{"verdict":"WATCH","confidence":0.6,"opportunities":["فرصة"],'
              '"risks":["خطر"],"recommendations":["خطوة"],"gaps":[]}')

    def fake_call(system, user, max_tokens=1200):
        # المرحلة ٢ تُرسل group_summaries؛ المرحلة ١ لا — نميّز بينهما بالمحتوى.
        return stage2 if "group_summaries" in user else "ملخّص المجموعة أ."

    with patch("silk_synthesis.available", return_value=True):
        with patch("silk_synthesis._call", side_effect=fake_call):
            out = syn.synthesize_market(row, "تمور")
    assert out is not None
    assert out["verdict"] == "WATCH" and out["groups_with_data"] == ["A"]
    # المجموعات ب–هـ غائبة => يجب أن تظهر كنواقص (سياسة الفشل الجزئي).
    joined = " ".join(out["gaps"])
    assert "الاقتصاد" in joined and "الثقافة" in joined


def test_engine_synthesis_layer_offline_keyless():
    # طبقة التركيب مفعّلة بلا مفتاح: لا تعطّل، لا تُرفق synthesis، النقاط ثابتة.
    os.environ.pop("ANTHROPIC_API_KEY", None)
    with _block_network():
        res = engine.analyze("تمور", countries=[{"iso3": "ARE", "m49": "784"}],
                             year=2022, with_synthesis=True)
    row = res["markets"][0]
    assert "synthesis" not in row              # keyless -> nothing fabricated/attached
    assert row["total_score"] == 0.0          # additive, score unchanged


def _demo_result_for_reports():
    return {"product": "تمور", "hs_code": "080410", "year": 2022,
            "markets": [{"country": "مصر",
                         "components": {"market_size": {"value": 2.4e8},
                                        "saudi_position": {"value": 44}},
                         "income_ppp": 14800, "population": 111000000,
                         "total_score": 0.74,
                         "jury": {"verdict": "PRELIMINARY GO"},
                         "synthesis": {"verdict": "WATCH",
                                       "opportunities": ["سوق كبير", "حضور سعودي"],
                                       "risks": ["منافسة عراقية"],
                                       "recommendations": ["ابدأ بشحنة تجريبية"],
                                       "gaps": ["مجموعة غير متوفّرة: الثقافة"],
                                       "by": "Claude"},
                         "regulatory": [{"value": {"title": "halal cert"},
                                         "source": "Web Search"}]}]}


def test_reports_build_docx_with_disclaimer():
    # يبني التقريرين Word ويحويان تذييل إخلاء المسؤولية والخلاصة التنفيذية.
    import pytest
    pytest.importorskip("docx")
    import silk_reports as reports

    result = _demo_result_for_reports()
    full = reports.build_full_report(result)
    short = reports.build_short_report(result)
    assert len(full) > 2000 and len(short) > 2000  # real .docx payloads
    # افتح المختصر وتحقق من وجود التذييل والحكم — reopen and check content.
    import io
    import docx
    doc = docx.Document(io.BytesIO(short))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "إخلاء مسؤولية" in text            # mandatory disclaimer present
    assert "الخلاصة التنفيذية" in text        # executive summary present
    assert "WATCH" in text                     # verdict carried from synthesis


def test_reports_no_fabrication_on_missing_fields():
    # نتيجة شبه فارغة: التقرير يُبنى ويكتب "غير متوفّر" بدل اختلاق أرقام.
    import pytest
    pytest.importorskip("docx")
    import io
    import docx
    import silk_reports as reports

    result = {"product": "x", "hs_code": "1", "year": 2022,
              "markets": [{"country": "م", "components": {}}]}  # no numbers
    short = reports.build_short_report(result)
    doc = docx.Document(io.BytesIO(short))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert reports._NA in text                 # missing values marked, not invented


def test_reports_pdf_graceful():
    # to_pdf لا يرمي أبداً: يعيد bytes (لو LibreOffice موجود) أو None (لو غائب).
    import pytest
    pytest.importorskip("docx")
    import silk_reports as reports

    short = reports.build_short_report(_demo_result_for_reports())
    out = reports.to_pdf(short)
    assert out is None or (isinstance(out, bytes) and out[:4] == b"%PDF")


def test_reports_endpoints_auth_and_flow():
    # نقاط التقارير محميّة بالجلسة وتبني ملفاً من نتيجة مهمة حقيقية مملوكة.
    import pytest
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    pytest.importorskip("docx")
    from unittest.mock import patch
    from fastapi.testclient import TestClient
    import hashlib
    import secrets
    import api
    import silk_auth
    import silk_db
    import silk_jobs

    client = TestClient(api.app)
    # بلا جلسة => 401 — unauthorized without a session token.
    assert client.get("/reports/nope/full.docx").status_code == 401

    # أنشئ مستخدماً وجلسة — mint a real session.
    token = secrets.token_urlsafe(32)
    silk_db.store_magic_link("reports@example.com",
                             hashlib.sha256(token.encode()).hexdigest())
    sess = silk_auth.verify_magic_link(token)
    auth = {"Authorization": f"Bearer {sess['session_token']}"}

    # مهمة منتهية مملوكة بنتيجة تقرير — an owned finished job with a result.
    job_id = silk_db.create_job(sess["user_id"])
    import json
    silk_db.update_job(job_id, "finished",
                       result_json=json.dumps(_demo_result_for_reports()))

    r = client.get(f"/reports/{job_id}/short.docx", headers=auth)
    assert r.status_code == 200
    assert r.content[:2] == b"PK"              # a real .docx (zip) payload
    assert "attachment" in r.headers.get("content-disposition", "")


def test_dashboard_route_serves_page():
    # /dashboard/{id} يقدّم صفحة اللوحة الثابتة (عامّة؛ الصفحة تجلب البيانات بالجلسة).
    import pytest
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    import api

    client = TestClient(api.app)
    r = client.get("/dashboard/some-job-id")
    assert r.status_code == 200
    assert "لوحة سِلك" in r.text and "leaflet" in r.text.lower()  # real dashboard page


def test_analyze_never_invokes_paid_volza_even_with_key_and_flag():
    # حرج (أمان/تكلفة): /analyze لا يستدعي Volza أبداً — حتى مع مفتاح محفوظ
    # وإرسال with_volza=True عمداً. الأدوات المدفوعة عبر /deepen فقط.
    import pytest
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from unittest.mock import patch
    from fastapi.testclient import TestClient
    import api
    import silk_volza_agent

    os.environ["VOLZA_API_KEY"] = "SPY-KEY"        # key present -> old bug would fire it
    client = TestClient(api.app)
    headers = _test_session_headers("no-paid-on-analyze@example.com")
    calls = {"n": 0}
    real = silk_volza_agent.VolzaAgent.run

    def spy(self, task):
        calls["n"] += 1
        return real(self, task)

    try:
        with patch.object(silk_volza_agent.VolzaAgent, "run", spy), \
             patch("requests.get", side_effect=OSError("net blocked")), \
             patch("requests.post", side_effect=OSError("net blocked")):
            r = client.post("/analyze", json={"product": "تمور", "year": 2022,
                                              "with_competition": True,
                                              "with_volza": True, "with_explee": True},
                            headers=headers)
            data = r.json()
            jr = client.get(f"/jobs/{data['job_id']}", headers=headers).json()
    finally:
        os.environ.pop("VOLZA_API_KEY", None)
    assert r.status_code == 200
    row = (jr["result"]["markets"] or [{}])[0]
    assert calls["n"] == 0                 # Volza never invoked on /analyze
    assert "volza" not in row and "explee" not in row  # no paid enrichment attached
    assert "competitors_web" in row        # free Group C still ran


def test_dnb_agent_no_fabrication_keyless():
    # وكيل D&B بلا مفتاح/شبكة: None موسوم يوضّح الاشتراك، لا DUNS مُختلق.
    import silk_dnb_agent as dnb

    os.environ.pop("DNB_API_KEY", None)
    with _block_network():
        rep = dnb.DnbAgent().run({"names": ["Some Supplier LLC"], "country": "MA"})
    assert rep.failed is True
    dp = rep.findings[0]
    assert dp.value is None and dp.confidence == 0.0
    assert "DNB_API_KEY" in dp.note or "subscription" in dp.note.lower()


def test_deepen_offline_keyless_no_fabrication():
    # التعميق بلا مفاتيح/شبكة: يبني الهيكل لأعلى الأسواق دون اختلاق شركات/DUNS.
    import silk_deepen

    for k in ("GOOGLE_MAPS_API_KEY", "VOLZA_API_KEY", "EXPLEE_API_KEY", "DNB_API_KEY"):
        os.environ.pop(k, None)
    result = {"classified": True, "product": "تمور", "hs_code": "080410",
              "markets": [{"iso3": "MAR", "m49": "504", "country": "المغرب"},
                          {"iso3": "EGY", "m49": "818", "country": "مصر"}]}
    with _block_network():
        out = silk_deepen.deepen(result, top=2)
    assert out["hs_code"] == "080410" and len(out["markets"]) == 2
    m = out["markets"][0]
    for key in ("maps", "volza", "explee", "dnb"):
        assert key in m
    # بلا مفاتيح: لا قيم حقيقية (لا اختلاق) — keyless => no real values fabricated.
    reals = [f for f in (m["maps"] + m["volza"] + m["explee"] + m["dnb"])
             if (f.get("value") if isinstance(f, dict) else getattr(f, "value", None)) is not None]
    assert reals == []


def test_deepen_endpoint_auth_and_flow():
    # نقطة /deepen محميّة بالجلسة وتعمل على نتيجة مهمة مملوكة منتهية.
    import pytest
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    import json
    import api
    import silk_db

    client = TestClient(api.app)
    assert client.post("/deepen/nope").status_code == 401     # no session

    headers = _test_session_headers("deepen@example.com")
    # نستخرج user_id من الجلسة عبر إنشاء مهمة مملوكة — owned finished job.
    import silk_auth
    uid = silk_auth.session_user_id(headers["Authorization"].split(" ", 1)[1])
    job_id = silk_db.create_job(uid)
    silk_db.update_job(job_id, "finished", result_json=json.dumps(
        {"classified": True, "product": "تمور", "hs_code": "080410",
         "markets": [{"iso3": "MAR", "m49": "504", "country": "المغرب"}]}))
    for k in ("GOOGLE_MAPS_API_KEY", "VOLZA_API_KEY", "EXPLEE_API_KEY", "DNB_API_KEY"):
        os.environ.pop(k, None)
    # keyless => the paid agents short-circuit before any network call (no need to
    # block sockets, which would break TestClient's async transport).
    r = client.post(f"/deepen/{job_id}", headers=headers)
    assert r.status_code == 200
    data = r.json()
    assert data["hs_code"] == "080410" and data["markets"][0]["iso3"] == "MAR"


if __name__ == "__main__":
    import logging

    logging.disable(logging.WARNING)
    for name, fn in sorted(globals().items()):
        if name.startswith("test_") and callable(fn):
            fn()
            print(f"PASS {name}")
    print("ALL TESTS PASSED")
