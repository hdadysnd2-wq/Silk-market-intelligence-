"""أقفال إصلاح تصدير /research (تدقيق حيّ: تمور × هولندا).

البلاغ الإنتاجي: اللوحة تعرض تقرير البحث العميق الغنيّ (مراجَع، دورتان، أسعار
رفّ، HHI، لوائح EU)، لكن `GET /analyses/{id}/report.md` كان يعيد قالب /analyze
الفارغ («تغطية 0.0%»، «سنة البيانات None»، «with_research غير مفعّلة»، «0 أسواق»،
SWOT فارغ) و`GET /analyses/{id}/report.docx` يفشل 501 — لأن كلا المُصدِّرَين
لم يقرأا `dr["report"]`.

الأقفال:
  - report.md لنتيجة بحث يُصيَّر من deep_research (السرد نفسه)، لا قالب /analyze.
  - أرقام الصدق: «ثقة التصنيف»/«سنة البيانات» تُعرَض «—» لا نصّ "None".
  - report.docx (تقرير العميل) لا يفشل 501 حين يحمل السرد لغة حكم مُعرَّبة
    («درجة الثقة») — تُحوَّل تجارياً بدل نسف التصدير.

**شكل المدوّنة**: يُعاد بناء المفاتيح الدقيقة من `api._run_research_pipeline`
(الشكل المخزَّن الموثوق: markets:[]، deep_research{missions,analyst,verdict,
report,trace_id,budget_status}، market{iso3,m49,iso2,name_en,name_ar}،
data_economics) — لا مدوّنة مثالية مبسّطة. قاعدة الإنتاج الحيّة على Railway غير
قابلة للوصول من مِعزل CI (البروكسي يمنع مضيف Railway)، والتتبّع غير محلّي؛
فالشكل مُعاد بناؤه من الكود الذي يكتب المدوّنة نفسها.
"""
from __future__ import annotations

import contextlib
import os
import sys
import tempfile
from unittest import mock

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pytest

# المدوّنة القانونية الحقيقية الشكل مُمركَزة الآن في مصدر واحد يستورده هذا
# الاختبار الهرمتي ورُتبتا الخادم/المتصفّح الحقيقيتان معاً (رُتب ٢–٣). الغلاف
# `_netherlands_research_blob` يبقى بنفس الاسم للتوافق الرجعي مع بقيّة الملف.
sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(
    os.path.abspath(__file__))), "tools"))
from canonical_netherlands import netherlands_research_blob  # noqa: E402


@contextlib.contextmanager
def _env(**vals):
    old = {k: os.environ.get(k) for k in vals}
    try:
        for k, v in vals.items():
            if v is None:
                os.environ.pop(k, None)
            else:
                os.environ[k] = v
        yield
    finally:
        for k, v in old.items():
            if v is None:
                os.environ.pop(k, None)
            else:
                os.environ[k] = v


def _dp(value, source="UN Comtrade", conf=0.8, note="", ra="2026-07-15"):
    # غلاف رجعي — المُنشئ القانوني في tools/canonical_netherlands._dp.
    from canonical_netherlands import _dp as _canon_dp
    return _canon_dp(value, source, conf, note, ra)


def _netherlands_research_blob():
    """غلاف للمدوّنة القانونية الوحيدة (tools/canonical_netherlands) — نفس
    الشكل الذي تبذر منه رُتبتا الخادم/المتصفّح الحقيقيتان، فلا تتشعّب النسخ."""
    return netherlands_research_blob()


# ═══ القفل ١ — report.md يُصيَّر من deep_research لا من قالب /analyze ═══════

def test_report_md_renders_deep_research_not_analyze_template():
    import silk_render
    from silk_reports import render_markdown
    view = silk_render.build_view(_netherlands_research_blob())
    md = render_markdown(view)
    # السرد الغنيّ الحقيقي موجود.
    assert "HHI" in md and "940" in md              # المشهد التنافسي
    assert "EU 2017/625" in md                       # لوائح الوصول
    assert "6.20" in md or "7.49" in md              # أسعار الرفّ
    assert "دورة تنقيح" in md                        # المراجعة (دورتان)
    # قالب /analyze الفارغ اختفى تماماً.
    assert "with_research" not in md
    assert "0 أسواق مرشّحة" not in md
    assert "حزمة وكلاء البحث غير مفعّلة" not in md


# ═══ القفل ٢ — لا نصّ "None" حرفي في الترويسة (صدق الفجوات) ════════════════

def test_report_md_never_prints_none_literal():
    import silk_render
    from silk_reports import render_markdown, _data_year_label
    view = silk_render.build_view(_netherlands_research_blob())
    md = render_markdown(view)
    assert "None" not in md                          # لا سنة/ثقة "None"
    assert "ثقة التصنيف —" in md                     # فجوة معلنة صريحة
    # وحدةً: سنة غائبة => «—» لا "None".
    assert _data_year_label({"year": None}) == "—"
    assert _data_year_label({"year": 2023}) == "2023"


# ═══ القفل ٣ — report.docx (تقرير العميل) لا يفشل 501 على لغة الحكم ════════

def test_report_docx_client_does_not_501_on_judgment_language():
    import silk_render
    from silk_reports import render_client_docx, _client_forbidden_hits
    view = silk_render.build_view(_netherlands_research_blob())
    # كان يرفع RuntimeError→501 على «درجة الثقة» (algorithm_language).
    path = render_client_docx(view, os.path.join(tempfile.mkdtemp(), "c.docx"))
    assert os.path.exists(path)
    # والمستند نظيف فعلاً (لغة الحكم حُوّلت تجارياً لا سُرِّبت).
    from docx import Document
    doc = Document(path)
    blob = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                blob += "\n" + c.text
    assert not _client_forbidden_hits(blob), _client_forbidden_hits(blob)


# ═══ القفل ٤ — نقطة النهاية GET /report.md تقدّم السرد فعلاً ═══════════════

def test_report_md_endpoint_serves_deep_research_narrative():
    from fastapi.testclient import TestClient
    import importlib
    import api
    with _env(SILK_API_KEY="x", SILK_RATE_LIMIT="0",
              SILK_USAGE_DB=os.path.join(tempfile.mkdtemp(), "u.db")):
        importlib.reload(api)
        client = TestClient(api.create_app())
        with mock.patch("silk_storage.get_analysis",
                        return_value=_netherlands_research_blob()):
            r = client.get("/analyses/5/report.md", headers={"X-API-Key": "x"})
        assert r.status_code == 200, r.text
        body = r.text
        assert "HHI" in body and "EU 2017/625" in body
        assert "with_research" not in body and "None" not in body
