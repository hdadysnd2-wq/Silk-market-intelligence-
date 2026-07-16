"""C5 (SPEC-v2, Command #5b) — جدول «قائمة مستوردين وموزعين قابلين للتواصل»
يُصيَّر في md **و**docx (المدقّق والعميل) من بنية النموذج الواحدة، بالأعمدة
السبعة + مستوى التوثيق «◐ مرصود عبر خرائط قوقل» + سطر الإفصاح، وبلا اختلاق.
كذلك: إيقاف المكشطة أثناء تشغيلة لا يُسقط التقرير (فجوة معلنة).

Run: python3 -m pytest tests/test_importer_leads_render_c5.py -q
"""
import os
import sys
import tempfile

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


@pytest.fixture(autouse=True)
def _isolated_leads_cache(monkeypatch):
    """عزل مخزن الروابط لكل اختبار — لا تلوّث ترتيبي."""
    monkeypatch.setenv("SILK_CACHE_DIR", tempfile.mkdtemp())

_LEADS = {
    "path": "scraper",
    "note": "مرصود عبر مكشطة خرائط قوقل (هاتف/إيميل)",
    "leads": [
        {"name": "Ajwa XL", "address": "Rotterdam", "phone": "+3110123",
         "email": "sales@ajwa.nl", "website": "ajwa.nl", "rating": 4.5,
         "review_count": 120, "maps_link": "https://maps.google/ajwa",
         "doc_level": "◐ مرصود عبر خرائط قوقل"},
        {"name": "NutsWorld", "address": "—", "phone": "—", "email": "—",
         "website": "—", "rating": None, "review_count": None,
         "maps_link": "—", "doc_level": "○ مرشّح ويب غير موثَّق"},
    ],
}


def _blob(leads=_LEADS):
    m = {"agent_name": "LLMMissionAgent", "summary": "قنوات", "failed": False,
         "findings": []}
    return {
        "product": "تمور", "hs_code": "080410", "markets": [],
        "market": {"iso3": "NLD", "m49": 528, "iso2": "NL",
                   "name_en": "Netherlands", "name_ar": "هولندا"},
        "deep_research": {
            "missions": {"channels_importers": m},
            "analyst": {"report": {"agent_name": "a", "summary": "s",
                                   "findings": [], "failed": False},
                        "missing_categories": [], "by_category": {}},
            "verdict": {"verdict": "WATCH", "ai": {"verdict": "WATCH"}},
            "report": {"report": "## 1. الخلاصة\nنص التقرير.", "review_cycles": 1,
                       "unresolved_notes": [], "failure_reason": ""},
            "importer_leads": leads, "trace_id": "nld-c5"},
    }


def test_md_renders_leads_table_with_columns_and_disclaimer():
    import silk_render
    from silk_reports import render_markdown
    md = render_markdown(silk_render.build_view(_blob()))
    assert "قائمة مستوردين وموزعين قابلين للتواصل" in md
    for col in ("الاسم", "العنوان", "الهاتف", "الإيميل", "التقييم",
                "مستوى التوثيق"):
        assert col in md, f"عمود مفقود: {col}"
    assert "Ajwa XL" in md and "sales@ajwa.nl" in md and "+3110123" in md
    assert "◐ مرصود عبر خرائط قوقل" in md          # مستوى توثيق الخرائط
    assert "لا أنه يستورد التمور السعودية" in md    # سطر الإفصاح C5


def test_client_docx_renders_leads_table_and_survives_guard():
    import silk_render
    from silk_reports import render_client_docx
    from docx import Document
    view = silk_render.build_view(_blob())
    path = render_client_docx(view, os.path.join(tempfile.mkdtemp(), "c.docx"))
    assert os.path.exists(path)          # الحارس لم يُسقط التصدير على جدول الروابط
    doc = Document(path)
    blob = "\n".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for c in row.cells:
                blob += "\n" + c.text
    assert "قائمة مستوردين وموزعين قابلين للتواصل" in blob
    assert "Ajwa XL" in blob and "sales@ajwa.nl" in blob
    assert "مرصود عبر خرائط قوقل" in blob
    assert "لا أنه يستورد التمور السعودية" in blob


def test_no_leads_declares_gap_never_fabricates_rows():
    import silk_render
    from silk_reports import render_markdown
    gap = {"leads": [], "path": "gap", "note": "تعذّر الرصد"}
    md = render_markdown(silk_render.build_view(_blob(gap)))
    assert "قائمة مستوردين وموزعين قابلين للتواصل" in md
    assert "فجوة معلنة" in md
    assert "Ajwa XL" not in md            # لا صف مختلَق


def test_scraper_kill_during_run_does_not_break_report():
    """إيقاف/تعذّر المكشطة أثناء التشغيلة => finalize يعيد فجوة، والتقرير
    يُصيَّر كاملاً (عزل C4/معيار القبول ٥)."""
    import silk_gmaps
    import silk_render
    from silk_reports import render_markdown
    from unittest.mock import patch
    from types import SimpleNamespace
    ref = SimpleNamespace(iso3="NLD", iso2="NL", m49=528,
                          name_en="Netherlands", name_ar="هولندا")
    # عزل مخزن الروابط (مجلّد مؤقّت) — لا روابط مخزّنة من اختبار سابق تلوّث.
    fresh_cache = tempfile.mkdtemp()
    # المكشطة «حيّة» لكنها تُقتَل: كل نداء HTTP يفشل، وPlaces أيضاً => فجوة.
    with patch.dict(os.environ, {"SILK_GMAPS_SCRAPER_URL": "http://dead:8080",
                                 "SILK_CACHE_DIR": fresh_cache}), \
            patch("requests.post", side_effect=OSError("killed")), \
            patch("requests.get", side_effect=OSError("killed")), \
            patch("silk_maps_agent.find_places", return_value=[]):
        fut = silk_gmaps.submit_scrape_async("تمور", ref)
        out = silk_gmaps.finalize_leads(fut, "تمور", ref, [], timeout_s=3)
    assert out["path"] == "gap" and out["leads"] == []
    md = render_markdown(silk_render.build_view(_blob({"leads": [], "path": "gap",
                                                       "note": out["note"]})))
    assert "## 1. الخلاصة" in md or "الخلاصة" in md   # التقرير كامل رغم القتل
