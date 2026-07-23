"""القالب الأكاديمي (قرار المالك 2026-07-22) — أقفال البنية والعقود.

البنية المعتمدة (نموذج v3 + تعقيب المالك): «ملخّص الدراسة» الشامل يفتتح
التقرير والتوصية أول سطر فيه؛ «٧. التوصيات» قسم ختامي مستقل قبل المراجع؛
وثلاثة أقسام شواهد مهيكلة تُعرض دائماً من بيانات البعثات (الديموغرافيا/
حجم السكان، ثقافة المستهلك، الاشتراطات الجمركية). حتمي بالكامل — صفر
نداء كلود؛ نفس بوابات العميل كلها (تطهير/نقاء/اتساق الحكم/بوابة نصّ
المُنتَج النهائي).

Run: python3 -m pytest tests/test_academic_report_style.py -q
"""
from __future__ import annotations

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pytest

from tools.gen_academic_sample import academic_sample_blob  # noqa: E402


def _build_doc(tmp_path, blob=None):
    pytest.importorskip("docx")
    from docx import Document
    from silk_render import build_view
    import silk_reports as R
    os.environ["SILK_HERMETIC"] = "1"
    view = build_view(blob or academic_sample_blob())
    view["test_run"] = True
    out = str(tmp_path / "academic.docx")
    R.render_academic_docx(view, out)
    return Document(out)


def _headings(doc) -> list[str]:
    return [p.text for p in doc.paragraphs
            if p.style.name.startswith("Heading") and p.text.strip()]


def test_summary_opens_and_recommendations_close_the_body(tmp_path):
    doc = _build_doc(tmp_path)
    heads = _headings(doc)
    level1 = [h for h in heads if not h.startswith("٣.")]
    assert level1.index("ملخّص الدراسة") < level1.index("١. المقدمة وأهداف الدراسة")
    # «التوصيات» آخر أقسام المتن — قبل المراجع/المسرد مباشرة.
    assert level1.index("٧. التوصيات") > level1.index("٦. حدود الدراسة والبحث المستقبلي")
    assert level1.index("المراجع") > level1.index("٧. التوصيات")


def test_summary_first_line_is_the_verdict_with_calibrated_band(tmp_path):
    doc = _build_doc(tmp_path)
    paras = [p.text for p in doc.paragraphs]
    i = paras.index("ملخّص الدراسة")
    # عزل اتجاه الأقواس (WP-5) يحقن RLM غير مرئية — تُجرَّد قبل المطابقة.
    first = paras[i + 1].replace("‏", "")
    assert first.startswith("التوصية الختامية: مراقبة السوق")
    assert "منخفضة (50%)" in first          # سُلَّم المعايرة الواحد (WP-1)


def test_owner_requested_evidence_sections_present(tmp_path):
    """طلب المالك: حجم السكان/الديموغرافيا + ثقافة المستهلك + الاشتراطات
    الجمركية تظهر أقساماً مهيكلة داخل النتائج."""
    doc = _build_doc(tmp_path)
    heads = "\n".join(_headings(doc))
    assert "السياق الديموغرافي وحجم السكان" in heads
    assert "ثقافة المستهلك وأنماط الطلب" in heads
    assert "الاشتراطات الجمركية ومتطلبات الدخول" in heads
    cells = [c.text for t in doc.tables for row in t.rows for c in row.cells]
    joined = "\n".join(cells)
    assert "4.3 مليون نسمة" in joined       # حجم السكان وصل فعلاً
    assert any("GSO" in c for c in cells)   # اشتراط جمركي (ضمن نصّ البند)
    # WS10 (قرار المالك): المتن نظيف — لا شارة قوة دليل ولا عمود «قوة الدليل»
    # ولا عمود مصدر لكل صف؛ الإسناد كلّه في قسم المراجع وحده.
    headers = [c.text for t in doc.tables for c in t.rows[0].cells]
    assert "قوة الدليل" not in headers
    assert "المصدر" not in headers          # لا عمود مصدر لكل صف في المتن
    for sym in ("✓ موثّق", "◐ ثانوي", "○ غير", "✓ موثق"):
        assert sym not in joined            # صفر شارة قوة دليل في أيّ خلية
    # جدول الشواهد الإلزامي صار عموداً واحداً «البند المرصود» (لا مصدر/شارة).
    assert "البند المرصود" in headers


def test_missing_evidence_missions_skip_sections_not_fabricate(tmp_path):
    """مدوّنة بلا بعثات الشواهد: الأقسام تُسقَط — لا جدول مختلَق."""
    from tools.canonical_netherlands import netherlands_research_blob
    doc = _build_doc(tmp_path, netherlands_research_blob())
    heads = "\n".join(_headings(doc))
    assert "السياق الديموغرافي وحجم السكان" not in heads
    assert "الاشتراطات الجمركية" not in heads


def test_academic_doc_passes_client_hygiene_and_verdict_gates(tmp_path):
    """نفس عقود العميل: صفر مصطلحات محظورة، حكم واحد متسق، لا «…»."""
    from silk_reports import _client_forbidden_hits, count_suspicious_brackets
    doc = _build_doc(tmp_path)
    text = "\n".join([p.text for p in doc.paragraphs]
                     + [c.text for t in doc.tables
                        for row in t.rows for c in row.cells])
    assert _client_forbidden_hits(text) == []
    assert "بند تقني غير قابل للعرض المباشر" not in text
    assert "إذن ماذا" not in text
    assert not [l for l in text.splitlines()
                if len(l.strip()) > 25 and l.strip().endswith(("...", "…"))]
    assert text.count("التوصية الختامية: مراقبة السوق") >= 1
    assert "التوصية بالدخول" not in text     # لا حكم ثانٍ


def test_academic_render_is_deterministic_and_offline(tmp_path):
    """حتمي بالكامل: تشغيلان متتاليان = نفس النص؛ وبلا شبكة (لا نداء كلود)."""
    import socket
    from unittest.mock import patch
    with patch.object(socket, "socket",
                      side_effect=OSError("network disabled for offline test")):
        d1 = _build_doc(tmp_path / "a" if False else tmp_path)
        t1 = "\n".join(p.text for p in d1.paragraphs)
        d2 = _build_doc(tmp_path)
        t2 = "\n".join(p.text for p in d2.paragraphs)
    assert t1 == t2


def test_api_style_param_wires_academic_renderer():
    src = open(os.path.join(os.path.dirname(os.path.dirname(
        os.path.abspath(__file__))), "api.py"), encoding="utf-8").read()
    assert src.count('style == "academic"') == 2      # docx + pdf
    assert "render_academic_docx" in src
    assert "render_academic_pdf" in src
    # التصدير الأكاديمي عميلٌ أيضاً — يمرّ من بوابة التسليم نفسها (§0):
    # البوابة تُستدعى قبل فرع الأسلوب في كلا المسارين.
    docx_handler = src.split('def report_docx(')[1].split("\n    @app.")[0]
    assert docx_handler.index("_block_client_export_if_gate_failed") \
        < docx_handler.index('style == "academic"')


# ── السجل الأكاديمي للكاتب (متابعة المالك: «اختلفت الصياغة») ────────────────

def test_academic_writer_contract_exists_with_register_rules():
    from silk_style_contract import (ACADEMIC_SECTION_CLOSER,
                                     ACADEMIC_WRITER_CONTRACT)
    assert ACADEMIC_SECTION_CLOSER == "دلالة هذه النتيجة:"
    for marker in ("تشير النتائج إلى", "تخلص الدراسة إلى",
                   "لا تحويل إلى الريال", "دلالة هذه النتيجة"):
        assert marker in ACADEMIC_WRITER_CONTRACT, marker


def test_deep_report_swaps_contract_when_style_academic(monkeypatch):
    """style=academic يحقن عقد السجل الأكاديمي بدل عقد التاجر — والحكم
    المعتمد وقواعد الصدق تبقى في البرومبت كما هي."""
    import silk_ai_judge as J
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key")
    captured = {}

    def fake_call(system, user, max_tokens=1600, model=None, timeout=None):
        captured["user"] = user
        return "## 1. الخلاصة التنفيذية\nنص."

    monkeypatch.setattr(J, "_call", fake_call)
    verdict = {"verdict": "WATCH", "confidence": 0.5}
    J.deep_report({}, "ملخص", verdict, "منتج", "Kuwait", style="academic")
    academic_prompt = captured["user"]
    assert "السجل الأكاديمي" in academic_prompt
    assert "دلالة هذه النتيجة" in academic_prompt
    assert "الحكم المعتمد" in academic_prompt          # قيد WP-1 باقٍ
    J.deep_report({}, "ملخص", verdict, "منتج", "Kuwait")
    commercial_prompt = captured["user"]
    assert "السجل الأكاديمي" not in commercial_prompt


def test_write_reviewed_report_passes_style_to_both_writer_calls():
    src = open(os.path.join(os.path.dirname(os.path.dirname(
        os.path.abspath(__file__))), "silk_ai_judge.py"),
        encoding="utf-8").read()
    body = src.split("def write_reviewed_report(")[1]
    assert body.count("style=style") == 2      # المسوّدة + التنقيح


def test_regenerate_endpoint_wires_style_and_persists_it():
    src = open(os.path.join(os.path.dirname(os.path.dirname(
        os.path.abspath(__file__))), "api.py"), encoding="utf-8").read()
    regen = src.split("def regenerate_report(")[1].split("\n    @app.")[0]
    assert "regen_style" in regen
    assert "style=regen_style" in regen
    assert '["report_style"] = regen_style' in regen
    # التصديرات تتبع الأسلوب المخزَّن افتراضياً.
    assert src.count('.get("report_style")') >= 2   # docx + pdf


def test_view_propagates_stored_report_style():
    from silk_render import build_view
    from tools.gen_academic_sample import academic_sample_blob
    os.environ["SILK_HERMETIC"] = "1"
    view = build_view(academic_sample_blob())
    assert view["deep_research"]["report_style"] == "academic"


def test_committed_sample_narrative_is_academic_register(tmp_path):
    doc = _build_doc(tmp_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "تخلص الدراسة إلى" in text
    assert "دلالة هذه النتيجة" in text
    assert "ماذا يعني هذا لقرارك" not in text


def test_sample_generator_and_committed_sample_exist():
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    assert os.path.exists(os.path.join(root, "tools", "gen_academic_sample.py"))
    assert os.path.exists(os.path.join(root, "samples",
                                       "academic_report_latest.docx"))
