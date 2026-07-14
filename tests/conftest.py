"""أدوات اختبار مشتركة — shared test helpers (M0).

يوفّر `block_network` القانوني الواحد بدل النسخ المكرَّرة في ملفات الموجات
(الأثر التاريخي يُنظَّف في M9). الاختبارات الجديدة تستورد من هنا حصراً.
Canonical network guard for hermetic tests; new tests import from here only.
"""
import contextlib
import os
import socket
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


@contextlib.contextmanager
def block_network():
    """اقطع الشبكة مؤقتاً — make outbound sockets fail so 'no data' paths hold.

    بلاغ حي (تسريب تسلسل اختبارات CI): جلسة requests المشتركة الدائمة
    (silk_data_layer._session — تجميع اتصالات keep-alive للأداء الإنتاجي)
    قد تحمل اتصالاً TCP حياً فعلياً تركه نداء سابق غير محظور في نفس عملية
    pytest (تشغيل تسلسلي واحد لكل ملفات tests/). إعادة استعمال اتصال
    مجمَّع قائم لا يستدعي socket.socket() من جديد، فيتجاوز الحجب أدناه
    صامتاً ويُرجع بيانات حقيقية رغم دخول هذا السياق — ظهر هذا حين أضاف
    ملف اختبار جديد بضعة نداءات فأزاح ترتيب التنفيذ فكشف اتصالاً مجمَّعاً
    كان يبقى خاملاً غير مستغَل سابقاً. إغلاق تجمّعات الاتصال المعروفة عند
    كل دخول يمنع نجاة اتصال حيّ لاختبار يُفترض به حجب كامل — Session.close()
    يُغلق التجمّع الحالي فقط لا الكائن نفسه، فيُعاد فتح اتصال جديد طبيعياً
    خارج هذا السياق حين تُستأنف الشبكة.
    """
    real = socket.socket

    def _no_net(*a, **k):  # noqa: ANN002, ANN003
        # صياغة بلا كلمة hermetic عمداً: حارس تقارير الإنتاج يرفض أي أثر يحمل
        # الكلمة (إصلاح مراجعة Stage 5) — قطع الشبكة حالة تشغيل صادقة لا بديل
        # بيانات، فلا يجوز أن تسمّم ملاحظاتُه تقريراً مشتقاً في اختبار.
        raise OSError("network disabled for offline test")

    try:
        import silk_data_layer
        silk_data_layer._session.close()
    except Exception:  # noqa: BLE001 — أفضل جهد؛ الحجب الأساسي (socket) نافذ بدونه
        pass

    socket.socket = _no_net
    try:
        yield
    finally:
        socket.socket = real


def docx_all_text(path: str) -> str:
    """كل نص مستند Word — فقرات + خلايا جداول (مراجعة المشروع: بعض أقسام
    render_docx صارت جداولاً حقيقية بدل نقاط سردية؛ `doc.paragraphs` وحدها
    لا تصل خلايا الجداول، فتفوّت اختباراتٌ محتوًى انتقل إليها بلا انحدار فعلي).
    """
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


import tempfile

import pytest


@pytest.fixture(autouse=True)
def _isolated_fact_store(monkeypatch):
    """عزل مخزن الحقائق لكل اختبار — كتابة M2 العابرة دفّأت المخزن الافتراضي
    فتسرّبت حقائق حقيقية بين الاختبارات (اكتُشف عبر test_engine_localprice_layer_offline
    بعد تشغيلات تدقيق Stage 1). Every test gets its own store unless it overrides."""
    monkeypatch.setenv("SILK_STORE_DB",
                       os.path.join(tempfile.mkdtemp(), "store.db"))
    # 1b: عطّل مباعدة النداءات في الاختبارات — الشبكة مقطوعة أصلاً، والمباعدة
    # 250ms × مئات النداءات الفاشلة كانت ستبطئ الحزمة بلا فائدة.
    monkeypatch.setenv("SILK_HTTP_MIN_GAP_MS", "0")
    # نافذة كومتريد الخاصة (بلاغ 429، افتراضي 1100ms) تُصفَّر أيضاً — بلا
    # هذا نامت الحزمة الهيرمتية ~ساعتين فعلياً (1.1ث × مئات نداءات كومتريد
    # المقطوعة الشبكة) — اكتُشف حياً عند إضافة النافذة.
    monkeypatch.setenv("SILK_COMTRADE_MIN_GAP_MS", "0")
    # الموجة ٦ (V5): عزل ملفات التتبّع أيضاً — بلا هذا، اختبارات /research
    # الحقيقية (TestClient) تكتب data/traces/*.jsonl فعلياً على القرص.
    monkeypatch.setenv("SILK_TRACE_DIR", tempfile.mkdtemp())
    # عزل عدّاد data_economics بين الاختبارات — contextvar بلا حدود عملية
    # مستقلة (pytest يُشغّل كل الاختبارات على نفس الخيط)، فاختبار سابق ترك
    # عدّاداً بأرقام عالية كان سيُفعِّل سقف silk_llm_runtime._run_loop
    # الكلي زوراً في اختبار لاحق لا علاقة له (انحدار اكتُشف فعلياً، الموجة
    # ٦). الإنتاج غير متأثر: كل طلب /research يستدعي begin_data_counter()
    # صراحة قبل أي استخدام.
    import silk_context
    silk_context._data_counter.set(None)
