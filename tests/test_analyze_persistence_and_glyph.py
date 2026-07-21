"""قفلٌ لحادثة LESSONS ٣١ — نتائج /analyze لم تكن تحليلاتٍ محفوظةً حقيقيةً.

الجذر (مراجعة شيفرة، file:line): `silk_engine.analyze` كان يثبّت
`db_path="data/silk.db"` (مسارٌ نسبيٌّ لمجلد العمل الجاري)، والمعالج `/analyze`
لا يمرّر `db_path` — فيكتب `_persist` عبر `save_analysis(result, "data/silk.db")`
بمسارٍ **صريح** يتجاوز `silk_storage._db_path()`. بينما كل القرّاء
(`get_analysis`/`list_analyses`/كل نقاط `report.*`) ومسار `/research` يحسمون
`path=None → _db_path()` الذي يحترم `SILK_DATA_DIR`. على النشر (`SILK_DATA_DIR=
/data`) كان `/analyze` يكتب `./data/silk.db` الفاني (يُمسح عند إعادة النشر
فيعود العدّاد ١) والقرّاء يقرؤون `/data/silk.db` — فالمعرّف «1» ثم 404. نفس عائلة
LESSONS ٤، لم تُلتقط لأن `/analyze` لم يُدخَّن حيًّا قط (Commands #1-6 غطّت
`/research` حصرًا).

يغطي هذا الملف:
  (أ) توجيه القاعدة: `analyze(persist=True)` يكتب لنفس قاعدة `_db_path()` التي
      يقرأ منها `/analyses/{id}` — مع `SILK_DATA_DIR` مضبوطًا.
  (ب) تدفّق حقيقي عبر TestClient: POST /analyze → analysis_id → GET
      /analyses/{id} ينجح → التصديرات (md/docx/pdf) لا تُرجِع 404 → المعرّف
      يظهر في «بحوثي السابقة».
  (ج) نفس التدفّق لـ«مسح الأسواق» (نداء /analyze بلا سوق محدّد) — جذرٌ مشترك.
  (د) قفل الزجاج: صفر «§» في أي سلسلة تواجه العميل (title=، تسميات الأزرار).

Run: python3 -m pytest tests/test_analyze_persistence_and_glyph.py -q
"""
import contextlib
import os
import re
import sys
import tempfile
from unittest.mock import patch

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


@contextlib.contextmanager
def _env(**vals):
    saved = {k: os.environ.get(k) for k in vals}
    try:
        for k, v in vals.items():
            if v is None:
                os.environ.pop(k, None)
            else:
                os.environ[k] = v
        yield
    finally:
        for k, old in saved.items():
            if old is None:
                os.environ.pop(k, None)
            else:
                os.environ[k] = old


@contextlib.contextmanager
def _no_network():
    """اقطع الشبكة على مستوى requests (نمط اختبارات TestClient الموثَّق) —
    قطع socket عالميًّا يكسر نقل TestClient نفسه."""
    err = OSError("network blocked in test")
    with patch("requests.get", side_effect=err), \
         patch("requests.post", side_effect=err), \
         patch("requests.Session.request", side_effect=err):
        yield


def _client():
    from fastapi.testclient import TestClient
    import api
    import importlib
    importlib.reload(api)   # يلتقط SILK_DATA_DIR الحالي في create_app
    return TestClient(api.app)


_KEYLESS = dict(SILK_API_KEY=None, ANTHROPIC_API_KEY=None, VOLZA_API_KEY=None,
                EXPLEE_API_KEY=None, LOCALPRICE_API_KEY=None)


# ── (أ) توجيه القاعدة على مستوى المحرّك — الحارس الضيّق للجذر ────────────────

def test_engine_persist_writes_to_canonical_db_path_not_relative_literal():
    """`analyze(persist=True)` بلا `db_path` صريح يكتب لقاعدة `_db_path()`
    (تحترم SILK_DATA_DIR) لا لمسارٍ نسبيٍّ ثابت. القرّاء بمسار افتراضي يجدون
    الصفّ — وهذا بالضبط ما فشل حيًّا (id «1» ثم 404)."""
    import importlib
    tmp = tempfile.mkdtemp()
    with _env(SILK_DATA_DIR=tmp, **_KEYLESS), _no_network():
        import silk_engine
        import silk_storage
        importlib.reload(silk_storage)
        importlib.reload(silk_engine)
        result = silk_engine.analyze("شاي أخضر", persist=True)
        aid = result.get("analysis_id")
        assert aid is not None, "لم يُرفَق analysis_id رغم persist=True"
        # القارئ بالمسار الافتراضي (نفس ما يفعله GET /analyses/{id}) يجده:
        assert silk_storage._db_path() == os.path.join(tmp, "silk.db")
        found = silk_storage.get_analysis(aid)     # path=None → _db_path()
        assert found is not None, (
            "الصفّ غير موجود في القاعدة القانونية — كُتب لقرصٍ آخر (الجذر)")
        assert found.get("product") == "شاي أخضر"
        # كُتب فعلًا في قاعدة SILK_DATA_DIR، وليس في ./data/silk.db النسبي:
        assert os.path.exists(os.path.join(tmp, "silk.db"))
        assert aid in {r["id"] for r in silk_storage.list_analyses()}


# ── (ب) تدفّق حقيقي كامل عبر TestClient — نفس ما لمسه المالك ─────────────────

def _run_and_assert_full_flow(body: dict):
    with _env(SILK_DATA_DIR=tempfile.mkdtemp(), **_KEYLESS), _no_network():
        c = _client()
        r = c.post("/analyze", json=body)
        assert r.status_code == 200, r.text
        res = r.json()
        aid = res.get("analysis_id")
        assert aid is not None, "POST /analyze لم يُعِد analysis_id"

        # 1) إعادة الفتح بالمعرّف تنجح (كانت 404):
        g = c.get(f"/analyses/{aid}")
        assert g.status_code == 200, (
            f"GET /analyses/{aid} = {g.status_code} (الجذر: 404 'not found')")

        # 2) التصديرات الثلاثة تجد التحليل — لا 404 لأيٍّ منها:
        md = c.get(f"/analyses/{aid}/report.md")
        assert md.status_code == 200 and md.text.strip(), md.status_code
        docx = c.get(f"/analyses/{aid}/report.docx")
        assert docx.status_code == 200, f"docx={docx.status_code} {docx.text[:200]}"
        pdf = c.get(f"/analyses/{aid}/report.pdf")
        # PDF: 200 إن توفّر محرّك التحويل، أو 503 نظيف إن غاب — لكن **ليس
        # 404 أبدًا** (404 = التحليل غير موجود = الجذر المُصلَح).
        assert pdf.status_code != 404, "report.pdf رجع 404 — التحليل غير موجود"
        assert pdf.status_code in (200, 503), pdf.status_code

        # 3) المعرّف يظهر في «بحوثي السابقة» (GET /analyses):
        listing = c.get("/analyses").json()
        assert aid in {row["id"] for row in listing}, (
            "التحليل المحفوظ لا يظهر في قائمة بحوثي السابقة")
        return aid


def test_quick_scan_analyze_full_persisted_flow_no_404():
    """مسحٌ سريع لسوقٍ محدّد (peanut butter/UAE-شكل): كل التدفّق يعمل.

    hs_confirmed=True: هذا الاسم يُصادف تعليم بوّابة تأكيد HS (الموجة ٢ —
    040510/«زبدة» لا يشمل «فول سوداني»، عمداً — نفس عيّنة الحادثة الحيّة)
    والاختبار يقصد فحص تدفّق الحفظ/التصدير لا البوّابة، فيتجاوزها كمستخدمٍ
    أكّد الرمز صراحةً (اختبار البوّابة نفسها في test_report_quality_upgrade.py)."""
    _run_and_assert_full_flow(
        {"product": "زبدة الفول السوداني", "persist": True,
         "markets": ["ARE"], "hs_confirmed": True})


def test_compare_all_markets_analyze_shares_the_same_fixed_flow():
    """«مسح الأسواق» = نداء /analyze بلا سوق محدّد. جذرٌ مشترك مع المسح
    السريع (كلاهما عبر نفس المعالج) — فيُصلَح بنفس الإصلاح."""
    _run_and_assert_full_flow({"product": "زبدة الفول السوداني",
                               "persist": True, "hs_confirmed": True})


def test_analyze_backend_honors_persist_false_contract():
    """المعالج يحترم persist=false (المتعاقَد عليه): لا صفّ يُحفَظ، لا
    analysis_id. الواجهة تُرسل persist:true دائمًا (buildBody) — العقد مُتّسِق."""
    with _env(SILK_DATA_DIR=tempfile.mkdtemp(), **_KEYLESS), _no_network():
        c = _client()
        res = c.post("/analyze",
                     json={"product": "زبدة الفول السوداني",
                           "persist": False, "markets": ["ARE"],
                           "hs_confirmed": True}).json()
        assert res.get("analysis_id") is None
        assert c.get("/analyses").json() == []


# ── (د) قفل الزجاج: صفر «§» في السلاسل التي تواجه العميل ─────────────────────

def test_no_section_glyph_in_client_facing_strings():
    """انضباط السلاسل اليتيمة (LESSONS ٩): «§» يُبدَّل بصريًّا في RTL — ممنوع
    في أي سلسلة يراها العميل (title=، نصّ الأزرار/التسميات). الاستبدال القياسي
    «القسم N». تعليقات JS/docstrings (غير مرئية للعميل) خارج النطاق."""
    html = open(os.path.join(_ROOT, "web", "index.html"), encoding="utf-8").read()
    # انزع أجسام <script>/<style> — تعليقاتها/سلاسلها منطقٌ داخليّ لا يراه
    # العميل (الرمز يُبدَّل بصريًّا فقط حين يُعرَض كنصّ RTL في الصفحة).
    markup = re.sub(r"<(script|style)\b.*?</\1>", "", html,
                    flags=re.S | re.I)

    # (١) أي قيمة سمة title="..." تحوي § :
    title_hits = [m.group(0)[:80]
                  for m in re.finditer(r'title="[^"]*§[^"]*"', markup)]
    assert not title_hits, f"«§» في تلميحات title=: {title_hits}"

    # (٢) أي نصّ ظاهر بين وسمَين — تسميات الأزرار/العناوين المرئية.
    visible_hits = [m.group(1).strip()[:80]
                    for m in re.finditer(r">([^<>]*§[^<>]*)<", markup)]
    assert not visible_hits, f"«§» في نصّ ظاهر للعميل: {visible_hits}"


def test_no_section_glyph_in_rendered_analyze_report():
    """بلاغ حي (2026-07-20، تمر سكري/NLD): «§4b» على تقرير العميل من ملاحظة
    محرّك القرار (silk_decision) و«§8» من stage — ترقيمُ أقسامٍ داخليّ يصل
    report.md + docx العميل. القفل: تشغيلة /analyze حتمية حقيقية => صفر «§»
    في report.md (render_markdown) وdocx العميل ووثيقة المشغّل."""
    import silk_store
    import silk_engine  # يُستورَد قبل قطع الشبكة (requests-level) لتفادي كسر ssl
    # بذر مخزن الحقائق كي يُنتِج المحرّك قراراً (لا شبكة) — نفس بذرة المولّد.
    with _env(SILK_DATA_DIR=tempfile.mkdtemp(), SILK_STORE_DB=None, **_KEYLESS):
        silk_store.migrate()
        silk_store.upsert_trade_flows([
            {"hs6": "080410", "reporter_iso3": "CHN", "partner_iso3": "WLD",
             "year": 2023, "flow": "M", "value_usd": 6.0e7, "qty_kg": 2.0e7},
            {"hs6": "080410", "reporter_iso3": "CHN", "partner_iso3": "SAU",
             "year": 2023, "flow": "M", "value_usd": 1.8e7, "qty_kg": 8.0e6},
        ])
        with _no_network():
            result = silk_engine.analyze(
                "تمور", countries=[{"iso3": "CHN", "m49": "156"}],
                year=2023, with_research=True, with_risk=True)
        from silk_render import build_view
        from silk_reports import render_markdown, render_docx
        view = build_view(result)
        md = render_markdown(view)
        assert "§" not in md, ("«§» في report.md: "
                               + "; ".join(l for l in md.splitlines() if "§" in l)[:200])
        # تقرير /analyze الكلاسيكي يُصدَّر عبر render_docx (لا render_client_docx،
        # الأخير للبحث العميق حصراً — مغطّى باختبار منفصل).
        p = os.path.join(tempfile.mkdtemp(), "r.docx")
        render_docx(view, p)
        from docx import Document
        doc = Document(p)
        txt = " ".join(par.text for par in doc.paragraphs)
        for tb in doc.tables:
            for row in tb.rows:
                for c in row.cells:
                    txt += " " + c.text
        assert "§" not in txt, "«§» في render_docx"


def test_no_section_glyph_in_deep_research_client_surfaces():
    """القفل يمتدّ لمسار البحث العميق: report.md + docx العميل للمدوّنة القانونية
    خاليان من «§» (شبكة أمان منقّي العميل تُحيّد أي «§Nx» متبقٍّ)."""
    import silk_render as R
    import silk_reports as RP
    sys.path.insert(0, os.path.join(_ROOT, "tools"))
    from canonical_netherlands import netherlands_research_blob
    # حقن ملاحظة قديمة الطراز بـ«§» في حقيقة معروضة — يجب أن ينقّيها العميل.
    blob = netherlands_research_blob()
    blob["deep_research"]["missions"]["trade_flow"]["findings"][0]["note"] = \
        "واردات من حزمة §4b المتحقَّق منها"
    view = R.build_view(blob)
    md = RP._md_deep_research(view, [])
    assert "§" not in md
    p = os.path.join(tempfile.mkdtemp(), "c.docx")
    RP.render_client_docx(view, p)
    from docx import Document
    doc = Document(p)
    txt = " ".join(par.text for par in doc.paragraphs)
    for tb in doc.tables:
        for row in tb.rows:
            for c in row.cells:
                txt += " " + c.text
    assert "§" not in txt


def test_client_sanitizer_strips_section_glyph_token():
    """منقّي العميل يزيل رمز القسم «§Nx» (belt-and-suspenders) بلا كسر الجملة."""
    from silk_reports import _client_sanitize
    assert "§" not in _client_sanitize("قرار من حزمة §4b المتحقَّق منها")
    assert "§" not in _client_sanitize("بوابة خطر (قاعدة §8 المعلنة)")
    assert "§" not in _client_sanitize("المحرك §10.3 الموزون")
    # مراجعة الشيفرة: لا يبتلع كلمةً عربية تالية حين لا يتبع «§» رقمٌ.
    assert _client_sanitize("انظر § المنهجية للتفصيل") == "انظر المنهجية للتفصيل"
    # «§» وحده (بلا رمز) يُزال أيضاً — لا يبقى على سطح العميل.
    assert "§" not in _client_sanitize("رمز § وحده")
    # لاحقة قسمٍ بشرطة (§11.5-2) تُزال كاملة.
    assert "§" not in _client_sanitize("ارجع §11.5-2 هنا")
