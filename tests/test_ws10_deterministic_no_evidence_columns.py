"""WS10 المرحلة ٢ (المسارات الحتمية) — لا عمود «مستوى التوثيق» في متن التقرير.

يقفل ثلاثة مسارات حتمية كانت تُظهر عمود الإسناد/التوثيق خارج برومبت الكاتب:
جدول الروابط (`_LEADS_HEADER`/`_lead_cells`)، ومُعرِّبات المُطهِّر (درجة الثقة/
confidence/«| الدليل |»). حتميّ وهرمتيّ بالكامل — صفر نموذج، صفر تكلفة.

Run:  python3 -m pytest tests/test_ws10_deterministic_no_evidence_columns.py -q
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


def test_leads_header_has_no_documentation_level_column():
    import silk_reports as R
    assert "مستوى التوثيق" not in R._LEADS_HEADER
    assert R._LEADS_HEADER == ["الاسم", "العنوان", "الهاتف", "الإيميل",
                               "الموقع", "التقييم"]


def test_lead_cells_match_header_and_drop_doc_level():
    import silk_reports as R
    lead = {"name": "شركة مثال", "address": "أمستردام", "phone": "+31",
            "email": "a@b.nl", "website": "https://ex.nl",
            "rating": 4.5, "review_count": 12, "doc_level": "✓ موثّق"}
    cells = R._lead_cells(lead)
    # الطول يطابق الترويسة (وإلا انزاح الجدول) ولا خلية «مستوى التوثيق».
    assert len(cells) == len(R._LEADS_HEADER) == 6
    assert "✓ موثّق" not in cells          # قيمة doc_level لم تُعرَض
    assert "مستوى التوثيق" not in cells


def test_leads_table_renders_without_doc_level_column(tmp_path):
    # جدول Word الفعلي لا يحمل ترويسة «مستوى التوثيق».
    from docx import Document
    import silk_reports as R
    dr = {"product": "زبدة الفول السوداني", "market_iso3": "NLD",
          "importer_leads": {"leads": [
              {"name": "Amsterdam Importers BV", "address": "Amsterdam, NL",
               "phone": "+31 20 1", "email": "sales@ai.nl",
               "website": "https://ai.nl", "rating": 4.6, "review_count": 20,
               "doc_level": "✓ موثّق"}]}}
    doc = Document()
    R._docx_leads(doc, dr)
    headers = [c.text for t in doc.tables for c in t.rows[0].cells]
    if headers:   # جدول مبنيّ (الرائد نجا من الفلترة)
        assert "مستوى التوثيق" not in headers
        cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
        assert "✓ موثّق" not in cells


def test_sanitizer_does_not_produce_documentation_level_label():
    import silk_reports as R
    for leaked in ("درجة الثقة", "confidence", "Confidence 0.8",
                   "| الدليل |"):
        out = R._client_sanitize(leaked)
        assert "مستوى التوثيق" not in out
