"""WP-5 (برنامج إصلاح جودة التقارير) — انعكاس الأقواس في PDF العربي.

بلاغ التدقيق (2026-07-22): الأقواس ظهرت معكوسة «) ... (» في الـPDF المُسلَّم
حول المقاطع اللاتينية/الرقمية داخل الفقرات العربية — `_finalize_rtl` كانت
تضبط bidi على مستوى الفقرة/الـrun لكن خط الأنابيب لا يحقن أي عزل اتجاه
(RLM) حول المقاطع مختلطة الاتجاه قبل تحويل LibreOffice headless. الأقفال:

1. `_bidi_isolate_brackets`: RLM بعد القوس الافتتاحي وقبل الختامي لمقطع
   لاتيني/رقمي في سياق عربي — بلا ازدواج، وبلا مساس بنصٍّ غير عربي.
2. تُطبَّق داخل `_finalize_rtl` فيحملها docx الفعلي المُصدَّر.
3. `count_suspicious_brackets` + فحص `_pdf_bracket_check` في `docx_to_pdf`:
   فوق العتبة يفشل التصدير بصوت عالٍ (لا PDF معكوس الأقواس يُسلَّم).

Run: python3 -m pytest tests/test_wp5_rtl_brackets.py -q
"""
from __future__ import annotations

import os
import sys
import types

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pytest

RLM = "‏"


def test_rlm_injected_after_opening_and_before_closing_bracket():
    from silk_reports import _bidi_isolate_brackets
    out = _bidi_isolate_brackets("الواردات (UN Comtrade) في نمو")
    assert f"({RLM}UN Comtrade{RLM})" in out


def test_rlm_injection_is_idempotent():
    from silk_reports import _bidi_isolate_brackets
    once = _bidi_isolate_brackets("متوسط السعر (6 USD/kg) مؤشر")
    assert _bidi_isolate_brackets(once) == once


def test_non_arabic_text_and_arabic_only_parentheticals_untouched():
    from silk_reports import _bidi_isolate_brackets
    latin = "See (UN Comtrade) for details"
    assert _bidi_isolate_brackets(latin) == latin       # لا سياق عربي
    arabic = "الواردات (مرتفعة) هذا العام"
    assert _bidi_isolate_brackets(arabic) == arabic     # لا مقطع لاتيني/رقمي


def test_finalize_rtl_applies_bracket_isolation_to_real_docx(tmp_path):
    pytest.importorskip("docx")
    from docx import Document
    from silk_reports import _finalize_rtl
    doc = Document()
    doc.add_paragraph("مؤشر التركز (HHI 2500) مرتفع نسبياً")
    _finalize_rtl(doc)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert f"({RLM}HHI 2500{RLM})" in text


def test_count_suspicious_brackets_detects_mirroring_signature():
    from silk_reports import count_suspicious_brackets
    mirrored = "التقرير يظهر ) قيمة مقلوبة (\nوسطر آخر معلق (\nونهاية ("
    assert count_suspicious_brackets(mirrored) == 3
    clean = "الواردات (UN Comtrade) مستقرة (HHI 2500) تماماً."
    assert count_suspicious_brackets(clean) == 0


def test_pdf_bracket_check_fails_export_above_threshold(tmp_path, monkeypatch):
    """فحص الـPDF النهائي: نصّ مستخرَج فوق العتبة => RuntimeError (لا تسليم)."""
    import silk_reports as R

    class _FakePage:
        def __init__(self, text):
            self._t = text

        def get_text(self):
            return self._t

    class _FakePdf:
        def __init__(self, text):
            self._pages = [_FakePage(text)]

        def __enter__(self):
            return self._pages

        def __exit__(self, *a):
            return False

    bad_text = "( \n( \n( \n( \n"          # ٤ أقواس معلّقة > العتبة 3
    fake_fitz = types.SimpleNamespace(open=lambda p: _FakePdf(bad_text))
    monkeypatch.setitem(sys.modules, "fitz", fake_fitz)
    with pytest.raises(RuntimeError, match="اتجاه الأقواس"):
        R._pdf_bracket_check("/tmp/any.pdf")
    # تحت العتبة: يمرّ بصمت.
    ok_text = "الواردات (UN Comtrade) سليمة."
    monkeypatch.setitem(sys.modules, "fitz",
                        types.SimpleNamespace(open=lambda p: _FakePdf(ok_text)))
    R._pdf_bracket_check("/tmp/any.pdf")


def test_docx_to_pdf_wires_the_bracket_check():
    src = open(os.path.join(os.path.dirname(os.path.dirname(
        os.path.abspath(__file__))), "silk_reports.py"),
        encoding="utf-8").read()
    body = src.split("def docx_to_pdf(")[1].split("\ndef ")[0]
    assert "_pdf_bracket_check(" in body


def test_calibration_tool_carries_bracket_fixtures():
    src = open(os.path.join(os.path.dirname(os.path.dirname(
        os.path.abspath(__file__))), "tools", "rtl_calibration.py"),
        encoding="utf-8").read()
    assert "_BRACKET_LINES" in src
    assert "build_bracket_fixture" in src
    assert "bracket_suspicious_count" in src
