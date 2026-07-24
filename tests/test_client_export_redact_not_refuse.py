"""PART A (أمر العمل الرئيس — عائلة 501 المتكرّرة): «نقِّ لا ترفض».

تكرّر فشل تصدير docx للعميل (501) ثلاث مرّات (#90/#103/#106)، وكلّ مرة كان
العلاج مطاردة مصطلح متسرّب. التغيير البنيوي هنا ينهي العائلة: مسار تنقية
نهائي يستبدل أيّ مصطلح ممنوع متبقٍّ بمفردة محايدة ويُسلّم المستند (مع سطر
إفصاح)، والحارس `_client_assert_clean` يبقى شبكة أمان أخيرة لِما يستحيل
تنقيته فقط — فلا يسقط التصدير كله على تسرّب واحد.

الحالة الحرجة المغطّاة: كلمة إنجليزية تشغيلية عارية (status/run/call) يتركها
المُطهِّر عمداً (تفادي تشويه أسماء المصادر) — مثلاً عنوان مصدر «Market Status
Report» — كانت تُسقِط التصدير؛ الآن تُنقّى.

Run: python3 -m pytest tests/test_client_export_redact_not_refuse.py -q
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pytest  # noqa: E402

pytest.importorskip("docx")


def test_redact_text_neutralizes_bare_english_operational_tokens():
    import silk_reports as R
    # هذه الكلمات لا يعرّبها _client_sanitize (تفادي تشويه أسماء مصادر) —
    # فكانت تصل الحارس وتُسقِط 501. التنقية تستبدلها بمحايد.
    out = R._client_redact_text("تقرير يستند إلى Market Status Report الرسمي.")
    assert not R._client_forbidden_hits(out), out
    out2 = R._client_redact_text("راجع the call log و run id في النظام.")
    assert not R._client_forbidden_hits(out2), out2


def test_redact_residual_over_a_built_doc_then_assert_passes():
    """مستند فيه مصطلح تشغيلي متسرّب => _client_redact_residual ينقّيه،
    فيمرّ _client_assert_clean بلا استثناء (لا 501)."""
    import silk_reports as R
    from docx import Document
    doc = Document()
    doc.add_paragraph("المنتج واعد بحسب Market Status Report.")
    doc.add_heading("قسم", level=1)
    t = doc.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "الحقل يحوي كلمة status تشغيلية."

    changed = R._client_redact_residual(doc)
    assert changed is True
    # لا استثناء — الحارس النهائي يمرّ بعد التنقية.
    R._client_assert_clean(doc)
    blob = "\n".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for c in row.cells:
                blob += "\n" + c.text
    assert "status" not in blob.lower()
    assert "Market الحالة Report" in blob or "الحالة" in blob


def test_clean_doc_is_not_changed_by_redaction():
    """مستند نظيف أصلاً => لا تغيير (المسار العادي لا يُفعِّل الإفصاح)."""
    import silk_reports as R
    from docx import Document
    doc = Document()
    doc.add_paragraph("واردات التمور بلغت 61 مليون دولار وفق UN Comtrade.")
    assert R._client_redact_residual(doc) is False


def test_render_client_docx_does_not_501_on_english_source_title():
    """التكامل: نتيجة بحث سرد الكاتب فيها عنوان مصدر إنجليزي يحمل كلمة
    تشغيلية عارية => التصدير يُنتِج ملفاً (لا RuntimeError/501)، ويُعاد فتحه."""
    import tempfile
    import silk_reports as R
    from silk_render import build_view
    from silk_agents import AgentReport
    from silk_data_layer import DataPoint
    from docx import Document

    rep = AgentReport(
        "LLMAgent:trade_flow",
        [DataPoint(61_000_000.0, "UN Comtrade", 0.9, "واردات 2023")],
        False, "واردات مرصودة")
    body = ("## 1. الخلاصة التنفيذية\n"
            "يستند التقييم إلى تقرير Market Status Report الرسمي وبيانات "
            "الجمارك. التوصية مراقبة السوق.\n"
            "## 2. منهجية البحث ونطاقه\nمصادر رسمية.\n")
    result = {
        "product": "تمور", "hs_code": "080410", "markets": [],
        "deep_research": {
            "missions": {"trade_flow": rep},
            "analyst": {"report": rep, "by_category": {},
                       "missing_categories": []},
            "verdict": {"verdict": "WATCH", "ai": {"verdict": "WATCH",
                                                   "reasoning": "سوق مراقَب"}},
            "report": {"report": body, "review_cycles": 1,
                      "unresolved_notes": []},
        },
    }
    path = os.path.join(tempfile.mkdtemp(), "client.docx")
    view = build_view(result)
    # قبل «نقِّ لا ترفض» كان هذا يرفع RuntimeError (يترجمه المسار إلى 501).
    out = R.render_client_docx(view, path)
    assert os.path.exists(out)
    reopened = Document(out)               # يُفتَح فعلياً (عقد البند ٣)
    text = "\n".join(p.text for p in reopened.paragraphs)
    assert "status" not in text.lower()    # نُقّي المصطلح التشغيلي
    # HF4.2 (بلاغ قطر): سطرُ إفصاح التنقية لا يصل العميل — سطحُ المدقّق فقط.
    assert "نُقّيت بعض المصطلحات" not in text
    iview = {**view, "internal": True}
    iout = R.render_client_docx(iview, os.path.join(tempfile.mkdtemp(), "audit.docx"))
    itext = "\n".join(p.text for p in Document(iout).paragraphs)
    assert "نُقّيت بعض المصطلحات" in itext   # المدقّق يرى الإفصاح
