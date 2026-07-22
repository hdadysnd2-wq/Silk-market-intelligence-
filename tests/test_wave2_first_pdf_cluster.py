"""Wave 2 — عنقود أوّل PDF حيّ (فيتوتشيني/إيطاليا): قفلٌ سلوكيّ هرمتيّ.

كل الإصلاحات تُتحقَّق على **المدوّنة القانونية الحقيقية الشكل**
(`tools/canonical_fettuccine.py`) لا نموذجٍ مثالي:
- البند ٤: رائدٌ بجغرافيا خاطئة (عنوان أمريكي، سوق إيطاليا) يُسقَط.
- البند ٥: رائدٌ من جملة نثرٍ إنجليزية يُسقَط (لا نثرٌ في خلية عميل).
- البند ٦: رائدٌ حشو (اسمٌ وكلُّ الاتصال «—») يُسقَط؛ صفرٌ => سطر فجوة صادق.
- البند ٨ + الطيّة E: «سِلك» متّصلة «سلك» بلا مِحرفٍ مُركَّب (docx + PDF).
- البند ٩: مقاس A4 لا Letter.
- البند ١٠ + الطيّة B: سطر الإخلاء بارامتري بالمنتج (لا «التمور السعودية»).

الرُتبة ٣ البصرية (docx→PDF→pdftotext) تُخطَّى محليًا إن غاب soffice/pdftotext
(تعمل في وظيفة e2e-live-shape على CI).
"""
import os
import re
import shutil
import subprocess
import sys
import tempfile

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(
    os.path.abspath(__file__))), "tools"))

import silk_render  # noqa: E402
import silk_reports  # noqa: E402
from canonical_fettuccine import fettuccine_research_blob  # noqa: E402

pytest.importorskip("docx")

_COMBINING = "".join(chr(c) for c in list(range(0x064B, 0x0653)) + [0x0670])
# مِحرفٌ مُركَّبٌ **منفصل** يشقّ كلمة: يظهر في بداية رمزٍ أو بعد فراغ (نمط «ِس لك»)
# لا ملتصقًا بحرفٍ (الشدّة الشرعية في «المرشّحة» ملتصقةٌ فلا تُعَدّ شقًّا).
_DETACHED_MARK = re.compile(r"(?:^|[\s])[" + re.escape(_COMBINING) + r"]")


def _view():
    return silk_render.build_view(fettuccine_research_blob())


def _md():
    return silk_reports.render_markdown(_view())


def _docx_all_text(path):
    from docx import Document
    d = Document(path)
    out = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            out += [c.text for c in row.cells]
    for s in d.sections:
        for hf in (s.header, s.footer):
            out += [p.text for p in hf.paragraphs]
    return "\n".join(out)


# ═══ البنود ٤/٥/٦ — جدول الروابط ═══════════════════════════════════════════

def test_wrong_geo_lead_dropped_valid_kept():
    """عنوانٌ أمريكيٌّ في دراسةِ إيطاليا => الرائد يُسقَط؛ الرائد الإيطالي يبقى."""
    md = _md()
    seg = md[md.find("قائمة مستوردين"):]
    assert "Pastificio Milano" in seg          # صالح (إيطاليا) — يبقى
    assert "NutsWorld" not in seg              # عنوان أمريكي — يُسقَط
    assert "United States" not in seg


def test_prose_leak_sentence_never_becomes_a_lead_row():
    """جملةُ نثرٍ إنجليزية لا تصير صفَّ رائد (تُوجَّه للسرد) — لا نثرٌ في الجدول."""
    seg = _md()
    seg = seg[seg.find("قائمة مستوردين"):]
    assert "Italy imports a significant" not in seg
    # unit: المميِّز يرفض الجملة ويقبل الاسم.
    from silk_gmaps import looks_like_name
    assert looks_like_name("Pastificio Milano Srl") is True
    assert looks_like_name(
        "Italy imports a significant volume of pasta from several "
        "European suppliers according to recent trade data.") is False


def test_filler_all_dash_lead_dropped():
    """رائدٌ اسمُه فقط وكلُّ الاتصال «—»/فارغ => يُسقَط (لا صفّ حشو)."""
    seg = _md()
    assert "Anonimo Distribuzione" not in seg[seg.find("قائمة مستوردين"):]


def test_all_leads_filtered_out_shows_one_honest_gap_line():
    """لو صُفّيت كلُّ الروابط => سطرُ فجوةٍ صادقٌ واحد، لا جدولٌ فارغ/حشو."""
    blob = fettuccine_research_blob()
    # اجعل كلَّ الروابط غير صالحة (كلها حشو/نثر/جغرافيا خاطئة).
    for lead in blob["deep_research"]["importer_leads"]["leads"]:
        lead["name"] = "x"          # حشو (بلا اتصال)
        for k in ("phone", "email", "website", "maps_link", "address"):
            lead[k] = ""
    md = silk_reports.render_markdown(silk_render.build_view(blob))
    seg = md[md.find("قائمة مستوردين"):]
    assert "فجوة معلنة" in seg
    assert "| الاسم |" not in seg   # لا جدول


def test_no_untranslated_english_sentence_in_client_docx_table():
    """تقرير العميل (docx): لا جملةٌ إنجليزية خام في خلايا جدول الروابط."""
    p = os.path.join(tempfile.mkdtemp(), "c.docx")
    silk_reports.render_client_docx(_view(), p)
    txt = _docx_all_text(p)
    assert "Italy imports a significant" not in txt
    assert "Pastificio Milano" in txt          # الصالح يبقى


# ═══ البند ١٠ + الطيّة B — لا ترميز منتج/دولة في القوالب ═══════════════════

def test_disclaimer_parametrized_by_study_product_not_dates():
    """سطر الإخلاء يُشتَقّ من المنتج المدروس («فيتوتشيني») — لا «التمور السعودية»."""
    md = _md()
    assert "فيتوتشيني" in md[md.find("قائمة مستوردين"):]
    assert "التمور السعودية" not in md
    from silk_gmaps import maps_disclaimer, MAPS_DISCLAIMER
    assert "عسل" in maps_disclaimer("عسل")     # منتجٌ آخر => يظهر هو
    assert "التمور" not in MAPS_DISCLAIMER      # الثابت العام بلا منتج مثبَّت


def test_no_hardcoded_product_word_in_client_facing_templates():
    """الطيّة B: القوالب/السلاسل القابلة لإعادة الاستعمال في مسار الروابط تخلو
    من أيّ اسم منتجٍ مثبَّت (تمور/dates/معكرونة). المنشأ «السعودية» يبقى (صحيح)."""
    import inspect
    from silk_gmaps import maps_disclaimer
    blob = inspect.getsource(maps_disclaimer)
    # الدالّة تشتقّ من الوسيط، لا تحمل اسم منتجٍ مثبَّت.
    for tok in ("التمور", "dates", "معكرونة", "pasta", "فيتوتشيني"):
        assert tok not in blob, f"اسم منتجٍ مثبَّت في سطر الإخلاء: {tok}"
    # الثابت العام كذلك.
    src = _read("silk_gmaps.py")
    m = re.search(r'MAPS_DISCLAIMER = .+', src)
    assert m and "التمور" not in m.group(0)


def _read(rel):
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    with open(os.path.join(root, rel), encoding="utf-8") as f:
        return f.read()


# ═══ البند ٨ + الطيّة E — تشكيل العلامة/العربية (مصدر + بصري) ══════════════

def test_docx_brand_is_shape_safe_no_combining_marks():
    """docx: «سلك» متّصلة، ولا «سِلك» بالكسرة، ولا أيّ مِحرفٍ مُركَّب يشقّ كلمة."""
    p = os.path.join(tempfile.mkdtemp(), "c.docx")
    silk_reports.render_client_docx(_view(), p)
    txt = _docx_all_text(p)
    assert "سلك" in txt                         # متّصلة
    assert "سِلك" not in txt                     # لا الكسرة المُركَّبة في العلامة
    assert not _DETACHED_MARK.search(txt), "مِحرفٌ مُركَّبٌ منفصلٌ يشقّ كلمة"


def test_docx_page_size_is_a4_not_letter():
    """كل مسارات docx بمقاس A4 (210×297مم) لا Letter الأمريكي."""
    from docx import Document
    for fn in ("render_client_docx", "render_docx"):
        p = os.path.join(tempfile.mkdtemp(), "d.docx")
        getattr(silk_reports, fn)(_view(), p)
        s = Document(p).sections[0]
        assert abs(s.page_width.mm - 210) < 1 and abs(s.page_height.mm - 297) < 1, \
            f"{fn}: ليس A4 ({s.page_width.mm:.1f}×{s.page_height.mm:.1f})"


# ── الرُتبة ٣ البصرية: docx→PDF→pdftotext (مُخطَّاة إن غاب المحرّك) ──────────

def _pdf_tools():
    return bool(silk_reports._find_soffice()) and bool(shutil.which("pdftotext"))


# ── الشرط ١: نسخةُ الـPDF المجرّدةُ تشمل كلَّ جزءٍ نصّيّ، والمجموعةُ حركاتٌ فقط ──

def test_pdf_diacritic_free_copy_strips_every_xml_part():
    """الشرط ١: بعد `_pdf_diacritic_free_copy`، صفرُ نقاطٍ في U+064B–U+0652/U+0670
    عبر **كلّ** جزء نصّيّ في الحزمة (document.xml + header*/footer* + الحواشي)."""
    import zipfile
    p = os.path.join(tempfile.mkdtemp(), "r.docx")
    silk_reports.render_client_docx(_view(), p)
    tmp = silk_reports._pdf_diacritic_free_copy(p)
    z = zipfile.ZipFile(tmp)
    parts = [n for n in z.namelist()
             if n.startswith("word/") and n.endswith(".xml")]
    assert any(n.startswith("word/header") for n in parts), "لا جزء ترويسة"
    assert any(n.startswith("word/footer") for n in parts), "لا جزء تذييل"
    marks = set(_COMBINING)
    for n in parts:
        text = z.read(n).decode("utf-8")
        leaked = [hex(ord(c)) for c in text if c in marks]
        assert not leaked, f"حركاتٌ باقيةٌ في {n}: {leaked}"


def test_strip_set_is_only_combining_marks_never_base_letters():
    """الشرط ١: مجموعةُ التجريد حركاتٌ مُركَّبةٌ (Mn) حصرًا — لا حرفَ أساسٍ ولا
    همزةَ ولا صيغةَ ألفٍ تُمَسّ أبدًا."""
    import unicodedata
    for c in silk_reports._AR_COMBINING:
        assert unicodedata.category(c) == "Mn", \
            f"{hex(ord(c))} ليس حركةً مُركَّبة (قد يكون حرفًا أساسيًا)"
    # صيغُ الهمزة والألف (حروفُ أساسٍ) خارج المجموعة صراحةً.
    for base in "اأإآىءؤئ":
        assert base not in silk_reports._AR_COMBINING, f"حرفُ أساسٍ في المجموعة: {base}"


@pytest.mark.skipif(not _pdf_tools(),
                    reason="soffice/pdftotext غير متاح (يعمل في e2e-live-shape)")
def test_visual_pdf_lock_production_entrypoint_bare_no_split_no_leaks():
    """قفلٌ بصريّ عبر **مسار الإنتاج** (البنود ٨/٩ + الطيّة E + الشروط ٢/٣/٤):
    `render_client_pdf` (المُنتِج الحقيقي: docx مُنهًى + تجريد نسخة التحويل) →
    pdftotext على PDF العميل: (٢) المسار إنتاجيّ لا اختباريّ؛ (٣) خطٌّ عربيّ حاضر
    (فشلٌ عالٍ)؛ (E) صفرُ مِحرفٍ مُركَّبٍ يشقّ كلمة و«سلك» متّصلة؛ (٢-تسريب) لا
    رموز تشغيل داخلية على النصّ العاري."""
    assert silk_reports.has_arabic_font(), \
        "لا خطّ عربيّ الشكل — الـPDF سيُصيَّر tofu (fc-list بلا Naskh/Arabic)"
    pdf = os.path.join(tempfile.mkdtemp(), "client.pdf")
    silk_reports.render_client_pdf(_view(), pdf)     # مسار الإنتاج نفسه
    assert os.path.exists(pdf)
    txt = subprocess.run(["pdftotext", "-enc", "UTF-8", pdf, "-"],
                         capture_output=True, timeout=60).stdout.decode(
                             "utf-8", "replace")
    # الطيّة E: نصٌّ عارٍ تمامًا — لا مِحرفَ مُركَّبٍ إطلاقًا => لا شقَّ كلمة.
    combining = re.compile("[" + re.escape(_COMBINING) + "]")
    assert not combining.search(txt), "مِحرفٌ مُركَّبٌ باقٍ في استخراج PDF"
    assert "سلك" in txt.replace("\n", "") and "س لك" not in txt
    assert "التمور السعودية" not in txt
    # الشرط ٤ (§2): كنسُ التسريب على النصّ العاري (أوثقُ بلا حركات).
    for leak in ("tool-use", "tool_use", "Claude", "anthropic", "⚠",
                 "importer_leads", "deep_research", "MagicMock"):
        assert leak not in txt, f"تسريبٌ تشغيليٌّ في PDF العميل: {leak}"


@pytest.mark.skipif(not _pdf_tools(),
                    reason="soffice/pdftotext غير متاح (يعمل في e2e-live-shape)")
@pytest.mark.xfail(
    strict=False,
    reason="§E (حزمة الفكس v2.1) قيد التحقيق NOT DONE: استخراج pdftotext من "
           "مخرَج LibreOffice يقلب ترتيب محارف العربية (مؤشر→مؤرش، غير→غري) — "
           "عطلٌ حقيقيّ في تسليم الـPDF يحتاج فكساً في التحويل نفسه (تضمين خط "
           "بـToUnicode CMap صحيح، أو خيارات تصدير LO بديلة) لم يُنفَّذ/يُتحقَّق "
           "منه بعد. هذا القفل detector: يُخفق حتى يُصلَح فعلاً (عندها يُصبح "
           "XPASS فيُرفَع الوسم). لا يُدمَج §E «جاهزاً» قبل ذلك — الدلو الثالث "
           "«no sufficient evidence — pending».")
def test_visual_pdf_lock_no_reversed_glyph_tokens():
    """§E (حزمة الفكس v2.1) — بلاغ حي: الـPDF المُسلَّم استخرج نصّاً مقلوب
    ترتيب المحارف («مؤرش» بدل «مؤشر»، «رشط» بدل «شرط»، «أكرث» بدل «تركّز»
    (يُختبَر عبر «مؤشر التركّز»)، «مرسد» بدل «مسرد»، «السوداين» بدل
    «السودان»، «$M» بدل «M$»/الصيغة الصحيحة) رغم أن مسار التحويل هو
    python-docx → LibreOffice (`docx_to_pdf`). يبني هذا القفل تقرير المدقّق
    الداخلي الكامل (`render_research_pdf`، مدوّنة هولندا القانونية الشكل)
    الذي يحمل كلا الكلمتين المستهدَفتين فعلياً («مؤشر التركّز HHI = 940»،
    و«شرطا قلب الحكم» لحكمٍ WATCH) — لا نموذج مثالي مصطنَع للاختبار فقط.

    **الحالة: xfail (§E غير منجَز).** أثبت CI (`e2e-live-shape`) أن الانقلاب
    حقيقيٌّ في بيئة التحويل الفعلية — القفل يرصده صحيحاً، لكن فكس التحويل
    نفسه مؤجَّل لتحقيقٍ لا يمكن التحقّق منه في صندوقٍ بلا soffice+pdftotext
    عاملَين. يُرفَع `xfail` فور شحن الفكس (سيصبح XPASS)."""
    import sys as _sys
    tools_dir = os.path.join(os.path.dirname(os.path.dirname(
        os.path.abspath(__file__))), "tools")
    if tools_dir not in _sys.path:
        _sys.path.insert(0, tools_dir)
    from canonical_netherlands import netherlands_research_blob
    assert silk_reports.has_arabic_font(), \
        "لا خطّ عربيّ الشكل — الـPDF سيُصيَّر tofu (fc-list بلا Naskh/Arabic)"
    view = silk_render.build_view(netherlands_research_blob())
    pdf = os.path.join(tempfile.mkdtemp(), "research.pdf")
    silk_reports.render_research_pdf(view, pdf)   # مسار الإنتاج (المدقّق)
    assert os.path.exists(pdf)
    txt = subprocess.run(["pdftotext", "-enc", "UTF-8", pdf, "-"],
                         capture_output=True, timeout=60).stdout.decode(
                             "utf-8", "replace")
    flat = txt.replace("\n", "")
    assert "مؤشر" in flat, "الكلمة الهدف «مؤشر» غائبة — تحقّق من محتوى المدوّنة"
    assert "شرط" in flat, "الكلمة الهدف «شرط» غائبة — تحقّق من محتوى المدوّنة"
    for reversed_token in ("مؤرش", "رشط", "أكرث", "مرسد", "السوداين", "$M"):
        assert reversed_token not in flat, (
            f"محرفٌ مقلوب الترتيب في استخراج PDF: «{reversed_token}» — "
            "أثر انقلاب اتجاه غير مُصلَح في تحويل docx→PDF")


def test_shape_safe_helper_strips_only_combining_marks():
    """`_shape_safe_ar`: يجرّد الحركات/التطويل فقط، ويُبقي الحروف/اللاتيني."""
    assert silk_reports._shape_safe_ar("سِلك") == "سلك"
    assert silk_reports._shape_safe_ar("Italy 2024 — سلك") == "Italy 2024 — سلك"
