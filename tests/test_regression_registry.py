"""سجلّ الانحدار الموحّد — one guard per real incident, one meta-test for coverage.

> **الغرض (أمر المُشرِف).** لكل حادثة إنتاجية حقيقية في هذا المستودع — صفوف
> `docs/LESSONS.md` **وفخاخ** `silk-operations` §2 (THE TRAPS) — حارسٌ واحد
> يُفشِل على عودة نفس العائلة، و**اختبار تغطية شامل** (meta) يثبت أنّ كل حادثة
> مُسجَّلة هنا فعلاً — فلا تسقط حادثة من الشبكة بصمت. يشمل ذلك **الحوادث الثلاث
> لعطل 501 في تصدير docx** (صفوف LESSONS ٣/١١/١٣) بحُرّاس سلوكيين فعليين.
>
> **Why a registry (not just per-file lock-tests).** Lock-tests live scattered
> across `tests/`; this file is the single index that maps EVERY known incident
> to a live guard and then proves — mechanically — that the index is complete
> against both incident ledgers. A new incident that lands in either ledger
> without a registry entry fails `test_meta_registry_covers_every_known_incident`.

هرمتي بالكامل (قراءة مصدر + سلوك محلي، بلا شبكة). الحُرّاس السلوكية (٣/١١/١٣)
تبني/تنقّي فعلياً من المدوّنة القانونية الحقيقية الشكل — لا نماذج مثالية.

Run: python3 -m pytest tests/test_regression_registry.py -q
"""
from __future__ import annotations

import os
import re
import sys
import tempfile

import pytest

_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, _ROOT)
sys.path.insert(0, os.path.join(_ROOT, "tools"))


def _read(rel: str) -> str:
    with open(os.path.join(_ROOT, rel), encoding="utf-8") as f:
        return f.read()


def _exists(rel: str) -> bool:
    return os.path.exists(os.path.join(_ROOT, rel))


def _needles(rel: str, *needles: str):
    """حارس وجود: كل إبرة حاضرة في الملف — يعيد callable للتسجيل."""
    def check():
        assert _exists(rel), f"ملف الإنفاذ مفقود: {rel}"
        src = _read(rel)
        missing = [n for n in needles if n not in src]
        assert not missing, f"{rel}: رموز/علامات إنفاذ مفقودة {missing}"
    return check


def _absent(rel: str, *forbidden: str):
    def check():
        src = _read(rel)
        present = [n for n in forbidden if n in src]
        assert not present, f"{rel}: رموز يجب أن تكون قد أُزيلت لا تزال {present}"
    return check


# ── حُرّاس سلوكية للحوادث الثلاث لعطل docx-501 (LESSONS ٣/١١/١٣) ──────────────

def _guard_docx501_row3():
    """LESSONS ٣ — 501 شُحن لأن الاختبارات نماذج مموّهة. الحارس: تصدير docx
    العميل يُنتِج ملفاً حقيقياً قابلاً للفتح من **المدوّنة القانونية الحقيقية
    الشكل** (لا نموذج)، بلا 501."""
    import silk_render
    from silk_reports import render_client_docx
    from canonical_netherlands import netherlands_research_blob
    view = silk_render.build_view(netherlands_research_blob())
    path = render_client_docx(view, os.path.join(tempfile.mkdtemp(), "c.docx"))
    assert os.path.exists(path)
    from docx import Document
    doc = Document(path)
    assert any(p.text.strip() for p in doc.paragraphs), "docx فارغ"


def _guard_docx501_row11():
    """LESSONS ١١ — 501 تكرّر لأن محفّزات الحارس العربية بلا استبدال مقابل.
    الحارس: مصطلح حكم عربي ممنوع («درجة الثقة») يُحيَّد فعلياً بمُطهِّر/منقِّي
    العميل فلا يبقى في مخرَج التصدير."""
    from silk_reports import _client_redact_text, _client_forbidden_hits
    leaked = "التقييم يعتمد درجة الثقة العالية على مصدر البيانات"
    cleaned = _client_redact_text(leaked)
    assert not _client_forbidden_hits(cleaned), (
        f"محفّز حارس عربي بقي بعد التنقية: {_client_forbidden_hits(cleaned)}")


def _guard_docx501_row13():
    """LESSONS ١٣ — docx يفشل حياً (501) بينما الهرمتي أخضر؛ الحارس كان يرفض
    على تسرّب واحد. القاعدة «نقِّ لا ترفض»: مصطلح إنجليزي عارٍ متسرّب يُستبدَل
    بمحايد ويُسلَّم المستند — لا 501. الحارس: `_client_redact_text` ينقّي
    مصطلحاً إنجليزياً تشغيلياً بدل رفعه."""
    from silk_reports import _client_redact_text, _client_forbidden_hits
    leaked = "mission status: successful run"
    assert _client_forbidden_hits(leaked), "التهيئة خاطئة — النص يجب أن يتسرّب أولاً"
    cleaned = _client_redact_text(leaked)
    assert not _client_forbidden_hits(cleaned), (
        f"مصطلح إنجليزي تشغيلي بقي بعد التنقية: {_client_forbidden_hits(cleaned)}")


# ── حُرّاس سلوكية لفخّي المسار (markets:[] + view بعد التخزين) ────────────────

def _guard_trap_markets_empty():
    """TRAP «markets:[] misroutes exporters» — نتيجة /research دوماً markets:[]؛
    أي مسار عرض يثق بـ`markets[0]` يحصل على {} فيُنتِج صفحة /analyze فارغة.
    الحارس: build_view للمدوّنة القانونية يُنتِج فرع deep_research (لا قالب
    فارغ) رغم markets:[]."""
    import silk_render
    from canonical_netherlands import netherlands_research_blob
    blob = netherlands_research_blob()
    assert blob["markets"] == [], "المدوّنة القانونية يجب أن تحمل markets:[]"
    view = silk_render.build_view(blob)
    assert view.get("deep_research"), "build_view لم يُنتِج فرع deep_research"


def _guard_trap_view_after_persist():
    """TRAP «view attached AFTER persist on one path» — البلوب المخزَّن قد لا
    يحمل view مشتقّاً؛ مسار القراءة يجب أن يعيد بناءه. الحارس: `GET /analyses/{id}`
    يبني view عند غيابه (`found["view"] = _view(found)`)."""
    src = _read("api.py")
    assert 'found["view"] = _view(found)' in src or 'found["view"]=_view(found)' in src, (
        "مسار /analyses/{id} لا يعيد بناء view الغائب")
    assert 'found.setdefault("analysis_id"' in src, (
        "مسار /analyses/{id} لا يضمن analysis_id للبلوبات الأقدم")


# ── حُرّاس تراب باقية (وجود رمز الإصلاح، file:line-anchored) ──────────────────

def _guard_trap_two_sanitizers():
    _needles("silk_reports.py", "_client_assert_clean", "_client_sanitize")()
    _needles("silk_render.py", "_strip_internal_plumbing")()


def _guard_trap_redaction_mangling():
    # الإصلاح الالتفافي: مرحلة التصعيد سُمّيت `..._escalate{attempt}` لتفادي
    # الجزء القصير الذي كان يُشوَّه؛ والمنقِّح `_redact` باقٍ. (بنيوياً مفتوح
    # — لا حارس طول أدنى بعد؛ مُتتبَّع في silk-operations §2.)
    _needles("silk_ai_judge.py", "_escalate{attempt}")()
    _absent("silk_ai_judge.py", "maxtok_retry")()
    _needles("silk_diagnostics.py", "def _redact")()


def _guard_trap_strip_plumbing_three_leaks():
    """TRAP «_strip_internal_plumbing leaked three raw forms» — الحارس يبني
    السلاسل الإنتاجية الحرفية الثلاث ويؤكّد تحييد كلٍّ منها (silk-operations §4:
    القفل بالسلسلة الحرفية). الثلاث: ريبر DataPoint(...) يُحيَّد كاملاً بلا
    نصف-ترجمة؛ JSON مضمَّن بمفاتيح score/summary يُحيَّد؛ رمز حكم بأي حالة أحرف."""
    from silk_render import _strip_internal_plumbing
    # (١) ريبر DataPoint(...) — كامل التحييد، لا نصف-ترجمة، تُستخرَج القيمة
    dp = ("مبنيّ على DataPoint(value='واردات 120 مليون دولار', source='UN "
          "Comtrade', confidence=0.9, note='n', retrieved_at='2026-07-01', "
          "status='')")
    o1 = _strip_internal_plumbing(dp)
    assert "DataPoint(" not in o1 and "confidence=" not in o1, o1
    assert "درجة الثقة=" not in o1, f"نصف-ترجمة: {o1}"
    assert "واردات 120 مليون دولار" in o1, o1
    # (٢) JSON مضمَّن بمفاتيح score/summary
    o2 = _strip_internal_plumbing('التوصية: {"score": 0.72, "summary": "سوق واعد"}')
    assert "{" not in o2 and '"summary"' not in o2, o2
    assert "سوق واعد" in o2, o2
    # (٣) رمز حكم بأي حالة أحرف
    o3 = _strip_internal_plumbing("الحكم go — مراقبة قبل الدخول")
    assert not re.search(r"\bgo\b", o3, re.I), f"رمز حكم خام بقي: {o3}"


def _guard_trap_parallel_cache_window():
    # فخّ معروف غير مُصلَح بعد (نافذة الذاكرة المؤقتة عبر ١٢ بعثة متوازية).
    # الحارس يُبقيه مُتتبَّعاً: آلية التوازي (ThreadPoolExecutor) لا تزال في
    # مُشغّل البعثات، والفخّ موسوم صراحةً «Known, not yet fixed» في المهارة.
    _needles("silk_missions.py", "ThreadPoolExecutor")()
    _needles(".claude/skills/silk-operations/SKILL.md", "not yet fixed")()


# كل مدخلة: Incident(key, source, match, check)
#   source: "LESSONS" (key=رقم الصفّ int) أو "trap" (key=slug، match=جزء من
#   اسم الفخّ العريض في §2). check: callable يُفشِل على عودة الانحدار.


def _guard_datapoint_repr_flexible():
    """LESSONS ١٧ — ريبر DataPoint المختصر/الشاذ كان يمرّ نصف مترجم (هجوم
    المشرف الحي). الحارس: النمط المرن + شبكة الأمان يمسكان كل العائلة."""
    import silk_render as _r
    cases = [
        "DataPoint(value=None, confidence=0.0)",
        "DataPoint(value='12.5', source='comtrade', confidence=0.9, "
        "note='ok (x)', retrieved_at='2026', status='ok')",
        "DataPoint(confidence=0.5, value=None)",
        "قبل DataPoint(value=None, confidence=0.0) بعد",
    ]
    for c in cases:
        out = _r._strip_internal_plumbing(c)
        assert "DataPoint" not in out and "confidence" not in out and \
               "درجة الثقة=" not in out, f"leak: {c!r} -> {out!r}"
    assert _r._strip_internal_plumbing(cases[1]).strip().startswith("12.5")


def _guard_vendor_name_leak():
    """LESSONS ١٨ — اسم مزوّد داخلي (Volza/Explee/…) تسرّب لسطح العميل (بلاغ
    UK الحي). الحارس السلوكي: الأسطر الحرفية المسرّبة (سطر «الخطوة التالية»
    القديم + ترجمة `silk_narrative` التي تُسمّي المزوّد) تُحيَّد فعلياً بمنقِّي
    العميل فلا يبقى اسم مزوّد في مخرَج التصدير؛ وسطر next_step المُولَّد لم
    يعُد يحمل اسم مزوّد أصلاً."""
    from silk_reports import (_client_redact_text, _client_forbidden_hits,
                              _client_sanitize)
    # (١) الأسطر الحرفية المسرّبة من البلاغ الحي — كلٌّ يتسرّب أولاً ثم يُحيَّد.
    leaked = [
        "فعّل خدمة التعميق المدفوعة للتحقق من المستوردين وجهات الاتصال (Volza/Explee)",
        "إكسبلي غير متاح حالياً",
        "فولزا: لا مستوردون بالاسم مرصودون لرمز 0804 في GBR",
        "buyers via Serper and SerpApi, priced by LocalPrice",
        "seasonality from pytrends; risk news from GDELT",
    ]
    for line in leaked:
        assert _client_forbidden_hits(line), (
            f"التهيئة خاطئة — يجب أن يتسرّب أولاً: {line!r}")
        cleaned = _client_redact_text(_client_sanitize(line))
        assert not _client_forbidden_hits(cleaned), (
            f"اسم مزوّد بقي بعد التنقية: {_client_forbidden_hits(cleaned)}")
    # (٢) الحارس الصارم يملك أسماء المزوّدين لاتينيةً وعربيةً معاً — لا يعتمد
    # على المُطهِّر وحده (متغيّر مستقبلي يفلت المُطهِّر يبقى يُرفَع بصوت عالٍ).
    for v in ("Volza", "Explee", "إكسبلي", "فولزا", "LocalPrice", "Serper",
              "SerpApi", "pytrends", "GDELT"):
        hits = _client_forbidden_hits(f"مبنيّ على {v} التجارية")
        assert any(h.startswith("vendor_name") for h in hits), (
            f"اسم مزوّد ليس في قائمة الرفض الصارم (_client_assert_clean): {v}")
    # (٣) سطر next_step المُولَّد لا يحمل اسم مزوّد إطلاقاً.
    import silk_render
    from canonical_netherlands import netherlands_research_blob
    blob = netherlands_research_blob()
    blob["deep_research"]["verdict"] = {"verdict": "GO", "confidence": 0.7,
                                        "ai": {"verdict": "GO"}}
    view = silk_render.build_view(blob)
    nxt = (view.get("deep_research") or {}).get("next_step") or ""
    assert nxt and not _client_forbidden_hits(nxt), (
        f"سطر الخطوة التالية يحمل اسم مزوّد: {nxt!r}")


def _guard_export_format_contract():
    """LESSONS ١٩ — عائلة export-format-contract (بلاغ المُشرِف عند السطر): زرّ
    «تصدير التقرير» الأساسي كان موصولاً بـ`dlReport("docx")` فينزّل Word بينما
    المُسلَّم النهائي للعميل PDF غير قابل للتحرير (§3، اتفاق المالك). الحارس:
    (١) زرّ PDF موصول بـ`dlReport("pdf")` لا docx؛ (٢) `dlReport` يملك فرع pdf
    بامتداد `.pdf` ورسالة 503 عربية صريحة؛ (٣) بطاقة الدردشة المصغّرة تُصدِّر
    PDF؛ (٤) الخادم يخدم report.pdf بنوع application/pdf؛ (٥) صورة النشر
    (Dockerfile) + وظيفة e2e تُثبّتان محرّك التحويل فلا يموت الزرّ حياً."""
    html = _read("web/index.html")
    # (١) الوصلة الأساسية: PDF لا docx (السطر المعطوب الأصلي غائب).
    assert '$("#pdfBtn").addEventListener("click",function(){dlReport("pdf")})' \
        in html, "زرّ PDF غير موصول بـdlReport(\"pdf\")"
    assert '$("#pdfBtn").addEventListener("click",function(){dlReport("docx")})' \
        not in html, "زرّ PDF لا يزال موصولاً بتنزيل docx (البلاغ الأصلي)"
    # زرّ Word ثانوي حاضر (النسخة القابلة للتحرير للمشغّل، لا العميل).
    assert 'id="wordBtn"' in html and \
        '$("#wordBtn").addEventListener("click",function(){dlReport("docx")})' \
        in html, "زرّ Word الثانوي غائب أو غير موصول"
    # (٢) فرع pdf في dlReport: امتداد .pdf + رسالة 503 العربية الصريحة.
    assert 'kind==="pdf"' in html, "dlReport بلا فرع pdf"
    assert '"سِلك_تقرير_"+id+".pdf"' in html, "اسم/امتداد ملف الـPDF خاطئ"
    assert "محرّك التحويل غير متاح — جرّب Word مؤقتاً" in html, \
        "رسالة 503 العربية الصريحة غائبة"
    assert "r.status===503" in html, "فرع pdf لا يعالج 503 صراحةً"
    # (٣) بطاقة الدردشة المصغّرة تُصدِّر PDF لا docx.
    assert 'data-act="pdf"' in html, "بطاقة الدردشة المصغّرة لا تُصدِّر PDF"
    assert 'this.dataset.act==="board"?nav("board"):dlReport("pdf")' in html, \
        "معالج بطاقة الدردشة لا يستدعي dlReport(\"pdf\")"
    # (٤) الخادم يخدم report.pdf بنوع application/pdf.
    api = _read("api.py")
    assert "/analyses/{analysis_id}/report.pdf" in api and \
        'media_type="application/pdf"' in api, "نقطة نهاية report.pdf غائبة/خاطئة"
    # (٥) محرّك التحويل مثبَّت على النشر (Dockerfile) وفي وظيفة e2e — كي يعمل
    # الزرّ حيّاً لا في CI فقط (البند ٦ من أمر العمل).
    dockerfile = _read("Dockerfile")
    assert "libreoffice-writer" in dockerfile, \
        "محرّك تحويل PDF غير مثبَّت في صورة النشر — الزرّ سيموت حياً بـ503"
    e2e = _read(".github/workflows/e2e-live-shape.yml")
    assert "libreoffice-writer" in e2e, \
        "وظيفة e2e لا تثبّت محرّك التحويل — تأكيد %PDF سيفشل"
    # التدفّق يؤكّد توقيع %PDF على المسار الأساسي.
    flow = _read("tests/e2e/live_shape_flow.cjs")
    assert 'pdfBuf[0] === 0x25 && pdfBuf[1] === 0x50' in flow, \
        "تدفّق e2e لا يؤكّد توقيع %PDF لزرّ PDF"


def _guard_world_tier2_no_fabrication():
    """LESSONS ٢٠ — عائلة tier2-fabrication (تصميم الميزة أ، قفل استباقي): توسيع
    الترتيب لكل دول العالم يجب ألّا يختلق قيمة فئة-٢ ولا يفجّر ميزانية كومتريد.
    الحارس (قراءة مصدر + سلوك حيّ): (١) وحدة الترتيب لا تقرأ أيّ CSV محلّي؛
    (٢) الفئة-٢ تحمل الوسم التعاقدي + فجوتَي موقع السعودية/المنافسة معلنتين؛
    (٣) نداء العالم الواحد + التدهور عند نفاد الميزانية موجودان؛ (٤) ملف القفل
    قائم."""
    src = _read("silk_market_ranker.py")
    # (١) لا CSV محلّي في وحدة الترتيب إطلاقاً.
    for forbidden in ("agreements_l1", "demographics_l1", "market_locale",
                      "muslim_share", "requirements_l1"):
        assert forbidden not in src, f"الترتيب يقرأ CSV محلّياً: {forbidden}"
    # (٢) الوسم التعاقدي + الفجوة المعلنة + المسجّل + الصمّام.
    for needle in ('TIER2_LABEL = "تغطية أساسية — بيانات محلية محدودة"',
                   'def _tier2_gather_row', 'status="tier2_gap"',
                   'def _world_markets_enabled', 'def world_import_totals',
                   'def _comtrade_budget_left'):
        assert needle in src, f"علامة إنفاذ الفئة-٢ مفقودة: {needle}"
    # (٣) نداء العالم الواحد (partner=0) مشترك للفئتين + تدهور الميزانية.
    assert 'flow="M", partner=0' in src, "نداء العالم الواحد (partner=0) غائب"
    assert '_comtrade_budget_left()' in src and '_WORLD_BUDGET_RESERVE' in src, \
        "فرع التدهور عند نفاد الميزانية غائب"
    # (٤) ملف القفل قائم بأقفاله السبعة.
    assert _exists("tests/test_world_coverage_tierA.py"), "ملف قفل الميزة أ مفقود"
    lock = _read("tests/test_world_coverage_tierA.py")
    for fn in ("test_tier_separation_and_labels",
               "test_tier2_never_carries_a_local_csv_value",
               "test_tier2_gather_makes_zero_comtrade_calls",
               "test_budget_exhausted_degrades_to_tier1_only",
               "test_ranking_is_deterministic_on_fixture"):
        assert f"def {fn}" in lock, f"قفل الميزة أ مفقود: {fn}"


def _guard_out_of_coverage_thin_study():
    """LESSONS ٢٢ — عائلة out-of-coverage-thin-study (مواصفة المالك، الميزة أ):
    سوقٌ خارج التغطية يجب ألّا يشغّل دراسةً هزيلة بل يُعاد برسالةٍ صادقة ويُسجَّل
    إشارةَ طلب. الحارس (قراءة مصدر): البوّابة + الرسالة الحرفية + التسجيل +
    تسطيح الواجهة + ملف القفل."""
    api = _read("api.py")
    assert "def _market_in_coverage" in api, "دالّة فحص التغطية غائبة"
    assert '"error": "out_of_coverage"' in api, "بوّابة خارج التغطية غائبة"
    assert "هذه السوق خارج التغطية الحالية" in api and \
        "تواصل معنا لإضافتها" in api, "الرسالة الصادقة الحرفية غائبة"
    assert '"out_of_coverage_demand"' in api, "تسجيل إشارة الطلب غائب"
    assert "_world_markets_enabled()" in api, "البوّابة غير مقيّدة بالصمّام"
    html = _read("web/index.html")
    assert "x.message||x.reason||x.error" in html, \
        "الواجهة لا تُسطّح رسالة detail (لن تظهر رسالة خارج التغطية)"
    assert _exists("tests/test_out_of_coverage_guard.py"), "ملف قفل البوّابة مفقود"
    lock = _read("tests/test_out_of_coverage_guard.py")
    for fn in ("test_out_of_coverage_market_returns_honest_message_and_logs_demand",
               "test_tier1_curated_market_is_always_covered",
               "test_flag_off_no_coverage_guard_any_country_works_todays_way"):
        assert f"def {fn}" in lock, f"قفل البوّابة مفقود: {fn}"


def _guard_intake_no_silent_guess():
    """LESSONS ٢١ — عائلة intake-silent-guess (تصميم الميزة ب، قفل استباقي):
    استقبال المنتج من صورة يجب ألّا يختلق اسماً ولا يبدأ تحليلاً قبل تأكيد
    المستخدم، والمحوّل أماميّ معزول عن طبقات التحليل. الحارس (قراءة مصدر):
    (١) عقد عدم الاختلاق (فرع readable/العتبة => تعذّر قراءة صادق)؛ (٢) حدود
    الصورة + التقييس + العزل؛ (٣) القياس (حجز واحد) في نقطة النهاية؛ (٤) المحوّل
    لا يستورد/يستدعي طبقات التحليل؛ (٥) ملف القفل قائم."""
    import ast as _ast
    src = _read("silk_product_intake.py")
    # (١) عقد عدم الاختلاق + الرسالة الموحّدة + العتبة.
    for needle in ('READ_FAILED_MSG = "تعذّرت القراءة — اكتب الاسم يدوياً"',
                   'def _read_failed', 'def intake_image', 'readable',
                   '_MIN_CONFIDENCE', 'def enabled'):
        assert needle in src, f"علامة إنفاذ الاستقبال مفقودة: {needle}"
    # (٢) حدود الصورة + التقييس + العزل.
    for needle in ('MAX_IMAGE_BYTES', 'ALLOWED_MEDIA_TYPES', 'def _decode_and_check',
                   'def _sanitize', 'def _isolate', '_MAGIC'):
        assert needle in src, f"علامة سلامة الصورة مفقودة: {needle}"
    # (٣) القياس — نقطة النهاية تحجز تفعيلة واحدة كأيّ نداء مدفوع.
    api = _read("api.py")
    assert 'def _intake_vision_allowed' in api and \
        'try_reserve_paid_calls(1)' in api, "قياس نداء الرؤية غائب"
    assert '@app.post("/products/intake")' in api, "نقطة نهاية الاستقبال غائبة"
    assert 'intake.enabled()' in api, "صمّام SILK_IMAGE_INTAKE غير مفحوص"
    # (٤) المحوّل أماميّ معزول — لا يستورد أيّ طبقة تحليل، ولا يستدعيها نصّاً.
    tree = _ast.parse(src)
    imported = {n.names[0].name.split(".")[0] for n in _ast.walk(tree)
                if isinstance(n, _ast.Import)}
    imported |= {(n.module or "").split(".")[0] for n in _ast.walk(tree)
                 if isinstance(n, _ast.ImportFrom)}
    forbidden = {"silk_engine", "silk_missions", "silk_market_analyst",
                 "silk_ai_judge", "silk_market_ranker", "correlation",
                 "silk_synthesis", "silk_llm_runtime"}
    assert imported.isdisjoint(forbidden), imported & forbidden
    for banned in ("analyze(", "deep_research(", "write_reviewed_report",
                   "ResearchManager", "rank_markets("):
        assert banned not in src, f"الاستقبال يمسّ مسار التحليل: {banned}"
    # (٥) ملف القفل قائم بأقفاله المركزية.
    assert _exists("tests/test_product_intake_featureB.py"), "ملف قفل الميزة ب مفقود"
    lock = _read("tests/test_product_intake_featureB.py")
    for fn in ("test_low_confidence_or_unreadable_never_fabricates",
               "test_intake_module_imports_no_pipeline_code",
               "test_endpoint_image_call_is_metered_from_the_cap",
               "test_image_validation_rejects_bad_inputs"):
        assert f"def {fn}" in lock, f"قفل الميزة ب مفقود: {fn}"


def _guard_unresolved_hs_silent_spend():
    """LESSONS ٢٣ — حادثة الفيتوتشيني: دراسةٌ مدفوعةٌ بدأت برمز HS غير محسوم.
    الحارس السلوكي: (١) المُصنِّف لا يختلق (منتجٌ مجهول => منتقٍ يدوي، hs6=None،
    ثقة 0.0)؛ (٢) `_validate` يرفض فصلًا مستبعَدًا وما ليس رمزًا (عقد عدم اختلاق)؛
    (٣) بوّابة `unresolved_hs` موجودةٌ وتسبق حجز الدولار في `/research`."""
    import silk_hs_classifier as hsc
    out = hsc.classify("qwxzptvbmzzz منتج لا وجود له", allow_claude=False)
    assert out["hs6"] is None and out["status"] == "manual" and \
        out["confidence"] == 0.0, out
    assert hsc._validate({"hs6": "270900", "confidence": 0.9}) is None  # فصل ٢٧
    assert hsc._validate({"hs6": "زائف", "confidence": 0.9}) is None    # ليس رمزًا
    api = _read("api.py")
    assert "def _require_hs6" in api and '"error": "unresolved_hs"' in api, \
        "بوّابة hs6 الصلبة غائبة"
    assert "def classify_hs" in api and "def _classify_ai_allowed" in api, \
        "نقطة/حارس التصنيف غائبة"
    gate = api.index('"error": "unresolved_hs"')
    reserve = api.index("try_reserve_usd(_expected_usd)")
    assert gate < reserve, "بوّابة hs6 يجب أن تسبق حجز الدولار (لا إنفاق على رمز مجهول)"


def _guard_hardcoded_product_rule():
    """LESSONS ٢٤ — الحارسان (مُصنِّف HS + استشارة بلد المنشأ) قاعدتان مبنيّتان
    على البيانات لا حالتا منتج (نفس عائلة «التمور السعودية»). الحارس: (١) منطقهما
    يخلو من أيّ منتج/ISO/HS من العيّنات، والعتبة config-driven؛ (٢) سلوكيًا القاعدة
    تُعمَّم من ترتيب البيانات — عيّنةٌ مُرقَّعةٌ صناعيّةٌ تُطلق/تصمت بالعتبة."""
    import inspect
    import unittest.mock as _mock
    import silk_hs_classifier as hsc
    import silk_market_ranker as ranker
    blob = inspect.getsource(hsc)
    for fn in (ranker.world_export_totals, ranker.top_world_exporters,
               ranker.is_top_world_exporter, ranker._producer_advisory_topn):
        blob += "\n" + inspect.getsource(fn)
    for tok in ("معكرونة", "pasta", "fettuccine", "تمور", "dates", "olive",
                "عسل", "honey", "ITA", "ESP", "GBR", "ARE",
                "190219", "150910", "080410", "040900"):
        if tok.isascii():
            assert not re.search(r"(?<![A-Za-z0-9])" + re.escape(tok)
                                 + r"(?![A-Za-z0-9])", blob), \
                f"ترميزٌ صلبٌ في منطق الحارس: {tok}"
        else:
            assert tok not in blob, f"ترميزٌ صلبٌ في منطق الحارس: {tok}"
    assert "SILK_PRODUCER_ADVISORY_TOPN" in blob, "العتبة ليست config-driven"

    # سلوكي: القاعدة من البيانات — رموزٌ صناعيّةٌ بحتة (لا اسم حقيقي).
    def _fake(hs_code, year):
        return [{"iso3": c, "m49": "0", "total_usd": 9 - i}
                for i, c in enumerate(["XXA", "XXB", "XXC"])]
    with _mock.patch.object(ranker, "world_export_totals", side_effect=_fake):
        top, _l = ranker.is_top_world_exporter("AAAAAA", "XXA", 2023, 2)
        bot, _l2 = ranker.is_top_world_exporter("AAAAAA", "XXC", 2023, 2)
    assert top is True and bot is False, "القاعدة لا تتبع ترتيب البيانات"


def _guard_wrong_direction_study():
    """LESSONS ٢٥ — عائلة wrong-direction-study (Wave 1.5، A): استشارةُ بلد
    المنشأ تُعمَّم لأشقّائها. الحارس السلوكي: (١) تصدير إلى بلد المنشأ نفسه =>
    self_origin (config-driven عبر env)؛ (٢) فصلٌ مقيَّد من مرجع المالك؛
    (٣) البوّابة في api؛ (٤) صفر ISO/HS مكتوب صلبًا في منطق المطابقة."""
    import silk_prerun as sp
    import os as _os
    old = _os.environ.get("SILK_ORIGIN_ISO3")
    _os.environ["SILK_ORIGIN_ISO3"] = "SAU"
    try:
        assert any(a["kind"] == "self_origin"
                   for a in sp.sibling_advisories("080410", "SAU"))
        assert not any(a["kind"] == "self_origin"
                       for a in sp.sibling_advisories("080410", "ITA"))
    finally:
        if old is None:
            _os.environ.pop("SILK_ORIGIN_ISO3", None)
        else:
            _os.environ["SILK_ORIGIN_ISO3"] = old
    # فصلٌ مقيَّد من المرجع (خنزير في سوقٍ خليجية) — عضوٌ من العائلة.
    assert any(a["kind"] == "restricted_chapter"
               for a in sp.sibling_advisories("020329", "SAU"))
    api = _read("api.py")
    assert '"error": "prerun_advisory"' in api and "advisories_ack" in api, \
        "بوّابة أشقّاء الاستشارة غائبة"
    # صفر رمز HS/دولة مكتوب صلبًا في منطق المطابقة.
    import inspect
    blob = "\n".join(inspect.getsource(fn) for fn in (
        sp.sibling_advisories, sp._restricted_hits))
    assert not re.search(r"(?<!\d)\d{4,6}(?!\d)", blob), "رمز HS صلب في المطابقة"
    assert not re.search(r'"[A-Z]{3}"', blob), "رمز دولة صلب في المطابقة"


def _guard_silent_external_failure():
    """LESSONS ٢٦ — عائلة silent-external-failure (Wave 1.5، C): فشلُ خدمةٍ
    خارجية مُهيَّأة يُعلَن للمشغّل. الحارس السلوكي: (١) record_service_failure
    يكتب صفَّ service_failure؛ (٢) المكشطة تُعلِن فشلها؛ (٣) جدول التدقيق قائم."""
    import silk_ops_log
    import tempfile as _tf
    import unittest.mock as _mock
    with _tf.TemporaryDirectory() as td:
        path = os.path.join(td, "ops.db")
        with _mock.patch.object(silk_ops_log, "_db_path", lambda: path):
            silk_ops_log.record_service_failure("comtrade", "429 rate limited")
            rows = silk_ops_log.last_errors(5, path)
    assert rows and rows[0]["kind"] == "service_failure" and \
        rows[0]["context"]["service"] == "comtrade"
    assert "record_service_failure" in _read("silk_gmaps.py"), \
        "المكشطة لا تُعلِن فشلها للمشغّل"
    assert _exists("docs/EXTERNAL_SERVICES_FAILURE_AUDIT.md"), "جدول التدقيق مفقود"


def _guard_readiness_before_spend():
    """LESSONS ٢٧ — عائلة spend-before-knowing (Wave 1.5، D): لوحةُ الجاهزية
    تعرض كلَّ تدهورٍ قبل الحجز. الحارس: نقطة `/research/readiness` + المُركِّب
    `_readiness_checks` (مع can_run/blocking) قائمان، والصمّام config-driven."""
    api = _read("api.py")
    assert "def _readiness_checks" in api and "def research_readiness" in api, \
        "لوحة الجاهزية (نقطة/مُركِّب) غائبة"
    assert '"/research/readiness"' in api and '"can_run"' in api and \
        '"blocking"' in api, "عقد لوحة الجاهزية غير مكتمل"
    import silk_prerun
    assert hasattr(silk_prerun, "advisories_enabled")


def _guard_leads_table_hygiene():
    """LESSONS ٢٨ — عنقود أوّل PDF: جدولُ روابط العميل نُقِّي عند الحدّ. الحارس
    السلوكي على المدوّنة القانونية (فيتوتشيني): جغرافيا خاطئة/نثر/حشو تُسقَط،
    الصالح يبقى، وسطر الإخلاء بارامتري بالمنتج (لا «التمور السعودية»)."""
    import silk_render
    import silk_reports
    from canonical_fettuccine import fettuccine_research_blob
    md = silk_reports.render_markdown(
        silk_render.build_view(fettuccine_research_blob()))
    seg = md[md.find("قائمة مستوردين"):]
    assert "Pastificio Milano" in seg          # صالح — يبقى
    assert "NutsWorld" not in seg              # جغرافيا أمريكية — يُسقَط
    assert "Italy imports a significant" not in seg   # نثر — يُسقَط
    assert "Anonimo Distribuzione" not in seg  # حشو — يُسقَط
    assert "فيتوتشيني" in seg and "التمور السعودية" not in md


def _guard_report_arabic_shape_a4():
    """LESSONS ٢٩ — العلامة «سِلك» كُسِرت «ِس لك» + الصفحة Letter لا A4. الحارس
    السلوكي: docx يحوي «سلك» متّصلة بلا كسرة، بمقاس A4 (210×297مم)."""
    import silk_render
    import silk_reports
    from canonical_fettuccine import fettuccine_research_blob
    import tempfile
    from docx import Document
    view = silk_render.build_view(fettuccine_research_blob())
    path = silk_reports.render_client_docx(
        view, os.path.join(tempfile.mkdtemp(), "r.docx"))
    doc = Document(path)
    txt = "\n".join(p.text for p in doc.paragraphs)
    for s in doc.sections:
        for hf in (s.header, s.footer):
            txt += "\n" + "\n".join(p.text for p in hf.paragraphs)
    assert "سلك" in txt and "سِلك" not in txt, "العلامة غير آمنة التشكيل"
    sec = doc.sections[0]
    assert abs(sec.page_width.mm - 210) < 1 and abs(sec.page_height.mm - 297) < 1, \
        "الصفحة ليست A4"


def _guard_client_template_no_hardcoded_product():
    """LESSONS ٣٠ — «التمور السعودية» كانت مثبَّتةً في تقرير أيّ منتج (عائلة
    hardcoded-product-rule موسَّعة للقوالب). الحارس: سطر الإخلاء بارامتري بالمنتج
    ولا يحمل اسم منتجٍ مثبَّت."""
    import inspect
    from silk_gmaps import maps_disclaimer, MAPS_DISCLAIMER
    src = inspect.getsource(maps_disclaimer)
    for tok in ("التمور", "dates", "معكرونة", "pasta"):
        assert tok not in src, f"اسم منتجٍ مثبَّت في سطر الإخلاء: {tok}"
    assert "التمور" not in MAPS_DISCLAIMER
    assert "عسل" in maps_disclaimer("عسل")   # يُشتَقّ من المنتج فعلًا


def _guard_analyze_persist_canonical_db():
    """LESSONS ٣١ — نتائج /analyze لم تكن محفوظةً في القاعدة القانونية: المحرّك
    ثبّت `db_path="data/silk.db"` النسبيّ فكتب لقرصٍ لا يقرأ منه أحد (المعرّف «1»
    ثم 404). الحارس السلوكي: مع SILK_DATA_DIR مضبوطًا، `analyze(persist=True)`
    يكتب لقاعدة `_db_path()` نفسها التي يقرأ منها `get_analysis` (بمسار افتراضي)."""
    import importlib
    tmp = tempfile.mkdtemp()
    saved = {k: os.environ.get(k) for k in ("SILK_DATA_DIR", "SILK_DB")}
    try:
        os.environ["SILK_DATA_DIR"] = tmp
        os.environ.pop("SILK_DB", None)
        import silk_engine
        import silk_storage
        importlib.reload(silk_storage)
        importlib.reload(silk_engine)
        # لا شبكة: المحرّك يتدهور لفجوات معلنة لكن الصفّ يُحفَظ ويُقرَأ.
        import unittest.mock as M
        with M.patch("requests.get", side_effect=OSError("blocked")), \
             M.patch("requests.post", side_effect=OSError("blocked")):
            result = silk_engine.analyze("شاي أخضر", persist=True)
        aid = result.get("analysis_id")
        assert aid is not None, "لم يُرفَق analysis_id رغم persist=True"
        assert silk_storage._db_path() == os.path.join(tmp, "silk.db")
        found = silk_storage.get_analysis(aid)   # path=None → _db_path()
        assert found is not None and found.get("product") == "شاي أخضر", (
            "الصفّ غير موجود في القاعدة القانونية — الجذر: كُتب لقرصٍ نسبيّ آخر")
    finally:
        for k, v in saved.items():
            if v is None:
                os.environ.pop(k, None)
            else:
                os.environ[k] = v
        import silk_storage
        importlib.reload(silk_storage)


def _guard_new_source_contracts():
    """LESSONS ٣٢ — مصدرٌ جديد = نفس العقود (فجوة معلنة/ops/مخزَّن/محكوم/نظيف
    الشروط). الحارس السلوكي: (أ) IMF/WTO دون الشبكة => فجوة معلنة None/0.0 لا
    اختلاق؛ (ب) WTO بلا مفتاح => فجوة معلنة بصفر نداء شبكة؛ (ج) سلسلة التراجع
    كلا-الفشلين تُبقي مصدر WITS؛ (د) البوّابة العربية للبنك الدولي تطابق تامّ
    (لا تُحوِّل WITS)؛ (هـ) كل نطاق مُفضَّل بعثته تملك web_search (لا إعداد ميت)."""
    from unittest.mock import patch
    import silk_imf_agent as imf
    import silk_wto_tariff as wto
    import silk_tariffs_agent as tar
    import silk_missions as M
    from silk_data_layer import DataPoint, public_source_url, WORLD_BANK_AR_PORTAL

    # (أ) لا اختلاق دون الشبكة
    with patch("silk_cache.cached_get", return_value=None):
        assert imf.imf_indicator("NLD", "gdp_growth").value is None
    # (ب) WTO بلا مفتاح => صفر نداء شبكة
    saved = {k: os.environ.get(k) for k in ("WTO_TTD_API_KEY", "WTO_API_KEY")}
    try:
        os.environ.pop("WTO_TTD_API_KEY", None)
        os.environ.pop("WTO_API_KEY", None)
        with patch("silk_cache.cached_get") as cg:
            dp = wto.wto_applied_tariff("080410", "NLD")
        cg.assert_not_called()
        assert dp.value is None
    finally:
        for k, v in saved.items():
            if v is None:
                os.environ.pop(k, None)
            else:
                os.environ[k] = v
    # (ج) سلسلة التراجع: كلا الفشلين => مصدر WITS يبقى
    with patch("silk_wto_tariff.wto_applied_tariff",
               return_value=DataPoint(None, "WTO TTD", 0.0, "x")), \
         patch("silk_tariffs_agent.applied_tariff",
               return_value=DataPoint(None, "World Bank WITS", 0.0, "y")):
        dp = tar.tariff_with_fallback("080410", "NLD")
    assert dp.value is None and dp.source == "World Bank WITS"
    # (د) البوّابة العربية للعميل: تطابق تامّ، WITS لا يُحوَّل
    assert public_source_url("World Bank", arabic=True) == WORLD_BANK_AR_PORTAL
    assert public_source_url("World Bank WITS", arabic=True) != WORLD_BANK_AR_PORTAL
    assert public_source_url("World Bank") == "https://data.worldbank.org/"
    # (هـ) لا نطاق مُفضَّل بعثته بلا web_search
    for key in M.PREFERRED_DOMAINS:
        assert "web_search" in M.MISSIONS[key]["allowed_tools"]


def _guard_report_quality_upgrade():
    """LESSONS ٣٢ — إصلاحُ المحرّك لا تحرير التقرير (تدقيق زبدة الفول السوداني/
    اليمن): كل عائلة عيبٍ تحريريّ صارت إنفاذًا حتميًّا. الحارس السلوكي على
    مدوّنة اليمن الإنتاجية الشكل: (١) عقد التأكيد يُعلِّم الرمز الخاطئ ولا
    يُعلِّم الصحيح؛ (٢) الرمز المُعلَّم يُعيد التأطير بملاحظةٍ واحدة + يسقف
    الثقة؛ (٣) شرطا قلب الحكم حقلان مهيكلان."""
    import silk_render as R
    from silk_hs_confirm import confirm_hs, is_flagged, CONTEXTUAL_TAG
    from tools.canonical_yemen import yemen_research_blob
    # (١) عقد التأكيد: الصفة المميّزة لا تخسر أمام كلمة ثانوية عارية.
    assert is_flagged(confirm_hs("زبدة الفول السوداني", "040510"))
    assert confirm_hs("تمور", "080410")["confirmed"] is not False
    # (٢) التأطير + سقف الثقة على المدوّنة، بملاحظةٍ واحدة (لا تكرار).
    dr = R.build_view(yemen_research_blob())["deep_research"]
    assert dr["hs_flagged"] is True
    assert dr["verdict"]["confidence"] <= 0.5
    assert sum(1 for l in dr["limits"] if CONTEXTUAL_TAG in l) == 1
    # (٣) شرطا قلب الحكم المهيكلان (حكم مراقبة).
    assert len(dr["flip_conditions"]) == 2
    assert all(c.get("closes_via") for c in dr["flip_conditions"])


def _guard_parse_provenance_not_prose():
    """LESSONS ٣٣ — حلِّل المصدر لا النثر: قاعدةُ إفصاح التقادُم تُرسى إلى
    بياناتٍ بنيوية. الحارس السلوكي: (١) `fact_year` يقرأ الوسم البنيويّ
    `year=YYYY`/`retrieved_at`؛ (٢) حقيقةٌ متقادِمة تُوسَم بأيّ صياغة؛
    (٣) رمز HS 2008 بلا حقيقة خلفه لا يُوسَم؛ (٤) «الطعام 2013» بلا حقيقة لا
    يُوسَم (لا false-positive نثريّ)."""
    import silk_render as R
    from silk_staleness import fact_year, stale_fact_years, is_stale_fact
    # (١) المصدر البنيويّ.
    assert fact_year({"value": 1, "note": "x year=2013", "retrieved_at": "2026"}) == 2013
    assert fact_year({"value": 1, "retrieved_at": "2018-12-31"}) == 2018
    assert not is_stale_fact({"value": 1, "retrieved_at": "2026-01-01"})
    # (٢) الوسم مستقلّ عن الصياغة.
    for s in ["في 2013 بلغ الدخل.", "عام 2013م.", "الدخل 2013 منخفض."]:
        assert R._STALE_TAG in R._tag_stale_years(s, {2013}), s
    # (٣) رمز HS 2008 لا يُوسَم (ليس سنة حقيقة، وليس في القائمة).
    assert R._STALE_TAG not in R._tag_stale_years("البند 2008 للمحضرات.", {2013})
    # (٤) «الطعام 2013» بلا حقيقة متقادِمة => بلا وسم (لا مطابقة داخل كلمة).
    assert R._STALE_TAG not in R._tag_stale_years("استهلاك الطعام 2013.", set())
    # (٥) القائمة تُشتَقّ من حقائق اليمن (2013/2018).
    from tools.canonical_yemen import yemen_research_blob
    ms = yemen_research_blob()["deep_research"]["missions"]
    allf = [f for v in ms.values() for f in v["findings"]]
    assert stale_fact_years(allf) == {2013, 2018}


def _guard_hs_gate_shared_choke_point_fail_safe():
    """LESSONS ٣٥ — تقرير الكويت الحيّ (زبدة الفول السوداني، 2026-07-21):
    بوّابة تأكيد HS كانت موصولة بـ/research وحده خلف صمّامٍ مُطفأ افتراضياً.
    الحارس السلوكي: (١) `gate_enabled` فشل-آمن — مفعّلة بلا أيّ متغيّر env؛
    (٢) `preflight_block` نقطة اختناق واحدة تحجب رمزاً غير مؤكَّد؛ (٣) كلا
    معالجَي `/analyze` و`/research` في api.py يستدعيانها فعلياً (لا نسخة
    مكرَّرة/مسار واحد فقط)."""
    import silk_hs_confirm as C
    saved = os.environ.pop("SILK_HS_CONFIRM_GATE", None)
    try:
        # (١) فشل-آمن: بلا أيّ ضبط => مفعّلة.
        assert C.gate_enabled() is True
        # إطفاءٌ صريح فقط يُعطّلها.
        os.environ["SILK_HS_CONFIRM_GATE"] = "0"
        assert C.gate_enabled() is False
        os.environ["SILK_HS_CONFIRM_GATE"] = "1"
        assert C.gate_enabled() is True
        del os.environ["SILK_HS_CONFIRM_GATE"]
        # (٢) نقطة الاختناق تحجب فعلياً — نفس عيّنة الحادثة الحية.
        blocked = C.preflight_block("زبدة الفول السوداني", "040510")
        assert blocked is not None and blocked["error"] == "hs_confirmation_needed"
        assert C.preflight_block("زبدة الفول السوداني", "040510",
                                 hs_confirmed=True) is None
    finally:
        if saved is None:
            os.environ.pop("SILK_HS_CONFIRM_GATE", None)
        else:
            os.environ["SILK_HS_CONFIRM_GATE"] = saved
    # (٣) كلا المعالجَين يستدعيان preflight_block — لا مسارٌ واحد فقط.
    api_src = _read("api.py")
    assert api_src.count("preflight_block(") >= 2, (
        "preflight_block يجب أن تُستدعى من كلا /analyze و/research")


def _guard_cross_market_checkpoint_leak():
    """LESSONS ٣٦ — تسرّب اليمن↔الكويت: نقاط تفتيش بعثات `/research` كانت
    تُقرأ بمفتاح analysis_id فقط بلا عمود سوق، واستئنافٌ بسوقٍ مختلف يُعيد
    استهلاكها بصمت. الحارس السلوكي: (١) نقطة تفتيش مختومة بسوقٍ (اليمن) لا
    تُعاد لطلبٍ بسوقٍ آخر (الكويت)؛ (٢) صفوفٌ قديمة بلا ختم لا تُحجَب؛
    (٣) بوّابة `/research`'s resume_market_mismatch (٤٠٩) موجودة في api.py
    **قبل** فرع «مكتملة => أعِدها كما هي» (لا إرجاعٌ صامتٌ يتجاهل الطلب)."""
    import silk_storage
    from silk_agents import AgentReport
    import tempfile as _tf
    db = os.path.join(_tf.mkdtemp(), "silk.db")
    yemen_report = AgentReport(agent_name="x", findings=[], failed=False,
                               summary="سوق عدن المركزي / ربوع")
    silk_storage.save_mission_checkpoint(1, "consumer_culture", yemen_report,
                                         path=db, market_iso3="YEM")
    # (١) طلبٌ بسوق آخر لا يستلم الصفّ.
    assert "consumer_culture" not in silk_storage.load_mission_checkpoints(
        1, path=db, market_iso3="KWT")
    assert "consumer_culture" in silk_storage.load_mission_checkpoints(
        1, path=db, market_iso3="YEM")
    # (٢) صفٌّ قديم بلا ختم (market_iso3=None) لا يُحجَب.
    old_report = AgentReport(agent_name="y", findings=[], failed=False, summary="s")
    silk_storage.save_mission_checkpoint(2, "tradeflow", old_report, path=db)
    assert "tradeflow" in silk_storage.load_mission_checkpoints(
        2, path=db, market_iso3="KWT")
    # (٣) بوّابة API تسبق فرع الإعادة الصامتة لتشغيلةٍ مكتملة.
    api_src = _read("api.py")
    assert "resume_market_mismatch" in api_src
    gate_idx = api_src.index("resume_market_mismatch")
    completed_shortcut_idx = api_src.index(
        'if run_row.get("status") == "completed"')
    assert gate_idx < completed_shortcut_idx, (
        "بوّابة تعارض السوق يجب أن تسبق فرع «مكتملة => أعِدها كما هي»")


def _guard_golden_contract_test_exists_and_covers_both_paths():
    """LESSONS ٣٧ — الاختبار الذهبي موجودٌ فعلياً ويفحص كِلا مسارَي الدخول
    على نفس سيناريو الحادثة (زبدة الفول السوداني/الكويت)، لا مساراً واحداً."""
    assert _exists("tools/canonical_kuwait_peanut_butter.py")
    assert _exists("tests/test_golden_deep_research_contract.py")
    golden_src = _read("tests/test_golden_deep_research_contract.py")
    assert '"/analyze"' in golden_src and '"/research"' in golden_src
    assert "resume_market_mismatch" in golden_src
    smoke_src = _read("tools/post_deploy_smoke.py")
    assert "hs_confirmation_needed" in smoke_src, (
        "فحص الدخان بعد النشر يجب أن يثبت بوّابة HS حياً (Wave 3.2)")


_LESSONS = {
    1: _needles("docs/LIVE_PROOF_RUNBOOK.md", "لا يُشغَّل هيرمتياً"),
    2: _needles("silk_render.py", "_deep_research_view"),
    3: _guard_docx501_row3,          # docx-501 (١)
    4: _needles("api.py", "SILK_REQUIRE_PERSISTENT_DATA_DIR",
                "SILK_DATA_DIR غير مضبوط"),
    5: _needles("silk_storage.py", "def create_research_run",
                "def load_mission_checkpoints"),
    6: _needles("silk_llm_runtime.py", "_JSON_PARSE_FAILURE_GAP"),
    7: _needles("silk_data_layer.py", "_WB_INDICATOR_SOURCE"),
    8: _needles("silk_data_layer.py", "class DataPoint"),
    9: _absent("web/index.html", 'id="snapBtn"'),
    10: _needles("docs/AUDIT_STATUS.md", "قراءة فقط", "غير موجود"),
    11: _guard_docx501_row11,         # docx-501 (٣)
    12: _needles("silk_render.py", "_reconcile_mission_limits", "_first_clause"),
    13: _guard_docx501_row13,         # docx-501 (٢، الفشل الحيّ)
    14: _needles("silk_quality_gate.py", "_check_confidentiality_leaks",
                 "_check_style"),
    15: _needles("tools/live_shape_server.py", "class LiveShapeServer",
                 "def seed_db"),
    16: _needles("silk_ai_judge.py", "_WRITER_MAX_TOKENS", "_MAX_TOKENS_CEILING",
                 "max_tokens=_MAX_TOKENS_CEILING"),
    17: _guard_datapoint_repr_flexible,  # هجوم المشرف — ريبر DataPoint المرن
    18: _guard_vendor_name_leak,         # بلاغ UK — تسريب اسم مزوّد للعميل
    19: _guard_export_format_contract,   # بلاغ المُشرِف — زرّ PDF كان ينزّل docx
    20: _guard_world_tier2_no_fabrication,  # الميزة أ — لا تلفيق فئة-٢/تفجّر ميزانية
    21: _guard_intake_no_silent_guess,      # الميزة ب — لا اختلاق منتج من صورة
    22: _guard_out_of_coverage_thin_study,  # الميزة أ — سوق خارج التغطية لا دراسة هزيلة
    23: _guard_unresolved_hs_silent_spend,  # Wave 1 — الفيتوتشيني: لا إنفاق برمز HS مجهول
    24: _guard_hardcoded_product_rule,      # Wave 1 — الحارسان قاعدتان مبنيّتان على البيانات
    25: _guard_wrong_direction_study,       # Wave 1.5 A — أشقّاء «الدراسة بالاتجاه الخاطئ»
    26: _guard_silent_external_failure,     # Wave 1.5 C — لا فشلٌ صامت لخدمةٍ خارجية
    27: _guard_readiness_before_spend,      # Wave 1.5 D — كلُّ تدهورٍ قبل الحجز
    28: _guard_leads_table_hygiene,         # Wave 2 — نقاء جدول الروابط (جغرافيا/نثر/حشو)
    29: _guard_report_arabic_shape_a4,      # Wave 2 — «سلك» متّصلة + A4
    30: _guard_client_template_no_hardcoded_product,  # Wave 2 — لا منتج مثبَّت في القوالب
    31: _guard_analyze_persist_canonical_db,   # /analyze — التخزين للقاعدة القانونية لا قرصٍ نسبيّ فانٍ
    32: _guard_report_quality_upgrade,         # ترقية جودة التقرير — إصلاحُ المحرّك لا تحرير التقرير
    33: _guard_parse_provenance_not_prose,     # التقادُم من المصدر لا النثر (قرار المالك)
    34: _guard_new_source_contracts,           # دمج مصادر جديدة — نفس العقود (فجوة/ops/مخزَّن/محكوم/نظيف)
    35: _guard_hs_gate_shared_choke_point_fail_safe,  # تقرير الكويت — بوّابة HS فشل-آمن + نقطة اختناق مشتركة
    36: _guard_cross_market_checkpoint_leak,          # تقرير الكويت — تسرّب يمن↔كويت عبر نقاط تفتيش بعثات
    37: _guard_golden_contract_test_exists_and_covers_both_paths,  # الاختبار الذهبي — كل العقود، كلا المسارين

}

_TRAPS = [
    ("mock_passes_real_fails", "Mock-passes / real-fails",
     _needles("silk_render.py", "_deep_research_view")),
    ("markets_empty_misroute", "misroutes exporters",
     _guard_trap_markets_empty),
    ("two_sanitizers", "Two different sanitizers",
     _guard_trap_two_sanitizers),
    ("redaction_mangling", "Redaction mangling",
     _guard_trap_redaction_mangling),
    ("parallel_cache_window", "Parallel missions and the cache window",
     _guard_trap_parallel_cache_window),
    ("cap_counts_not_dollars", "Cap counted operations, not dollars",
     _needles("silk_usage.py", "def try_reserve_usd")),
    ("styled_not_wired", "Styled-but-never-wired UI affordance",
     _needles("web/index.html", 'data-id="',
              '$("#histList").addEventListener("click"')),
    ("silent_noop_family", "silent no-op has three forms",
     _needles("web/index.html", "function openStoredAnalysis",
              # حزمة الإغلاق، البند ٤: سلسلة /markets في بناء نيّة الدردشة
              # اكتسبت .catch عربياً (كانت ترفض صامتةً فتُعلّق مؤشّر الانتظار).
              "تعذّر تحميل قائمة الأسواق")),
    ("view_after_persist", "view attached AFTER persist",
     _guard_trap_view_after_persist),
    ("strip_plumbing_three_leaks", "leaked three raw forms",
     _guard_trap_strip_plumbing_three_leaks),
    ("orphan_reservation_leak", "Orphaned runs leak their USD reservation",
     lambda: (
         _needles("silk_storage.py", "def reap_orphan_research_runs",
                  "reconcile_usd", "SILK_ORPHAN_STALE_MINUTES")(),
         _needles("api.py", "reap_orphan_research_runs")(),
         _needles("silk_collectors.py", "reap_orphan_research_runs")())),
]


# ── حارس واحد لكل حادثة (تُوسَّع برمجياً لتقارير pytest واضحة) ────────────────

@pytest.mark.parametrize("row", sorted(_LESSONS), ids=[f"lessons-{n}" for n in sorted(_LESSONS)])
def test_lessons_incident_guard_holds(row):
    """كل صفّ في docs/LESSONS.md له حارس حيّ يُفشِل على عودة عائلته."""
    _LESSONS[row]()


@pytest.mark.parametrize("slug,match,check", _TRAPS,
                         ids=[t[0] for t in _TRAPS])
def test_operations_trap_guard_holds(slug, match, check):
    """كل فخّ في silk-operations §2 (THE TRAPS) له حارس حيّ."""
    check()


# ── الاختبار الشامل: التغطية كاملة ضدّ كلا السِّجلّين ─────────────────────────

def _lessons_rows_in_ledger() -> set[int]:
    ledger = _read("docs/LESSONS.md")
    return {int(m.group(1))
            for m in re.finditer(r"^\|\s*(\d+)\s*\|", ledger, re.M)}


def _trap_rows_in_skill() -> list[str]:
    """صفوف بيانات جدول §2 (THE TRAPS) — الخلية الأولى (اسم الفخّ العريض)."""
    skill = _read(".claude/skills/silk-operations/SKILL.md")
    m = re.search(r"## 2\. THE TRAPS(.*?)\n## 3\.", skill, re.S)
    assert m, "قسم §2 THE TRAPS غير موجود في مهارة silk-operations"
    rows = []
    for line in m.group(1).splitlines():
        line = line.strip()
        if not line.startswith("| **"):        # صفوف البيانات فقط
            continue
        first_cell = line.split("|")[1].strip()
        rows.append(first_cell)
    return rows


def test_meta_registry_covers_every_known_incident():
    """التغطية الشاملة: كل صفّ LESSONS وكل فخّ §2 مُسجَّل هنا بحارس، ولا مدخلة
    يتيمة بلا صفّ مقابل. حادثة جديدة تسقط في أي سِجلّ بلا حارس تُحمِّر هنا."""
    # (أ) LESSONS: مفاتيح السجلّ = أرقام الصفوف بالضبط، متتابعة ١..N.
    ledger_rows = _lessons_rows_in_ledger()
    assert ledger_rows == set(range(1, max(ledger_rows) + 1)), (
        f"أرقام صفوف LESSONS غير متتابعة: {sorted(ledger_rows)}")
    registry_rows = set(_LESSONS)
    assert registry_rows == ledger_rows, (
        f"صفوف LESSONS بلا حارس في السجلّ: {sorted(ledger_rows - registry_rows)}؛ "
        f"حُرّاس بلا صفّ: {sorted(registry_rows - ledger_rows)}")

    # (ب) TRAPS: كل صفّ فخّ في §2 يطابقه حارس واحد بالضبط عبر إبرة `match`،
    # وكل حارس فخّ يطابق صفّاً واحداً على الأقل (لا يتيم).
    trap_rows = _trap_rows_in_skill()
    assert trap_rows, "لم تُقرَأ صفوف فخاخ من المهارة"
    for row in trap_rows:
        matched = [slug for slug, match, _ in _TRAPS if match in row]
        assert len(matched) == 1, (
            f"صفّ الفخّ «{row[:60]}…» يطابقه {len(matched)} حُرّاس (المتوقّع ١): "
            f"{matched}")
    for slug, match, _ in _TRAPS:
        assert any(match in row for row in trap_rows), (
            f"حارس الفخّ «{slug}» (match={match!r}) لا يطابق أيّ صفّ في §2 — "
            "يتيم؛ حدِّث السجلّ أو المهارة")


def test_meta_docx501_trio_all_have_behavioral_guards():
    """الحوادث الثلاث لعطل docx-501 (LESSONS ٣/١١/١٣) لها حُرّاس **سلوكية**
    (تبني/تنقّي فعلياً)، لا مجرّد وجود رمز — أمر المُشرِف الصريح."""
    behavioral = {3: _guard_docx501_row3, 11: _guard_docx501_row11,
                  13: _guard_docx501_row13}
    for row, guard in behavioral.items():
        assert _LESSONS[row] is guard, (
            f"صفّ docx-501 رقم {row} ليس مربوطاً بحارسه السلوكي")
        guard()  # يجب أن يمرّ فعلياً الآن
