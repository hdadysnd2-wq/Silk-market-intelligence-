"""Master Prompt Part 2 §B — بوابة اتساق الحكم عند التسليم.

الحكم حقلٌ واحدٌ (`verdict_tone`) يجب أن تشتقّ منه شارة الغلاف وصفّ الجدول
وسطر «الحكم:»/«التوصية:» في قسم القرار معاً — لا نصّاً مستقلاً في كل موضع.
هذا الملف يقفل الآلية (`silk_reports._assert_verdict_consistency_doc`/
`_assert_verdict_consistency_text`) على حالتين: (أ) تعارضٌ مصطنَع تلتقطه
الدالة مباشرةً (وحدة)، و(ب) المدوّنة القانونية الحقيقية (Kuwait/Netherlands)
تمرّ بلا رفعٍ — الإصلاح لا يكسر مساراً حياً صحيحاً.

Run: python3 -m pytest tests/test_master_prompt_part2_verdict.py -q
"""
from __future__ import annotations

import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pytest

from tools.canonical_kuwait_peanut_butter import kuwait_research_blob  # noqa: E402


def test_declared_verdict_labels_catches_badge_table_mismatch():
    """وحدة: شارة تقول «مراقبة السوق» بينما صفّ الجدول يقول «التوصية
    بالدخول» — يجب أن يُرفَع RuntimeError يسمّي كلا التسميتين."""
    pytest.importorskip("docx")
    from docx import Document
    from silk_reports import _assert_verdict_consistency_doc

    doc = Document()
    doc.add_paragraph("  مراقبة السوق  ")  # شارة الغلاف
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "الحكم"
    table.rows[0].cells[1].text = "التوصية بالدخول"  # صفّ جدول متعارض

    with pytest.raises(RuntimeError, match="تناقض حكمٍ"):
        _assert_verdict_consistency_doc(doc, "WATCH", "اختبار")


def test_declared_verdict_labels_passes_when_all_sites_agree():
    """لا رفع حين تتّفق كل المواضع الثلاثة على نفس الحكم."""
    pytest.importorskip("docx")
    from docx import Document
    from silk_reports import _assert_verdict_consistency_doc

    doc = Document()
    doc.add_paragraph("  مراقبة السوق  ")
    doc.add_paragraph("الحكم: مراقبة السوق")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "التوصية"
    table.rows[0].cells[1].text = "مراقبة السوق"

    _assert_verdict_consistency_doc(doc, "WATCH", "اختبار")  # لا استثناء


def test_flip_condition_narrative_does_not_false_positive():
    """نقاش «شرطا قلب الحكم» (LESSONS ٣٢) قد يذكر تسمية حكمٍ أخرى ضمن فقرةٍ
    حرّة («لو تحقّق كذا لتحوّل الحكم إلى دخول مشروط») — هذا **ليس** تناقضاً؛
    الفحص تصريحيّ محض (شارة/جدول/سطر «الحكم:») لا مسحٌ لكامل السرد."""
    pytest.importorskip("docx")
    from docx import Document
    from silk_reports import _assert_verdict_consistency_doc

    doc = Document()
    doc.add_paragraph("  مراقبة السوق  ")
    doc.add_paragraph("الحكم: مراقبة السوق")
    doc.add_paragraph(
        "شرطا قلب الحكم: لو تأكّد موزّع محلي وبيانات استيراد دقيقة، لتحوّل "
        "الحكم إلى دخول مشروط خلال ٩٠ يوماً.")

    _assert_verdict_consistency_doc(doc, "WATCH", "اختبار")  # لا استثناء


def test_kuwait_client_and_research_docx_pass_verdict_gate(monkeypatch):
    """المدوّنة القانونية (زبدة الفول السوداني/الكويت، حكم WATCH) تُصدَّر
    docx للعميل وللمشغّل معاً بلا تناقض حكمٍ — الإصلاح لا يكسر مساراً حياً."""
    pytest.importorskip("docx")
    from silk_render import build_view
    from silk_reports import render_docx, render_client_docx
    monkeypatch.setenv("SILK_HERMETIC", "1")
    view = build_view(kuwait_research_blob())
    tmp = tempfile.mkdtemp()
    render_docx(view, os.path.join(tmp, "kwt_research.docx"))
    render_client_docx(view, os.path.join(tmp, "kwt_client.docx"))


def test_markdown_verdict_consistency_catches_mismatch():
    """معادل Markdown: صفّ جدول «| الحكم |» يخالف سطر «- التوصية:»."""
    from silk_reports import _assert_verdict_consistency_text
    blob = "\n".join([
        "| الحكم | مراقبة السوق |",
        "## الحكم وأساسه",
        "- التوصية: **التوصية بالدخول**",
    ])
    with pytest.raises(RuntimeError, match="تناقض حكمٍ"):
        _assert_verdict_consistency_text(blob, "WATCH", "اختبار")


def test_markdown_verdict_consistency_passes_kuwait(monkeypatch):
    from silk_render import build_view
    from silk_reports import render_markdown
    monkeypatch.setenv("SILK_HERMETIC", "1")
    view = build_view(kuwait_research_blob())
    render_markdown(view)  # لا استثناء


def test_verdict_tone_recognizes_arabic_labels_not_only_english_codes():
    """LESSONS ٤٤: بعض مسارات الحكم (نداء كلود المرحلة الثانية، أو مدوّناتٌ
    يضبطها المتصل) تضع التسمية العربية مباشرةً في `ai["verdict"]` بدل الرمز
    الإنجليزي (`"دخول مشروط"` لا `"CONDITIONAL-GO"`) — هذه المدخلة كانت
    تنهار إلى tone="unknown" فتعرض الشارة «تعذّر إصدار توصية» بينما الجدول/
    المتن يذكران التسمية العربية الصحيحة، وهو بالضبط تناقض الشارة/المتن
    الذي صُمِّمت `_verdict_tone` أصلاً لمنعه. اكتُشف عبر بوابة اتساق الحكم
    الجديدة (§B) على مدوّنة اختبارٍ حقيقية الشكل (`test_client_report_export.py`)."""
    from silk_render import _verdict_tone
    assert _verdict_tone("دخول مشروط") == "conditional"
    assert _verdict_tone("مراقبة السوق") == "watch"
    assert _verdict_tone("عدم الدخول حالياً") == "nogo"
    assert _verdict_tone("التوصية بالدخول") == "go"
    # الرموز الإنجليزية تبقى كما كانت (لا انحدار).
    assert _verdict_tone("CONDITIONAL-GO") == "conditional"
    assert _verdict_tone("WATCH") == "watch"
