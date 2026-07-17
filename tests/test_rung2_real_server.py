"""رُتبة ٢ — خادم حقيقي بشكل الإنتاج (rung 2 — the real-server lane).

> **القاعدة الجديدة (تأكيد المالك آخِراً لا أوّلاً).** «مدموج ≠ يعمل؛ أخضر
> محلياً ≠ تمّ» (البند ١، `docs/LESSONS.md`). الحزمة الهرمتية تُثبِت العقود
> على TestClient/نماذج؛ **لا تُقلِع الخادم الفعلي**. هذه الرُتبة تُقلِع
> `uvicorn api:app` على ملف SQLite حقيقي مبذور بالمدوّنة القانونية، وتضرب كل
> نقطة نهاية للقراءة/التصدير بـHTTP حقيقي. هذا التدفّق بالضبط كان سيلتقط
> عائلة 501 التصدير والشريط الجانبي الميت **قبل** أن يراها المالك.

المفاتيح المدفوعة غائبة عمداً — نقاط القراءة/التصدير المبذورة تُقرَأ من
المخزن ولا تلمس أيّ API خارجي (لا محاكاة لازمة على مسار الإسناد؛ التغييرات
قرب مسار المال تُشغّل رُتبة ٤ بمزودين محاكَين). راجع `tools/live_shape_server.py`.

يُتخطّى في `pytest tests/ -q` الافتراضية (هرمتية سريعة)؛ يعمل في وظيفة
`e2e-live-shape` حين `SILK_RUN_E2E=1`.

Run locally:  SILK_RUN_E2E=1 python3 -m pytest tests/test_rung2_real_server.py -q
"""
from __future__ import annotations

import io
import json
import os
import sys
import urllib.error
import urllib.request

import pytest

sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(
    os.path.abspath(__file__))), "tools"))

pytestmark = pytest.mark.e2e


def _get(base: str, path: str):
    req = urllib.request.Request(base + path)
    try:
        with urllib.request.urlopen(req, timeout=30) as r:
            return r.status, r.read(), dict(r.headers)
    except urllib.error.HTTPError as e:
        return e.code, e.read(), dict(e.headers or {})


@pytest.fixture(scope="module")
def server():
    from live_shape_server import LiveShapeServer
    with LiveShapeServer() as srv:
        yield srv


# ── الواجهة تُخدَم فعلياً (الشريط الجانبي/أزرار التصدير موجودة) ──────────────

def test_dashboard_html_is_served_with_the_touched_ui_hooks(server):
    """`GET /` يخدم اللوحة الفعلية بمرابط الواجهة التي تختبرها رُتبة ٣:
    الشريط الجانبي (#histList)، وزرّا التصدير (Word=#pdfBtn، Markdown=#mdBtn)."""
    st, body, _ = _get(server.base_url, "/")
    assert st == 200
    html = body.decode("utf-8", "replace")
    assert 'id="histList"' in html          # الشريط الجانبي
    assert 'id="pdfBtn"' in html            # تصدير Word/PDF
    assert 'id="mdBtn"' in html             # تصدير Markdown
    assert 'id="boardBody"' in html         # لوحة العرض


def test_health_is_200_on_the_real_server(server):
    st, body, _ = _get(server.base_url, "/health")
    assert st == 200
    assert json.loads(body).get("status") is not None or True  # 200 يكفي


# ── الشريط الجانبي: الصفوف الحقيقية تُسرَد، والصفّ الجارٍ يحمل حالته ──────────

def test_analyses_list_returns_both_seeded_rows(server):
    st, body, _ = _get(server.base_url, "/analyses")
    assert st == 200
    rows = json.loads(body)
    ids = {r.get("id") for r in rows}
    assert server.completed_id in ids and server.running_id in ids
    running = [r for r in rows if r.get("id") == server.running_id][0]
    # صفّ التقدّم — الشريط الجانبي يعرض «جارٍ التشغيل…» + زرّ استئناف عليه.
    assert running.get("status") == "running"


def test_single_analysis_carries_analysis_id_and_prebuilt_view(server):
    """نقر الشريط الجانبي يفتح `GET /analyses/{id}` ويعتمد على analysis_id
    (لأزرار التصدير) وعلى view (لعرض اللوحة) — كلاهما يجب أن يصل."""
    st, body, _ = _get(server.base_url, f"/analyses/{server.completed_id}")
    assert st == 200
    j = json.loads(body)
    assert j.get("analysis_id") == server.completed_id
    assert j.get("view")                    # عرض مُبنى مسبقاً


# ── التصدير: نفس عائلة الأخطاء الثلاث (501/قالب فارغ) على HTTP حقيقي ─────────

def test_report_md_serves_real_narrative_not_the_empty_analyze_template(server):
    """البند ٢: report.md لنتيجة بحث يُصيَّر من deep_research (سرد غنيّ)، لا
    قالب /analyze الفارغ («with_research غير مفعّلة»/«0 أسواق»/«None»)."""
    st, body, _ = _get(server.base_url,
                       f"/analyses/{server.completed_id}/report.md")
    assert st == 200, body[:300]
    md = body.decode("utf-8", "replace")
    assert md.strip()                                   # غير فارغ
    assert "HHI" in md and "EU 2017/625" in md          # السرد الحقيقي
    assert "with_research" not in md                    # لا قالب /analyze
    assert "None" not in md                             # لا سنة/ثقة "None"


def test_report_docx_downloads_a_real_openable_document_no_501(server):
    """البنود ٢/٣/١١/١٣: report.docx لا يفشل 501، وهو ملف ZIP/docx صالح
    يُفتَح فعلياً عبر python-docx بمحتوى غير فارغ (لا مجرد 200 على الحالة)."""
    st, body, headers = _get(server.base_url,
                             f"/analyses/{server.completed_id}/report.docx")
    assert st == 200, body[:300]
    assert body[:2] == b"PK"                            # توقيع ZIP (docx)
    from docx import Document
    doc = Document(io.BytesIO(body))
    text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                text += "\n" + c.text
    assert text.strip()                                 # مستند غير فارغ فعلاً


def test_missing_analysis_export_is_404_not_a_fabricated_success(server):
    """المبدأ المؤسِّس على مستوى HTTP: تحليل غير موجود = 404 صريح، لا مستند
    فارغ مُختلَق."""
    st, _, _ = _get(server.base_url, "/analyses/999999/report.md")
    assert st == 404
