"""Wave 3 §6 — سجل الأدلة يحمل رابطًا عموميًّا **حقيقيًّا** من سجلّ المصادر.

أمر العمل الرئيس §6: لكل حقيقةٍ اسمُ مصدرها العمومي وتاريخُ جمعها ورابطُها —
من طبقة البيانات لا بتخمينٍ عند التصيير. كان عمودُ «الرابط» يُكشَط من النصّ
فيعود «—» لكلّ حقائق Comtrade/World Bank الحقيقية؛ الآن رابطٌ رسميٌّ من
`silk_data_layer.SOURCE_PUBLIC_URL`. لا اختلاق: مصدرٌ مجهول/مدفوع => «—».

هرمتيّ بالكامل — لا شبكة، لا مفتاح.
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pytest  # noqa: E402

pytest.importorskip("docx")

import silk_data_layer as dl  # noqa: E402
import silk_reports as sr  # noqa: E402


# ── وحدة السجلّ: public_source_url ─────────────────────────────────────────
def test_public_source_url_real_for_known_public_sources():
    """المصادر العمومية المعروفة تُصيب روابطها الرسمية الحقيقية — حتى مع
    اللواحق العربية بين قوسين (كما تظهر فعلاً في DataPoint.source)."""
    cases = {
        "UN Comtrade (مخزن الحقائق)": "https://comtradeplus.un.org/",
        "UN Comtrade (تقرير سعودي مباشر — مرآة)": "https://comtradeplus.un.org/",
        "World Bank (لقطة مضمّنة)": "https://data.worldbank.org/",
        "Google Trends (تكميلية — بحث لا شراء)": "https://trends.google.com/trends/",
        "OpenAlex": "https://openalex.org/",
        "FAOSTAT": "https://www.fao.org/faostat/en/",
        "Eurostat (مسح ميزانية الأسرة)": "https://ec.europa.eu/eurostat/",
        "GDELT": "https://www.gdeltproject.org/",
    }
    for label, url in cases.items():
        got = dl.public_source_url(label)
        assert got == url, (label, got, url)
        assert got.startswith("https://")


def test_public_source_url_empty_for_unknown_no_fabrication():
    """لا اختلاق: مصدرٌ مدفوع/بحثٌ/فارغ/مجهول => «» (لا رابط مخترَع)."""
    for label in ("Volza", "Explee", "Web Search (Serper)", "Local retail",
                  "", None, "مصدر غير معروف", "Some Vendor Ltd"):
        assert dl.public_source_url(label) == ""


# ── وحدة المتصل: _evidence_url (مصدرٌ واحدٌ للأولوية) ───────────────────────
def test_evidence_url_prefers_specific_scraped_url():
    """رابطٌ محدّدٌ في نصّ الحقيقة (نتيجة بحث) يفوز على رابط السجلّ العامّ."""
    got = sr._evidence_url("انظر https://example.com/report", "World Bank", "")
    assert got == "https://example.com/report"


def test_evidence_url_falls_back_to_registry_homepage():
    """بلا رابطٍ محدّد، الرابطُ الرسميُّ للمصدر المسمّى من السجلّ."""
    assert sr._evidence_url("واردات 2023", "UN Comtrade", "") == \
        "https://comtradeplus.un.org/"


def test_evidence_url_dash_when_no_specific_and_no_public_source():
    """لا رابطَ محدّدٌ ولا مصدرٌ عموميّ => «—» صادقة، لا اختلاق."""
    assert sr._evidence_url("قيمة", "Volza", "") == "—"
    assert sr._evidence_url("", "", "") == "—"


# ── تكامل: سجل الأدلة (المدقّق + العميل) يعرض رابطًا حقيقيًّا ────────────────
def _dr_with_comtrade_fact():
    return {"missions": {"trade_flow": {"findings": [{
        "value": "بلغت واردات السوق 61 مليون دولار في 2023",
        "source": "UN Comtrade (مخزن الحقائق)",
        "note": "قيمة الواردات السنوية",
        "retrieved_at": "2026-07-01", "confidence": 0.9}]}}}


def _tables_text(doc):
    out = []
    for t in doc.tables:
        hdr = "|".join(c.text for c in t.rows[0].cells)
        body = ["|".join(c.text for c in r.cells) for r in t.rows[1:]]
        out.append((hdr, body))
    return out


def test_auditor_evidence_log_has_real_url_for_comtrade_fact():
    """ملحق المدقّق (التقني): عمود «الرابط» يحمل رابط Comtrade الرسمي الحقيقي."""
    from docx import Document
    doc = Document()
    sr._docx_technical_appendix(doc, _dr_with_comtrade_fact())
    tables = [t for t in _tables_text(doc)
              if "الحقيقة" in t[0] and "الرابط" in t[0]]
    assert tables, "لا جدول سجل أدلة بعمود رابط"
    joined = "\n".join("\n".join(b) for _, b in tables)
    assert "comtradeplus.un.org" in joined, joined


def test_client_references_section_has_real_url_for_comtrade_fact():
    """§A (حزمة الفكس v2.1): «المراجع» تحلّ محلّ جدول سجل الأدلة القديم —
    مصدر عمومي معروف يظهر باسمه ورابطه الرسمي الحقيقي."""
    from docx import Document
    doc = Document()
    sr._client_references_section(doc, _dr_with_comtrade_fact())
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "المراجع" in text
    assert "UN Comtrade" in text
    assert "comtradeplus.un.org" in text


def test_client_references_paid_source_dropped_not_leaked():
    """لا اختلاق ولا تسريب: مصدرٌ مدفوع بلا رابط عمومي حقيقي لا يظهر في
    «المراجع» إطلاقاً (لا اسم مورّد، لا رابط مخترَع) — يبقى في الملحق
    الداخلي فقط (§A-2/٥)."""
    from docx import Document
    dr = {"missions": {"buyers": {"findings": [{
        "value": "ثلاثة مستوردين نشطين في السوق",
        "source": "Volza", "note": "لا رابط عموميّ",
        "retrieved_at": "2026-07-01", "confidence": 0.6}]}}}
    doc = Document()
    sr._client_references_section(doc, dr)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "volza" not in text.lower()
    assert "لا مصادر عمومية موثّقة" in text
