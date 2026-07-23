"""اختبارات الموجة ١١ (V5) — هوية سِلك البصرية + تصليب الوكلاء الضعيفين +
جاهزية الإصدار.

يغطي: ١١.١ (قالب سِلك البصري في silk_reports.py)، ١١.٢أ (تصليب وكيل
المنافسين)، ١١.٢ب (تصليب وكيل المخاطر/الأخبار)، ١١.٣ (فحص جاهزية الإصدار).
Run:  python3 -m pytest tests/test_wave11_identity_and_hardening.py -q
"""
import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from conftest import docx_all_text  # noqa: E402


# ── ١١.١: هوية سِلك البصرية ────────────────────────────────────────────

def test_branding_config_file_exists_and_loads_defaults():
    import silk_reports
    b = silk_reports._load_branding()
    assert b["primary_color"]
    assert b["contact_footer"]


def test_branding_config_missing_file_falls_back_to_defaults():
    import silk_reports
    b = silk_reports._load_branding("no/such/path.yaml")
    assert b == silk_reports._BRANDING_DEFAULTS


def test_hex_to_rgbcolor_parses_and_degrades_gracefully():
    import silk_reports
    rgb = silk_reports._hex_to_rgbcolor("1B3B6F")
    assert bytes(rgb) == bytes([0x1B, 0x3B, 0x6F])
    bad = silk_reports._hex_to_rgbcolor("not-a-color")
    assert bad is not None  # لا استثناء — رجوع للكحلي الافتراضي


def _deep_research_result():
    from silk_agents import AgentReport
    from silk_data_layer import DataPoint
    from silk_market_resolver import resolve_market

    ref, _ = resolve_market("Spain")
    analyst_report = AgentReport(
        "LLMAgent:market_analyst",
        [DataPoint("طلب استدلالي", "x", 0.6, "[demand] ...")], False, "تحليل")
    report_text = (
        "## 1. الخلاصة التنفيذية\nنص.\n"
        "## 6. المشهد التنافسي\n"
        "| الدولة | الحصة |\n| --- | --- |\n| فرنسا | 40% |\n| المغرب | 30% |\n"
    )
    return {
        "product": "تمور", "hs_code": "080410", "year": None,
        "market": {"iso3": ref.iso3, "m49": ref.m49, "iso2": ref.iso2,
                  "name_en": ref.name_en, "name_ar": ref.name_ar},
        "markets": [],
        "deep_research": {
            "missions": {
                "trade_flow": AgentReport(
                    "LLMAgent:trade_flow",
                    [DataPoint(950000.0, "UN Comtrade", 0.9, "n")], False, "ok")},
            "analyst": {"report": analyst_report,
                       "by_category": {"demand": analyst_report.findings,
                                      "entry_cost": [], "price_competitiveness": [],
                                      "entry_door": [], "swot": []},
                       "missing_categories": ["entry_cost"]},
            # WP-1: الحكم المعروض من الحقل الحتمي حصراً.
            "verdict": {"verdict": "WATCH", "confidence": 0.5,
                       "ai": {"verdict": "WATCH", "confidence": 0.5,
                             "reasoning": "سبب"}},
            "report": {"report": report_text, "review_cycles": 1,
                      "unresolved_notes": []},
        },
    }


def test_cover_wordmark_is_real_not_bracketed_placeholder(monkeypatch):
    """§7 (أمر العمل الرئيس): الغلاف بلا نصّ نائب مُقوَّس «[شعار سِلك]» —
    علامة اسمية حقيقية «سِلك» بلون العلامة، لا قوس نائب."""
    from silk_render import build_view
    from silk_reports import render_docx
    monkeypatch.setenv("SILK_HERMETIC", "1")
    view = build_view(_deep_research_result())
    path = os.path.join(tempfile.mkdtemp(), "brand.docx")
    render_docx(view, path)
    text = docx_all_text(path)
    assert "[شعار سِلك]" not in text and "شعار سِلك" not in text
    assert "سلك" in text  # Wave 2: تشكيل آمن (بلا كسرة مُركَّبة)


def test_page_header_and_footer_present(monkeypatch):
    from silk_render import build_view
    from silk_reports import render_docx
    from docx import Document
    monkeypatch.setenv("SILK_HERMETIC", "1")
    view = build_view(_deep_research_result())
    path = os.path.join(tempfile.mkdtemp(), "header.docx")
    render_docx(view, path)
    doc = Document(path)
    section = doc.sections[0]
    header_text = "\n".join(p.text for p in section.header.paragraphs)
    footer_text = "\n".join(p.text for p in section.footer.paragraphs)
    assert "تمور" in header_text
    assert "سلك لذكاء الأسواق" in footer_text  # Wave 2


def test_table_header_row_shaded_with_primary_color(monkeypatch):
    from silk_render import build_view
    from silk_reports import render_docx, _TABLE_HEADER_FILL
    from docx import Document
    from docx.oxml.ns import qn
    monkeypatch.setenv("SILK_HERMETIC", "1")
    view = build_view(_deep_research_result())
    path = os.path.join(tempfile.mkdtemp(), "shaded.docx")
    render_docx(view, path)
    doc = Document(path)
    # §7 (ترقية الطباعة): رأس الجدول أخضرُ سِلك #166534 (كان لون العلامة الأساس).
    header_fill = _TABLE_HEADER_FILL.upper()
    found = False
    for table in doc.tables:
        hdr_cell = table.rows[0].cells[0]
        shd = hdr_cell._tc.get_or_add_tcPr().find(qn("w:shd"))
        if shd is not None and shd.get(qn("w:fill")).upper() == header_fill:
            found = True
            break
    assert found


def test_markdown_table_has_no_machine_caption(monkeypatch):
    """§7 (أمر العمل الرئيس): لا تعليق آليّ «جدول: المؤشر · القيمة» قبل
    الجداول — الأعمدة تعرّف نفسها بترويستها. عنوان حقيقي أو لا شيء."""
    from silk_render import build_view
    from silk_reports import render_docx
    monkeypatch.setenv("SILK_HERMETIC", "1")
    view = build_view(_deep_research_result())
    path = os.path.join(tempfile.mkdtemp(), "caption.docx")
    render_docx(view, path)
    text = docx_all_text(path)
    assert "جدول: " not in text


def test_sample_docx_regenerated_reflects_new_branding_and_structure():
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    path = os.path.join(root, "samples", "research_report_latest.docx")
    assert os.path.exists(path)
    text = docx_all_text(path)
    assert "[شعار سِلك]" not in text  # §7: لا نصّ نائب مُقوَّس
    assert "منهجية البحث ونطاقه" in text
    assert "التوصيات الاستراتيجية" in text
    assert "سلك لذكاء الأسواق" in text  # Wave 2: تشكيل آمن


# ── ١١.٢أ: تصليب وكيل المنافسين — تحقّق ضد إسبانيا ────────────────────

def _spain_ref():
    from silk_market_resolver import resolve_market
    ref, _ = resolve_market("Spain")
    return ref


def test_competitors_mission_dry_run_against_spain_calls_comtrade_competitors(
        tmp_path):
    import json
    from unittest.mock import patch
    import silk_missions as sm
    from silk_data_layer import DataPoint

    def fake_market_competitors(hs, m49, year):
        return [
            DataPoint({"partner": "تونس", "code": "788", "value_usd": 1e6,
                      "share": 40.0}, "UN Comtrade", 0.9, "n"),
            DataPoint({"partner": "المغرب", "code": "504", "value_usd": 6e5,
                      "share": 24.0}, "UN Comtrade", 0.9, "n"),
        ]

    calls = {"n": 0}

    def fake_call_tools(system, messages, tools=None, max_tokens=1600,
                        model=None, timeout=None):
        calls["n"] += 1
        if tools and calls["n"] == 1:
            return {"stop_reason": "tool_use", "content": [
                {"type": "tool_use", "id": "t1", "name": "comtrade_competitors",
                 "input": {"year": 2023}}]}
        return {"stop_reason": "end_turn", "content": [
            {"type": "text", "text": json.dumps(
                {"findings": [{"claim": "تونس أكبر مورّد بحصة 40%",
                              "datapoint_ids": ["dp1"], "confidence": 0.9}],
                 "gaps": [], "summary": "dry-run إسبانيا"})}]}

    with patch("silk_data_layer_v2.market_competitors",
              side_effect=fake_market_competitors), \
         patch("silk_llm_runtime._call_tools", side_effect=fake_call_tools):
        out = sm.deep_research(_spain_ref(), product="تمور", hs_code="080410",
                               dry_run=True, only_agent="competitors",
                               trace_dir=str(tmp_path))

    assert out["mode"] == "dry_run"
    assert not out["report"].failed
    tool_names = [e.get("tool") for e in out["events"] if e.get("kind") == "tool_call"]
    assert "comtrade_competitors" in tool_names


# ── ١١.٢ب: تصليب وكيل المخاطر/الأخبار — تحقّق ضد إسبانيا ──────────────

def test_risk_news_mission_dry_run_against_spain_calls_exchange_rate(tmp_path):
    import json
    from unittest.mock import patch
    import silk_missions as sm
    from silk_data_layer import DataPoint

    calls = {"n": 0}

    def fake_call_tools(system, messages, tools=None, max_tokens=1600,
                        model=None, timeout=None):
        calls["n"] += 1
        if tools and calls["n"] == 1:
            return {"stop_reason": "tool_use", "content": [
                {"type": "tool_use", "id": "t1", "name": "worldbank_indicator",
                 "input": {"indicator": "exchange_rate", "year": 2023}}]}
        return {"stop_reason": "end_turn", "content": [
            {"type": "text", "text": json.dumps(
                {"findings": [{"claim": "سعر الصرف اليورو/دولار 0.92",
                              "datapoint_ids": ["dp1"], "confidence": 0.95}],
                 "gaps": [], "summary": "dry-run إسبانيا"})}]}

    with patch("silk_llm_runtime.world_bank",
              return_value=DataPoint(0.92, "World Bank", 0.95,
                                     "PA.NUS.FCRF year=2023", "2026-01-01")), \
         patch("silk_llm_runtime._call_tools", side_effect=fake_call_tools):
        out = sm.deep_research(_spain_ref(), product="تمور", hs_code="080410",
                               dry_run=True, only_agent="risk_news",
                               trace_dir=str(tmp_path))

    assert out["mode"] == "dry_run"
    assert not out["report"].failed
    tool_names = [e.get("tool") for e in out["events"] if e.get("kind") == "tool_call"]
    assert "worldbank_indicator" in tool_names


# ── ١١.٣: جاهزية الإصدار — أداة التقييم + وثائق ────────────────────────

def _root():
    return os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def test_eval_harness_runs_citation_axis_against_sample_result(monkeypatch):
    """محور الاستشهاد البرمجي يعمل بلا مفتاح كلود ولا شبكة — يعيد نتيجة
    محسوبة (لا استثناء) حتى بلا حالة ذهبية مسجَّلة."""
    monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
    import runpy
    ns = runpy.run_path(os.path.join(_root(), "tools", "gen_research_sample.py"))
    import silk_evals
    out = silk_evals.evaluate_report(ns["result"])
    assert out is not None
    assert out["axes"]["citation_correctness"] in (0, 100)
    assert out["note"]  # فجوة معلنة صراحة: محاور كلود غير محسوبة بلا مفتاح


def test_golden_cases_still_declared_empty_not_fabricated():
    import json
    path = os.path.join(_root(), "evals", "golden_cases.json")
    cases = json.load(open(path, encoding="utf-8"))
    assert cases == []  # فجوة معلنة — لا حالة مُختلَقة لسدّ الفراغ


def test_release_notes_document_exists_with_required_sections():
    text = open(os.path.join(_root(), "docs", "RELEASE_NOTES_v1.md"),
               encoding="utf-8").read()
    for marker in ("الضمانات", "الحدود المعلنة", "التكلفة لكل تشغيلة",
                  "دليل المالك التشغيلي"):
        assert marker in text


def test_execution_plan_marks_deep_research_waves_delivered():
    text = open(os.path.join(_root(), "docs", "EXECUTION_PLAN.md"),
               encoding="utf-8").read()
    assert "موجات ١-١١ منفَّذة ومدموجة بالكامل" in text
    assert "DEEP_RESEARCH_DECISIONS.md" in text
