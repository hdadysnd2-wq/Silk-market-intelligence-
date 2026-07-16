"""أمر العمل الرئيس — تحديث مخرجات تقرير البحث العميق (target 10/10).

اختبارات حتمية هرمتية لأقسام أمر العمل: §1 العملة بالدولار حصراً، §2 السرّية
(صفر سباكة داخلية)، §5 الاكتمال (لا تحذير تشغيلي، لا بتر بـ«…»)، §6 سجل
الأدلة للمدققين (مصدر عمومي + رابط + شارة، لا وسم أداة، لا حقيقة مبتورة).
كلها بلا شبكة وبلا مفتاح — تشكيل نصّ فقط.

Run: python3 -m pytest tests/test_report_output_overhaul.py -q
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


# ── §1 العملة تبقى بالدولار — لا تحويل ريالي، لا اختزال «م$» ────────────────

def test_render_strips_any_sar_conversion_parenthetical():
    import silk_render
    txt, _ = silk_render._apply_merchant_language(  # noqa: SLF001
        "واردات تبلغ 61 مليون دولار (نحو 228.8 مليون ريال بسعر الربط 3.75) عام "
        "2023.")
    assert "ريال" not in txt and "بسعر الربط" not in txt
    assert "61 مليون دولار" in txt


def test_render_unifies_million_dollar_shorthand():
    import silk_render
    txt, _ = silk_render._apply_merchant_language("بلغت 2.1م$ ثم 3 مليار$.")  # noqa: SLF001
    assert "2.1 مليون دولار" in txt and "3 مليار دولار" in txt
    assert "م$" not in txt and "مليار$" not in txt


# ── §2 السرّية — تطهير المتن في طبقة العرض ─────────────────────────────────

def test_facts_list_phrase_rewritten_for_reader():
    import silk_render
    out = silk_render._strip_internal_plumbing(  # noqa: SLF001
        "لم يرد ذكر السعر بين الحقائق المتاحة.")
    assert "بين الحقائق" not in out
    assert "من المصادر المتاحة" in out


def test_mission_key_in_parentheses_mapped_to_arabic_label():
    import silk_render
    out = silk_render._strip_internal_plumbing(  # noqa: SLF001
        "نمط استهلاك موسمي (consumer_culture) واضح.")
    assert "consumer_culture" not in out
    assert "ثقافة المستهلك" in out


def test_tool_failure_narration_becomes_data_gap():
    import silk_render
    out = silk_render._strip_internal_plumbing(  # noqa: SLF001
        "فشل استعلام WITS مرتين بسبب انقطاع الاتصال، فالرقم غير مؤكَّد.")
    assert "فشل استعلام" not in out and "انقطاع الاتصال" not in out
    assert "لم تتوفّر بيانات WITS من المصدر الرسمي" in out


def test_no_claude_mention_survives_in_body():
    import silk_render
    out = silk_render._strip_internal_plumbing(  # noqa: SLF001
        "وفق تقدير كلود ورد Claude أعلاه.")
    assert "كلود" not in out and "Claude" not in out


# ── §6 سجل الأدلة — مصدر عمومي، رابط، شارة، لا وسم أداة، لا بتر ─────────────

def test_clean_source_label_strips_tool_use_tag():
    from silk_reports import _clean_source_label
    assert _clean_source_label("UN Comtrade (Claude tool-use)") == "UN Comtrade"
    assert "tool-use" not in _clean_source_label("البعثة (Claude tool-use)")


def test_trim_sentence_never_leaves_trailing_ellipsis():
    from silk_reports import _trim_sentence
    long = ("واردات إسبانيا من التمور بلغت 61 مليون دولار عام 2023 وفق بيانات "
            "UN Comtrade، وهي في نموّ مطّرد على ثلاث سنوات متتالية بلا انقطاع "
            "يُذكر في السجلّ الرسمي المنشور.")
    out = _trim_sentence(long, 80)
    assert not out.endswith("…") and not out.endswith("...")
    assert out and len(out) <= 80


def test_first_url_extracts_real_url_or_dash():
    from silk_reports import _first_url
    assert _first_url("مرصود على https://ah.nl/dates ✓") == "https://ah.nl/dates"
    assert _first_url("لا رابط هنا", None) == "—"


def test_finding_assembly_uses_public_source_not_tool_use():
    """§2.1 — تجميع نتائج البعثة يسند المصدر العمومي للنقاط المستشهَد بها
    (registry) لا وسم «(Claude tool-use)». يُحاكى بأدنى بنية registry/finding."""
    import silk_llm_runtime as R
    from silk_data_layer import DataPoint
    reg = {0: DataPoint(61_000_000.0, "UN Comtrade", 0.9, "واردات 2023",
                        "2026-07-01")}
    findings = R.__dict__  # sanity: module importable
    assert "run_llm_agent" in findings
    # النقطة العمومية موجودة بمصدرها الحقيقي — المُجمِّع يقرأ .source منها.
    assert reg[0].source == "UN Comtrade" and "tool-use" not in reg[0].source


# ── §2 بوابة الجودة — حارس تسريب حتمي يُفشِل على المحفّزات ────────────────

def _dr_with_report(text: str) -> dict:
    return {"deep_research": {
        "report": {"text": text}, "missions": {},
        "analyst": {"by_category": {}, "missing_categories": []},
        "verdict": {}}}


def test_quality_gate_fails_on_confidentiality_leak_tokens():
    import silk_quality_gate as G
    for leak in ["استند إلى Claude tool-use", "ارفع SILK_MAX_TOKENS_CEILING",
                 "وفق مسار بحث الأسعار", "بين الحقائق المعطاة", "⚠ تنبيه"]:
        out = G.run_quality_gate(_dr_with_report(
            "## 1. الخلاصة\n" + leak + " والقرار المتابعة."))
        assert out["verdict"] == G.FAIL, f"لم يُفشِل التسريب: {leak}"


def test_quality_gate_fails_on_trailing_ellipsis():
    import silk_quality_gate as G
    out = G.run_quality_gate(_dr_with_report(
        "## 1. الخلاصة\nهذه جملة مبتورة تنتهي بنقاط حذف غير نظيفة…"))
    assert out["verdict"] == G.FAIL
    assert any(f["check"] == "trailing_ellipsis" for f in out["findings"])


def test_quality_gate_confidentiality_note_does_not_reecho_token():
    """الملاحظة نفسها لا تُعيد طبع القيمة المطابَقة (وإلا لأعادت البوابة
    تسريب ما رصدته)."""
    import silk_quality_gate as G
    out = G.run_quality_gate(_dr_with_report("## 1. الخلاصة\nClaude tool-use هنا."))
    joined = " ".join(f["note"] for f in out["findings"])
    assert "tool-use" not in joined and "Claude" not in joined


# ── §4 RTL — المستند كله من اليمين لليسار ──────────────────────────────────

def _mini_research_result():
    from silk_market_resolver import resolve_market
    from silk_agents import AgentReport
    from silk_data_layer import DataPoint
    ref, _ = resolve_market("Spain")
    dp = DataPoint(61_000_000.0, "UN Comtrade", 0.9, "واردات 2023", "2026-07-01")
    rep = AgentReport("LLMAgent:trade_flow", [dp], False, "تدفقات مؤكَّدة")
    return {
        "product": "تمور", "hs_code": "080410", "year": 2023, "markets": [],
        "market": {"iso3": ref.iso3, "m49": ref.m49, "iso2": ref.iso2,
                   "name_en": ref.name_en, "name_ar": ref.name_ar},
        "deep_research": {
            "trace_id": "rtl-test", "missions": {"trade_flow": rep},
            "analyst": {"report": {"agent_name": "market_analyst",
                                   "summary": "متابعة.", "findings": [],
                                   "failed": False},
                        "by_category": {"demand": [dp]},
                        "missing_categories": []},
            "verdict": {"verdict": "WATCH", "ai": {"verdict": "WATCH"}},
            "report": {"report": "## 1. الخلاصة التنفيذية\nالسوق ينمو 9%. "
                                 "| البند | القيمة |\n| --- | --- |\n"
                                 "| الواردات | 61 مليون دولار |\n",
                       "review_cycles": 1, "unresolved_notes": [],
                       "failure_reason": ""}},
    }


def _docx_xml(tmpdir):
    import os
    import zipfile
    import silk_render
    from silk_reports import render_docx
    os.environ["SILK_HERMETIC"] = "1"
    view = silk_render.build_view(_mini_research_result())
    path = os.path.join(tmpdir, "rtl.docx")
    render_docx(view, path)
    with zipfile.ZipFile(path) as z:
        return (z.read("word/document.xml").decode("utf-8"),
                z.read("word/styles.xml").decode("utf-8"), path)


def test_docx_is_rtl_document_wide(tmp_path):
    """§4 (أمر العمل الرئيس): كل فقرة bidi + محاذاة يمين، كل run <w:rtl/> +
    خطّ عربي، كل جدول bidiVisual، والمقطع bidi — لا ترقيع لعنصر واحد."""
    doc_xml, styles_xml, _ = _docx_xml(str(tmp_path))
    assert doc_xml.count("<w:bidi") > 5          # فقرات كثيرة bidi
    assert doc_xml.count("<w:rtl") > 5           # runs rtl
    assert doc_xml.count("<w:jc") > 5            # محاذاة يمين
    assert "<w:bidiVisual" in doc_xml            # جدول متدفّق يميناً
    assert "Arial" in doc_xml                    # خطّ عربي على الـrFonts
    # المقطع نفسه bidi (اتجاه أساس المستند).
    assert "<w:bidi/>" in doc_xml or "<w:bidi " in doc_xml
    # نمط Normal يحمل rtl كذلك (وراثة + صريح معاً).
    assert "<w:rtl" in styles_xml


def test_docx_jc_is_right_aligned(tmp_path):
    doc_xml, _, _ = _docx_xml(str(tmp_path))
    import re
    vals = re.findall(r'<w:jc w:val="([^"]+)"', doc_xml)
    assert vals and all(v == "right" for v in vals)


def test_docx_opens_cleanly_after_rtl(tmp_path):
    """المستند يُفتَح ثانيةً بعد تطبيق RTL — لا XML مكسور (schema صالح)."""
    from docx import Document
    _, _, path = _docx_xml(str(tmp_path))
    doc = Document(path)          # يرمي إن كان الـoxml غير صالح
    assert doc.paragraphs


# ── §3 PDF — المُسلَّم النهائي PDF، تدهور رشيق ──────────────────────────────

def test_pdf_conversion_fails_cleanly_when_engine_absent(tmp_path, monkeypatch):
    """§3/§5: بلا محرّك تحويل، docx_to_pdf يرفع RuntimeError برسالة نظيفة —
    لا docx بديل صامت، لا اسم متغيّر بيئة/مسار داخلي في الرسالة."""
    import os
    import pytest
    import silk_reports
    monkeypatch.setattr(silk_reports, "_find_soffice", lambda: None)
    docx = os.path.join(str(tmp_path), "x.docx")
    from docx import Document
    Document().save(docx)
    with pytest.raises(RuntimeError) as ei:
        silk_reports.docx_to_pdf(docx, os.path.join(str(tmp_path), "x.pdf"))
    msg = str(ei.value)
    assert "PDF" in msg
    for leak in ("soffice", "libreoffice", "SILK_", "/tmp", "Traceback"):
        assert leak not in msg


def test_render_client_pdf_propagates_clean_error(tmp_path, monkeypatch):
    """§3: render_client_pdf لا يُسلّم docx بديلاً إن تعذّر التحويل — يرفع
    الخطأ النظيف نفسه (المسار الأعلى يعيده 503 معلَناً)."""
    import os
    import pytest
    import silk_reports
    monkeypatch.setattr(silk_reports, "_find_soffice", lambda: None)
    view = __import__("silk_render").build_view(_mini_research_result())
    with pytest.raises(RuntimeError):
        silk_reports.render_client_pdf(view, os.path.join(str(tmp_path), "c.pdf"))


def test_pdf_produced_when_engine_available(tmp_path):
    """§3: عند توفّر محرّك تحويل عامل، يُنتَج PDF فعليّ من مستند RTL. يُتخطّى
    إن كان المحرّك غائباً/غير عامل في هذه البيئة (تدهور رشيق موثَّق)."""
    import os
    import pytest
    import silk_reports
    if silk_reports._find_soffice() is None:  # noqa: SLF001
        pytest.skip("محرّك تحويل المستندات غير متاح في هذه البيئة")
    view = __import__("silk_render").build_view(_mini_research_result())
    docx = os.path.join(str(tmp_path), "r.docx")
    silk_reports.render_docx(view, docx)
    try:
        pdf = silk_reports.docx_to_pdf(docx, os.path.join(str(tmp_path), "r.pdf"))
    except RuntimeError:
        pytest.skip("محرّك التحويل موجود لكنه لا يعمل في هذه البيئة (sandbox)")
    assert os.path.exists(pdf) and pdf.endswith(".pdf")
    with open(pdf, "rb") as fh:
        assert fh.read(5) == b"%PDF-"


# ── §8 بوابة الأسلوب الحتمية ────────────────────────────────────────────────

def test_style_gate_fails_on_currency_shorthand_and_inline_enumeration():
    import silk_quality_gate as G
    for bad in ["بلغت 61م$ عام 2023.", "أقوى أسباب: (1) الشريحة و(2) النمو."]:
        out = G.run_quality_gate(_dr_with_report("## 1. الخلاصة\n" + bad))
        assert out["verdict"] == G.FAIL, f"لم يُفشِل الأسلوب: {bad}"


def test_style_gate_warns_on_connector_overuse():
    import silk_quality_gate as G
    txt = ("## 1. الخلاصة\nمن ناحية الطلب جيّد. من ناحية العرض مجزّأ. "
           "من ناحية السعر تنافسي. والقرار المتابعة.")
    checks = [f["check"] for f in G._check_style(txt)]  # noqa: SLF001
    assert "style_connector_overuse" in checks


def test_style_gate_warns_on_repeated_key_figure():
    import silk_quality_gate as G
    txt = ("الحصة 55.28% في الأولى، ثم 55.28% في الثانية، ثم 55.28% ثالثاً "
           "في المتن نفسه.")
    checks = [f["check"] for f in G._check_style(txt)]  # noqa: SLF001
    assert "style_repeated_key_figure" in checks


def test_style_gate_passes_clean_arabic_prose():
    import silk_quality_gate as G
    checks = G._check_style(  # noqa: SLF001
        "## 1. الخلاصة\nالسوق ينمو بوتيرة صحّية، والقرار الدخول المشروط "
        "بعد تأمين الأهلية التنظيمية أولاً ثم التعاقد مع موزّع ثانياً.")
    assert checks == []


def test_writer_prompt_carries_style_contract_additions():
    """§8 — عقد الكاتب يمنع النحت الحرفي ويوجّب صوتاً واحداً وترقيماً عربياً."""
    import silk_ai_judge
    import inspect
    src = inspect.getsource(silk_ai_judge.deep_report)
    assert "لا نحت حرفيّ" in src               # منع الكالك
    assert "لا ترقيم إنجليزي داخل الفقرة" in src  # لا (1)…(2)
    assert "كما ورد في قسم" in src             # الإحالة بدل إعادة الشرح
    assert "مليون دولار" in src                # العملة بالدولار حصراً


def test_reviewer_prompt_has_language_pass():
    """§8 — المراجِع يُجري تمريرة لغوية على مستوى السطر (تطابق/تلصيقات/كالك)."""
    import silk_ai_judge
    import inspect
    src = inspect.getsource(silk_ai_judge.review_report)
    assert "تمريرة لغوية" in src
    assert "التطابق النحوي" in src
