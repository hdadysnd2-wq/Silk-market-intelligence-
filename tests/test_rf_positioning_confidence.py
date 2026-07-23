"""اختبارات PR-F (توصياتي): S2 موقع سعر المصدّر ضمن أسعار المنافسين (سطر
كاتب)، وS3 مؤشّر تغطية المصادر (عدّ داخليّ، عرضٌ محايد الصياغة بعد WS10).

كلاهما يمرّ عبر build_view → render_client_docx وحارس المصطلحات المحظورة.
Run:  python3 -m pytest tests/test_rf_positioning_confidence.py -q
"""
import inspect
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from conftest import docx_all_text  # noqa: E402


# ── S3: مؤشّر ثقة الدراسة ──────────────────────────────────────────────────

def test_confidence_section_tallies_badges_correctly():
    """WS10: التجميع الداخلي كما هو (2 رسمي من 0.9+0.95، الفاشلة مُتخطّاة)،
    لكن العرض محايد الصياغة — لا شارات ✓/◐/○ ولا مصطلح «موثّق/ثانوي»."""
    from docx import Document
    import silk_reports as R
    dr = {
        "missions": {
            "m1": {"failed": False, "findings": [
                {"confidence": 0.9}, {"confidence": 0.6}, {"confidence": 0.2}]},
            "m2": {"failed": True, "findings": [{"confidence": 0.95}]}},
        "analyst": {"by_category": {"demand": [{"confidence": 0.95}]}},
    }
    doc = Document()
    R._client_confidence_section(doc, dr)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "مؤشّر تغطية المصادر" in text     # العنوان المحايد الجديد
    assert "إجمالي 4 مؤشراً" in text        # 3 من m1 + 1 محلل (m2 فاشلة مُتخطّاة)
    assert "50%" in text                     # 2 من 4 مصادر رسمية أوّلية
    tbl = doc.tables[0]
    cells = [c.text for row in tbl.rows for c in row.cells]
    assert "مصدر رسمي أوّلي" in cells         # صفٌّ محايد بعدّه (2)
    assert "2" in cells
    # WS10: صفر شارة/مصطلح قوة دليل في المتن — الرُتبة داخلية، الإسناد للمراجع.
    for sym in ("✓", "◐", "○", "موثّق", "ثانوي", "غير متحقق"):
        assert sym not in text
        assert all(sym not in c for c in cells)


def test_confidence_section_absent_when_no_findings():
    from docx import Document
    import silk_reports as R
    doc = Document()
    R._client_confidence_section(doc, {"missions": {}, "analyst": {}})
    assert all("تغطية المصادر" not in p.text for p in doc.paragraphs)


# ── S2: سطر عقد برومبت الكاتب ─────────────────────────────────────────────

def test_writer_prompt_requires_price_positioning_line():
    import silk_ai_judge as J
    src = inspect.getsource(J.deep_report)
    assert "موقع سعرك ضمن أسعار المنافسين" in src
    assert "مئين" in src                     # تموضع مئيني إن أمكن


# ── العيّنة المحفوظة تُظهر الاثنين وتبقى نظيفة ─────────────────────────────

def test_committed_sample_shows_positioning_and_confidence_index():
    """§A (حزمة الفكس v2.1): جدول مزيج الثقة (✓/◐/○، S3) أُسقط من بناء
    العميل النهائي (يبقى في ?internal=1 فقط) — النموذج المحفوظ لم يعد
    يحمله؛ التموضع السعري (S2) غير متأثر ويبقى مفحوصاً."""
    import silk_reports as R
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    text = docx_all_text(os.path.join(root, "samples", "client_report_latest.docx"))
    assert "مؤشّر ثقة الدراسة" not in text                 # §A: أُسقط من العميل
    # S2: التموضع السعري مقروء من داخل قسم المنافسة (صياغة الكاتب التوضيحية)
    assert "الموقع السعري" in text
    assert "ضمن أسعار المنافسين المرصودة" in text
    assert R._client_forbidden_hits(text) == []          # يبقى خالياً من التِلِمِتري
