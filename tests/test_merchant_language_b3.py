"""B1/B3 (SPEC-v2) — لغة التاجر لا الاقتصادي. التقرير لصاحب قرار تجاري غير
متخصص: كل مصطلح تقني عند **أول** ورود يُشرح بين قوسين بالعربية، ويُذيَّل
التقرير بمسرد من المصطلحات المستعملة فعلاً، والأرقام الكبيرة بالدولار
تُسيَّق بالريال. مُنفَّذ حتمياً في طبقة العرض الواحدة
(silk_render._apply_merchant_language) فيظهر على md **و**docx معاً حتى لو
انحرف النموذج — لا HHI/CAGR/LPI/MFN مستقلّة عند أول ورود.

Run: python3 -m pytest tests/test_merchant_language_b3.py -q
"""
import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

_TERMS = ["HHI", "CAGR", "LPI", "MFN"]

# سرد فيه المصطلحات **مستقلّة** (لا بين قوسين مسبقاً) كي يُختبَر الشرح الحتمي.
_REPORT_TEXT = (
    "## 1. الخلاصة التنفيذية\n"
    "السوق مفتّت ومؤشر التركّز HHI منخفض، والنمو السنوي CAGR بلغ 8%. "
    "مؤشر الشحن LPI مرتفع، والتعريفة العادية MFN صفرية على التمر.\n\n"
    "## 2. حجم الطلب\n"
    "الطلب السنوي يتجاوز 129.6 مليون دولار.\n\n"
    "## 11. الخلاصة والقرار\nالقرار: المتابعة قبل الالتزام.\n")


def _dp(value, source="UN Comtrade", note=""):
    from silk_data_layer import DataPoint
    return DataPoint(value, source, 0.9, note)


def _blob():
    """مدوّنة بحث عميق بالشكل المخزَّن — المصطلحات تظهر في السرد فقط (ملاحظات
    البعثات بلا اختصارات خام كي يكون أول ورود لكل مصطلح هو السرد المشروح)."""
    m = {"agent_name": "LLMMissionAgent",
         "summary": "المشهد التنافسي مفتّت", "failed": False,
         "findings": [_dp(940, note="مؤشر التركّز للسوق")]}
    return {
        "product": "تمور", "hs_code": "080410", "markets": [],
        "market": {"iso3": "NLD", "m49": 528, "iso2": "NL",
                   "name_en": "Netherlands", "name_ar": "هولندا"},
        "deep_research": {
            "missions": {"competition": m},
            "analyst": {"report": {"agent_name": "market_analyst",
                                   "summary": "هولندا — متابعة.",
                                   "findings": [], "failed": False},
                        "missing_categories": [], "by_category": {}},
            "verdict": {"verdict": "WATCH", "ai": {"verdict": "WATCH"}},
            "report": {"report": _REPORT_TEXT, "review_cycles": 1,
                       "unresolved_notes": [], "failure_reason": ""},
            "trace_id": "nld-b3"},
    }


def _first_is_glossed(text: str, term: str) -> bool:
    """أول ورود للمصطلح متبوعٌ بشرح بين قوسين (تجاهل مسافة واحدة)."""
    i = text.find(term)
    assert i >= 0, f"{term} غائب تماماً"
    tail = text[i + len(term): i + len(term) + 3].lstrip()
    return tail.startswith("(")


def test_md_has_glossary_and_no_standalone_jargon_at_first_use():
    import silk_render
    from silk_reports import render_markdown
    md = render_markdown(silk_render.build_view(_blob()))
    assert "مسرد المصطلحات" in md, "المسرد غائب عن md"
    for t in _TERMS:
        assert _first_is_glossed(md, t), f"{t} ورد مستقلّاً عند أول ورود في md"


def test_docx_has_glossary_and_no_standalone_jargon_at_first_use():
    import silk_render
    from silk_reports import render_client_docx
    from docx import Document
    view = silk_render.build_view(_blob())
    path = render_client_docx(view, os.path.join(tempfile.mkdtemp(), "c.docx"))
    assert os.path.exists(path)
    doc = Document(path)
    blob = "\n".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for c in row.cells:
                blob += "\n" + c.text
    assert "مسرد المصطلحات" in blob, "المسرد غائب عن docx العميل"
    for t in _TERMS:
        assert _first_is_glossed(blob, t), f"{t} ورد مستقلّاً عند أول ورود في docx"


def test_glossary_lists_only_terms_actually_used():
    """المسرد يُبنى من المصطلحات المستعملة فعلاً — EORI (غير مذكور) لا يظهر."""
    import silk_render
    from silk_reports import render_markdown
    md = render_markdown(silk_render.build_view(_blob()))
    glossary = md.split("مسرد المصطلحات", 1)[1]
    assert "HHI" in glossary and "CAGR" in glossary
    assert "EORI" not in glossary and "TRACES" not in glossary


def test_usd_amount_stays_usd_no_sar_conversion():
    """§1 (أمر العمل الرئيس): العملة تبقى بالدولار كما وردت — لا تحويل إلى
    الريال ولا مقابل مُقوَّس. «129.6 مليون دولار» تظهر كما هي، ولا «ريال»
    ولا «بسعر الربط» في أي مخرَج (يُلغى تسييق B1 الريالي)."""
    import silk_render
    from silk_reports import render_markdown
    md = render_markdown(silk_render.build_view(_blob()))
    assert "129.6 مليون دولار" in md, "المبلغ الدولاري تغيّرت صياغته"
    assert "ريال" not in md, "تسرّب تحويل ريالي رغم إلغائه في §1"
    assert "بسعر الربط" not in md and "486" not in md


def test_million_dollar_shorthand_is_unified():
    """§1: الاختزال «م$» يُوحَّد إلى الصيغة الكاملة «مليون دولار» ولا يبقى."""
    import silk_render
    txt, _ = silk_render._apply_merchant_language(  # noqa: SLF001
        "بلغت الواردات 61م$ عام 2023.")
    assert "61 مليون دولار" in txt and "م$" not in txt


def test_no_fabrication_gloss_only_annotates_never_invents_numbers():
    """الشرح والمسرد لا يغيّران أي رقم — 940/8%/129.6 تبقى كما وردت."""
    import silk_render
    view = silk_render.build_view(_blob())
    txt = view["deep_research"]["report"]["text"]
    assert "8%" in txt and "129.6" in txt  # الأرقام الأصلية محفوظة
