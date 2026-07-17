"""البند ١٨ (بلاغ UK الحي، أمر العمل الرئيس ITEM 1) — لا اسم مزوّد داخلي على
سطح العميل. lock-test قبل الفكس: يتحقّق أن تصدير العميل يُحيّد كل اسم مزوّد
(Volza/Explee/إكسبلي/فولزا/Serper/SerpApi/LocalPrice/pytrends/GDELT) عبر
المدوّنة القانونية الحقيقية الشكل **وشكل تقرير UK**، وأن حارس التصدير يرفض
بصوت عالٍ اسم مزوّد يستحيل تنقيته — تماماً كحارس المصطلحات التشغيلية القائم.

هرمتي بالكامل (بلا شبكة): يبني/يفتح docx فعلياً من الشكل الحقيقي، لا نموذج.

Run: python3 -m pytest tests/test_vendor_name_leak_item1.py -q
"""
from __future__ import annotations

import os
import sys
import tempfile

import pytest

_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, _ROOT)
sys.path.insert(0, os.path.join(_ROOT, "tools"))

pytest.importorskip("docx")

# قائمة المزوّدين الحرفية من عقد المالك (تُطابِق _CLIENT_VENDOR_NAMES_*).
_VENDORS = ("Volza", "Explee", "إكسبلي", "فولزا", "Serper", "SerpApi",
            "LocalPrice", "pytrends", "GDELT")


def _inject_vendor_leaks(blob: dict) -> dict:
    """احقن أسماء مزوّدين في كل حقل يصل تقرير العميل — بصيغ خام (يترجمها
    silk_narrative لأسماء عربية) وصيغ علامة تجارية عارية معاً."""
    dr = blob["deep_research"]
    # سرد الكاتب — صيغ خام يعرّبها silk_narrative + علامات عارية.
    dr["report"]["report"] += (
        "\n\n## ملاحظة تشغيلية مسرّبة\n"
        "Explee unavailable: timeout. "
        "Volza: no named importers parsed for HS0804 into GBR. "
        "المشترون عبر Serper وSerpApi، والتسعير من LocalPrice، "
        "والموسمية من pytrends وأخبار المخاطر من GDELT.")
    # جدول الروابط — ملاحظة الفجوة كثيراً ما تحمل اسم المزوّد.
    dr["importer_leads"] = {
        "leads": [], "path": "gap",
        "note": "فولزا/إكسبلي غير متاحين — فجوة معلنة"}
    # تقاطعات المحلل (المسار الاحتياطي لمتن العميل).
    dr["analyst"]["by_category"] = {
        "price_competitiveness": [
            {"value": "تسعير مرجعي من LocalPrice وExplee", "source": "Explee",
             "confidence": 0.5, "note": "", "retrieved_at": "2026-07-15"}]}
    return blob


def _uk_research_blob() -> dict:
    """شكل تقرير UK (عسل × بريطانيا) — مشتقّ من المدوّنة القانونية بتبديل
    السوق/المنتج، بأسماء مزوّدين محقونة في كل حقل عميل."""
    from canonical_netherlands import netherlands_research_blob
    blob = netherlands_research_blob()
    blob["product"] = "عسل"
    blob["hs_code"] = "040900"
    blob["market"] = {"iso3": "GBR", "m49": 826, "iso2": "GB",
                      "name_en": "United Kingdom", "name_ar": "بريطانيا"}
    blob["deep_research"]["market"] = dict(blob["market"])
    # حكم GO كي يُولَّد سطر «الخطوة التالية» (كان يحقن (Volza/Explee)).
    blob["deep_research"]["verdict"] = {
        "verdict": "GO", "confidence": 0.7,
        "ai": {"verdict": "GO", "reasoning": "سوق واعد يحتاج تحقّق المستوردين."}}
    return _inject_vendor_leaks(blob)


def _client_export_text(blob: dict) -> str:
    """ابنِ تقرير العميل docx من الشكل الحقيقي وأعِد كل نصّه (فقرات + جداول)."""
    import silk_render
    from silk_reports import render_client_docx
    from conftest import docx_all_text
    view = silk_render.build_view(blob)
    path = render_client_docx(view, os.path.join(tempfile.mkdtemp(), "c.docx"))
    return docx_all_text(path)


def test_client_export_names_no_vendor_across_canonical_and_uk_shapes():
    """كنس (d): المدوّنة القانونية + شكل UK — بعد الفكس، صفر ظهور لأيّ اسم
    مزوّد داخلي في تقرير العميل المُنتَج فعلياً؛ والحارس النهائي لا يرفعه
    501 (نُقِّي قبله)."""
    from silk_reports import _client_forbidden_hits
    from canonical_netherlands import netherlands_research_blob

    canonical = _inject_vendor_leaks(netherlands_research_blob())
    for name, blob in (("canonical", canonical), ("uk", _uk_research_blob())):
        text = _client_export_text(blob)
        hits = _client_forbidden_hits(text)
        vendor_hits = [h for h in hits if h.startswith("vendor_name")]
        assert not vendor_hits, f"[{name}] اسم مزوّد تسرّب للعميل: {vendor_hits}"
        # تأكيد مباشر على الأسماء الحرفية (شبكة أمان فوق الحارس).
        for v in _VENDORS:
            assert v not in text, f"[{name}] «{v}» ظهر في تقرير العميل"


def test_next_step_carries_no_vendor_name():
    """سطر «الخطوة التالية» المُولَّد (GO) لغة أعمال عامة بلا اسم مزوّد."""
    import silk_render
    from silk_reports import _client_forbidden_hits
    view = silk_render.build_view(_uk_research_blob())
    nxt = (view.get("deep_research") or {}).get("next_step") or ""
    assert nxt, "سطر الخطوة التالية غائب رغم حكم GO"
    assert not _client_forbidden_hits(nxt), f"تسرّب في next_step: {nxt!r}"
    assert "خدمة التعميق المدفوعة" in nxt   # لغة أعمال عامة محفوظة


def test_client_vendor_guard_fails_loud_on_injected_vendor_name():
    """(b) الحارس يرفض بصوت عالٍ اسم مزوّد يستحيل تنقيته — نفس نمط الحارس
    القائم (_client_assert_clean يرفع RuntimeError)."""
    from docx import Document
    from silk_reports import _client_assert_clean, _client_forbidden_hits
    doc = Document()
    doc.add_paragraph("التقرير مبنيّ على بيانات Volza وإكسبلي التجارية.")
    assert _client_forbidden_hits("\n".join(p.text for p in doc.paragraphs)), (
        "التهيئة خاطئة — يجب أن يكتشف الحارس اسم المزوّد أولاً")
    with pytest.raises(RuntimeError):
        _client_assert_clean(doc)


def test_vendor_names_survive_as_legit_public_sources():
    """المصادر البشرية العمومية (UN Comtrade/World Bank/Eurostat) ليست
    مزوّدين ممنوعين — لا يلتقطها حارس أسماء المزوّدين."""
    from silk_reports import _client_forbidden_hits
    clean = ("المصدر UN Comtrade وWorld Bank وEurostat — استشهاد مشروع.")
    vendor_hits = [h for h in _client_forbidden_hits(clean)
                   if h.startswith("vendor_name")]
    assert not vendor_hits, f"مصدر عمومي مشروع صُنّف مزوّداً: {vendor_hits}"
