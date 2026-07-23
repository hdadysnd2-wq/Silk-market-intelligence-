"""WP-4 (برنامج إصلاح جودة التقارير) — اتساق الفجوات: الختام يرى كل الفجوات.

بلاغ التدقيق (2026-07-22): القسم الختامي طبع «لا فجوة جوهرية تمنع اتخاذ
القرار» بينما قسم المخاطر أعلن ثلاث فجوات بيانات حقيقية (فشل واجهة حوكمة
البنك الدولي، الموسمية، سعر الصرف الاسمي) — فجوات البعثات المعلنة لم تكن
مدخلاً للقسم الختامي إطلاقاً (مساران للحقيقة). الأقفال:

1. `_client_gap_inputs` هو المصدر الواحد (٤ مدخلات) للقسم الختامي ولحارس
   البوابة معاً؛ فجوات البعثات («فجوات: …» في الملخّصات) مدخله الرابع.
2. السطر الإيجابي يُطبَع فقط حين تخلو القوائم الأربع معاً؛ فجوات غير حاجبة
   وحدها => صياغة «لا تمنع القرار الحالي لكنها تقيّد يقينه».
3. حارس بوابة: «فجوة بيانات» في المتن + ختام سيطبع «لا فجوة جوهرية» = FAIL.

Run: python3 -m pytest tests/test_wp4_gaps_consistency.py -q
"""
from __future__ import annotations

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pytest


def _dr(mission_summary: str = "", flip_met: bool = True,
        report_text: str = "## 1. الخلاصة التنفيذية\nنص.") -> dict:
    return {
        "report": {"text": report_text},
        "analyst": {"by_category": {"demand": [{"value": "x",
                                                "confidence": 0.9}]},
                   "missing_categories": []},
        "missions": {"macro_context": {
            "label": "السياق الكلي", "failed": False,
            "summary": mission_summary,
            "findings": [{"value": "v", "confidence": 0.8}]}},
        "flip_conditions": [{"condition": "شرط", "met": flip_met,
                             "closes_via": "خطوة"}],
    }


def test_mission_declared_gaps_are_a_fourth_input():
    """فجوة بعثة معلنة («فجوات: …») تظهر في مدخلات القسم الختامي —
    الحالة المُسلَّمة: فشل واجهة حوكمة البنك الدولي كان مرئياً في المخاطر
    وغائباً عن الختام."""
    from silk_reports import _client_gap_inputs
    critical, informational = _client_gap_inputs(_dr(
        "اكتمل الجمع | فجوات: تعذّر جلب مؤشرات الحوكمة من البنك الدولي؛ "
        "لا بيانات موسمية"))
    assert critical == []
    assert len(informational) == 2
    assert any("الحوكمة" in g for g in informational)
    assert all("لا تمنع القرار الحالي" in g for g in informational)


def test_mission_gaps_are_deduped_and_capped_per_mission():
    from silk_reports import _client_gap_inputs
    dr = _dr("فجوات: أ؛ أ؛ ب؛ ج")   # مكرّر + ثلاث فجوات فريدة
    _, informational = _client_gap_inputs(dr)
    assert len(informational) == 2   # سقف بندين لكل بعثة، بلا تكرار


def test_positive_sentence_only_when_all_four_lists_empty(tmp_path):
    pytest.importorskip("docx")
    from docx import Document
    from silk_reports import _client_gaps_section

    def _render(dr):
        doc = Document()
        _client_gaps_section(doc, dr)
        return "\n".join(p.text for p in doc.paragraphs)

    # (أ) كل المدخلات خالية => السطر الإيجابي.
    assert "لا فجوة جوهرية" in _render(_dr())
    # (ب) فجوة بعثة معلنة وحدها => الصياغة المقيِّدة، لا النفي.
    t = _render(_dr("فجوات: تعذّر جلب مؤشرات الحوكمة"))
    assert "لا فجوة جوهرية" not in t
    assert "تقيّد يقينه" in t
    assert "الحوكمة" in t
    # (ج) شرط قلب غير محقَّق (حرج) => قائمة الفجوات الحرجة.
    t = _render(_dr(flip_met=False))
    assert "لا فجوة جوهرية" not in t
    assert "لم يتحقّق بعد" in t


def test_gate_fails_on_closing_contradiction():
    """«فجوة بيانات» في المتن + مدخلات فجوات خالية (الختام سينفي) = FAIL."""
    import silk_quality_gate as G
    dr = _dr(report_text="## 9. تقييم المخاطر\nتوجد فجوة بيانات في مؤشرات "
                         "الحوكمة لهذا السوق.")
    f = G._check_gaps_closing_contradiction(dr)
    assert f and f[0]["check"] == "gaps_closing_contradiction"
    assert f[0]["repairable"] is False
    # وحين توجد فجوة معلنة (الختام لن ينفي) — لا تناقض.
    dr2 = _dr(mission_summary="فجوات: نقص بيانات",
              report_text="## 9. تقييم المخاطر\nتوجد فجوة بيانات معلنة.")
    assert G._check_gaps_closing_contradiction(dr2) == []
    src = open(os.path.join(os.path.dirname(os.path.dirname(
        os.path.abspath(__file__))), "silk_quality_gate.py"),
        encoding="utf-8").read()
    fail_block = src.split('if any(f["check"] in (')[1].split(
        'for f in non_repairable')[0]
    assert '"gaps_closing_contradiction"' in fail_block


def test_governance_failure_fixture_lists_the_gap_in_closing(tmp_path):
    """قبول WP-4: مدوّنة بفشل واجهة الحوكمة تُنتج ختاماً يذكر الفجوة."""
    pytest.importorskip("docx")
    from docx import Document
    from silk_reports import _client_gaps_section
    doc = Document()
    _client_gaps_section(doc, _dr(
        "اكتمل | فجوات: فشل واجهة مؤشرات الحوكمة العالمية للبنك الدولي"))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "الحوكمة" in text
    assert "لا فجوة جوهرية" not in text
