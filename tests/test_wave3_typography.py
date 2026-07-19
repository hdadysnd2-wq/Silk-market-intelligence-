"""Wave 3 §7 — ترقيةُ طباعة تقارير العميل (قرار المالك، مواصفةٌ مُثبَّتة).

هرمتيّ على مستوى XML الـdocx (لا soffice): الخطّ IBM Plex، أحجام/ألوان الأنماط،
وتنسيقُ الجداول (رأسٌ أخضر، شريطٌ متناوب، حدودٌ شعريّة، هوامشُ خلايا، خطٌّ أصغر
للجداول الكثيفة). قبولُ التضمين (pdffonts) والهندسة (§4) يقعان في CI بالخطّ
مثبَّتًا — لا يُتحقَّق منهما محلّيًّا (لا soffice/لا Plex).
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pytest  # noqa: E402

pytest.importorskip("docx")

import silk_reports as sr  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402


def test_body_font_is_ibm_plex_sans_arabic():
    assert sr._RTL_BODY_FONT == "IBM Plex Sans Arabic"  # noqa: SLF001


def test_typography_styles_sizes_colors():
    from docx import Document
    doc = Document()
    sr._apply_typography(doc)  # noqa: SLF001
    normal = doc.styles["Normal"]
    assert normal.font.name == "IBM Plex Sans Arabic"
    assert abs(normal.font.size.pt - 11.0) < 0.01
    assert abs(normal.paragraph_format.line_spacing - 345 / 240) < 0.001
    h1 = doc.styles["Heading 1"]
    assert h1.font.bold and abs(h1.font.size.pt - 14.0) < 0.01
    assert str(h1.font.color.rgb) == "166534"
    assert str(doc.styles["Heading 3"].font.color.rgb) == "333333"
    assert str(doc.styles["Title"].font.color.rgb) == "166534"
    assert abs(doc.styles["Title"].font.size.pt - 18.0) < 0.01
    hf = doc.styles["Header"]
    assert str(hf.font.color.rgb) == "555555" and abs(hf.font.size.pt - 9.0) < 0.01


def _cell_fill(cell):
    shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    return shd.get(qn("w:fill")) if shd is not None else None


def test_table_header_zebra_borders_and_margins():
    from docx import Document
    doc = Document()
    sr._add_table(doc, ["الحقيقة", "المصدر"],
                  [["أ", "ب"], ["ج", "د"], ["هـ", "و"]])
    table = doc.tables[-1]
    # رأسٌ أخضر سِلك بخطٍّ أبيض عريض
    assert _cell_fill(table.rows[0].cells[0]) == "166534"
    hdr_run = table.rows[0].cells[0].paragraphs[0].runs[0]
    assert hdr_run.bold and str(hdr_run.font.color.rgb) == "FFFFFF"
    # شريطٌ متناوبٌ أخضرُ خفيف على الصفّ الزوجيّ (١-مرتكز => الصف الثاني بيانات)
    assert _cell_fill(table.rows[2].cells[0]) == "F2F7F3"
    # حدودٌ شعريّةٌ خافتة #BBBBBB
    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    assert borders is not None
    assert borders.find(qn("w:top")).get(qn("w:color")).upper() == "BBBBBB"
    # هوامشُ خليّة داخلية مضبوطة
    mar = table.rows[1].cells[0]._tc.get_or_add_tcPr().find(qn("w:tcMar"))
    assert mar is not None and mar.find(qn("w:left")).get(qn("w:w")) == "120"


def test_dense_table_drops_to_smaller_font():
    from docx import Document
    doc = Document()
    rows = [[str(i), "قيمة"] for i in range(25)]      # >20 => كثيف
    sr._add_table(doc, ["رقم", "قيمة"], rows)
    table = doc.tables[-1]
    body_run = table.rows[5].cells[0].paragraphs[0].runs[0]
    assert abs(body_run.font.size.pt - 9.0) < 0.01


def test_has_plex_arabic_font_is_callable_bool():
    # لا نفترض تثبيت Plex محلّيًّا — نؤكّد فقط أنّ الفحص يعمل ويُرجِع bool.
    assert isinstance(sr.has_plex_arabic_font(), bool)


def test_pdffonts_embeds_plex_regular_and_bold(tmp_path):
    """§7 قبول (قرار المالك): pdffonts يُظهِر IBM Plex Sans Arabic **Regular
    و**Bold مُضمَّنَين في الـPDF المُصيَّر — لا تبديلَ صامتٍ من LibreOffice.

    مُلزَم تحت SILK_PDF_ACCEPTANCE=1 (CI بالخطّ مثبَّتًا). يُتخطّى محلّيًّا حيث
    لا soffice/pdffonts/Plex — كبقية مسار القبول."""
    import os
    import shutil
    import subprocess
    require = os.environ.get("SILK_PDF_ACCEPTANCE") == "1"

    def _gate(reason):
        if require:
            pytest.fail(f"بوابة قبول §7 مُلزَمة وفشل شرطها: {reason}")
        pytest.skip(reason)

    if sr._find_soffice() is None:  # noqa: SLF001
        _gate("محرّك تحويل PDF غائب")
    if not shutil.which("pdffonts"):
        _gate("pdffonts (poppler-utils) غائب")
    if not sr.has_plex_arabic_font():
        _gate("IBM Plex Sans Arabic غير مثبَّت — LibreOffice سيبدّله صامتًا")

    from docx import Document
    doc = Document()
    doc.add_heading("عنوان القسم الرئيسي", level=1)          # Bold Plex
    doc.add_paragraph("هذا نصّ المتن العربي في التقرير للتحقق من تضمين الخطّ.")
    sr._finalize_rtl(doc)  # noqa: SLF001
    docx = os.path.join(str(tmp_path), "t.docx")
    doc.save(docx)
    try:
        pdf = sr.docx_to_pdf(docx, os.path.join(str(tmp_path), "t.pdf"))
    except RuntimeError as e:
        _gate(f"تعذّر تحويل PDF: {e}")
    out = subprocess.run(["pdffonts", pdf], capture_output=True,
                         text=True, timeout=60).stdout
    plex = [ln for ln in out.splitlines() if "plexsansarabic" in ln.lower()]
    assert any("regular" in ln.lower() for ln in plex), out
    assert any("bold" in ln.lower() for ln in plex), out
    # مُضمَّن: عمود emb = yes على أسطر Plex (لا تبديل، لا إحالة خارجية).
    assert all("yes" in ln.lower() for ln in plex), out
