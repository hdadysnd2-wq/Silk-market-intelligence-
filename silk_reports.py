"""مشتقا التقارير لسِلك — Silk report derivatives (wave 5c, vision §10.3-§10.4).

مشتقان جديدان من القالب الموحّد (`silk_render.build_view`) — لا مسار عرض
مستقلاً جديداً، بل اشتقاقان فوق النموذج القانوني نفسه:

  - **التقرير الكامل (Word)** — §10.3: الخلاصة التنفيذية أولاً، ثم
    "موقعك التنافسي"، ثم الأسواق **بسطر مصدر تحت كل رقم** (مبني في القالب
    نفسه — `components_detail` — فيستحيل بنيوياً رقم بلا نسب)، ثم قسم
    **"حدود هذا التقرير" قبل التوصيات**.
  - **المختصر** — §10.4: منتج مختلف لا نسخة مصغرة — صفحة "رسالة جوال":
    القرار + ٣ أرقام حاسمة + سطرا الموقع التنافسي + إحالة اللوحة.

python-docx تبعية اختيارية: غيابها = RuntimeError واضحة (الاستيراد كسول)،
والمختصر نصّ خالص يعمل دوماً. صفر شبكة، صفر تعديل أرقام — عرض صرف.
"""
from __future__ import annotations

import logging
import re

log = logging.getLogger(__name__)

_DOCX_HINT = "python-docx غير مثبتة — pip install python-docx"

# آثار برهانية لا يجوز أن تبلغ تقرير إنتاج أبداً — hermetic-only markers.
_HERMETIC_MARKERS = ("MagicMock", "example.org", "hermetic", "demo double",
                     "بدائل موسومة")


def _assert_production_clean(view: dict) -> None:
    """حارس الإنتاج (إصلاح مراجعة Stage 5): أي أثر برهاني في تشغيلة غير موسومة
    SILK_HERMETIC = رفض توليد التقرير بصوت عالٍ — لا تقرير مسموماً بصمت.
    التشغيلات البرهانية الموسومة تمرّ وتحمل لافتة TEST RUN الظاهرة بدلاً من ذلك.
    """
    import os as _os
    if _os.environ.get("SILK_HERMETIC") or view.get("test_run"):
        return
    import json as _json
    blob = _json.dumps(view, ensure_ascii=False, default=str)
    for marker in _HERMETIC_MARKERS:
        if marker in blob:
            raise RuntimeError(
                f"hermetic artifact '{marker}' found in a production report "
                "view — رفض التوليد: أثر برهاني في تقرير إنتاجي (اضبط "
                "SILK_HERMETIC=1 للتشغيلات البرهانية)")


def _degraded_banner_text(view: dict) -> str | None:
    """نص لافتة التدهور — None إن لم يكن التشغيل متدهوراً (بلاغ حي: بوابة
    ما قبل التشغيل في api.py توسم `degraded=true` فقط عند allow_degraded=
    true صريح أو سباق نادر على حجز الميزانية؛ هذه اللافتة تجعل التدهور
    مرئياً في كل مشتق بدل هيكل يبدو كالمنتج النهائي)."""
    if not view.get("degraded"):
        return None
    reason = view.get("degraded_reason") or "خدمة التحليل الآلي غير متاحة"
    return f"⚠ DEGRADED — نظام الذكاء الاصطناعي غير متاح ({reason})"


def _stamp_degraded_banner(doc, view: dict) -> None:
    """اطبع لافتة تدهور حمراء بارزة — تُستدعى عند الغلاف وأعلى كل قسم رئيسي
    من قسم البحث العميق. لا أثر إن لم يكن التشغيل متدهوراً."""
    text = _degraded_banner_text(view)
    if not text:
        return
    from docx.shared import RGBColor
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)


def _fmt(v: object) -> str:
    """تنسيق قيمة للعرض — display formatting (None = فجوة معلنة)."""
    if v is None:
        return "—"
    if isinstance(v, (int, float)) and not isinstance(v, bool) and abs(v) >= 1000:
        return f"{v:,.0f}"
    return str(v)


def _truncate_at_word(text: str, max_len: int) -> str:
    """قصّ نص عند حدّ كلمة كاملة — بلاغ حي (الموجة ٩): "لا تتوفر من أد"
    (كلمة مقطوعة منتصفها) كانت ناتجة عن قصّ حرفي بلا مراعاة حدود الكلمات
    (هنا، وفي ملاحظة استشهاد `run_llm_agent` — راجع الإصلاح المقابل هناك).
    لا يقصّ أبداً منتصف كلمة؛ يتراجع لآخر مسافة قبل الحد."""
    if len(text) <= max_len:
        return text
    cut = text[:max_len]
    sp = cut.rfind(" ")
    if sp > max_len * 0.5:  # لا تراجُع مفرط إن كانت الكلمة الأخيرة طويلة جداً
        cut = cut[:sp]
    return cut.rstrip() + "…"


def _clean_report_text(s: object, max_len: int = 300) -> str:
    """نص آمن للعرض — بلاغ حي (الموجة ٨): لا يجوز أن يتسرّب رد كلود الخام
    غير المفسَّر (كتلة JSON/سياج ```) لواجهة تقرير — يُستبدَل بملخّص نظيف
    بدل عرضه حرفياً؛ التفاصيل الكاملة تبقى في أثر التتبّع (data/traces/).
    القصّ عند حدّ كلمة (الموجة ٩) — لا كلمة مقطوعة منتصفها بعد الآن."""
    text = str(s or "").strip()
    if not text:
        return text
    if text.lstrip().startswith(("{", "```")):
        return "بند تقني غير قابل للعرض المباشر — التفاصيل الكاملة في أثر التتبّع."
    return _truncate_at_word(text, max_len)


# §5/§6 (أمر العمل الرئيس): قصّ نظيف بلا «…» متدلٍّ + رابط عمومي + مصدر نظيف.
_URL_RE = re.compile(r"https?://[^\s)>\]،؛]+")
_TOOLUSE_RE = re.compile(r"\s*\((?:Claude\s*)?tool[-\s]?use\)\s*", re.I)


def _first_url(*texts: object) -> str:
    """أول رابط http(s) في أيٍّ من النصوص — لعمود «الرابط» في سجل الأدلة.
    لا اختلاق: «—» إن لم يُرصَد رابط فعلي (DataPoint لا يحمل حقل رابط مستقلّ)."""
    for t in texts:
        m = _URL_RE.search(str(t or ""))
        if m:
            return m.group(0).rstrip(".,،؛)")
    return "—"


def _evidence_url(note: object, source: object, value: object) -> str:
    """رابط عمود «الرابط» في سجل الأدلة (§6، أمر العمل الرئيس):
    ١) الأولوية لرابطٍ **محدّد** مرصودٍ في نصّ الحقيقة/الملاحظة (نتيجة بحثٍ مثلاً)؛
    ٢) وإلا الرابطُ العموميُّ الرسميُّ للمصدر المسمّى من سجلّ `silk_data_layer`
       (UN Comtrade → comtradeplus.un.org …) — رابطٌ حقيقيٌّ للتحقق من المجموعة؛
    ٣) وإلا «—».
    لا اختلاق: «—» صادقةٌ حين لا رابطَ محدّدٌ ولا مصدرٌ عموميٌّ معروف. لا نبحث
    في اسم المصدر عن رابط (لن يحمله)، فالبحثُ عن الرابط المحدّد في النصّ فقط."""
    scraped = _first_url(note, value)
    if scraped != "—":
        return scraped
    from silk_data_layer import public_source_url
    return public_source_url(source) or "—"


def _clean_source_label(s: object) -> str:
    """اسم مصدر نظيف لسجل الأدلة (§2/§6) — يُزال وسم «(Claude tool-use)»
    وأي بادئة وكيل داخلية، فيبقى اسم المصدر العمومي (UN Comtrade …) وحده.
    شبكة أمان أخيرة فوق الإصلاح في المصدر (silk_llm_runtime finding assembly)."""
    txt = _TOOLUSE_RE.sub(" ", str(s or ""))
    txt = re.sub(r"LLM(?:Mission)?Agent:\s*[A-Za-z_]+", "", txt)
    return re.sub(r"\s{2,}", " ", txt).strip(" ،-—") or "—"


def _trim_sentence(s: object, max_len: int = 240) -> str:
    """قصّ نظيف عند حدّ جملة (§5/§6 — «سجل الأدلة» لا يعرض حقيقة مبتورة
    بـ«…»): النص كاملاً إن كان ضمن الحدّ، وإلا يُقصّ عند آخر علامة ترقيم
    ختامية قبل الحدّ (. ! ؟ ؛ ،) بلا نقاط حذف؛ فإن غابت فعند آخر كلمة كاملة."""
    text = str(s or "").strip()
    if not text:
        return text
    if text.lstrip().startswith(("{", "```")):
        return "بند تقني غير قابل للعرض المباشر — التفاصيل في أثر التتبّع."
    if len(text) <= max_len:
        return text
    cut = text[:max_len]
    best = max((cut.rfind(ch) for ch in ".!؟؛،"), default=-1)
    if best > max_len * 0.4:
        return cut[:best + 1].rstrip()
    sp = cut.rfind(" ")
    return (cut[:sp] if sp > 0 else cut).rstrip()


# WP-2 §1 — النصّ النائب التقني (يبقى في المسارات الداخلية/?internal=1 فقط).
_UNRENDERABLE_NOTE = "بند تقني غير قابل للعرض المباشر"


def _client_prose(s: object, max_len: int = 400) -> str:
    """نص آمن لمتن **العميل** — WP-2 §1: لا نصّ نائب يصل العميل أبداً.
    كتلة JSON/سياج من ردٍّ خام: يُحاوَل استخلاص المضمون (نفس مستخلص
    `silk_render._strip_raw_json_leak`)، وإن تعذّر تُسقَط الكتلة ("") —
    فيخلو القسم وتلتقطه بوابة الجودة (FAIL يمنع التسليم) بدل تسليم نائبٍ
    مثل «بند تقني غير قابل للعرض المباشر — التفاصيل في أثر التتبع»."""
    text = str(s or "").strip()
    if not text:
        return ""
    if text.lstrip().startswith(("{", "```")):
        try:
            from silk_render import _strip_raw_json_leak
            extracted = str(_strip_raw_json_leak(text) or "").strip()
        except Exception:  # noqa: BLE001 — تعذّر الاستخلاص = إسقاط
            extracted = ""
        if (not extracted or extracted.lstrip().startswith(("{", "```"))
                or "تعذّر تفسير" in extracted):
            return ""
        text = extracted
    return _trim_sentence(text, max_len)


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")
_CODE_SPAN_RE = re.compile(r"`([^`]*)`")


def _strip_inline_markdown(line: str) -> str:
    """أزل تنسيق Markdown من فقرة عادية — بلاغ حي (الموجة ٩): "**عريض**"/
    "`كود`"/"#" خام كانت تظهر حرفياً على وجه التقرير. العناوين '## '/'### '
    تتحوّل لعناوين حقيقية في موضع آخر (قبل بلوغ هذه الدالة) — هذه للفقرات
    العادية التي تحوي تنسيقاً inline فقط."""
    line = _BOLD_RE.sub(r"\1", line)
    line = _ITALIC_RE.sub(r"\1", line)
    line = _CODE_SPAN_RE.sub(r"\1", line)
    if line.lstrip().startswith("#"):
        line = line.lstrip("#").strip()
    return line


def _is_markdown_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def _is_markdown_table_separator(line: str) -> bool:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return bool(cells) and all(c and set(c) <= {"-", ":"} for c in cells)


# شارة الأدلة رُحِّلت إلى silk_narrative.evidence_badge (P2) — لتُستعمل من
# نموذج العرض (silk_render._deep_research_view) لا من طبقة عرض النص وحدها.
# الاسم القديم يبقى هنا كتوافق خلفي لكل مستدعي هذا الملف.
from silk_narrative import (  # noqa: E402
    EVIDENCE_SECONDARY_MIN as _EVIDENCE_SECONDARY_MIN,
    EVIDENCE_VERIFIED_MIN as _EVIDENCE_VERIFIED_MIN,
    evidence_badge as _evidence_badge,
)


# سدّ تسريب (الطبقة ٦): _verdict_tone/_VERDICT_LABELS_AR انتقلتا إلى
# silk_render.py — مصدر واحد يستهلكه غلاف/خلاصة docx هنا ولوحة الويب معاً
# (view["deep_research"]["verdict_tone"/"verdict_label"])، بدل نسخة بايثون
# ونسخة JS منفصلتين قد تختلفان لنفس الرمز.
from silk_render import _VERDICT_LABELS_AR, _verdict_tone  # noqa: E402


# ── هوية سِلك البصرية (الموجة ١١، §11.1) — config/branding.yaml ─────────

_BRANDING_PATH = "config/branding.yaml"
_BRANDING_DEFAULTS = {
    "logo_path": "", "primary_color": "1B3B6F", "secondary_color": "C9A227",
    "contact_footer": "سِلك لذكاء الأسواق",
}


def _load_branding(path: str = _BRANDING_PATH) -> dict:
    """اقرأ هوية سِلك البصرية — محلّل مسطّح خفيف (key: value فقط، بلا
    تعشيش) بلا PyYAML (ليست ضمن requirements.txt؛ المشروع stdlib-first —
    راجع تعليق أعلى config/branding.yaml). ملف غائب/سطر مشوَّه = القيمة
    الافتراضية لذلك المفتاح فقط، لا فشل كامل."""
    out = dict(_BRANDING_DEFAULTS)
    try:
        with open(path, encoding="utf-8") as f:
            for line in f:
                s = line.strip()
                if not s or s.startswith("#") or ":" not in s:
                    continue
                k, v = s.split(":", 1)
                k, v = k.strip(), v.strip()
                if k in out and v:
                    out[k] = v
    except FileNotFoundError:
        pass
    except Exception as e:  # noqa: BLE001 — الهوية تحسين عرض لا شرط توليد
        log.warning("branding config unavailable (%s): %s", path, e)
    return out


def _hex_to_rgbcolor(hex_str: str):
    from docx.shared import RGBColor
    h = (hex_str or "").lstrip("#")
    try:
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    except Exception:  # noqa: BLE001 — لون مشوَّه = الكحلي الافتراضي
        return RGBColor(0x1B, 0x3B, 0x6F)


def _set_cell_shading(cell, hex_color: str) -> None:
    """ظلّل خلفية خلية جدول — python-docx لا يعرض هذا كخاصية عليا؛ عنصر
    w:shd عبر oxml مباشرة (نمط موثَّق شائع لتلوين رؤوس/أشرطة الجداول)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), (hex_color or "").lstrip("#") or "FFFFFF")
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_margins(cell, top=100, bottom=100, left=120, right=120) -> None:
    """§7: هوامشُ خليّةٍ داخلية (twips) — تنفّسٌ حول النصّ كالمنصّات المرجعية."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcpr = cell._tc.get_or_add_tcPr()
    mar = tcpr.find(qn("w:tcMar"))          # عناصرُ lxml الفارغةُ falsy — نستعمل
    if mar is None:                          # is None صراحةً لا `or` (فخّ معروف)
        mar = OxmlElement("w:tcMar")
        tcpr.append(mar)
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        el = mar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            mar.append(el)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")


def _set_table_borders(table, hex_color: "str | None" = None) -> None:
    """§7: حدودٌ شعريّةٌ ناعمة (single, size 4 = ½pt) بلونٍ خافت بدل أسود
    Table-Grid الافتراضي — مظهرٌ احترافيّ هادئ."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    hex_color = hex_color or _TABLE_BORDER
    tblpr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), (hex_color or "").lstrip("#"))
        borders.append(el)
    existing = tblpr.find(qn("w:tblBorders"))
    if existing is not None:
        tblpr.remove(existing)
    tblpr.append(borders)


def _apply_typography(doc) -> None:
    """§7 (قرار المالك): نظامُ الطباعة على مستوى الأنماط — متنٌ IBM Plex،
    تباعدُ أسطرٍ ١٫٤٤، عناوينُ خضراءُ مُدرَّجة، ترويسة/تذييلٌ رماديّ. يُطبَّق
    مرّةً على المستند (الأنماطُ تُشار إليها بالاسم، فيسري على كل المحتوى)."""
    from docx.shared import Pt, RGBColor
    styles = doc.styles

    def _style(name):
        try:
            return styles[name]
        except KeyError:
            return None

    normal = _style("Normal")
    if normal is not None:
        normal.font.name = _RTL_BODY_FONT
        normal.font.size = Pt(_TYPO["body_pt"])
        pf = normal.paragraph_format
        pf.line_spacing = _TYPO["line_spacing"]
        pf.space_after = Pt(_TYPO["space_after_pt"])
    for name, key in (("Title", "title"), ("Heading 1", "h1"),
                      ("Heading 2", "h2"), ("Heading 3", "h3"),
                      ("Header", "header_footer"), ("Footer", "header_footer")):
        st = _style(name)
        if st is None:
            continue
        pt, hexc = _TYPO[key]
        st.font.name = _RTL_BODY_FONT
        st.font.size = Pt(pt)
        st.font.color.rgb = RGBColor.from_string(hexc)
        if key in ("title", "h1", "h2", "h3"):
            st.font.bold = True


# §7 (قرار المالك — ترقيةُ الطباعة، مُثبَّتٌ على تقريرٍ حقيقيّ): العائلةُ
# الرسميةُ IBM Plex Sans Arabic (OFL) — Regular للمتن، Bold للعناوين/رؤوس
# الجداول. تُضبط على ascii+cs فتُشكَّل الحروفُ موصولةً بلا مربّعات tofu. يجب
# أن تكون مثبَّتةً في بيئة التصيير (فحصُ الحضور يفشل بصوتٍ عالٍ إن غابت — لا
# نترك LibreOffice يبدّلها صامتًا).
_RTL_BODY_FONT = "IBM Plex Sans Arabic"

# §7 نظامُ الطباعة (وحدات docx). أنصافُ نقاطٍ للحجم؛ الألوانُ HEX بلا #.
_TYPO = {
    "body_pt": 11.0,          # 22 نصفَ نقطة
    "line_spacing": 345 / 240,  # 1.4375 (سطرٌ ≈١٫٤٤)
    "space_after_pt": 9.0,    # 180 twip
    "title": (18.0, "166534"),          # 36 نصفَ نقطة، عريض، موسَّط
    "h1": (14.0, "166534"), "h2": (13.0, "166534"), "h3": (12.0, "333333"),
    "header_footer": (9.0, "555555"),   # 18 نصفَ نقطة، رماديّ
}
_TABLE_HEADER_FILL = "166534"
_TABLE_ZEBRA_FILL = "F2F7F3"
_TABLE_BORDER = "BBBBBB"
_TABLE_DENSE_ROWS = 20        # جداولٌ >٢٠ صفًّا تنزل لـ٩ نقاط (18 نصفَ نقطة)
_TABLE_DENSE_PT = 9.0


def _set_rtl_paragraph(ppr) -> None:
    """§4 (مُصحَّح، مُتحقَّق تجريبياً): أضِف <w:bidi/> + محاذاة **منطقية**
    START على خصائص فقرة (pPr) عبر oxml.

    **حرِج**: مع فقرة bidi يفسّر OOXML قيمة w:jc **منطقياً** — فـ`right`
    تُصيَّر بصرياً **يساراً** في العربية (انقلاب). الصحيح `start` (يُصيَّر
    محاذاةً لليمين في RTL؛ وهو الافتراضي أصلاً). لا تستعمل right/end/both
    أبداً. العناوين/التذييلات الموسَّطة (`center`) تبقى موسَّطةً كما هي."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    if ppr.find(qn("w:bidi")) is None:
        ppr.append(OxmlElement("w:bidi"))
    jc = ppr.find(qn("w:jc"))
    if jc is not None and jc.get(qn("w:val")) == "center":
        return  # عنوان/تذييل موسَّط — يبقى موسَّطاً
    if jc is None:
        jc = OxmlElement("w:jc")
        ppr.append(jc)
    # START (لا right/end/both) — الانقلاب المنطقي يجعل right تُصيَّر يساراً.
    jc.set(qn("w:val"), "start")


def _set_rtl_run_fonts(rpr, font: str = _RTL_BODY_FONT) -> None:
    """§4: أضِف <w:rtl/> + خطّ عربي (ascii+hAnsi+cs مع szCs) على rPr — كي
    يُشكَّل النصّ العربي فعلاً (python-docx لا يضبط w:cs)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    if rpr.find(qn("w:rtl")) is None:
        rpr.append(OxmlElement("w:rtl"))
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), font)
    # szCs بجانب sz (المطلوب في §4) — إن وُجد sz انسخه إلى szCs.
    sz = rpr.find(qn("w:sz"))
    if sz is not None and rpr.find(qn("w:szCs")) is None:
        szcs = OxmlElement("w:szCs")
        szcs.set(qn("w:val"), sz.get(qn("w:val")))
        rpr.append(szcs)


def _set_table_rtl(table) -> None:
    """§4: اجعل الجدول من اليمين لليسار — <w:bidiVisual/> على tblPr (تتدفّق
    الأعمدة يميناً) + محاذاة كل خلية يميناً + عرض أعمدة حقيقي بوحدات DXA."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Twips
    tblPr = table._tbl.tblPr
    if tblPr.find(qn("w:bidiVisual")) is None:
        tblPr.append(OxmlElement("w:bidiVisual"))
    # عرض أعمدة حقيقي (DXA): توزيع عرض الصفحة المتاح (~9360 twips لـA4 بهوامش)
    # بالتساوي — الجداول لم تكن تحدّد أيّ عرض (autofit فقط).
    cols = len(table.columns)
    if cols:
        table.autofit = False
        total = 9360
        each = Twips(total // cols)
        for col in table.columns:
            for cell in col.cells:
                cell.width = each
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _set_rtl_paragraph(p._p.get_or_add_pPr())


def _apply_rtl(doc, font: str = _RTL_BODY_FONT) -> None:
    """§4 (أمر العمل الرئيس): اجعل المستند كلّه من اليمين لليسار على مستوى
    الافتراضات لا ترقيعاً لكل عنصر — (١) كل مقطع يحمل <w:bidi/>؛ (٢) نمط
    Normal وكل أنماط الفقرات/العناوين: bidi + محاذاة يمين + <w:rtl/> + خطّ
    عربي (ascii+cs). الرموز اللاتينية (الدولار، HS، الروابط، أرقام اللوائح)
    تبقى LTR داخل الفقرة تلقائياً عبر خوارزمية bidi (اتجاه الأساس RTL)."""
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    for section in doc.sections:
        sectPr = section._sectPr
        if sectPr.find(qn("w:bidi")) is None:
            sectPr.append(OxmlElement("w:bidi"))   # <w:jc> غير صالح على sectPr
    for style in doc.styles:
        if getattr(style, "type", None) != WD_STYLE_TYPE.PARAGRAPH:
            continue
        try:
            _set_rtl_paragraph(style.element.get_or_add_pPr())
            _set_rtl_run_fonts(style.element.get_or_add_rPr(), font)
        except Exception:  # noqa: BLE001 — نمط بلا pPr/rPr صالح يُتخطّى
            continue


# WP-5 — انعكاس الأقواس في PDF: فقرة عربية bidi تحوي مقطعاً لاتينياً/رقمياً
# بين قوسين كانت تُصيَّر «) ... (» بعد تحويل LibreOffice headless — خط
# الأنابيب لم يكن يحقن أي عزل اتجاه حول المقاطع مختلطة الاتجاه. الفكس:
# علامة RLM (U+200F) بعد القوس الافتتاحي وقبل الختامي حين يحوي المقطع
# لاتينية/أرقاماً في سياق عربي — فيلتصق القوسان بجارٍ قويّ الاتجاه RTL
# ويُصيَّران باتجاههما الصحيح. لا تغيير في أي محتوى مرئي (RLM غير مرئية).
_RLM = "\u200f"
_BRACKET_SPAN_RE = re.compile(r"([(\[])([^()\[\]\n]*[A-Za-z0-9][^()\[\]\n]*)([)\]])")
_ARABIC_CHAR_RE = re.compile("[\\u0600-\\u06ff]")


def _bidi_isolate_brackets(text: str) -> str:
    """احقن RLM داخل أقواس المقاطع اللاتينية/الرقمية في نصٍّ عربي السياق —
    تُطبَّق قبل `_finalize_rtl` على كل run. نص بلا عربية يمرّ كما هو."""
    if not text or not _ARABIC_CHAR_RE.search(text):
        return text

    def _wrap(m: "re.Match") -> str:
        inner = m.group(2)
        if inner.startswith(_RLM) and inner.endswith(_RLM):
            return m.group(0)   # مُعالَج سلفاً — لا ازدواج
        return f"{m.group(1)}{_RLM}{inner}{_RLM}{m.group(3)}"

    return _BRACKET_SPAN_RE.sub(_wrap, text)


# WP-5 §2 — الفحص الآلي على نصّ الـPDF المستخرَج: قوس افتتاحي يتبعه فراغ/
# نهاية سطر أثرُ انعكاسٍ نموذجي («) ... (» تُستخرَج "(" معلّقةً قبل فراغ).
_SUSPICIOUS_OPEN_BRACKET_RE = re.compile(r"\((?=\s|$)", re.M)


def count_suspicious_brackets(text: str) -> int:
    """عدد الأقواس الافتتاحية المعلّقة (يتبعها فراغ/نهاية سطر) في نصٍّ
    مستخرج من PDF — مقياس انعكاس الأقواس، صفر على مستند سليم تقريباً."""
    return len(_SUSPICIOUS_OPEN_BRACKET_RE.findall(text or ""))


def _pdf_bracket_check(pdf_path: str) -> None:
    """WP-5: افحص الـPDF النهائي — فوق العتبة (SILK_PDF_BRACKET_FAIL_MAX،
    افتراضياً 3) يفشل التصدير بصوت عالٍ بدل تسليم مستند بأقواس معكوسة.
    بلا pymupdf (بيئة بلا أداة استخراج) يُتخطّى الفحص — لا ادعاء فحصٍ
    لم يحدث (يُسجَّل سطر تشخيص فقط)."""
    try:
        import fitz  # pymupdf — حاضرة على بيئة e2e/الإنتاج
    except ImportError:
        log.info("pdf bracket check skipped: pymupdf غير مثبّتة")
        return
    try:
        with fitz.open(pdf_path) as pdf:
            text = "\n".join(page.get_text() for page in pdf)
    except Exception as e:  # noqa: BLE001 — تعذّر الاستخراج ≠ مستند معكوس
        log.warning("pdf bracket check extraction failed: %s", e)
        return
    import os
    n = count_suspicious_brackets(text)
    # مراجعة شيفرة PR #147: قيمة بيئة مشوَّهة كانت تُفجِّر ValueError خاماً
    # فتقتل كل تصدير PDF — تراجُع آمن للافتراضي بدل الانهيار.
    try:
        limit = int(os.environ.get("SILK_PDF_BRACKET_FAIL_MAX", "3"))
    except (TypeError, ValueError):
        limit = 3
    if n > limit:
        raise RuntimeError(
            f"فشل فحص اتجاه الأقواس في الـPDF النهائي: {n} قوساً افتتاحياً "
            f"معلّقاً (العتبة {limit}) — مؤشر انعكاس أقواس RTL؛ لا يُسلَّم "
            "مستند معكوس الأقواس")


def _finalize_rtl(doc, font: str = _RTL_BODY_FONT) -> None:
    """§4: تمريرة ختامية تضمن أن **كل** فقرة (bidi+محاذاة يمين) و**كل** run
    (<w:rtl/> + خطّ عربي ascii+cs) تحمل الاتجاه صراحةً — لا اعتماداً على وراثة
    النمط وحدها (بعض إصدارات Word لا تُورِّث rtl للتنسيق المباشر). تشمل المتن
    والجداول (المتداخلة) وترويسات/تذييلات المقاطع."""
    def _do_paragraph(p):
        _set_rtl_paragraph(p._p.get_or_add_pPr())
        for run in p.runs:
            # WP-5: عزل اتجاه أقواس المقاطع اللاتينية/الرقمية (RLM) قبل
            # ضبط rtl — يمنع انعكاس «) ... (» في تحويل PDF.
            if run.text:
                isolated = _bidi_isolate_brackets(run.text)
                if isolated != run.text:
                    run.text = isolated
            _set_rtl_run_fonts(run._r.get_or_add_rPr(), font)
            # Wave 2 (البند ٨): العلامة «سِلك» => «سلك» متّصلة في **الوورد**
            # (وثيقة المشغّل القابلة للتحرير). أمّا التجريد الكامل لكلّ الحركات
            # فيقع عند حدّ **الـPDF** (المُسلَّم النهائي للعميل) في `docx_to_pdf`،
            # كي تبقى وثيقةُ الوورد مشكّلةً كمصدرٍ، ويخرج الـPDF مجرّدًا (لا كلمةَ
            # يشقّها مِحرفٌ مُركَّب في الاستخراج — تعريفُ عائلة E). الطيّة E تُثبَّت
            # على استخراج الـPDF لا على الوورد.
            if run.text and "سِلك" in run.text:
                run.text = run.text.replace("سِلك", "سلك")

    def _walk_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _do_paragraph(p)
                    _walk_tables(cell.tables)   # جداول متداخلة (SWOT وغيرها)

    _apply_typography(doc)   # §7: نظامُ الطباعة (خطّ/حجم/لون/تباعد) مرّةً
    for p in doc.paragraphs:
        _do_paragraph(p)
    _walk_tables(doc.tables)
    from docx.shared import Mm
    for section in doc.sections:
        # Wave 2 (البند ٩): مقاس A4 صراحةً (كان الافتراضي Letter الأمريكي، بينما
        # حساب عرض الأعمدة يفترض A4 ~9360 twips — تصحيح عدم التطابق). 210×297مم.
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                _do_paragraph(p)


def _add_page_number_field(paragraph) -> None:
    """أدرج حقل رقم الصفحة الديناميكي (PAGE) — python-docx لا يعرض حقول
    Word كخاصية عليا؛ نمط oxml موثَّق شائع (w:fldChar begin/instrText/end)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


# Wave 2 (البند ٨ + الطيّة E): «سِلك» كانت تُكسَر «ِس لك» في PDF (٥٠×، الغلاف +
# كل تذييل)، وكذلك **أشقّاؤها** — أيّ كلمةٍ عربيةٍ مشكّلةٍ في النصّ (مثل «مُوصًى»
# => « ُموًصى»، «يُتَّخذ» => « ُيَّتخذ») — تنفصل حركاتُها عند تشكيل الخط في
# LibreOffice فتُشقّ الكلمة في الاستخراج. الإصلاح **قاعدةُ حدٍّ لا حالة**: تجريدُ
# الحركات (U+064B–U+0652، U+0670) من **مخرَج** runs العميل عند حدّ العرض — القوالب
# المصدرية تبقى مشكّلةً، والمخرَج مجرّد. القفل البصري (E) هو تعريفُ العائلة: لا
# كلمةَ عربيةٍ يشقّها مِحرفٌ مُركَّب في استخراج PDF.
_AR_COMBINING = "".join(chr(c) for c in list(range(0x064B, 0x0653)) + [0x0670])
_AR_COMBINING_RE = re.compile("[" + re.escape(_AR_COMBINING) + "]")


def _shape_safe_ar(text: str) -> str:
    """جرّد الحركات المُركَّبة (U+064B–U+0652، U+0670) فيخرج النصّ العربيّ متّصلًا
    في الاستخراج — قاعدةُ حدِّ العرض للمخرَج (المصادر تبقى مشكّلة)."""
    return _AR_COMBINING_RE.sub("", text or "")


def _add_cover_wordmark(doc, branding: dict) -> None:
    """شعار سِلك على الغلاف — صورة فعلية إن وُجد `logo_path` صالح، وإلا
    علامة اسمية نصّية حقيقية «سِلك» بلون العلامة الأساس (§7، أمر العمل
    الرئيس: لا نصّ نائب بين أقواس «[شعار سِلك]»). لا استثناء يُسقِط التوليد
    إن تعذّرت قراءة الصورة (مسار خاطئ/صيغة غير مدعومة) — رجوع للعلامة النصّية."""
    logo_path = branding.get("logo_path")
    if logo_path:
        try:
            import os
            if os.path.exists(logo_path):
                doc.add_picture(logo_path, width=_docx_inches(1.5))
                return
        except Exception as e:  # noqa: BLE001 — الشعار تحسين عرض لا شرط
            log.warning("cover logo unavailable (%s): %s", logo_path, e)
    from docx.shared import Pt
    p = doc.add_paragraph()
    run = p.add_run(_shape_safe_ar("سِلك"))   # «سلك» متّصلة (لا «ِس لك»)، لا نائب
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = _hex_to_rgbcolor(branding["primary_color"])


def _docx_inches(n: float):
    from docx.shared import Inches
    return Inches(n)


def _add_page_header_footer(doc, title: str) -> None:
    """رأس/تذييل ثابتان لكل صفحة — عنوان التقرير أعلى، ورقم صفحة ديناميكي
    + سطر هوية سِلك أسفل (الموجة ١١، §11.1: تناسق بصري عبر كل صفحات
    التقرير، لا الغلاف فقط)."""
    branding = _load_branding()
    section = doc.sections[0]
    hp = section.header.paragraphs[0] if section.header.paragraphs \
        else section.header.add_paragraph()
    hp.text = _shape_safe_ar(title)
    fp = section.footer.paragraphs[0] if section.footer.paragraphs \
        else section.footer.add_paragraph()
    fp.add_run(_shape_safe_ar(branding["contact_footer"]) + " — صفحة ")
    _add_page_number_field(fp)


# لون شارة الحكم بالـtone — conditional (دخول مشروط) لون مستقل (أخضر مزرقّ)
# يميّزه عن go الأخضر وwatch الكهرماني (بلاغ مراجعة المالك على النموذج).
_VERDICT_TEXT_COLORS = {"go": (0x1E, 0x7D, 0x32), "conditional": (0x00, 0x69, 0x5C),
                        "watch": (0xB8, 0x86, 0x0B),
                        "nogo": (0xC0, 0x00, 0x00), "unknown": (0x60, 0x60, 0x60)}
_VERDICT_HIGHLIGHTS = {"go": "BRIGHT_GREEN", "conditional": "TEAL",
                       "watch": "YELLOW", "nogo": "RED"}


def _add_verdict_badge(doc, vtxt: str) -> None:
    """شارة حكم ملوّنة على الغلاف — بلاغ حي (الموجة ٩، P0-A): "درجات تبدو
    بلا سند" — الحكم النصي وحده مدفون بين حقول الجدول؛ شارة بصرية بارزة
    (تظليل + لون + حجم) تجعل التوصية أول ما تراه العين، كتقرير احترافي."""
    from docx.shared import Pt, RGBColor
    tone = _verdict_tone(vtxt)
    p = doc.add_paragraph()
    run = p.add_run(f"  {_VERDICT_LABELS_AR[tone]}  ")
    run.bold = True
    run.font.size = Pt(16)
    rgb = _VERDICT_TEXT_COLORS[tone]
    run.font.color.rgb = RGBColor(*rgb)
    try:
        from docx.enum.text import WD_COLOR_INDEX
        name = _VERDICT_HIGHLIGHTS.get(tone)
        if name:
            run.font.highlight_color = getattr(WD_COLOR_INDEX, name)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    except Exception:  # noqa: BLE001 — التظليل تحسين عرض لا شرط توليد
        pass


# ═══ Master Prompt Part 2 §B — بوابة اتساق الحكم عند التسليم ═══
#
# الحكم حقلٌ واحدٌ (verdict_tone) تشتقّ منه كل مواضع العرض الأربعة: شارة
# الغلاف، صفّ الجدول («الحكم»/«التوصية»)، وسطر «الحكم:»/«التوصية:» في قسم
# القرار. الفحص هنا **تصريحيّ محض** — يلتقط فقط هذه المواضع الأربعة، لا
# يمسح كامل السرد: نقاش «شرطا قلب الحكم» (LESSONS ٣٢) قد يذكر تصنيفاً آخر
# افتراضياً بصياغة "لو تحقّق كذا لتحوّل الحكم إلى دخول مشروط" بلا أيّ تناقضٍ
# حقيقي — مسحٌ عامٌّ كان سيُبلِغ زائفاً هنا (نفس عائلة أخطاء false-positive
# التي عضّت CAGR/العملة سابقاً، LESSONS ٤٢). هذه بوابة تسليم صلبة: تُشغَّل
# بعد بناء المستند مباشرة وقبل الحفظ، وترفع RuntimeError على أي تعارض حقيقي
# — لا تسجيل صامت.
_VERDICT_DECL_RE = re.compile(r"^(?:الحكم|التوصية)\s*:\s*(.+)$")


def _match_known_verdict_label(s: str) -> "str | None":
    """صنِّف نصّاً كتسمية حكمٍ معروفة إن كان **يبدأ** بإحدى تسميات
    `_VERDICT_LABELS_AR` (مطابقة بادئة لا مساواة سطر كامل) — سطر «مختصر»
    قد يُلحِق سياقاً بعد التسمية («مراقبة السوق — سوق الكويت (بحث عميق
    شامل)») بلا أن يكون هذا تناقضاً؛ نصٌّ لا يبدأ بأيّ تسمية معروفة يُهمَل
    (ليس تصريحاً بحكمٍ أصلاً) بدل تصنيفه تعارضاً زائفاً."""
    s = s.strip()
    for lbl in _VERDICT_LABELS_AR.values():
        if s == lbl or s.startswith(lbl + " ") or s.startswith(lbl + "—") \
                or s.startswith(lbl + "،"):
            return lbl
    return None


def _resolve_vtxt(dr: dict) -> str:
    """سلسلة الحكم الخام الواحدة (GO/WATCH/...) — نقطة اشتقاقٍ مشتركة
    (Master Prompt Part 2 §B، البند ٤: لا يُشتقّ الحكم من نصٍّ منفصل في أكثر
    من موضع). WP-1: **الحكم الحتمي أولاً** عبر المصدر الواحد
    `silk_narrative.authoritative_verdict` — `ai.verdict` قراءة استشارية
    داخلية لا توصية معروضة (كان الترتيب معكوساً فتناقضت الشارة مع المتن)."""
    from silk_narrative import authoritative_verdict
    raw, _ = authoritative_verdict((dr or {}).get("verdict"))
    return raw or ""


def _declared_verdict_labels(doc) -> list[str]:
    """كل ذكرٍ تصريحيّ لتسمية الحكم في المواضع الثلاثة القابلة للفحص داخل
    كائن docx: شارة الغلاف (فقرة نصّها بالضبط إحدى تسميات _VERDICT_LABELS_AR)،
    سطر «الحكم:»/«التوصية:» في متن الفقرات، وصفّ جدول أول عموده «الحكم» أو
    «التوصية»."""
    out: list[str] = []
    labels = set(_VERDICT_LABELS_AR.values())
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if t in labels:  # شارة الغلاف: تطابق تامّ (فقرة قصيرة مستقلة) لا بادئة
            out.append(t)
            continue
        m = _VERDICT_DECL_RE.match(t)
        if m:
            lbl = _match_known_verdict_label(m.group(1))
            if lbl:
                out.append(lbl)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and cells[0] in ("الحكم", "التوصية"):
                lbl = _match_known_verdict_label(cells[1])
                if lbl:
                    out.append(lbl)
    return out


def _assert_verdict_consistency_doc(doc, vtxt: str, where: str) -> None:
    """بوابة تسليم: كل ذكرٍ تصريحيّ للحكم في المستند المُولَّد فعلياً يجب أن
    يطابق الحكم القانوني الواحد (`_verdict_tone(vtxt)`). تعارضٌ = رفضُ
    التوليد قبل وصول المستند للعميل/المشغّل، لا بعد بلاغٍ حي (Master Prompt
    Part 2 §B، البندان ٤-٥)."""
    canonical_label = _VERDICT_LABELS_AR[_verdict_tone(vtxt)]
    conflicting = sorted({d for d in _declared_verdict_labels(doc)
                          if d and d != canonical_label})
    if conflicting:
        raise RuntimeError(
            f"تناقض حكمٍ في {where}: الحكم القانوني الواحد "
            f"'{canonical_label}' بينما مواضع عرضٍ أخرى (الشارة/الجدول/"
            f"سطر القرار) تذكر '{'، '.join(conflicting)}' — كل موضع عرضٍ "
            "للحكم يُشتَقّ من الحقل الواحد verdict_tone، لا نصّاً منفصلاً "
            "قد يتباعد عنه")


def _declared_verdict_labels_text(blob: str) -> list[str]:
    """معادل نصّي (Markdown لا docx) لـ`_declared_verdict_labels` — يفحص
    سطر «- التوصية: **س**» وصفّ جدول «| الحكم | س |» فقط، لا كامل السرد."""
    out: list[str] = []
    for line in blob.splitlines():
        s = line.strip()
        m = re.match(r"^-\s*(?:الحكم|التوصية)\s*:\s*\*{0,2}(.+?)\*{0,2}\s*$", s)
        if m:
            lbl = _match_known_verdict_label(m.group(1))
            if lbl:
                out.append(lbl)
            continue
        m2 = re.match(r"^\|\s*(?:الحكم|التوصية)\s*\|\s*(.+?)\s*\|$", s)
        if m2:
            lbl = _match_known_verdict_label(m2.group(1))
            if lbl:
                out.append(lbl)
    return out


def _assert_verdict_consistency_text(blob: str, vtxt: str, where: str) -> None:
    """معادل Markdown لـ`_assert_verdict_consistency_doc`."""
    canonical_label = _VERDICT_LABELS_AR[_verdict_tone(vtxt)]
    conflicting = sorted({d for d in _declared_verdict_labels_text(blob)
                          if d and d != canonical_label})
    if conflicting:
        raise RuntimeError(
            f"تناقض حكمٍ في {where}: الحكم القانوني الواحد "
            f"'{canonical_label}' بينما مواضع عرضٍ أخرى تذكر "
            f"'{'، '.join(conflicting)}'")


def _render_markdown_table(doc, table_lines: list[str]) -> None:
    """حوّل جدول Markdown (| عمود | عمود |) لجدول Word حقيقي — بلاغ حي
    (الموجة ٩): سلاسل رقمية (تدفقات استيراد، سلّم أسعار، اشتراطات،
    ديموغرافيا) كانت تُعرض نقاطاً سردية مبعثرة بدل جدول واحد يقارَن بنظرة."""
    rows = []
    for ln in table_lines:
        if _is_markdown_table_separator(ln):
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:  # رأس فقط أو فارغ — لا جدول ذو معنى
        return
    headers, *data = rows
    n = len(headers)
    norm = [(r + [""] * n)[:n] for r in data]
    # §7 (أمر العمل الرئيس): لا تعليق آليّ («جدول: المؤشر · القيمة») قبل
    # الجداول — عناوينها من ترويستها. عنوان حقيقي أو لا شيء، لا وصف آلة.
    _add_table(doc, headers, norm, caption=None)


def _narrative_exec_summary(view: dict) -> list[str]:
    """خلاصة تنفيذية — التحليل الاحترافي (ai_report) إن توفر، وإلا exec_summary.

    `view["ai_report"]` (silk_ai_judge.ai_report، كلود) فقرات سردية احترافية
    مبنية على حزمة البحث الكاملة للسوق الأول + الأسواق المرتّبة؛ عند غيابه
    (لا مفتاح/فشل النداء) يُرجَع لـ `exec_summary` الحتمية في silk_narrative
    (النسخة الأولى هنا كانت تعيد صياغة الحقول لكنها ما تزال تمرّر «السبب
    التجاري: score 0.636 في النطاق الشرطي» حرفياً — شرط كود على وجه التقرير؛
    exec_summary تبني الفقرات الثلاث من الأرقام المرصودة نفسها بعربية بشرية،
    بلا درجة معيارية ولا اسم وكيل ولا شرط كود). كلا المسارين مبنيّ حصراً على
    حقول محسوبة فعلاً — لا اختلاق.
    """
    ai_text = view.get("ai_report")
    if ai_text:
        return [p.strip() for p in str(ai_text).split("\n") if p.strip()]
    from silk_narrative import exec_summary
    paras = exec_summary(view)
    note = view.get("ai_report_note")
    return paras + [str(note)] if note else paras


def _market_scope_paragraph(view: dict) -> str:
    """فقرة تعريف السوق ونطاقه (مواصفة تقرير عالمي §3) — ما يشمله التقرير
    وما يستثنيه بجملة واحدة، من حقول محسوبة — لا حكم جديد يُضاف."""
    h = view.get("header") or {}
    product = h.get("product") or view.get("product") or "—"
    hs = view.get("hs_code") or h.get("hs_code") or "—"
    n_markets = len(view.get("markets") or [])
    top_market = h.get("target_market") or "—"
    scope_txt = ("سوق واحد مرشّح" if n_markets == 1
                else f"{n_markets} أسواق مرشّحة")
    return (f"يشمل هذا التقرير: صادرات {product} (رمز HS {hs}) من المملكة "
            f"العربية السعودية، مُقيَّمة عبر {scope_txt}، "
            f"ومفصَّلة بعمق للسوق الأعلى ترتيباً ({top_market}). يستثني "
            "هذا التقرير: التسعير التفصيلي بالتجزئة والملفات المالية "
            "للمنافسين (تتطلبان خدمة التعميق المدفوعة)، وتحليل "
            "الشرائح السلوكية للعملاء واستخبارات الطلب المباشرة (تتطلبان "
            "بحثاً أولياً — مقابلات أو استبيانات — لم يُجرَ بعد).")


def _methodology_lines(view: dict) -> list[str]:
    """أسطر قسم المنهجية (مواصفة تقرير عالمي §2) — من حقول محسوبة فعلاً.

    P1: تحفّظ «قراءة أولية» يُقال هنا مرة واحدة هادئة — لا يتكرر في كل قسم؛
    وسطر كفاية البيانات يُترجم (لا اسم صنف وكيل على وجه التقرير).
    """
    from silk_narrative import translate_gaps
    h = view.get("header") or {}
    prov = view.get("provenance") or []
    sources = sorted({str(b.get("source")) for b in prov
                      if b.get("contributed")})
    lines = [f"سنة البيانات المعتمدة: {_data_year_label(view)}.",
             f"تغطية البيانات الإجمالية لهذا التحليل: {h.get('coverage_pct')}%."]
    if sources:
        lines.append("المصادر التي أسهمت فعلياً في هذا التشغيل: "
                     + "، ".join(sources) + ".")
    lines.append("كل رقم في هذا التقرير يُعرض مع مصدره وتاريخ سحبه؛ "
                 "القيمة غير المتاحة تُعرض «—».")
    suff = (view.get("decision") or {}).get("sufficiency")
    if suff:
        lines.append(translate_gaps([suff])[0])
    lines.append("هذه قراءة أولية مبنية على البيانات العامة المتاحة "
                 "وقت الإعداد.")
    return lines


def _gap_ar(g: object) -> str:
    """فجوة داخلية → عربية هادئة للعرض — display-only translation (5b)."""
    from silk_narrative import translate_gaps
    return translate_gaps([g])[0]


def _gap_list_ar(gaps: list) -> list[str]:
    """قائمة فجوات مترجمة للعرض — display-only batch translation (5b)."""
    from silk_narrative import translate_gaps
    return translate_gaps(gaps)


def _data_year_label(view: dict) -> str:
    """سنة البيانات الفعلية — the year actually used, flagging the declared fallback.

    عند تراجُع المحرّك إلى أحدث سنةٍ منشورة (السنة المطلوبة لم تُنشر بعد) نعرض السنة
    الفعلية مع سبب التراجُع — لا يظهر رقمُ سنةٍ فارغةٍ كأنه سنة التحليل.
    """
    dy = view.get("data_year", view.get("year"))
    if dy is None:
        return "—"   # فجوة معلنة — لا يُطبع نص "None" الحرفي (تدقيق: /research بلا سنة)
    if view.get("year_fell_back") and dy != view.get("year"):
        return f"{dy} (أحدث سنة منشورة؛ {view.get('year')} لم تُنشر بعد)"
    return str(dy)


# ── مساعدو حزمة البحث وقرار الدخول (§7) — research/decision display helpers ──
# عرض صرف فوق view.markets[i]: لا شبكة، لا تعديل أرقام، الفجوات تُعلن بنصّها.
# Pure display over the canonical view; gaps are declared verbatim, never filled.

_MODELED_TAG = "مُقدَّر — نموذج بافتراضات معلنة"

_PILLAR_AR = {"market": "جاذبية السوق", "competition": "المنافسة",
              "regulatory": "التنظيم", "profit": "الربحية",
              "risk": "أمان السوق (المخاطر)"}

_SEC_AR = {"market_size": "حجم السوق والمنافسة", "demand": "الطلب والقدرة",
           "regulatory": "الاشتراطات والتعريفة", "competitors": "المنافسون بالاسم",
           "pricing": "الأسعار", "risk": "المخاطر", "trend": "الاتجاه"}


def _research_bundle(m: dict) -> tuple[dict | None, str]:
    """حزمة البحث أو غيابها المعلّل — the research bundle, or a declared absence."""
    r = (m or {}).get("research")
    if isinstance(r, dict) and r.get("error"):
        return None, f"حزمة البحث فشلت — research bundle error: {r['error']}"
    if not isinstance(r, dict) or not r.get("agents"):
        return None, ("حزمة وكلاء البحث غير مفعّلة (with_research) — "
                      "لا اكتشافات مرصودة لهذا القسم")
    return r, ""


def _ragent(m: dict, name: str) -> dict:
    """وكيل بحث بالاسم من الحزمة — one research agent's block ({} if absent)."""
    r, _ = _research_bundle(m)
    return ((r or {}).get("agents") or {}).get(name) or {}


def _rfind(agent: dict, metric: str) -> dict | None:
    """اكتشاف بالاسم — the finding dict for a metric, or None."""
    for f in agent.get("findings") or []:
        if f.get("metric") == metric:
            return f
    return None


def _f_text(f: dict) -> str:
    """سطر الاكتشاف — metric: value unit [+ وسم «مُقدَّر» ونص المعادلة إن نموذجاً].

    اسم المقياس الداخلي (snake_case) يُعرَّب دوماً — لا يصل وجه العميل
    خاماً (بلاغ تدقيق: "tam_usd"/"hhi" ظهرت حرفياً في جداول docx/markdown)."""
    from silk_narrative import internal_ar
    unit = f" {f.get('unit')}" if f.get("unit") else ""
    txt = f"{internal_ar(f.get('metric'))}: {_fmt(f.get('value'))}{unit}"
    if f.get("modeled"):
        txt += f" [{_MODELED_TAG}]"
        if f.get("formula"):
            txt += f" — المعادلة: {f['formula']}"
    return txt


def _f_src_bare(f: dict | None) -> str:
    """نص المصدر بلا بادئة — source text without the "المصدر: " label.

    لخلايا الجداول (`_add_table`) التي تحمل عمود "المصدر" بذاتها — البادئة
    هناك تتكرر ("المصدر: المصدر: ..."). `_f_srcline` تبقى للاستخدام السردي.
    """
    from silk_narrative import confidence_phrase
    parts = []
    for s in (f or {}).get("sources") or []:
        seg = str(s.get("source") or "غير مرصود")
        if s.get("retrieved_at"):
            seg += f" | سُحب: {s['retrieved_at']}"
        if s.get("confidence") is not None:
            seg += f" | ثقة: {confidence_phrase(s['confidence'])}"
        if s.get("url"):
            seg += f" | {s['url']}"
        parts.append(seg)
    return "؛ ".join(parts) if parts else "غير مرصود"


def _f_srcline(f: dict | None) -> str:
    """سطر المصدر لاكتشاف — source line (source | retrieved_at | confidence | url)."""
    return "المصدر: " + _f_src_bare(f)


def _entry_text(e: object) -> str:
    """سطر مرشّح (شركة/مورّد) — one candidate line: name/url/address/via/date.

    إصلاح مراجعة Stage 5 (ثغرة ٢): بند kind=reference عنوانُ صفحة ويب لا اسم
    كيان — يُطبع دائماً بوسم «مرجع للمراجعة اليدوية»، لا كمنافس بالاسم.
    """
    if not isinstance(e, dict):
        return str(e)
    is_ref = _candidate_kind(e) == "reference"
    label = str(e.get("name") or e.get("title") or "؟")
    if is_ref:
        bits = [f"مرجع للمراجعة اليدوية: {label}"]
    elif e.get("business_hint") == "retail_or_food_service":
        # بلاغ مالك حقيقي: محلات عصير/مطاعم ظهرت كأنها مستوردون — تصنيف
        # جوجل الفعلي (types) يكشف ذلك، فنُعلِنه بدل تسميته موزّعاً بالجملة.
        bits = [f"⚠ {label} — يبدو محل تجزئة/مطعماً حسب تصنيف Google Maps، "
               "ليس بالضرورة موزّعاً بالجملة"]
    else:
        bits = [label]
    if e.get("url"):
        bits.append(str(e["url"]))
    if e.get("address"):
        bits.append(f"العنوان: {e['address']}")
    if e.get("via"):
        bits.append(f"عبر {e['via']}")
    if e.get("retrieved_at"):
        bits.append(f"سُحب: {e['retrieved_at']}")
    # كياناتٌ مُستخلَصة (كلود) تحمل ملاحظةَ «غير موثَّق، أكّده» — تُظهَر لا تُخفى.
    if not is_ref and e.get("note") and "مُستخلَص" in str(e.get("note")):
        bits.append(str(e["note"]))
    return " — ".join(bits)


def _candidate_kind(e: object) -> str:
    """نوع المرشّح — entity (اسم Google Places) أم reference (عنوان بحث ويب)."""
    if not isinstance(e, dict):
        return "reference"
    if e.get("kind") in ("entity", "reference"):
        return e["kind"]
    # توافُق خلفي: بنود قديمة بلا kind — Maps كيان، وكل ما جاء من بحث الويب مرجع.
    if e.get("via") == "Google Maps":
        return "entity"
    return "reference"


def _split_candidates(items: list) -> tuple[list, list]:
    """افصل الكيانات عن المراجع — entities first, web references second."""
    ents = [e for e in (items or []) if _candidate_kind(e) == "entity"]
    refs = [e for e in (items or []) if _candidate_kind(e) == "reference"]
    return ents, refs


def _listing_text(v: object) -> str:
    """سطر قائمة سعر تجزئة — one retail listing line (title/price/currency/store)."""
    if not isinstance(v, dict):
        return str(v)
    t = str(v.get("title") or v.get("name") or "قائمة")
    if v.get("price") is not None:
        t += f": {_fmt(v['price'])}" + (f" {v['currency']}" if v.get("currency") else "")
    if v.get("store"):
        t += f" — {v['store']}"
    return t


def _req_text(it: object) -> str:
    """سطر بند اشتراطات — one requirements-checklist item (بنده وجهته ورابطه)."""
    if not isinstance(it, dict):
        return str(it)
    bits = [str(it.get("item") or it.get("requirement") or "؟")]
    if it.get("authority"):
        bits.append(f"الجهة: {it['authority']}")
    if it.get("direction"):
        bits.append(f"الاتجاه: {it['direction']}")
    if it.get("source_url"):
        bits.append(str(it["source_url"]))
    return " — ".join(bits)


def _entry_decision_of(m: dict) -> tuple[dict | None, str]:
    """قرار الدخول §8 أو غيابه المعلّل — the §8 decision, or a declared absence."""
    ed = (m or {}).get("entry_decision")
    if isinstance(ed, dict) and ed.get("error"):
        return None, f"محرك القرار (§8) أبلغ خطأً — decision error: {ed['error']}"
    if not isinstance(ed, dict) or not ed.get("schema"):
        return None, ("قرار الدخول غير متاح — المحرك الموزون (§8) لم يعمل لهذا "
                      "التحليل (يتطلب with_research)")
    return ed, ""


def render_brief(view: dict, dashboard_url: str = "/") -> str:
    """المختصر (§10.4) — صفحة واحدة بتصميم "رسالة جوال"، من القالب حصراً.

    P1 (طبقة السرد): أسماء المقاييس الداخلية (market_size) ورمز الحكم الآلي
    وشعارات النزاهة أُزيلت من الوجه — سطر المصدر مع كل رقم حاضر هو إشارة
    النزاهة الوحيدة، والقيم نفسها بلا أي تغيير.
    """
    from silk_narrative import (GAP_WORD, competition_phrase, country_ar,
                                fmt_money, fmt_pct, internal_ar, verdict_ar)
    _assert_production_clean(view)
    d = view.get("decision") or {}
    cp = view.get("competitive_position") or {}
    top = (view.get("markets") or [{}])[0]
    dr = view.get("deep_research")
    numbers = []
    if dr:  # الموجة ٤: أرقام من تقاطعات المحلل الشامل لا components_detail
        by_cat = (dr.get("analyst") or {}).get("by_category") or {}
        for cat in ("demand", "price_competitiveness", "entry_cost"):
            items = by_cat.get(cat) or []
            if items:
                numbers.append(f"• {_CATEGORY_AR.get(cat, cat)}: "
                               f"{items[0].get('value')} "
                               f"[{items[0].get('source', '؟')}]")
    else:
        for c in (top.get("components_detail") or []):
            if c.get("value") is None:
                continue
            name = c.get("name")
            # كل مقياس بصيغته البشرية — مؤشر HHI الخام لا يصل وجه المستخدم أبداً.
            val = (fmt_money(c["value"]) if name in ("market_size",
                                                     "demand_capacity")
                   else fmt_pct(c["value"]) if name == "saudi_position"
                   else competition_phrase(c["value"]) if name == "competition"
                   else _fmt(c["value"]))
            numbers.append(f"• {internal_ar(name)}: {val} "
                           f"[{c.get('source', '؟')}]")
            if len(numbers) == 3:
                break
    if not numbers:
        numbers = [f"• {GAP_WORD}"]
    L = ([] if not view.get("test_run") else
         ["⚠ TEST RUN — تشغيل برهاني ببدائل موسومة، ليس تقريراً إنتاجياً"])
    market_ar = (((dr or {}).get("market") or {}).get("name_ar")
                if dr else country_ar(top.get("iso3"), d.get("market")))
    L += [f"سِلك | {view.get('product')} — سوق {market_ar} | "
          f"قراءة أولية {view.get('year')}",
         "",
         view.get("brief", [""])[0] if view.get("brief") else
         f"التوصية: {verdict_ar(d.get('verdict'))}",
         "", "أبرز الأرقام:", *numbers, ""]
    if len(view.get("brief") or []) > 1:
        L += view["brief"][1:3]
    else:
        L.append(cp.get("note", ""))
    L += ["", f"التفاصيل الكاملة باللوحة: {dashboard_url}"]
    return "\n".join(L)


# ── بناة أقسام Word الجديدة (§7) — new docx section builders (pure display) ──

def _add_table(doc, headers: list[str], rows: list[list],
               caption: str | None = None) -> None:
    """جدول Word موحّد سِلك — رأس بلون سِلك الأساس وخط أبيض، أشرطة متناوبة
    خفيفة، تسمية اختيارية أعلاه؛ no rows => no-op.

    مراجعة المشروع: النسخة الحية من التقرير (docx، ما يُرسله المستخدم فعلياً)
    كانت نقاطاً سردية بحتة بينما نظيرتها Markdown تستخدم جداول فعلية لنفس
    البيانات (قرار الدخول مثلاً) — تناقضٌ بين الصيغتين، وواحدة من أوضح علامات
    "تقرير غير احترافي" بمقارنة أي منصة أبحاث سوق مرجعية (Country Commercial
    Guides وITC Trade Map وEuromonitor). لا بيانات جديدة، عرضٌ صرفٌ فقط —
    "Table Grid" نمطٌ مدمج في python-docx (بلا قالب خارجي). الموجة ١١
    (§11.1): رأس/أشرطة ملوّنة من `config/branding.yaml` — تقرير موحَّد
    الهوية بصرياً بدل جدول Word افتراضي عادي.
    """
    if not rows:
        return
    if caption:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.bold = True
    from docx.shared import Pt
    dense = len(rows) > _TABLE_DENSE_ROWS   # §7: جدولٌ كثيفٌ => خطٌّ أصغر
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        _set_cell_shading(hdr[i], _TABLE_HEADER_FILL)   # §7: أخضر سِلك
        _set_cell_margins(hdr[i])
        if hdr[i].paragraphs[0].runs:
            r = hdr[i].paragraphs[0].runs[0]
            r.bold = True
            r.font.color.rgb = _hex_to_rgbcolor("FFFFFF")
            if dense:
                r.font.size = Pt(_TABLE_DENSE_PT)
    for row_idx, vals in enumerate(rows):
        cells = table.add_row().cells
        for i, v in enumerate(vals):
            cells[i].text = str(v) if v is not None else "—"
            _set_cell_margins(cells[i])
            if dense and cells[i].paragraphs[0].runs:
                cells[i].paragraphs[0].runs[0].font.size = Pt(_TABLE_DENSE_PT)
        if row_idx % 2 == 1:  # §7: شريطٌ متناوبٌ أخضرُ خفيف كل صفٍّ زوجيّ
            for c in cells:
                _set_cell_shading(c, _TABLE_ZEBRA_FILL)
    _set_table_borders(table)   # §7: حدودٌ شعريّةٌ خافتة بدل أسود Grid
    _set_table_rtl(table)   # §4: تدفّق أعمدة يميناً + محاذاة خلايا + عرض DXA


def _docx_entry_strategy(doc, m: dict) -> None:
    """استراتيجية دخول السوق — فصلٌ مستقل يركّب توصية الدخول من أرقامٍ مرصودة.

    مراجعة المشروع: التوصية («ادخل عبر موزّع قائم») كانت مدفونة كسطرٍ واحد
    داخل «الخطوات الأولى» — أي منصة مرجعية (Country Commercial Guides) تُفرد
    لها فصلاً. يُركِّب هنا فقط ما هو مرصود فعلاً (تركّز الموردين + بوابة
    الأهلية + عدد المرشّحين بالاسم) في فقرة واحدة مسبَّبة؛ لا رقم جديد، لا
    اختلاق — تجميعٌ لحقائق موجودة أصلاً في أقسام أخرى من نفس التقرير.
    """
    doc.add_heading("استراتيجية دخول السوق", level=2)
    comp = _ragent(m, "competitor")
    reg = _ragent(m, "regulatory")
    if not comp and not reg:
        doc.add_paragraph("بيانات غير كافية لتركيب استراتيجية دخول — "
                          "يتطلب with_research")
        return
    hhi_f = _rfind(comp, "hhi")
    top_f = _rfind(comp, "top_supplier_share_pct")
    gate_f = _rfind(reg, "eligibility_gate")
    sd = m.get("supplier_directory") or {}
    n_target = len(sd.get("target") or [])
    n_saudi = len(sd.get("saudi") or [])
    hhi = hhi_f.get("value") if hhi_f else None
    top_share = top_f.get("value") if top_f else None
    concentrated = hhi is not None and hhi > 0.25
    model = ("عبر موزّع/مستورد محلي قائم" if (concentrated or n_target)
             else "مباشر أو عبر موزّع محلي — لا تفضيل واضح من بيانات التركّز")
    reasons = []
    if hhi is not None:
        reasons.append(f"مؤشر تركّز الموردين HHI={hhi} "
                       + ("(سوق مركّز)" if concentrated else "(سوق مفتّت)"))
    if top_share is not None:
        reasons.append(f"أكبر حصة مورّد واحد {top_share}%")
    if n_target:
        reasons.append(f"{n_target} موزّعاً/مستورداً محتملاً مرصوداً بالاسم "
                       "بالسوق المستهدف")
    if n_saudi:
        reasons.append(f"{n_saudi} مورّداً سعودياً مرصوداً بالاسم كمرشّح تعاقد")
    if gate_f and gate_f.get("value"):
        reasons.append("بوابة أهلية أمامية مفتوحة — لا خطوة لاحقة تُعتبر "
                       "سالكة قبل اجتيازها")
    doc.add_paragraph(f"النموذج الموصى به: {model}")
    if reasons:
        doc.add_paragraph("الأساس (من الأقسام أعلاه، لا رقم جديد):")
        for r in reasons:
            doc.add_paragraph(r, style="List Bullet")
    else:
        doc.add_paragraph("لا مؤشرات تركّز أو مرشّحين مرصودة — التوصية أعلاه "
                          "افتراضية حتى تتوفر بيانات المنافسة/الموردين")


def _docx_entry_decision(doc, m: dict) -> None:
    """قرار الدخول — verdict/score/pillars/conditions/risks.

    سدّ تسريب (الطبقة ٨، قرار المالك): العنوان والنص كانا يحملان رطانة
    مسار العمل الداخلي ("المحرك الموزون §8"، "بوابة GATE 3") ورمز مفتاح
    خام (A/B) — عنوان تجاري صرف الآن، وخيار الأوزان يُعرض باسمه الوصفي
    (`weights_label`) لا رمزه (`weights_option` يبقى كما هو في المصدر
    لأغراض داخلية/تتبّع فقط)."""
    doc.add_heading("قرار الدخول", level=2)
    ed, absent = _entry_decision_of(m)
    if ed is None:
        doc.add_paragraph(absent)
        return
    from silk_decision import _WEIGHT_LABEL_AR
    from silk_narrative import confidence_phrase, verdict_ar
    doc.add_paragraph(f"الحكم: {verdict_ar(ed.get('verdict'))} | "
                      f"النقاط: {_fmt(ed.get('score'))}"
                      f" | الثقة: {confidence_phrase(ed.get('confidence'))}")
    doc.add_paragraph(f"أساس الثقة: {ed.get('confidence_basis')}")
    sbo = ed.get("scores_by_option") or {}
    weights_label = ed.get("weights_label") or ed.get("weights_option") or ""
    other_label = _WEIGHT_LABEL_AR.get(
        "B" if ed.get("weights_option") == "A" else "A", "")
    other_score = sbo.get("B" if ed.get("weights_option") == "A" else "A")
    doc.add_paragraph(f"منهجية الترجيح المعتمدة: {weights_label} — النقاط "
                      f"{_fmt(ed.get('score'))} (بديل مقارنة، {other_label}: "
                      f"{_fmt(other_score)})")
    if ed.get("weights_note"):
        doc.add_paragraph(f"ملاحظة الأوزان: {ed['weights_note']}")
    # الأعمدة الأربعة كجدول (توازي render_markdown) — لا نقاط سردية مبعثرة.
    _add_table(
        doc, ["العمود", "القيمة", "الأساس", "مكوّنات غائبة"],
        [[_PILLAR_AR.get(key, key), _fmt(p.get("value")), p.get("basis"),
          "، ".join(map(str, p.get("missing") or [])) or "—"]
         for key, p in (ed.get("pillars") or {}).items()])
    if ed.get("missing_pillars"):
        doc.add_paragraph("أعمدة غائبة كلياً: " + "، ".join(
            _PILLAR_AR.get(k, k) for k in ed["missing_pillars"]))
    if ed.get("critical_risk"):
        doc.add_paragraph("تحذير: خطر حرج مرصود — راجع سجل المخاطر أدناه.")
    doc.add_paragraph("الشروط:")
    for c in ed.get("conditions") or ["لا شروط مفتوحة"]:
        doc.add_paragraph(str(c), style="List Bullet")
    doc.add_paragraph("سجل المخاطر:")
    for r in ed.get("risks") or []:
        doc.add_paragraph(f"{r.get('risk')} (الشدة: {r.get('severity')}) — "
                          f"الدليل: {r.get('evidence')}", style="List Bullet")
    if not ed.get("risks"):
        doc.add_paragraph("لا مخاطر مسجّلة في محرك القرار", style="List Bullet")
    doc.add_paragraph("الخطوات الأولى:")
    for s in ed.get("first_steps") or ["لا خطوات مقترحة"]:
        doc.add_paragraph(str(s), style="List Bullet")
    doc.add_paragraph(f"لماذا: {ed.get('why')}")
    if ed.get("note"):
        doc.add_paragraph(str(ed["note"]))


def _docx_market_size(doc, m: dict) -> None:
    """حجم السوق TAM/SAM/SOM — كل اكتشاف بمصدره؛ المنمذج موسوم بمعادلته."""
    doc.add_heading("حجم السوق — TAM/SAM/SOM", level=2)
    r, absent = _research_bundle(m)
    if r is None:
        doc.add_paragraph(absent)
        return
    ag = _ragent(m, "market_size")
    valued = [f for f in (ag.get("findings") or []) if f.get("value") is not None]
    if valued:
        from silk_narrative import internal_ar
        _add_table(
            doc, ["المؤشر", "القيمة", "النوع", "المصدر"],
            [[internal_ar(f.get("metric")),
              f"{_fmt(f.get('value'))}{(' ' + f['unit']) if f.get('unit') else ''}",
              (_MODELED_TAG if f.get("modeled") else "رصد مباشر"),
              _f_src_bare(f)]
             for f in valued])
        for f in valued:
            if f.get("modeled") and f.get("formula"):
                doc.add_paragraph(
                    f"معادلة «{internal_ar(f.get('metric'))}»: {f['formula']}",
                    style="Intense Quote")
    else:
        doc.add_paragraph("لا اكتشافات مرصودة لحجم السوق")
    for g in ag.get("gaps") or []:
        doc.add_paragraph(f"غير متوفر: {_gap_ar(g)}", style="List Bullet")


def _docx_competition_research(doc, m: dict) -> None:
    """طبقتا المنافسة من حزمة البحث — شركات بالاسم + أرقام الطبقة الدولية."""
    ag = _ragent(m, "competitor")
    if not ag:
        return
    doc.add_heading("شركات بالاسم (مرشّحون غير موثَّقين)", level=3)
    named = (_rfind(ag, "named_companies") or {}).get("value") or []
    ents, refs = _split_candidates(named)
    if ents:
        _add_table(
            doc, ["الاسم", "العنوان", "المصدر", "ملاحظة"],
            [[e.get("name") or e.get("title") or "؟", e.get("address"),
              e.get("via"),
              ("مُستخلَص، غير موثَّق" if "مُستخلَص" in str(e.get("note") or "")
               else "غير موثَّق")]
             for e in ents[:10]])
        from silk_narrative import confidence_phrase
        doc.add_paragraph(f"كيانات غير موثَّقة — الثقة {confidence_phrase(0.4)} "
                          "— أكّدها قبل أي تعاقد.")
    else:
        doc.add_paragraph("لا شركات مرصودة بالاسم في هذا التشغيل "
                          "(أسماء الأعمال تأتي من Google Places حصراً)")
    if refs:
        doc.add_heading("مراجع ويب للمراجعة اليدوية (ليست أسماء منافسين)",
                        level=3)
        for n in refs[:8]:
            doc.add_paragraph(_entry_text(n), style="List Bullet")
    doc.add_paragraph("الطبقة الدولية (تركّز الموردين — UN Comtrade):")
    from silk_narrative import internal_ar
    metrics_rows = []
    for metric in ("hhi", "top_supplier_share_pct", "saudi_share_pct"):
        f = _rfind(ag, metric)
        metric_ar = internal_ar(metric)
        if f and f.get("value") is not None:
            # وحدة الحقيقة (%) كانت تُسقَط هنا — نفس البند في نسخة الماركداون
            # (_f_text) يعرضها صحيحة؛ إصلاح تناسق: كلاهما يقرأ f['unit'] الآن.
            unit = f" {f['unit']}" if f.get("unit") else ""
            metrics_rows.append([metric_ar, f"{_fmt(f.get('value'))}{unit}",
                                 _f_src_bare(f)])
            if f.get("note"):
                metrics_rows[-1].append(_gap_ar(f["note"]))
        else:
            metrics_rows.append([metric_ar, "—", "—"])
    max_cols = max(len(r) for r in metrics_rows)
    for r in metrics_rows:
        while len(r) < max_cols:
            r.append("")
    _add_table(doc, ["المؤشر", "القيمة", "المصدر", "ملاحظة"][:max_cols],
              metrics_rows)


def _docx_pricing_layers(doc, m: dict) -> None:
    """التسعير بطبقتيه — الحدودية (نماذج موسومة بمعادلاتها) ثم التجزئة ومراجعها."""
    doc.add_heading("التسعير بطبقتيه", level=2)
    r, absent = _research_bundle(m)
    if r is None:
        doc.add_paragraph(absent)
        return
    ag = _ragent(m, "pricing")
    doc.add_paragraph("الطبقة الحدودية (قيم وحدة كومتريد):")
    from silk_narrative import internal_ar
    border_rows, border_formulas = [], []
    for metric in ("border_unit_value_usd_kg", "saudi_border_unit_value_usd_kg",
                   "margin_at_border_pct"):
        f = _rfind(ag, metric)
        metric_ar = internal_ar(metric)
        if f and f.get("value") is not None:
            unit = f" {f['unit']}" if f.get("unit") else ""
            border_rows.append([
                metric_ar, f"{_fmt(f.get('value'))}{unit}",
                (_MODELED_TAG if f.get("modeled") else "رصد مباشر"),
                _f_src_bare(f)])
            if f.get("modeled") and f.get("formula"):
                border_formulas.append(f"معادلة «{metric_ar}»: {f['formula']}")
        else:
            border_rows.append([metric_ar, "—", "—", "—"])
    _add_table(doc, ["المؤشر", "القيمة", "النوع", "المصدر"], border_rows)
    for line in border_formulas:
        doc.add_paragraph(line, style="Intense Quote")
    doc.add_paragraph("طبقة التجزئة:")
    rp = _rfind(ag, "retail_prices")
    vals = (rp or {}).get("value") or []
    if vals:
        for v in vals[:8]:
            doc.add_paragraph(_listing_text(v), style="List Bullet")
        doc.add_paragraph(_f_srcline(rp), style="Intense Quote")
    else:
        # فجوة القسم المعلنة تُطبع مترجمةً للعرض (_gap_ar) — لا سباكة
        # إنجليزية (ملاحظة الحارس المدفوع/أسماء المفاتيح) على وجه التقرير.
        gap = next((g for g in (ag.get("gaps") or []) if "retail_prices" in g),
                   "أسعار التجزئة: غير متوفرة في هذا التشغيل")
        doc.add_paragraph(_gap_ar(gap), style="List Bullet")
    # نقاطُ الأسعار المُستخلَصة (كلود) — أرقامٌ مذكورةٌ صراحةً، لا روابط.
    points = (_rfind(ag, "retail_price_points") or {}).get("value") or []
    if points:
        doc.add_paragraph("أسعار مُستخلَصة (مذكورة صراحةً في عناوين الويب — مؤشِّر "
                          "لا سعرَ رفٍّ مؤكَّد):")
        for p in points[:8]:
            p = p or {}
            doc.add_paragraph(
                f"{p.get('price')} {p.get('currency')}/{p.get('unit')}".rstrip("/")
                + (f" — {p.get('url')}" if p.get("url") else ""),
                style="List Bullet")
    refs = (_rfind(ag, "retail_references") or {}).get("value") or []
    if refs:
        doc.add_paragraph("مصادر الأسعار (للاستشهاد):")
        for ref in refs[:3]:
            doc.add_paragraph(f"{(ref or {}).get('title')} — {(ref or {}).get('url')}"
                              f" — سُحب: {(ref or {}).get('retrieved_at')}",
                              style="List Bullet")
    elif not points:
        doc.add_paragraph("مراجع الأسعار: غير مرصودة")


def _docx_swot(doc, m: dict) -> None:
    """SWOT قاعدي — شبكة ٢×٢ حقيقية (لا أربع عناوين متتالية)؛ كل خلية بدليلها.

    مراجعة المشروع: شكل SWOT كأربعة عناوين متسلسلة لا يُقرأ كمصفوفة عند
    الطباعة، بينما كل مرجعية بحث سوق (IBISWorld، Euromonitor) تعرضه شبكةً
    ٢×٢ فعلية — نفس المحتوى، عرضٌ مطابق للتعارف الصناعي.
    """
    doc.add_heading("تحليل SWOT (قاعدي من حقائق مرصودة)", level=2)
    sw = m.get("swot") or {}
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    layout = [(("S", "القوة"), ("W", "الضعف")),
             (("O", "الفرص"), ("T", "التهديدات"))]
    for row_idx, pair in enumerate(layout):
        for col_idx, (key, title) in enumerate(pair):
            cell = table.rows[row_idx].cells[col_idx]
            cell.paragraphs[0].add_run(title).bold = True
            items = sw.get(key) or []
            if not items:
                cell.add_paragraph("—")
            for it in items:
                cell.add_paragraph(f"{it.get('text')} — الدليل: "
                                   f"{it.get('evidence')}", style="List Bullet")
    _set_table_rtl(table)   # §4: شبكة SWOT ٢×٢ من اليمين لليسار كذلك
    if sw.get("note"):
        doc.add_paragraph(str(sw["note"]))


def _docx_segments(doc, m: dict) -> None:
    """شرائح العملاء — segment + basis؛ الفارغ يُعلن لا يُخترع."""
    doc.add_heading("شرائح العملاء", level=2)
    segs = m.get("segments") or []
    if not segs:
        doc.add_paragraph("بيانات غير كافية للشرائح — التقسيم السلوكي/"
                          "الديموغرافي يتطلب بحثاً أولياً (مقابلات أو "
                          "استبيانات) لم يُجرَ بعد؛ لا يُشتق من بيانات ثانوية")
        return
    for s in segs:
        doc.add_paragraph(f"{s.get('segment')} — الأساس: {s.get('basis')}",
                          style="List Bullet")


def _docx_supplier_directory(doc, m: dict) -> None:
    """دليل المورّدين والمصنّعين — قائمتا السعودية والسوق المستهدف + الملاحظة."""
    doc.add_heading("دليل المورّدين والمصنّعين", level=2)
    sd = m.get("supplier_directory") or {}
    for key, title in (("saudi", "مورّدون ومصنّعون سعوديون:"),
                       ("target", "موزّعون ومستوردون في السوق المستهدف:")):
        doc.add_paragraph(title)
        items = sd.get(key) or []
        if not items:
            doc.add_paragraph("—",
                              style="List Bullet")
        for e in items[:10]:
            doc.add_paragraph(_entry_text(e), style="List Bullet")
    if sd.get("note"):
        doc.add_paragraph(str(sd["note"]))


def _docx_regulatory(doc, m: dict) -> None:
    """الاشتراطات التنظيمية — بوابة الأهلية أولاً ثم بنود L1 بجهاتها وروابطها."""
    doc.add_heading("الاشتراطات التنظيمية", level=2)
    r, absent = _research_bundle(m)
    ag = _ragent(m, "regulatory")
    if r is None or not ag:
        doc.add_paragraph(absent or "وكيل الاشتراطات لم يعمل في هذا التحليل")
        return
    gate_f = _rfind(ag, "eligibility_gate")
    if gate_f and gate_f.get("value"):
        doc.add_paragraph("تحذير — بوابة أهلية أمامية: هذا السوق يتطلب منشأة "
                          "معتمدة (EU 2017/625) قبل أي بند لاحق؛ لا بند أدناه "
                          "يُعتبر سالكاً قبل اجتيازها.")
    checklist_f = _rfind(ag, "requirements_checklist")
    checklist = (checklist_f or {}).get("value") or []
    if checklist:
        _add_table(
            doc, ["البند", "الجهة", "الاتجاه", "الرابط"],
            [[it.get("item") or it.get("requirement"), it.get("authority"),
              it.get("direction"), it.get("source_url")]
             for it in checklist if isinstance(it, dict)])
        doc.add_paragraph(_f_srcline(checklist_f), style="Intense Quote")
    else:
        doc.add_paragraph("لا بنود اشتراطات في مرجع سِلك لهذا السوق — غير متوفر")
    tf = _rfind(ag, "tariff_applied_pct")
    if tf and tf.get("value") is not None:
        doc.add_paragraph(_f_text(tf))
        doc.add_paragraph(_f_srcline(tf), style="Intense Quote")
    for g in ag.get("gaps") or []:
        doc.add_paragraph(f"غير متوفر: {_gap_ar(g)}", style="List Bullet")


_CATEGORY_AR = {
    "demand": "الطلب الفعلي القابل للتوجيه",
    "entry_cost": "تكلفة وصعوبة الدخول",
    "price_competitiveness": "التنافسية السعرية",
    "entry_door": "أبواب الدخول الأكثر أماناً",
    "swot": "SWOT من منظور المصدّر السعودي",
}


def _docx_deep_research(doc, view: dict) -> None:
    """قسم البحث العميق (الموجة ٤، V5) — ١٢ بعثة + محلل + حكم + تقرير مراجَع.

    يُستدعى إضافياً فقط عند `view["deep_research"]` (نتيجة /research) — لا
    يمسّ الأقسام الأربعة عشر القائمة لتقرير /analyze الكلاسيكي.
    """
    dr = view.get("deep_research")
    if not dr:
        return
    doc.add_heading("قسم البحث العميق — التقاطعات الخمسة والمحلل الشامل",
                    level=1)
    _stamp_degraded_banner(doc, view)
    # WP-7 §1: ختم التجاوز على النسخ الداخلية — نسخة عميل سُلِّمت بتجاوز
    # مالكٍ لبوابة الجودة تُوثَّق هنا مع ملاحظات البوابة المرفقة.
    for ov in (view.get("owner_override_history") or []):
        doc.add_paragraph(
            f"⚠ سُلِّمت نسخة عميل من هذا التقرير بتجاوز مالكٍ لبوابة "
            f"الجودة بتاريخ {ov.get('created_at')} — ملاحظات البوابة مرفقة:",
            style="Intense Quote")
        for gf in (ov.get("gate_findings") or [])[:6]:
            doc.add_paragraph(f"• {gf.get('check')}: {gf.get('note')}",
                              style="Intense Quote")

    # نفس تصنيف/تعريب الحكم المستعمَل في الغلاف (_VERDICT_LABELS_AR عبر
    # _verdict_tone) — لا مصدر عرض ثانٍ قد يختلف نصّه عن الأول لنفس الرمز
    # (اختبار اتساق: الحكم يظهر متطابقاً حرفياً في الخلاصة وهنا معاً).
    market = dr.get("market") or {}
    verdict = dr.get("verdict") or {}
    ai = verdict.get("ai") or {}
    v_raw = _resolve_vtxt(dr)
    doc.add_paragraph(
        f"السوق: {market.get('name_ar') or market.get('name_en')} "
        f"({market.get('iso3')}) — الحكم: "
        f"{_VERDICT_LABELS_AR[_verdict_tone(v_raw)]}")
    # WP-1 §2: قراءة كلود حقل استشاري في التصدير الداخلي فقط — تُعرَض
    # موسومة «قراءة تحليلية للذكاء الاصطناعي»، لا توصيةً ثانية.
    if ai.get("verdict"):
        from silk_narrative import verdict_ar as _var
        doc.add_paragraph(
            f"قراءة تحليلية للذكاء الاصطناعي (استشارية — ليست التوصية): "
            f"{_var(ai.get('verdict'))}", style="Intense Quote")
    if ai.get("reasoning"):
        doc.add_paragraph(str(ai["reasoning"]), style="Intense Quote")

    doc.add_heading("ملخّص مصادر البحث", level=2)
    _stamp_degraded_banner(doc, view)
    missions = dr.get("missions") or {}
    # الاسم التجاري العربي (view label) بدل مفتاح snake_case الداخلي —
    # بلاغ مالك: "pricing_scout" وأخواتها ظهرت حرفياً في جدول العميل.
    _add_table(doc, ["البعثة", "الحالة", "الملخّص"], [
        [m.get("label") or key,
         "فشلت/بلا أدلة" if m.get("failed") else "ناجحة",
         _clean_report_text(m.get("summary"))]
        for key, m in missions.items()])

    # Phase 3 (بلاغ حي: "تقرير احترافي لا تفريغ بيانات"): التقرير السردي
    # (كاتب التقرير) يُعرض أولاً — فقرات تحليلية تُفسّر الأرقام لقرار الدخول،
    # بما فيها التقاطعات الخمسة كأقسام فرعية '### ' سردية — قبل ملحق الأدلة
    # الرقمية الخام (كان الترتيب معكوساً: نقاط خام تسبق أي سرد).
    if dr.get("report", {}).get("text"):
        doc.add_heading("التقرير الكامل (كاتب التقرير، مراجَع)", level=2)
        _stamp_degraded_banner(doc, view)
        lines = str(dr["report"]["text"]).splitlines()
        i, n = 0, len(lines)
        while i < n:
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            if line.startswith("### "):
                # عناوين فرعية عامة (الموجة ١٠: "خارطة طريق الدخول" داخل
                # "التوصيات الاستراتيجية" مثلاً) — بلا هذا كانت "### السطر"
                # تظهر نصاً خاماً بثلاث علامات # حرفية بدل عنوان فعلي.
                doc.add_heading(line[4:].strip(), level=4)
                i += 1
                continue
            if line.startswith("## "):
                heading_text = line[3:].strip()
                doc.add_heading(heading_text, level=3)
                # الموجة ١٠: نتائج بوابة الجودة غير القابلة للإصلاح تُلحَق
                # هنا برمجياً داخل قسم "منهجية البحث ونطاقه" — مهنياً، لا
                # لافتة تحذير على الغلاف ولا صمت (silk_quality_gate.py).
                if "منهجية" in heading_text:
                    notes = ((dr.get("quality_gate") or {})
                            .get("methodology_notes") or [])
                    if notes:
                        doc.add_heading("حدود المنهجية وجودة البيانات",
                                        level=4)
                        for note in notes:
                            doc.add_paragraph(str(note), style="List Bullet")
                i += 1
                continue
            if _is_markdown_table_row(line):
                # بلاغ حي (الموجة ٩): سلاسل رقمية (استيراد/أسعار/اشتراطات/
                # ديموغرافيا) كانت نقاطاً سردية مبعثرة — تُجمَع أسطر الجدول
                # المتتالية وتُحوَّل لجدول Word حقيقي واحد.
                block = []
                while i < n and _is_markdown_table_row(lines[i].strip()):
                    block.append(lines[i].strip())
                    i += 1
                _render_markdown_table(doc, block)
                continue
            stripped = _strip_inline_markdown(line)
            # سطر الخلاصة الإلزامي لكل قسم (P0-C، الموجة ٩) — يُكتشَف نصياً
            # (لا بالاعتماد على ** الحرفية من كلود، التي تُزال أعلاه أصلاً)
            # ويُغمَّق برمجياً — إخلاص شكلي لا يعتمد على انضباط كلود بالسياج.
            if stripped.startswith("ماذا يعني هذا لقرارك:"):
                p = doc.add_paragraph()
                p.add_run(stripped).bold = True
            else:
                doc.add_paragraph(stripped)
            i += 1

    _docx_glossary(doc, dr)  # B1: مسرد المصطلحات بعد السرد قبل ملحق الأدلة
    _docx_leads(doc, dr)     # C5: قائمة المستوردين القابلين للتواصل

    doc.add_heading("ملحق — الأدلة الرقمية الداعمة للتقاطعات الخمسة", level=2)
    _stamp_degraded_banner(doc, view)
    doc.add_paragraph("كل نقطة أدناه هي الحقيقة الخام (بمصدرها) التي بُني "
                      "عليها السرد أعلاه — للتحقق المباشر لا كبديل عنه.",
                      style="Intense Quote")
    analyst = dr.get("analyst") or {}
    by_cat = analyst.get("by_category") or {}
    for cat, ar_label in _CATEGORY_AR.items():
        doc.add_heading(ar_label, level=3)
        items = by_cat.get(cat) or []
        if not items:
            doc.add_paragraph("دليل غير كافٍ لهذا التقاطع في هذا التشغيل — "
                              "فجوة معلنة")
            continue
        for f in items:
            from silk_narrative import evidence_badge_for
            doc.add_paragraph(str(f.get("value")), style="List Bullet")
            doc.add_paragraph(f"[{_clean_source_label(f.get('source'))} — "
                              f"{evidence_badge_for(f)}] "
                              f"{f.get('note') or ''}", style="Intense Quote")

    # سدّ خلل (الطبقة ٨): كان هذا الشرط متداخلاً داخل حلقة التقاطعات
    # أعلاه فيتكرّر عنوان "ملاحظات مراجعة لم تُحلّ" وقائمتها مرة لكل تقاطع
    # له أدلة — خارج الحلقة الآن فيظهر مرة واحدة فقط. النص مُعرَّب أصلاً
    # (clean_unresolved في _deep_research_view، الطبقة ٢) — لا حاجة لمُطهِّر
    # إضافي هنا، القيمة نظيفة عند وصولها.
    if dr["report"].get("unresolved_notes"):
        doc.add_heading("ملاحظات مراجعة لم تُحلّ", level=3)
        for n in dr["report"]["unresolved_notes"]:
            doc.add_paragraph(str(n), style="List Bullet")

    # Wave 6.1: شرطا قلب الحكم (حكم مراقبة/مشروط) — حقل مهيكل من نموذج العرض.
    _flips = dr.get("flip_conditions") or []
    if _flips:
        from silk_render import FLIP_CONDITIONS_HEADING
        doc.add_heading(FLIP_CONDITIONS_HEADING, level=2)
        for c in _flips:
            _mark = "✓ محقَّق" if c.get("met") else "○ غير محقَّق"
            doc.add_paragraph(
                f"{c.get('condition')} — {_mark}؛ يُغلَق عبر: "
                f"{c.get('closes_via')}", style="List Bullet")

    # §F-5 (حزمة الفكس v2.1): دعوة التعميق المدفوع («next_step») أُزيلت من
    # متن الدراسة — تبقى دعوةً على سطح التسليم (شارة/زرّ اللوحة، web/
    # index.html) لا داخل المستند المُصدَّر نفسه.

    if dr.get("limits"):
        doc.add_heading("حدود قسم البحث العميق", level=2)
        for x in dr["limits"][:12]:
            # PART B1: حدود البعثات جُمل واحدة مطهَّرة أصلاً (silk_render)؛ قصّها
            # عند ٣٠٠ كان يُنهي سطراً منتصفَ جملة بـ«…». سقف أوسع (٦٠٠) +
            # مصدر الحدّ صار الجملة الأولى لا الملخّص كاملاً => بلا بتر وسط جملة.
            doc.add_paragraph(_clean_report_text(x, max_len=600),
                              style="List Bullet")

    _docx_technical_appendix(doc, dr)


def _docx_technical_appendix(doc, dr: dict) -> None:
    """سجل الأدلة للمدققين (§6، أمر العمل الرئيس) — كل استشهاد بحقيقته
    الكاملة ومصدره العمومي الحقيقي ورابطه (إن رُصد) وتاريخ جمعه وقوة دليله.
    لا اسم بعثة داخلي، لا وسم «(Claude tool-use)»، لا رقم ثقة خام (شارة
    ✓/◐/○ فقط، §7)، ولا حقيقة مبتورة بـ«…» (§5 — قصّ عند حدّ جملة)."""
    # §D-3 (حزمة الفكس v2.1): إزالة التكرار بـ(القيمة، المصدر) المُطبَّعين —
    # بلاغ حي: GCC/GAFTA/حلال/GSO/شهادة المنشأ/SFDA شُحنت كلٌّ مرّتين
    # (البعثات المختلفة تستشهد بنفس الحقيقة العامة). أول ورودٍ فقط يُعرَض.
    rows = []
    seen: set = set()
    for _key, m in (dr.get("missions") or {}).items():
        findings = m.get("findings") if isinstance(m, dict) else None
        for f in (findings or []):
            value_txt = _trim_sentence(f.get("value"), 240)
            source_txt = _clean_source_label(f.get("source"))
            dedup_key = (re.sub(r"\s+", " ", value_txt).strip().lower(),
                        re.sub(r"\s+", " ", source_txt).strip().lower())
            if dedup_key in seen:
                continue
            seen.add(dedup_key)
            # WP-3: شارة واعية بالمنشأ والمصالحة — بند رفضته المصالحة يعرض
            # «متعارض — مستبعد» لا «✓ موثّق»؛ بند وكيل بحث غير مسانَد يُسقَف.
            from silk_narrative import evidence_badge_for
            rows.append([
                value_txt, source_txt,
                _evidence_url(f.get("note"), f.get("source"), f.get("value")),
                f.get("retrieved_at") or "—",
                evidence_badge_for(f)])
    if not rows:
        return
    doc.add_heading("سجل الأدلة للمدققين", level=2)
    doc.add_paragraph("لكل حقيقة مصدرها العمومي ورابطها (إن توفّر) وتاريخ "
                      "جمعها وقوة دليلها (✓ موثّق، ◐ مُقدَّر، ○ غير متحقّق).",
                      style="Intense Quote")
    _add_table(doc,
               ["الحقيقة", "المصدر", "الرابط", "تاريخ الجمع", "قوة الدليل"],
               rows[:80])
    # WP-3 §2: الإفصاح الواحد عن كل تعارض رقمي حُسم في ممرّ المصالحة.
    for c in ((dr.get("reconciliation") or {}).get("conflicts") or []):
        doc.add_paragraph(str(c.get("note") or ""), style="Intense Quote")


def _render_research_docx(doc, view: dict) -> None:
    """تقرير /research الكامل — بلاغ حي (الموجة ٨): التقرير السردي (كاتب
    التقرير) هو متن التقرير من الصفحة الأولى، لا ملحقاً بعد هيكل كلاسيكي
    غير مُغذّى. الأقسام الأربعة عشر الكلاسيكية (render_docx العادي) مبنية
    لتحليل /analyze متعدد الأسواق (`top_m` من `markets[0]`) — /research
    يحلّل سوقاً واحداً بعمق عبر البعثات الاثنتي عشر فتبقى `markets=[]`
    بنيوياً دوماً هنا، فبناء تلك الأقسام كان يُنتج هيكلاً فارغاً ("التغطية
    0.0%"، "0 أسواق"، "تعذّر إصدار توصية") يسبق المحتوى الحقيقي ويحمل حكماً
    غير مُغذّى (محرك موزون بلا مدخلات) يناقض حكم التوليف الفعلي — حكم واحد
    لا حكمان، ومتن واحد لا متنان (b). لذا: تُحذَف كلياً هنا، لا تُبنى فارغة.
    """
    dr = view.get("deep_research") or {}
    h = view.get("header") or {}
    market = dr.get("market") or {}
    verdict = dr.get("verdict") or {}
    ai = verdict.get("ai") or {}
    # vtxt يبقى الرمز الخام (GO/WATCH/...) لتصنيف اللون (_verdict_tone) —
    # لا يُطبَع حرفياً على وجه العميل أبداً؛ كل موضع عرض يمرّ عبر
    # _VERDICT_LABELS_AR[_verdict_tone(vtxt)] العربي الكامل (بلاغ تدقيق:
    # كانت الشارة تُترجَم بينما سطر «الحكم:» أسفل الخلاصة التنفيذية يطبع
    # الرمز الخام مباشرة — نفس التصنيف، مصدر عرض واحد لا اثنان).
    vtxt = _resolve_vtxt(dr)
    verdict_label = _VERDICT_LABELS_AR[_verdict_tone(vtxt)]

    # ٠) الغلاف وبطاقة التعريف — هوية سِلك (الموجة ١١، §11.1)
    branding = _load_branding()
    _add_page_header_footer(doc, f"سِلك — تقرير بحث عميق: {view.get('product')}")
    _add_cover_wordmark(doc, branding)
    doc.add_heading(f"سِلك — تقرير بحث عميق: {view.get('product')}", 0)
    doc.add_paragraph("أُعد بواسطة منصة سِلك لذكاء الأسواق", style="Intense Quote")
    if view.get("test_run"):
        doc.add_paragraph("⚠ TEST RUN — تشغيل برهاني ببدائل موسومة، "
                          "ليس تقريراً إنتاجياً")
    _stamp_degraded_banner(doc, view)
    _add_verdict_badge(doc, vtxt)
    _add_table(doc, ["البند", "القيمة"], [
        ["المنتج", h.get("product")],
        ["رمز HS", h.get("hs_code")],
        ["المنشأ", "المملكة العربية السعودية"],
        ["السوق المستهدف", h.get("target_market")
         or market.get("name_ar") or market.get("name_en")],
        ["تاريخ الإعداد", h.get("date")]])

    # جدول محتويات مطابق لبنية البحث العميق الفعلية — لا الأقسام الكلاسيكية
    # التي لا تُبنى هنا أصلاً (لا جدول محتويات يَعِد بأقسام غائبة).
    doc.add_heading("المحتويات", level=1)
    for i, ttl in enumerate((
            "الخلاصة التنفيذية", "ملخّص مصادر البحث",
            "التقرير الكامل (كاتب التقرير، مراجَع)",
            "ملحق — الأدلة الرقمية الداعمة للتقاطعات الخمسة",
            "حدود هذا التقرير"), 1):
        doc.add_paragraph(f"{i}. {ttl}")

    # ١) الخلاصة التنفيذية — من حكم التوليف حصراً (المصدر الوحيد للحكم في
    # هذا التقرير) — لا محرك §8 الموزون غير المُغذّى، ولا نص JSON خام.
    doc.add_heading("١. الخلاصة التنفيذية", level=1)
    doc.add_paragraph(f"الحكم: {verdict_label}")
    reasoning = ai.get("reasoning") or verdict.get("note") or ""
    if reasoning:
        doc.add_paragraph(_clean_report_text(reasoning, max_len=600))
    missions = dr.get("missions") or {}
    n_ok = sum(1 for m in missions.values()
              if not (m.get("failed") if isinstance(m, dict) else
                      getattr(m, "failed", False)))
    doc.add_paragraph(f"مبني على {n_ok}/{len(missions)} بعثة بحث نجحت في "
                      "جمع نتائج مبنية على استشهاد، إضافة إلى المحلل "
                      "الشامل وكاتب التقرير (راجع «حدود هذا التقرير» "
                      "أدناه لكل فجوة معلنة).")

    _docx_deep_research(doc, view)


# ══════════════════════════════════════════════════════════════════════════
# تقرير العميل (Word) — القالب الثاني (فصل الجمهور)
# ══════════════════════════════════════════════════════════════════════════
# البصيرة الجوهرية (بلاغ المالك): التصدير الحالي (_render_research_docx) يعرض
# تِلِمِتري النظام لقارئ نهائي — خطأ جمهور لا خطأ كود. اللوحة (web/index.html)
# تبقى جمهور المشغّل (بعثات، حالات، اقتصاد بيانات). هذا القالب هو جمهور
# العميل: مفردات تجارية بحتة، بلا أي مصطلح تشغيلي. حارس نصّي يرفض التصدير
# إن تسرّب أيّ مصطلح ممنوع (نفس نمط _assert_production_clean).
#
# لا مسار عرض جديد: هذا مُشتقّ آخر فوق build_view (كـrender_docx/render_brief/
# render_markdown) — يستهلك view["deep_research"] نفسه، لا يحسب رقماً جديداً.

# المصطلحات الممنوعة في تصدير العميل — الطبقة ١: قائمة المالك الصريحة
# (mission/status/successful/run/call/declared gap/tool names). الطبقة ٢:
# لغة الخوارزمية (score/confidence/verdict-token/الدرجة) — ممنوعة أيضاً في
# متن العميل (متطلَّب المرحلة ١: لا لغة نظام تصنيف، الأحكام الرقمية للملحق).
_CLIENT_TOOL_NAMES = (
    "comtrade_imports", "comtrade_competitors", "worldbank_indicator",
    "wits_tariff", "trends_interest", "faostat_supply", "web_search",
    "gdelt_news", "openalex_search", "channels_importers", "lookup_reference",
    "eurostat_eu_signals")

# أسماء المزوّدين الداخليين — عقد المالك الصريح (بلاغ UK الحي، 2026-07-17):
# لا يُسمّى مزوّد بيانات داخلي/مدفوع لعميل نهائي إطلاقاً (Volza/Explee/Serper/
# SerpApi/LocalPrice/pytrends/GDELT وصيغها العربية إكسبلي/فولزا). أسماء
# المصادر البشرية العمومية (UN Comtrade/World Bank/Eurostat…) تبقى مسموحة —
# استشهاد مشروع لا سباكة داخلية. السطح التشغيلي (?internal=1) قد يُبقيها.
# الطبقة اللاتينية بحدود كلمة (لا تخترق كلمة أطول)، والعربية مطابقة مباشرة.
_CLIENT_VENDOR_NAMES_LATIN = (
    "Volza", "Explee", "Serper", "SerpApi", "LocalPrice", "pytrends", "GDELT")
_CLIENT_VENDOR_NAMES_AR = ("إكسبلي", "فولزا")
# اللغة التجارية العامة التي تحلّ محلّ أيّ اسم مزوّد على سطح العميل.
_CLIENT_VENDOR_GENERIC = "خدمة التحقق المدفوعة من المشترين وجهات الاتصال"

# تدقيق v2 (الموجة ١، تسريبات المشرف المؤكَّدة): بدل تعداد كل صيغة تشويش، **نطبّع
# قبل المطابقة** ببناء نمطٍ متسامح — يسمح بمحارف تشويش بين كل حرفين (فراغ/
# محارف عرض-صفري/تطويل/حركات). اللاتينية تتسامح مع الفراغ أيضاً («S e r p A p i»)،
# والعربية تتسامح مع الحركات/التطويل فقط («إكْسبِلي») لا الفراغ — كي لا يطابق
# النمطُ العربيُّ تسلسلَ حروفٍ عبر كلماتٍ مشروعةٍ مفصولةٍ بفراغ (إيجابية كاذبة).
_ZW = "\u200b-\u200f\ufeff"                # zero-width + bidi marks
_HARAKAT = "\u0610-\u061a\u064b-\u065f\u0670\u06d6-\u06ed"
_OBF_LATIN = "[\\s\u0640" + _ZW + "]*"          # space + tatweel + zero-width
_OBF_AR = "[\u0640" + _HARAKAT + _ZW + "]*"    # tatweel + harakat + zero-width (no space)


def _tolerant(word: str, sep: str) -> str:
    """نمطٌ يتسامح مع محارف التشويش (sep) بين كل حرفين من الكلمة."""
    return sep.join(re.escape(ch) for ch in word)


_CLIENT_VENDOR_RE = re.compile(
    r"(?<![A-Za-z])(?:"
    + "|".join(_tolerant(w, _OBF_LATIN) for w in _CLIENT_VENDOR_NAMES_LATIN)
    + r")(?![A-Za-z])|"
    + "|".join(_tolerant(w, _OBF_AR) for w in _CLIENT_VENDOR_NAMES_AR), re.I)

# أنماط الرفض — كلٌّ يُطابَق ضد نص التصدير المُجمَّع كاملاً (فقرات + خلايا
# جداول). عربية تشغيلية + أسماء أدوات snake_case + كلمات إنجليزية تشغيلية
# ككلمات كاملة (لا تظهر في نثر عربي إلا تسرّباً). أسماء المصادر البشرية
# (UN Comtrade/World Bank/Eurostat...) مسموحة — استشهاد مشروع لا أداة داخلية.
_CLIENT_FORBIDDEN_PATTERNS = [
    # بلاغ حي إنتاجي (تمور/هولندا، 501 التصدير): "بعثتي"/"بعثتا"/"بعثتها"
    # (المثنى والصيغ ذات ضمير متصل تُبدِّل تاء التأنيث المربوطة ة إلى ت
    # المفتوحة نحوياً) تُفلِت تماماً من هذا الحارس — الصيغة القديمة
    # `بعث(?:ة|ات)` تطابق حرفياً فقط الحالتين المفردة والجمع. الفرع "ت"
    # الجديد يطابق أيّ كلمة تبدأ بـ"بعثت" (مثنى/ضمير) دون تعداد كل صيغة.
    ("mission", re.compile(r"بعث(?:ة|ت|ات)|\bmissions?\b", re.I)),
    ("status", re.compile(r"\bstatus\b", re.I)),
    ("successful", re.compile(r"ناجحة|نجحت|\bsuccessful\b", re.I)),
    ("run", re.compile(r"تشغيلة|\brun\b", re.I)),
    ("call", re.compile(r"نداء(?:ات)?\s+أدوات?|استدعاء\s+أداة|\bcall\b", re.I)),
    # "الفجوات المعلنة" (بأداة التعريف على الكلمتين) لم تكن تطابق الصيغة غير
    # المعرَّفة "فجوات معلنة" — أداة التعريف الاختيارية تسدّ هذه الفجوة.
    ("declared_gap", re.compile(
        r"(?:ال)?فجوة\s+(?:ال)?معلنة|(?:ال)?فجوات\s+(?:ال)?معلنة")),
    ("tool_name", re.compile(
        r"\b(?:" + "|".join(_CLIENT_TOOL_NAMES) + r")\b")),
    # عقد المالك (بلاغ UK الحي): اسم مزوّد داخلي/مدفوع = تسريب لجمهور العميل.
    ("vendor_name", _CLIENT_VENDOR_RE),
    ("agent_role", re.compile(
        r"المحلل الشامل|كاتب التقرير|بوابة الجودة|LLMMissionAgent|LLMAgent")),
    ("citation_plumbing", re.compile(r"مبنيّ?ة?\s+على استشهاد|بلا استشهاد|"
                                     r"\bdatapoint\b|\bdp\d+\b", re.I)),
    # الطبقة ٢ — لغة الخوارزمية في المتن (الأحكام الرقمية تعيش في الملحق فقط).
    ("algorithm_language", re.compile(
        r"\bverdict\b|\bconfidence\b|\bscore\b|الدرجة الرقمية|درجة الثقة|"
        r"النتيجة الرقمية", re.I)),
    # عائلة تسريب «§» (M-9): رمز ترقيمٍ داخليٍّ للأقسام (§4b/§8/§10.3…) لا
    # يخصّ العميل. `_CLIENT_SANITIZE` أعلاه يزيله فعلياً في المسار العادي —
    # هذا الحارس النهائي شبكة أمانٍ إضافية (نفس نمط كل بند آخر في هذه
    # القائمة: تطهيرٌ استباقي + رفضٌ نهائي لِما فات) لمدوّنةٍ مخزَّنة قديمة
    # أو مسار عرضٍ لا يمرّ عبر `_client_sanitize` أولاً. الوثائق الداخلية/
    # تعليقات الشيفرة/أسماء الاختبارات لا تخضع لهذا الحارس (نطاقه مقصورٌ على
    # سطح تصدير العميل فقط).
    ("section_marker", re.compile(r"§(?:\s*[0-9][0-9A-Za-z.\-]*)?")),
]

# استبدالات التطهير — تُطبَّق على كل كتلة نص من سرد الكاتب قبل عرضها، فتحوّل
# أيّ مصطلح تشغيلي تسرّب من الكاتب إلى مفردة تجارية. الحارس النهائي يلتقط ما
# فات. الترتيب مهم (الأطول أولاً).
_CLIENT_SANITIZE = [
    # الدرس ٣٣ (belt-and-suspenders — التقادُم من الحقل البنيويّ لا النثر):
    # وسمُ «year=YYYY» لم يعد يُكتَب (جامعو DataPoint يضبطون data_year)، لكن
    # مدوّنةً مخزَّنةً قديمة قد تحمله في ملاحظة — يُزال من سطح العميل قبل أن يُرى.
    (re.compile(r"\s*\byear\s*=\s*\d{4}\b"), ""),
    # عائلة تسريب «§» (M-9 + بلاغ حي 2026-07-20، §4b على تقرير العميل): ترقيمُ
    # أقسامٍ داخليّ (§4b/§8/§10.3/§11.5-2…) لا يخصّ العميل ويُبدَّل بصريًّا في
    # RTL — يُزال من سطح العميل. المصدر مُصلَح (silk_decision/silk_render) وهذه
    # شبكة أمان لأيّ تشغيلة مخزَّنة قديمة أو صياغةٍ يردّدها النموذج.
    # رمز القسم = «§» + (اختيارياً) رقمٌ فلاحقة ASCII (أرقام/لاتيني/نقطة/شرطة)
    # فقط — مُرسًى برقمٍ ومقصورٌ على ASCII كي لا يبتلع كلمةً عربية تالية مثل
    # «§ المنهجية» (مراجعة الشيفرة)؛ و«§» وحده يُزال أيضاً.
    (re.compile(r"\s*§(?:\s*[0-9][0-9A-Za-z.\-]*)?"), ""),
    # تمرير النثر (R1): مصطلحات مترجَمة حرفياً عن الإنجليزية التقنية → لغة
    # الأعمال الخليجية. المصدر مُصلَح في الكاتب/البعثة، وهذه شبكة أمان تلتقط
    # أيّ مخرَج حيّ (أو تشغيلة مخزَّنة قديمة) لا يزال يحمل الصياغة الحرفية.
    (re.compile(r"سعر جملة مرجعي"), "متوسط سعر الاستيراد الرسمي"),
    (re.compile(r"كومتريد"), "UN Comtrade"),        # لا تعريب لاسم مصدر
    (re.compile(r"تكلفة\s+هبوط"), "التكلفة الواصلة"),  # landed cost
    (re.compile(r"\|\s*الدليل\s*\|"), "| مستوى التوثيق |"),  # خلية عمود الجدول
    (re.compile(r"مطبَّع(?:ةً|ة|اً|ًا)?\s+لكل\s+(?:كيلوغرام|كجم|كيلو|وحدة)"),
     "محسوبةً بسعر الكيلوغرام الواحد للمقارنة العادلة"),
    (re.compile(r"\b(?:" + "|".join(_CLIENT_TOOL_NAMES) + r")\b"),
     "السجلّات الرسمية"),
    (re.compile(r"من\s+السجلّات الرسمية"), "من السجلّات الرسمية"),
    # عقد المالك (بلاغ UK الحي): أيّ اسم مزوّد داخلي متسرّب (من سرد الكاتب أو
    # ملاحظة مُعرَّبة مثل «إكسبلي غير متاح حالياً») يُحوَّل للغة أعمال عامة قبل
    # أن يراه العميل. الحارس النهائي (vendor_name) شبكة أمان لِما فات.
    (_CLIENT_VENDOR_RE, _CLIENT_VENDOR_GENERIC),
    (re.compile(r"LLMMissionAgent|LLMAgent"), "مصدر البيانات"),
    # §2.2 (أمر العمل الرئيس): لا تُنسَب الحقائق لمسارات بحث داخلية — الصياغة
    # المحايدة «جمع البيانات» تحلّ محل «بعثة»/«مسار بحث». صيغ المثنى/الضمير
    # المتصل لـ"بعثة" (ة→ت نحوياً) تُحوَّل قبل الصيغتين المفردة/الجمع أدناه.
    (re.compile(r"بعثت\w*"), "جمع البيانات"),
    (re.compile(r"بعثات"), "عمليات جمع البيانات"),
    (re.compile(r"بعثة"), "جمع البيانات"),
    (re.compile(r"نجحت في جمع"), "أنتجت"),
    (re.compile(r"ناجحة"), "مكتملة"),
    (re.compile(r"نجحت"), "اكتملت"),
    (re.compile(r"فشلت"), "لم تكتمل"),
    # أداة التعريف الاختيارية (نفس تمديد الحارس أعلاه) قبل الصيغتين الأصليتين.
    (re.compile(r"(?:ال)?فجوات\s+(?:ال)?معلنة"), "بنود تحتاج تحققاً"),
    (re.compile(r"(?:ال)?فجوة\s+(?:ال)?معلنة"), "بند يحتاج تحققاً"),
    (re.compile(r"تشغيلة"), "دورة تحليل"),
    (re.compile(r"مبنيّ?ة?\s+على استشهاد"), "موثّقة بمصادرها"),
    # البلاغ الحي الثالث لعائلة 501 (2026-07-16، «فشل التنزيل: HTTP 501 وما
    # زال»): ثلاثة محفّزات عربية في الحارس كانت بلا أي استبدال مقابل هنا —
    # أي نثر كاتب يلمسها يُسقِط التصدير كله. القفل الميكانيكي الذي يمنع
    # تكرار العائلة بأكملها (كل محفّز حارس عربي مغطّى إجبارياً):
    # tests/test_client_sanitizer_covers_guard.py — LESSONS.md البند ١١.
    (re.compile(r"نداء(?:ات)?\s+(?:ال)?أدوات?|استدعاء(?:ات)?\s+(?:ال)?أدا(?:ة|وات)"),
     "عمليات جمع البيانات"),
    (re.compile(r"بوابة الجودة"), "مراجعة الجودة"),
    (re.compile(r"بلا استشهاد"), "دون توثيق مصدر"),
    # معرّفات الاستشهاد الداخلية dpN — مجموعة بين قوسين تُحذف كلياً (لا
    # «()» مبتورة)، والمفردة العارية تصير وصفاً تجارياً.
    (re.compile(r"\(\s*dp\d+(?:\s*[،,;\s]\s*dp\d+)*\s*\)"), ""),
    (re.compile(r"\bdp\d+\b"), "مصدر موثّق"),
    (re.compile(r"\bdatapoints?\b", re.I), "معلومة موثّقة"),
    (re.compile(r"المحلل الشامل"), "فريق التحليل"),
    (re.compile(r"كاتب التقرير"), "فريق الإعداد"),
    # بلاغ حي إنتاجي (تدقيق تمور/هولندا، 501 تصدير العميل): سرد الكاتب يمرّ
    # عبر `_strip_internal_plumbing` (silk_render) الذي يعرّب حقول الحكم
    # الإنجليزية (confidence→«درجة الثقة»، verdict→«الحكم»، score→«الدرجة»)،
    # فتصل الصيغة العربية «درجة الثقة»/«الدرجة الرقمية» متن العميل وتُسقِطه
    # حارسُ لغة الخوارزمية (_client_assert_clean) بـ501. تُحوَّل هنا لمفردة
    # توثيق تجارية — الأحكام الرقمية تبقى للملحق/التصدير التشغيلي فقط.
    (re.compile(r"الدرجة الرقمية|النتيجة الرقمية"), "التقييم"),
    (re.compile(r"درجة الثقة"), "مستوى التوثيق"),
    (re.compile(r"\bconfidence\b", re.I), "مستوى التوثيق"),
    (re.compile(r"\bverdict\b", re.I), "التوصية"),
    (re.compile(r"\bscore\b", re.I), "التقييم"),
]


def _client_sanitize(text: object) -> str:
    """طهّر كتلة نص من سرد الكاتب لجمهور العميل — يحوّل المصطلحات التشغيلية
    المتسرّبة لمفردات تجارية. الحارس النهائي (_client_assert_clean) يرفض أي
    بقايا. لا يمسّ الأرقام ولا المصادر البشرية."""
    s = str(text or "")
    for pat, repl in _CLIENT_SANITIZE:
        s = pat.sub(repl, s)
    return re.sub(r"[ \t]{2,}", " ", s).strip()


def _client_forbidden_hits(blob: str) -> list[str]:
    """أرجع كل مصطلح ممنوع ظهر في نص التصدير — للاختبار وللحارس."""
    hits = []
    for label, pat in _CLIENT_FORBIDDEN_PATTERNS:
        m = pat.search(blob or "")
        if m:
            hits.append(f"{label}: «{m.group(0)}»")
    return hits


# PART A (أمر العمل الرئيس — عائلة 501 المتكرّرة): بديل محايد آمن لكل نمط
# ممنوع، تُستعمَل في **مسار التنقية النهائي** (لا الرفض). فلسفة «نقِّ لا
# ترفض»: مصطلح تشغيلي متبقٍّ يُستبدَل بمفردة محايدة ويُسلَّم المستند مع سطر
# إفصاح، بدل إسقاط التصدير كله بـ501 على تسرّب واحد. الحارس يبقى شبكة أمان
# أخيرة لِما يستحيل تنقيته فقط.
_CLIENT_REDACT_PLACEHOLDER = {
    "mission": "مصدر البيانات",
    "status": "الحالة",
    "successful": "مكتملة",
    "run": "دورة تحليل",
    "call": "عملية جمع بيانات",
    "declared_gap": "بند يحتاج تحققاً",
    "tool_name": "السجلّات الرسمية",
    "agent_role": "فريق البحث",
    "citation_plumbing": "مصدر موثّق",
    "algorithm_language": "التقييم",
    "vendor_name": _CLIENT_VENDOR_GENERIC,
    # يُزال لا يُستبدَل — رمز قسمٍ داخليٍّ ليس معلومةً تُستبدَل بمرادف، مطابقةً
    # لسلوك `_CLIENT_SANITIZE` (استبدالٌ بفراغ) لا حشو placeholder.
    "section_marker": "",
}


def _client_redact_text(text: str) -> str:
    """استبدل كل مطابقة نمط ممنوع بمفردة محايدة — مسار التنقية النهائي.
    يُطبَّق بعد `_client_sanitize`، فيلتقط ما تبقّى (كلمات إنجليزية تشغيلية
    عارية لم يُعرّبها المُطهِّر تفادياً لتشويه أسماء المصادر)."""
    s = str(text or "")
    for label, pat in _CLIENT_FORBIDDEN_PATTERNS:
        s = pat.sub(_CLIENT_REDACT_PLACEHOLDER.get(label, "—"), s)
    return re.sub(r"[ \t]{2,}", " ", s)


def _redact_paragraph(par) -> bool:
    """نقِّ فقرة python-docx في مكانها — يعيد True إن غُيِّر شيء. يعيد بناء
    نصّ الفقرة في الجرية الأولى (يفقد تنسيقاً داخلياً نادراً، مقبول لمسار
    شبكة أمان لا يُفعَّل إلا عند تسرّب)."""
    if not _client_forbidden_hits(par.text):
        return False
    redacted = _client_redact_text(par.text)
    if par.runs:
        par.runs[0].text = redacted
        for r in par.runs[1:]:
            r.text = ""
    else:
        par.add_run(redacted)
    return True


def _client_redact_residual(doc) -> bool:
    """مسار «نقِّ لا ترفض» (PART A): امسح كل فقرة/خلية، واستبدل أيّ مصطلح
    ممنوع متبقٍّ بمحايد. يعيد True إن نُقِّي شيء (فيُلحَق سطر إفصاح)."""
    changed = False
    for par in doc.paragraphs:
        changed = _redact_paragraph(par) or changed
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    changed = _redact_paragraph(par) or changed
    return changed


def _client_assert_clean(doc) -> None:
    """حارس تصدير العميل — شبكة الأمان الأخيرة بعد التنقية. يرفض فقط إن
    استحال تنقية مصطلح (لا يقع في المسار العادي بعد _client_redact_residual).
    يمسح كل الفقرات وخلايا الجداول المُجمَّعة."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    blob = "\n".join(parts)
    hits = _client_forbidden_hits(blob)
    if hits:
        raise RuntimeError(
            "تصدير العميل يحوي مصطلحات ممنوعة تعذّرت تنقيتها (تسريب تِلِمِتري "
            "لجمهور العميل) — رُفض التوليد: " + "؛ ".join(hits[:8]))


# خريطة أقسام الكاتب (البنية الدولية بأحد عشر قسماً) → أقسام تقرير العميل
# السبعة. أقسام "منهجية البحث ونطاقه" و"الملاحق" تُسقَط (تُستبدَل بفقرة
# منهجية مضبوطة وملحق أدلة). الترتيب النهائي للعميل يُبنى صراحة في المُصيّر.
_CLIENT_SECTION_MAP = {
    "الخلاصة التنفيذية": "القرار وأساسه",
    "التوصيات الاستراتيجية": "القرار وأساسه",
    "نظرة عامة على السوق وحجمه": "السوق بالأرقام",
    "ديناميكيات السوق": "السوق بالأرقام",
    "تحليل المستهلك والطلب": "السوق بالأرقام",
    "المشهد التنافسي": "المنافسة والتسعير والهامش",
    "التنظيم والوصول للسوق": "مسار الدخول والمتطلبات",
    "اللوجستيات وسلسلة الإمداد": "مسار الدخول والمتطلبات",
    "تقييم المخاطر": "المخاطر",
}
# ترتيب أقسام العميل النهائي (بلاغ المالك، البنية المطلوبة).
_CLIENT_SECTION_ORDER = (
    "القرار وأساسه", "السوق بالأرقام", "المنافسة والتسعير والهامش",
    "مسار الدخول والمتطلبات", "المخاطر")

_WRITER_HEADING_RE = re.compile(r"^##\s+\d+\.\s*(.+?)\s*$")
_ROADMAP_SUBHEAD_RE = re.compile(r"^###\s+خارطة طريق الدخول")


def _split_at_roadmap(body: list[str]) -> "tuple[list[str], list[str]]":
    """قسّم متن التوصيات عند العنوان الفرعي "### خارطة طريق الدخول" — الجزء
    قبله (تعليل الحكم) والجزء منه فصاعداً (الخارطة). نقطة قسم حتمية (عنوان
    فرعي ثابت يفرضه كاتب التقرير) لا تحليل نصّي هشّ."""
    for idx, line in enumerate(body):
        if _ROADMAP_SUBHEAD_RE.match(line.strip()):
            return body[:idx], body[idx:]
    return body, []


def _parse_writer_sections(report_text: str) -> "list[tuple[str, list[str]]]":
    """قسّم سرد الكاتب إلى (عنوان القسم، أسطر متنه) — يعيد استعمال بنية
    '## N. العنوان' نفسها التي يفرضها كاتب التقرير (silk_ai_judge._REPORT_
    SECTIONS). العناوين الفرعية '### ' والجداول تبقى ضمن متن قسمها."""
    sections: list[tuple[str, list[str]]] = []
    cur_title, cur_body = None, []
    for line in str(report_text or "").splitlines():
        m = _WRITER_HEADING_RE.match(line.strip())
        if m:
            if cur_title is not None:
                sections.append((cur_title, cur_body))
            cur_title, cur_body = m.group(1).strip(), []
        elif cur_title is not None:
            cur_body.append(line)
    if cur_title is not None:
        sections.append((cur_title, cur_body))
    return sections


def _client_render_body_block(doc, lines: list[str]) -> None:
    """اعرض أسطر متن قسم (فقرات + جداول Markdown + سطر الخلاصة الغامق) بعد
    التطهير — نفس ميكانيكا _docx_deep_research لكن عبر _client_sanitize،
    ويُسقِط أيّ جدول بقي فيه مصطلح ممنوع بعد التطهير (جداول الدرجات الرقمية
    verdict/confidence — لغة خوارزمية لا تُعرَض للعميل)."""
    i, n = 0, len(lines)
    while i < n:
        raw = lines[i].strip()
        if not raw:
            i += 1
            continue
        if raw.startswith("### "):
            doc.add_heading(_client_sanitize(raw[4:]), level=3)
            i += 1
            continue
        if raw.startswith("## "):  # عنوان قسم فرعي غير متوقّع — كفقرة مطهَّرة
            doc.add_paragraph(_client_sanitize(raw[3:]))
            i += 1
            continue
        if _is_markdown_table_row(raw):
            block = []
            while i < n and _is_markdown_table_row(lines[i].strip()):
                block.append(lines[i].strip())
                i += 1
            joined = _client_sanitize("\n".join(block))
            if _client_forbidden_hits(joined):
                continue  # جدول درجات/تِلِمِتري — يُسقَط من متن العميل
            _render_markdown_table(doc, joined.splitlines())
            continue
        stripped = _strip_inline_markdown(_client_sanitize(raw))
        if stripped.startswith("ماذا يعني هذا لقرارك:"):
            p = doc.add_paragraph()
            p.add_run(stripped).bold = True
        else:
            doc.add_paragraph(stripped)
        i += 1


def _client_methodology_paragraph(dr: dict) -> str:
    """فقرة المنهجية المضبوطة (§2.5، أمر العمل الرئيس) — صياغة عامة لا تكشف
    بنية النظام: المصادر المُستشارة، تاريخ الجمع، وأسلوب التحقّق. **لا ذكر
    لوكلاء أو أدوات أو مسارات بحث أو عدد مكوّنات داخلية** — من حقول محسوبة
    فعلاً (المصادر العمومية وتواريخها) لا اختلاق."""
    missions = dr.get("missions") or {}
    # المصادر البشرية الفريدة الظاهرة فعلاً في النتائج (لا أسماء أدوات).
    # WP-3 §3: (أ) مفتاح التفريد مُطبَّع (casefold + إسقاط المحارف غير
    # المرئية + طيّ الفراغات) — «GAFTA secretariat» لا تتكرّر بصيغتين؛
    # (ب) مصدرٌ كل بنوده أخطاء (value=None — خدمة فشلت في هذه التشغيلة)
    # يُستبعَد من سطر «اعتمد هذا التقرير…» ويُذكَر في الحدود فقط
    # (silk_render._deep_research_view).
    contributed: dict[str, bool] = {}
    display: dict[str, str] = {}
    for m in missions.values():
        for f in (m.get("findings") or []) if isinstance(m, dict) else []:
            src = _client_sanitize(_clean_source_label(f.get("source")))
            src = str(src or "").strip()
            if not src or src == "—" or _client_forbidden_hits(src):
                continue
            # اسم مصدر بشري فقط (قبل أيّ شرطة توضيحية).
            name = re.split(r"\s+[—\-(]", src)[0].strip()
            norm = re.sub(r"\s+", " ",
                          re.sub("[\\u200b-\\u200f\\ufeff]", "", name)
                          ).strip().casefold()
            if not norm:
                continue
            display.setdefault(norm, name)
            contributed[norm] = contributed.get(norm, False) or (
                f.get("value") is not None)
    src_list = "، ".join(sorted(
        display[n] for n, ok in contributed.items() if ok)[:6]) \
        or "مصادر رسمية عامة"
    dates = sorted({str(f.get("retrieved_at"))
                    for m in missions.values()
                    if isinstance(m, dict)
                    for f in (m.get("findings") or [])
                    if f.get("retrieved_at")})
    date_txt = (f"أحدث تاريخ جمع بيانات: {dates[-1]}" if dates
                else "تواريخ الجمع مسجّلة في سجل الأدلة أدناه")
    return (
        f"اعتمد هذا التقرير على مصادر رسمية عامة ({src_list})، مع تحليل "
        "تقاطعي بينها ومراجعة للاتساق قبل الاعتماد. خضعت كل معلومة للتحقّق "
        f"من مصدرها العمومي المباشر. {date_txt}. قائمة المراجع الكاملة "
        "بروابطها الرسمية ختام التقرير.")


# صياغة تجارية لكل فجوة في قسم "ما لم يكتمل للقرار" (بلاغ المالك، النقطة ٣):
# لا عناوين فارغة متتالية، بل جملة تجارية موحّدة القالب لكل تقاطع بلا أدلة.
_CLIENT_GAP_TEMPLATE = ("لم نتمكّن من توثيق {what} من مصدر موثّق ضمن هذا "
                        "التقرير — إغلاق هذه الفجوة يتطلّب {how}.")
_CLIENT_GAP_WHAT = {
    "demand": ("الحجم الدقيق للطلب الفعلي القابل للتوجيه",
               "بحثاً ميدانياً أوّلياً (مقابلات موزّعين أو استبيان طلب)"),
    "entry_cost": ("تكلفة الدخول الكاملة بما فيها الشحن الفعلي",
                   "عرض أسعار شحن ملزماً من وكيل لوجستي"),
    "price_competitiveness": ("موقعك السعري الدقيق مقابل المنافسين",
                              "بطاقة منتجك (تكلفتك للكيلوغرام) وخدمة رصد "
                              "أسعار مدفوعة"),
    "entry_door": ("قائمة موزّعين/مستوردين مؤكَّدين بالاسم",
                   "خدمة تحقّق جهات اتصال مدفوعة (قواعد بيانات تجارية)"),
    "swot": ("الموقف التنافسي المؤكَّد من منظور مصدّر سعودي",
             "دمج نتائج التحقّق الميداني أعلاه في تقييم موحّد"),
}


def _client_confidence_section(doc, dr: dict) -> None:
    """مؤشّر ثقة الدراسة (S3) — عدّ شارات الأدلة (✓ موثّق/◐ ثانوي/○ غير متحقق)
    عبر بعثات الدراسة وتقاطعات المحلل، فيرى العميل شفافياً كم من الدراسة
    مرصود بمصدر رسمي مقابل مُقدَّر أو فجوة. تجميع بحت من درجات ثقة قائمة —
    لا رقم جديد ولا حكم، ويمرّ بحارس المصطلحات كأيّ قسم عميل."""
    # WP-3: العدّ بنفس الشارة الواعية بالمنشأ والمصالحة المعروضة في سجلّ
    # الأدلة — «متعارض — مستبعد» يُحسَب غير متحقق، لا موثّقاً.
    from silk_narrative import evidence_badge_for
    counts = {"verified": 0, "secondary": 0, "unverified": 0}

    def _tally(f) -> None:
        badge = evidence_badge_for(f)
        if badge.startswith("✓"):
            counts["verified"] += 1
        elif badge.startswith("◐"):
            counts["secondary"] += 1
        else:
            counts["unverified"] += 1

    for m in (dr.get("missions") or {}).values():
        if not isinstance(m, dict) or m.get("failed"):
            continue
        for f in (m.get("findings") or []):
            if _dp_conf(f) is not None:
                _tally(f)
    for dps in ((dr.get("analyst") or {}).get("by_category") or {}).values():
        for f in (dps or []):
            if _dp_conf(f) is not None:
                _tally(f)

    total = sum(counts.values())
    if not total:
        return   # لا مؤشرات معدودة — لا قسم فارغ
    verified_pct = round(100 * counts["verified"] / total)
    doc.add_heading("مؤشّر ثقة الدراسة", level=1)
    doc.add_paragraph(
        f"من إجمالي {total} مؤشراً مرصوداً في هذه الدراسة، "
        f"{counts['verified']} موثّق بمصدر رسمي (نحو {verified_pct}%)، "
        f"{counts['secondary']} من مصدر ثانوي، و{counts['unverified']} غير "
        "متحقق. كلّ رقم في التقرير يحمل شارة توثيقه ومصدره في سجل الأدلة، "
        "فالقرار يُتَّخذ بمعرفة درجة اليقين خلف كل رقم لا على ثقة عمياء.")
    _add_table(doc, ["درجة التوثيق", "عدد المؤشرات"], [
        ["✓ موثّق (مصدر رسمي)", str(counts["verified"])],
        ["◐ ثانوي (مصدر واحد غير رسمي)", str(counts["secondary"])],
        ["○ غير متحقق", str(counts["unverified"])]])


def _client_gap_inputs(dr: dict) -> "tuple[list[str], list[str]]":
    """(الفجوات الحرجة للقرار، الفجوات المعلنة غير الحاجبة) — WP-4: المصدر
    الواحد الذي يقرأه قسم «ما لم يكتمل للقرار» **وحارس تناقض الختام** في
    بوابة الجودة معاً، فلا يوجد مساران للحقيقة (كان الختام يطبع «لا فجوة
    جوهرية» بينما قسم المخاطر يعلن ثلاث فجوات بيانات حقيقية — فجوات
    البعثات المعلنة لم تكن مدخلاً لهذا القسم إطلاقاً).

    الحرجة (تمنع اكتمال القرار): تقاطعات المحلل الغائبة، باب الدخول الأول
    غير المحقَّق، شروط قلب الحكم غير المحقَّقة. غير الحاجبة (تقيّد اليقين):
    فجوات البعثات المعلنة (missions[*] «فجوات: …» — مطهَّرة، مقصوصة عند حدّ
    جملة، بلا تكرار، بسقف بندين لكل بعثة)."""
    analyst = dr.get("analyst") or {}
    missing = analyst.get("missing_categories") or []
    by_cat = analyst.get("by_category") or {}

    critical: list[str] = []
    for cat in missing:
        what, how = _CLIENT_GAP_WHAT.get(
            cat, ("بند تحليلي إضافي", "بحثاً تكميلياً موجّهاً"))
        critical.append(_CLIENT_GAP_TEMPLATE.format(what=what, how=how))
    if "entry_door" not in missing:
        from silk_narrative import EVIDENCE_SECONDARY_MIN
        unverified_doors = [
            f for f in (by_cat.get("entry_door") or [])
            if _dp_conf(f) is not None and _dp_conf(f) < EVIDENCE_SECONDARY_MIN]
        if unverified_doors:
            names = "، ".join(
                n for n in (_client_sanitize(_client_prose(f.get("value"), 80))
                            for f in unverified_doors[:2]) if n)
            _named = f" ({names})" if names else ""
            critical.append(
                f"لم نتمكّن من تأكيد قناة الدخول الأولى{_named} من مصدر "
                "موثّق — إغلاق هذه الفجوة يتطلّب خدمة تحقّق جهات اتصال مدفوعة "
                "(قواعد بيانات تجارية) قبل الالتزام بالموزّع.")
    for c in (dr.get("flip_conditions") or []):
        if not c.get("met"):
            critical.append(
                f"لم يتحقّق بعد: {c.get('condition')} — إغلاق هذا الشرط "
                f"يتطلّب {c.get('closes_via')}.")

    # WP-4 §1: المدخل الرابع — فجوات البعثات المعلنة داخل ملخّصاتها.
    from silk_render import _mission_gap_lines
    informational: list[str] = []
    seen: set[str] = set()
    for k, m in (dr.get("missions") or {}).items():
        if not isinstance(m, dict):
            continue
        label = m.get("label") or str(k)
        per_mission = 0
        for line in _mission_gap_lines(label, m.get("summary") or ""):
            g = _client_sanitize(_trim_sentence(line, 200)).rstrip(".؛،")
            if not g or g in seen or per_mission >= 2:
                continue
            seen.add(g)
            per_mission += 1
            informational.append(
                f"فجوة بيانات معلنة — {g}؛ لا تمنع القرار الحالي لكنها "
                "تقيّد يقين الاستنتاجات المتصلة بها.")
    return critical, informational


def _client_gaps_section(doc, dr: dict) -> None:
    """قسم "ما لم يكتمل للقرار والخطوة التالية" — يحوّل كل تقاطع بلا أدلة
    كافية لصياغة تجارية موحّدة (بلاغ المالك النقطة ٣)، ثم الخطوة التالية.
    لا عناوين فارغة متتالية: إن اكتمل كل شيء، سطر إيجابي واحد. WP-4: يقرأ
    كل المدخلات الأربعة من `_client_gap_inputs` (المصدر الواحد المشترك مع
    حارس البوابة) — السطر الإيجابي يُطبَع فقط حين تخلو القوائم الأربع معاً."""
    doc.add_heading("ما لم يكتمل للقرار، والخطوة التالية", level=1)
    gap_lines, mission_gap_lines = _client_gap_inputs(dr)

    if gap_lines:
        doc.add_paragraph(
            "النقاط التالية لم تكتمل توثيقاً ضمن هذا التقرير؛ هي ما يفصل "
            "التوصية الحالية عن قرار نهائي كامل، وكلٌّ منها قابل للإغلاق "
            "بخطوة محدّدة:")
        for line in gap_lines + mission_gap_lines:
            doc.add_paragraph(line, style="List Bullet")
    elif mission_gap_lines:
        # WP-4 §2: فجوات معلنة غير حاجبة — لا سطر «لا فجوة جوهرية» بجانبها.
        n = len(mission_gap_lines)
        count_txt = ("فجوة معلنة واحدة لا تمنع" if n == 1
                     else f"{n} فجوات معلنة لا تمنع")
        doc.add_paragraph(
            f"توجد {count_txt} القرار الحالي لكنها تقيّد يقينه — مفصّلة "
            "أدناه:")
        for line in mission_gap_lines:
            doc.add_paragraph(line, style="List Bullet")
    else:
        # WP-4 §2: السطر الإيجابي فقط حين تخلو القوائم الأربع معاً.
        doc.add_paragraph(
            "اكتملت التقاطعات التحليلية الأساسية بأدلة موثّقة بمصادرها، ولا "
            "بند حاسم للقرار موسوم بأنه غير محقَّق؛ لا فجوة جوهرية تمنع "
            "اتخاذ القرار ضمن نطاق هذا التقرير.")
    # §F-5 (حزمة الفكس v2.1): دعوة التعميق المدفوع («next_step») أُزيلت من
    # متن الدراسة المُسلَّم للعميل — تبقى دعوةً على سطح التسليم (شارة/زرّ
    # اللوحة، web/index.html) لا داخل المستند نفسه.


_CAT_TAG_RE = re.compile(
    r"^\[(?:demand|price_competitiveness|entry_cost|entry_door|swot|[a-z_]+)\]\s*")


def _dp_conf(f: object) -> "float | None":
    """درجة الثقة لبند أدلة (dict أو DataPoint) — None إن تعذّر."""
    if isinstance(f, dict):
        c = f.get("confidence")
    else:
        c = getattr(f, "confidence", None)
    try:
        return float(c)
    except (TypeError, ValueError):
        return None


def _readable_number(v: object) -> str:
    """رقم بصيغة عربية مقروءة — 38000000.0 → «38 مليون» (بلاغ مراجعة المالك:
    عشريات خام في سجل الأدلة). الوحدة (دولار/نسمة/…) تأتي من الملاحظة لا
    تُختلَق هنا."""
    try:
        n = float(v)
    except (TypeError, ValueError):
        return str(v)
    a = abs(n)
    if a >= 1e9:
        return f"{n / 1e9:.1f} مليار".replace(".0 ", " ")
    if a >= 1e6:
        return f"{n / 1e6:.1f} مليون".replace(".0 ", " ")
    if a >= 1e3:
        return f"{n:,.0f}"
    return f"{n:g}"


def _client_readable_fact(value: object, note: object) -> "str | None":
    """حوّل قيمة أدلة إلى حقيقة مقروءة لعمود «الحقيقة» — بلاغ مراجعة المالك
    (النقطة ٤): (أ) البنود التي لا تُعرَض مباشرة (قيمة dict غير معروفة/رد
    خام) تُسقَط (None) بدل سطر «بند تقني غير قابل للعرض» عديم المعنى؛ (ب)
    العشريات الخام تُنسَّق مقروءةً مع سياق ملاحظتها. لا اختلاق وحدة."""
    clean_note = _client_sanitize(_CAT_TAG_RE.sub("", str(note or "")).strip())
    if isinstance(value, dict):
        if value.get("partner") is not None and value.get("share") is not None:
            return f"{value['partner']}: حصة {value['share']}%"
        if value.get("hhi") is not None:
            return f"مؤشر تركّز المورّدين HHI={value['hhi']}"
        return None  # بنية غير معروفة — لا سطر عديم المعنى
    if isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        num = _readable_number(value)
        return f"{num} — {clean_note}" if clean_note else num
    # القالب الأكاديمي أعاد إحياء هذه الدالة لخلايا جداول الشواهد — القصّ
    # عبر `_client_prose` (حدّ جملة، بلا «…»، يُسقِط الكتل الخام) لا
    # `_clean_report_text` (كانت تبتر بـ«…» فتُفشِل بوابة نصّ المُنتَج).
    txt = _client_sanitize(_client_prose(value, 140))
    if not txt:
        return None  # بند لا يُعرَض — يُسقَط بدل ضجيج للمدقّق
    return txt


# §A (حزمة الفكس v2.1) — «المراجع» تحلّ محلّ «سجل الأدلة للمدققين» في
# بناء العميل: قائمة مصادر عمومية فريدة (لا سجلّ تدقيقي ببند لكل حقيقة)،
# مبنية من حقول DataPoint البنيوية (source/value/retrieved_at) — لا تخمين
# بالكلمات المفتاحية. سجلّ الأدلة الكامل يبقى في الملحق الداخلي فقط
# (`_docx_technical_appendix`، عبر `?internal=1`).
_INTERNAL_SOURCE_LABEL_RE = re.compile(
    r"silk\s*l1|مرجع\s*سلك|سجلات?\s*رسمية|silk\s*requirements", re.I)


def _reference_row_from_finding(source_raw: object, value: object,
                                note: object) -> "tuple[str, str] | None":
    """(اسم مصدر عمومي، رابطه الحقيقي) لبند واحد، أو None إن تعذّر تحديد
    مصدر عمومي حقيقي — عندها يبقى البند في الملحق الداخلي فقط (§A-2). لا
    تخمين بالكلمات المفتاحية: إمّا مصدر معروف في `SOURCE_PUBLIC_URL`، أو
    (لِمراجع طبقة ١ التي تحمل تسمية مصدر داخلية) حقلا `authority`/
    `source_url` البنيويّان المخزَّنان فعلاً داخل قيمة البند نفسه
    (`silk_requirements_agent._row_dp`)."""
    from silk_data_layer import public_source_url
    label = _clean_source_label(source_raw)
    if _INTERNAL_SOURCE_LABEL_RE.search(str(source_raw or "")) or \
            _INTERNAL_SOURCE_LABEL_RE.search(label):
        if isinstance(value, dict) and value.get("source_url"):
            name = str(value.get("authority") or "").strip()
            return (name, str(value["source_url"])) if name else None
        return None  # لا حقل مصدر عمومي بنيوي — يبقى داخلياً فقط
    if not label or label == "—" or _client_forbidden_hits(label):
        return None
    url = _first_url(note, value)
    if url == "—":
        url = public_source_url(label, arabic=True)
    if not url:
        return None
    return (label, url)


def _client_references_section(doc, dr: dict) -> None:
    """المراجع (§A) — مصادر عمومية فريدة فقط، سطر واحد لكل مصدر: الاسم +
    الرابط الرسمي الحقيقي + تاريخ آخر جمع بيانات منه. لا جدول أدلة تدقيقي
    ببند لكل حقيقة، ولا بند بشارة ○ غير متحقَّق (§A-4: لا يُعرَض كمرجعٍ
    مصدرٌ لم يُتحقَّق منه)."""
    doc.add_heading("المراجع", level=1)
    doc.add_paragraph(_client_methodology_paragraph(dr))
    refs: dict[str, dict] = {}
    for m in (dr.get("missions") or {}).values():
        if not isinstance(m, dict):
            continue
        for f in (m.get("findings") or []):
            from silk_narrative import RECONCILED_OUT_TAG, evidence_badge_for
            badge = evidence_badge_for(f)
            # WP-3: بند مستبعد بالمصالحة أو غير متحقَّق (بعد سقف المنشأ) —
            # لا يُعرَض كمرجع للعميل (§A-4).
            if badge.startswith("○") or badge.startswith(RECONCILED_OUT_TAG):
                continue
            row = _reference_row_from_finding(
                f.get("source"), f.get("value"), f.get("note"))
            if row is None:
                continue
            name, url = row
            key = name.strip().lower()
            ra = str(f.get("retrieved_at") or "")
            existing = refs.get(key)
            if existing is None or ra > existing["retrieved_at"]:
                refs[key] = {"name": name, "url": url, "retrieved_at": ra}
    if not refs:
        doc.add_paragraph("لا مصادر عمومية موثّقة قابلة للعرض هنا في هذه "
                          "التشغيلة.")
        return
    for key in sorted(refs, key=lambda k: refs[k]["name"]):
        r = refs[key]
        line = f"{r['name']} — {r['url']}"
        if r["retrieved_at"]:
            line += f" (تاريخ الجمع: {r['retrieved_at']})"
        doc.add_paragraph(line, style="List Bullet")


def render_client_docx(view: dict, path: str) -> str:
    """تقرير العميل (Word) — القالب الثاني، جمهور العميل (بلاغ المالك: فصل
    الجمهور). يستهلك view["deep_research"] نفسه (لا مسار عرض جديد) وينتج
    بنية العميل السبع: القرار وأساسه → السوق بالأرقام → المنافسة والتسعير →
    مسار الدخول والمتطلبات → المخاطر → ما لم يكتمل للقرار → المنهجية وسجل
    الأدلة. حارس نصّي يرفض التصدير إن تسرّب أيّ مصطلح تشغيلي/خوارزمي.

    نتيجة /analyze الكلاسيكية (بلا deep_research) ليست ضمن نطاق هذا القالب
    — تُوجَّه لـrender_docx العادي من المستدعي. يعيد المسار عند النجاح.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(_DOCX_HINT) from exc

    _assert_production_clean(view)
    dr = view.get("deep_research") or {}
    if not dr:
        raise RuntimeError("render_client_docx يتطلّب نتيجة بحث عميق "
                           "(view['deep_research']) — استخدم render_docx "
                           "لتقارير /analyze الكلاسيكية")

    doc = Document()
    _apply_rtl(doc)   # §4: المستند كله من اليمين لليسار (تقرير العميل)
    h = view.get("header") or {}
    market = dr.get("market") or {}
    verdict = dr.get("verdict") or {}
    ai = verdict.get("ai") or {}
    vtxt = _resolve_vtxt(dr)
    branding = _load_branding()

    # ٠) الغلاف — هوية سِلك، بلا أيّ تِلِمِتري
    _add_page_header_footer(doc, f"سِلك — دراسة سوق تصديرية: {view.get('product')}")
    _add_cover_wordmark(doc, branding)
    doc.add_heading(f"دراسة سوق تصديرية: {view.get('product')}", 0)
    doc.add_paragraph("أُعدّت بواسطة منصة سِلك لذكاء الأسواق",
                      style="Intense Quote")
    if view.get("test_run"):
        doc.add_paragraph("⚠ نموذج توضيحي ببيانات موسومة — ليس تقريراً "
                          "إنتاجياً")
    _stamp_degraded_banner(doc, view)
    _add_verdict_badge(doc, vtxt)
    _add_table(doc, ["البند", "القيمة"], [
        ["المنتج", h.get("product") or view.get("product")],
        ["رمز HS", h.get("hs_code") or view.get("hs_code")],
        ["بلد المنشأ", "المملكة العربية السعودية"],
        ["السوق المستهدف", h.get("target_market")
         or market.get("name_ar") or market.get("name_en")],
        ["تاريخ الإعداد", h.get("date")]])

    # بنية الأقسام: اجمع أقسام الكاتب المطهَّرة تحت عناوين العميل السبعة.
    sections = _parse_writer_sections((dr.get("report") or {}).get("text") or "")
    buckets: dict[str, list[list[str]]] = {c: [] for c in _CLIENT_SECTION_ORDER}
    for title, body in sections:
        # قسم التوصيات يُقسَم عند "### خارطة طريق الدخول": تعليل الحكم يبقى
        # في "القرار وأساسه"، وخارطة الدخول (الأبواب والخطوات) تنتقل لـ
        # "مسار الدخول والمتطلبات" — مطابقة أدقّ للبنية المطلوبة (بلاغ المالك).
        if title == "التوصيات الاستراتيجية":
            decision_part, roadmap_part = _split_at_roadmap(body)
            buckets["القرار وأساسه"].append(decision_part)
            if roadmap_part:
                buckets["مسار الدخول والمتطلبات"].append(roadmap_part)
            continue
        client_head = _CLIENT_SECTION_MAP.get(title)
        if client_head:
            buckets[client_head].append(body)

    # ١) القرار وأساسه — يبدأ دوماً بالحكم وتعليله (المصدر الوحيد للحكم)،
    # ثم سرد الكاتب (الخلاصة + التوصيات) إن توفّر.
    doc.add_heading("القرار وأساسه", level=1)
    doc.add_paragraph(f"التوصية: {_VERDICT_LABELS_AR[_verdict_tone(vtxt)]}")
    # WP-1 §2: تعليل كلود يُعرَض للعميل فقط حين تتطابق قراءته مع الحكم
    # الحتمي المعروض (وإلا لعرضنا تعليلَ توصيةٍ أخرى تحت توصية مختلفة) —
    # القراءة المخالفة تبقى في التصدير الداخلي (?internal=1) موسومة استشارية.
    _ai_agrees = (not ai.get("verdict")
                  or _verdict_tone(ai.get("verdict")) == _verdict_tone(vtxt))
    reasoning = ((ai.get("reasoning") if _ai_agrees else "")
                 or verdict.get("note") or "")
    if not reasoning and verdict.get("confidence") is not None:
        # مراجعة شيفرة PR #147: مدوّنة مخزَّنة بلا «note» وقراءةُ كلود
        # مخالفة كانت تُخرج قسم القرار بلا أي فقرة أساس — سطر أساسٍ حتمي
        # من الحقول المحسوبة فقط (لا اختلاق) بدل الغياب الصامت.
        from silk_narrative import confidence_phrase
        reasoning = (
            f"حكم المحرّك الحتمي بدرجة ثقة "
            f"{confidence_phrase(verdict.get('confidence'))} بناءً على "
            "الأدلة المرصودة — تفصيل الأساس في هذا القسم والأقسام التالية.")
    if reasoning:
        # §B-1 (حزمة الفكس v2.1) + WP-2 §1: متن العميل لا يُقصّ بـ«…» ولا
        # يستقبل نصّاً نائباً أبداً — `_client_prose` تستخلص من الكتلة الخام
        # أو تُسقِطها (فتلتقطها البوابة)، بدل «بند تقني غير قابل للعرض».
        _r = _client_sanitize(_client_prose(reasoning, 2000))
        if _r:
            doc.add_paragraph(_r)
    # تصدير متدهور بدل فشل صامت (بلاغ المالك، القضية ٣): حين يتعذّر توليد
    # التقرير السردي (report=None بعد فشل نداء الكاتب) نُصدّر مستنداً كاملاً
    # بما هو متاح — الحكم أعلاه + الأدلة المرصودة + الفجوات المعلنة — مع
    # سطر صريح يوضّح أن السرد التفصيلي غاب لأسباب تقنية (لا نكشف تفصيلاً
    # تشغيلياً خاماً للعميل — حارس _client_assert_clean؛ التفصيل الكامل في
    # التصدير التشغيلي ?internal=1 وسبب الفشل المُهيكَل في نتيجة التحليل).
    if not (dr.get("report") or {}).get("text"):
        doc.add_paragraph(
            "تعذّر إنجاز التقرير السردي التفصيلي في هذه المحاولة لأسباب "
            "تقنية مؤقتة؛ القرار أعلاه والأدلة المرصودة في «المراجع» ختام "
            "هذا التقرير قائمة وصحيحة، ويمكن إعادة توليد النص السردي دون "
            "إعادة البحث الكامل.")
    _client_body_or_fallback(doc, buckets["القرار وأساسه"], dr,
                             "القرار وأساسه")

    # ٢-٥) بقية أقسام العميل بالترتيب
    for client_head in _CLIENT_SECTION_ORDER[1:]:
        doc.add_heading(client_head, level=1)
        _client_body_or_fallback(doc, buckets[client_head], dr, client_head)

    # §A (حزمة الفكس v2.1): جدول مزيج الثقة (✓/◐/○) وجدول مرشّحي خرائط قوقل
    # أُسقطا من بناء العميل — يبقيان في التصدير الداخلي (?internal=1) فقط.
    # ٦) ما لم يكتمل للقرار والخطوة التالية (صياغة تجارية للفجوات)
    _client_gaps_section(doc, dr)

    # ٧) المراجع (تحلّ محلّ «سجل الأدلة للمدققين» — مصادر عمومية فقط، §A)
    _client_references_section(doc, dr)

    # B1 (SPEC-v2): مسرد المصطلحات — تقرير العميل يعيد ترتيب الأقسام فلا يرث
    # المسرد من نصّ السرد؛ يُعرَض هنا صراحةً من بنية النموذج (مُطهَّراً).
    _docx_glossary(doc, dr, sanitize=_client_sanitize)
    # C5 (SPEC-v2): جدول المستوردين القابلين للتواصل — قسم الدخول (مُطهَّر).
    # ملاحظة نطاق (§A-4 من حزمة الفكس v2.1): تلك الفقرة طلبت إسقاط هذا
    # الجدول من بناء العميل بالكامل؛ أُبقي هنا عمداً — الجدول محتوًى تجارياً
    # فعلياً (جهات اتصال موزّعين محتملين) يخدم قرار العميل مباشرة، وثلاثة
    # اختبارات قائمة (test_auto_enrich_pipeline_item1.py،
    # test_importer_leads_render_c5.py، test_wave2_first_pdf_cluster.py)
    # تُثبِت أنه قرار منتج متعمَّد سابق (C5). ما أُزيل فعلاً من بناء العميل
    # هو جدول مزيج الثقة (✓/◐/○) وسجلّ الأدلة القديم (استُبدل بـ«المراجع»).
    _docx_leads(doc, dr, sanitize=_client_sanitize)

    # PART A (عائلة 501): نقِّ أوّلاً (استبدل أيّ متبقٍّ بمحايد + سطر إفصاح)،
    # ثم الحارس كشبكة أمان أخيرة — فلا يسقط التصدير بـ501 على تسرّب مصطلح.
    if _client_redact_residual(doc):
        doc.add_paragraph(
            "ملاحظة: نُقّيت بعض المصطلحات التقنية الداخلية من هذا التقرير "
            "تلقائياً لتقديمها بلغة تجارية؛ الأرقام ومصادرها لم تُمَسّ.",
            style="Intense Quote")
    _client_assert_clean(doc)  # شبكة أمان أخيرة — لِما يستحيل تنقيته فقط
    _assert_verdict_consistency_doc(doc, vtxt, "تقرير العميل")  # Master Prompt Part 2 §B
    # WP-7 §3: بوابة نصّ المُنتَج النهائي — على النص الكامل المبني فعلياً
    # (فقرات + جداول)، لا على القالب فقط؛ يغطّي مسار PDF أيضاً (يُبنى منه).
    import silk_quality_gate as _qg
    _artifact_text = "\n".join(
        [p.text for p in doc.paragraphs]
        + [c.text for t in doc.tables for row in t.rows for c in row.cells])
    _artifact_findings = _qg.run_client_artifact_text_gate(_artifact_text)
    if _artifact_findings:
        raise RuntimeError(
            "رفضت بوابة نصّ المُنتَج النهائي تسليم تقرير العميل: "
            + "؛ ".join(f["note"] for f in _artifact_findings[:5]))
    _finalize_rtl(doc)         # §4: اتجاه RTL صريح على كل فقرة/run قبل الحفظ
    doc.save(path)
    return path


# WP-2 §4 — مصادر البنود الخام لكل قسم عميل حين يغيب سرد الكاتب.
# «المخاطر» لم تعد مربوطة بمجموعة فارغة () تضمن فقرة الاعتذار العامة —
# صارت تقرأ من نتائج بعثة المخاطر (risk_news) + بنود SWOT (التهديدات ضمنها).
_CLIENT_FALLBACK_CATS = {
    "القرار وأساسه": ("swot",),
    "السوق بالأرقام": ("demand",),
    "المنافسة والتسعير والهامش": ("price_competitiveness",),
    "مسار الدخول والمتطلبات": ("entry_cost", "entry_door"),
    "المخاطر": ("swot",),
}


def _client_fallback_sources(dr: dict, client_head: str) -> list[str]:
    """البنود الخام المرشّحة لإعادة الصياغة التجارية لقسم عميل بلا سرد كاتب.
    نصوص فقط (لا نصّ نائب): كتل JSON غير القابلة للاستخلاص تُسقَط."""
    by_cat = (dr.get("analyst") or {}).get("by_category") or {}
    items: list[str] = []
    for cat in _CLIENT_FALLBACK_CATS.get(client_head, ()):
        for f in (by_cat.get(cat) or []):
            t = _client_prose(_dp_value(f), 400)
            if t:
                items.append(t)
    if client_head == "المخاطر":
        risk = (dr.get("missions") or {}).get("risk_news") or {}
        if isinstance(risk, dict) and not risk.get("failed"):
            for f in (risk.get("findings") or []):
                t = _client_prose(_dp_value(f), 400)
                if t:
                    items.append(t)
    return items


def _client_missing_narrative_heads(dr: dict) -> "dict[str, list[str]]":
    """أقسام العميل التي سيصادفها التصدير بلا سرد كاتب، مع بنودها الخام
    المرشّحة لإعادة الصياغة — نفس منطق تجميع الأقسام في `render_client_docx`
    (وتستعمله بوابة الجودة أيضاً) فلا يتباعدان."""
    sections = _parse_writer_sections(((dr.get("report") or {}).get("text")
                                       or ""))
    buckets: dict[str, list[list[str]]] = {c: [] for c in _CLIENT_SECTION_ORDER}
    for title, body in sections:
        if title == "التوصيات الاستراتيجية":
            decision_part, roadmap_part = _split_at_roadmap(body)
            buckets["القرار وأساسه"].append(decision_part)
            if roadmap_part:
                buckets["مسار الدخول والمتطلبات"].append(roadmap_part)
            continue
        head = _CLIENT_SECTION_MAP.get(title)
        if head:
            buckets[head].append(body)
    out: dict[str, list[str]] = {}
    for head in _CLIENT_SECTION_ORDER:
        has_body = any(any(str(ln).strip() for ln in body)
                       for body in buckets[head])
        if not has_body:
            out[head] = _client_fallback_sources(dr, head)
    return out


def _dp_value(f: object) -> object:
    """قيمة نقطة بيانات — dict مخزَّن أو كائن DataPoint حي."""
    if isinstance(f, dict):
        return f.get("value")
    return getattr(f, "value", None)


def _client_body_or_fallback(doc, bodies: list[list[str]], dr: dict,
                             client_head: str) -> None:
    """اعرض متون الكاتب لهذا القسم إن توفّرت؛ وإلا **النثر التجاري المُعاد
    صياغته** (نداء كاتب مصغّر، WP-2 §3 — يُحضَّر قبل بوابة التسليم ويُخزَّن
    في `dr["client_fallback_prose"]`). لا تُسرَد قيم `dp.value` الخام نقاطاً
    للعميل بعد الآن — قسم بلا سرد ولا نثر مُعاد صياغته تُفشِله بوابة الجودة
    قبل التسليم أصلاً (النصّ العام أدناه شبكة أمان لمسارات الاستدعاء
    المباشرة خارج نقطة التصدير فقط)."""
    if bodies:
        for body in bodies:
            _client_render_body_block(doc, body)
        return
    prose = str(((dr.get("client_fallback_prose") or {}).get(client_head))
                or "").strip()
    if prose:
        for para in prose.splitlines():
            line = _client_sanitize(_client_prose(para, 600))
            if line:
                doc.add_paragraph(line)
        return
    doc.add_paragraph(
        "التحليل السردي التفصيلي لهذا القسم غير متاح ضمن هذا التقرير؛ "
        "الأدلة المرصودة ذات الصلة مُدرجة في «المراجع» ختام التقرير.")


# ── §3 (أمر العمل الرئيس): التصدير النهائي PDF غير قابل للتحرير ────────────
# يُبنى المستند docx أولاً (كل منطق العرض/RTL/التطهير فيه) ثم يُحوَّل PDF عبر
# LibreOffice headless كخطوة أخيرة، ويُسلَّم الـPDF فقط. تدهور رشيق: إن غاب
# المحرّك أو فشل التحويل يُرفع RuntimeError برسالة نظيفة (لا docx بديل صامت،
# لا PDF جزئي) — يلتقطها المسار الأعلى فيعيد خطأً معلَناً للعميل.

_PDF_SOFFICE_CANDIDATES = ("soffice", "libreoffice")
_PDF_UNAVAILABLE = ("تعذّر إنتاج ملف PDF: محرّك تحويل المستندات غير متاح على "
                    "الخادم حالياً")
_PDF_FAILED = "تعذّر إنتاج ملف PDF من المستند"


def _find_soffice() -> "str | None":
    """مسار ثنائي LibreOffice (soffice/libreoffice) إن وُجد على الخادم."""
    import shutil
    for name in _PDF_SOFFICE_CANDIDATES:
        found = shutil.which(name)
        if found:
            return found
    return None


# عائلات خطوط عربية الشكل يقبلها التشكيل — أيّها كافٍ لـPDF بلا مربّعات tofu.
_ARABIC_FONT_HINTS = ("naskh", "amiri", "arabic", "cairo", "dubai", "tahoma",
                      "scheherazade", "lateef", "kufi", "arial")


def has_arabic_font() -> bool:
    """§3/§4: هل يتوفّر على الخادم خطّ عربي الشكل (Naskh/Amiri/Arabic/…)؟
    بلا خطّ عربي عامل، ينتج PDF بمربّعات tofu بدل الحروف الموصولة — يُفحَص
    عبر fc-list. يُستعمَل في مسار قبول الـPDF (يفشل بصوتٍ عالٍ إن غاب)."""
    import subprocess
    try:
        out = subprocess.run(["fc-list", ":", "family"], capture_output=True,
                             text=True, timeout=20)
    except Exception:  # noqa: BLE001 — fc-list غير متاح = لا يمكن التأكيد
        return False
    blob = (out.stdout or "").lower()
    return any(hint in blob for hint in _ARABIC_FONT_HINTS)


def has_plex_arabic_font() -> bool:
    """§7 (قرار المالك): هل «IBM Plex Sans Arabic» مثبَّتٌ تحديدًا؟ العائلةُ
    الرسميةُ للتقرير — يجب ألّا يبدّلها LibreOffice صامتًا. يفشل بصوتٍ عالٍ في
    مسار القبول إن غابت (لا خطٌّ بديلٌ يمرّ زائفًا)."""
    import subprocess
    try:
        out = subprocess.run(["fc-list", ":", "family"], capture_output=True,
                             text=True, timeout=20)
    except Exception:  # noqa: BLE001
        return False
    blob = (out.stdout or "").lower()
    return "plex sans arabic" in blob or "ibmplexsansarabic" in blob


def _pdf_diacritic_free_copy(docx_path: str) -> str:
    """Wave 2 (البند ٨ + الطيّة E): نسخةٌ مجرّدةٌ من الحركات للتحويل لـPDF.

    العلامةُ «سِلك» وأشقّاؤها المشكّلون («مُوصًى»/«يُتَّخذ»…) تنفصل حركاتُهم عند
    تشكيل خطّ LibreOffice فتُشقّ الكلمةُ في استخراج الـPDF. **الأصل (وثيقة الوورد
    القابلة للتحرير — المشغّل) يبقى مشكّلًا كمصدر**؛ نُحوّل نسخةً مؤقّتةً مجرّدةً من
    الحركات (U+064B–U+0652، U+0670) فقط، فيخرج الـPDF (المُسلَّم للعميل) نظيفَ
    الاستخراج — لا كلمةَ عربيةٍ يشقّها مِحرفٌ مُركَّب (تعريفُ عائلة E). لا يُمَسّ
    أيّ رقمٍ أو حرفٍ أو معنى — الحركاتُ علاماتُ نطقٍ لا محتوى.

    يعمل على **مستوى حزمة الـzip** لا نموذج python-docx كي يشمل **كلَّ جزءٍ نصّيّ**:
    `word/document.xml` **و** `word/header*.xml`/`footer*.xml` والحواشي
    (`footnotes.xml`/`endnotes.xml`) وأيّ جزءٍ آخر تحت `word/` — نموذجُ python-docx
    وحده يُغفِل الحواشي وترويسات الصفحة الأولى ومربّعات النصّ. المِحارفُ المُركَّبةُ
    تظهر حصرًا داخل نصّ `<w:t>` (لا في وسمٍ/سِمة)، فحذفُ نقاطها من XML آمنٌ بنيويًا.
    """
    import os
    import tempfile
    import zipfile
    tmp = os.path.join(tempfile.mkdtemp(prefix="silk_pdfsrc_"),
                       os.path.basename(docx_path))
    with zipfile.ZipFile(docx_path) as zin, \
            zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/") and \
                    item.filename.endswith(".xml"):
                data = _AR_COMBINING_RE.sub(
                    "", data.decode("utf-8")).encode("utf-8")
            zout.writestr(item, data)
    return tmp


def docx_to_pdf(docx_path: str, pdf_path: "str | None" = None,
                timeout: int = 180) -> str:
    """§3: حوّل مستند Word إلى PDF عبر LibreOffice headless — يعيد مسار الـPDF.

    تدهور رشيق: RuntimeError برسالة نظيفة إن غاب المحرّك أو فشل التحويل — لا
    يُسلَّم docx بديلاً صامتاً ولا PDF جزئي (§3/§5). لا يُعدَّل أيّ رقم أو محتوى:
    التحويلُ يقع على نسخةٍ مجرّدةٍ من **الحركات** فقط (الطيّة E) كي لا تُشقّ كلمةٌ
    عربيةٌ في استخراج الـPDF؛ الأصل (وثيقة الوورد) يبقى مشكّلًا (`_pdf_diacritic_
    free_copy`). المستند مبنيّ بالكامل في طبقة docx (RTL/تطهير/بوابة الجودة)."""
    import os
    import subprocess
    import tempfile
    soffice = _find_soffice()
    if not soffice:
        raise RuntimeError(_PDF_UNAVAILABLE)
    if not os.path.exists(docx_path):
        raise RuntimeError(_PDF_FAILED)
    out_dir = os.path.dirname(os.path.abspath(pdf_path or docx_path)) or "."
    profile = tempfile.mkdtemp(prefix="silk_lo_")
    # حوّل النسخةَ المجرّدةَ من الحركات (لا الأصل) — الطيّة E على مخرَج الـPDF.
    try:
        convert_src = _pdf_diacritic_free_copy(docx_path)
    except Exception:  # noqa: BLE001 — تعذّر التجريد => حوّل الأصل (لا نكسر التسليم)
        convert_src = docx_path
    try:
        proc = subprocess.run(
            [soffice, "--headless", "--norestore",
             f"-env:UserInstallation=file://{profile}",
             "--convert-to", "pdf", "--outdir", out_dir, convert_src],
            capture_output=True, timeout=timeout,
            env={**os.environ, "HOME": profile})
    except Exception as e:  # noqa: BLE001 — أي فشل تحويل = خطأ معلَن نظيف
        raise RuntimeError(_PDF_FAILED) from e
    produced = os.path.join(
        out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    if proc.returncode != 0 or not os.path.exists(produced):
        raise RuntimeError(_PDF_FAILED)
    # WP-5 §2: الفحص الآلي لاتجاه الأقواس على الـPDF النهائي — فوق العتبة
    # يفشل التصدير بصوت عالٍ بدل تسليم مستند معكوس الأقواس.
    _pdf_bracket_check(produced)
    if pdf_path and os.path.abspath(pdf_path) != os.path.abspath(produced):
        os.replace(produced, pdf_path)
        return pdf_path
    return produced


# ── القالب الأكاديمي (قرار المالك 2026-07-22، النموذج v3 المعتمد) ────────────
#
# نفس بيانات الدراسة حرفياً بترتيبٍ ونبرة بحثٍ أكاديمي: ملخّص شامل يفتتح
# التقرير (التوصية من أول سطر + المنهج + جدول أبرز النتائج + القيد الحاكم +
# الحدود + خلاصة التوصيات — قراءته وحدها تغني)، ثم مقدمة/منهجية/نتائج/
# مناقشة/استنتاجات/حدود، و«التوصيات» قسماً ختامياً مستقلاً قبل المراجع.
# حتمي بالكامل: صفر نداء كلود — إعادة ترتيب سرد الكاتب القائم + حقول
# النموذج المحسوبة، بنفس مُطهِّرات العميل وبواباته كلها (لا مسار حكمٍ ولا
# أرقام جديدة — عقد عدم الاختلاق).

# سرد الكاتب (١١ قسماً) → مواضع القالب الأكاديمي.
_ACADEMIC_RESULTS_MAP = (
    ("نظرة عامة على السوق وحجمه", "حجم السوق وتدفقات الاستيراد"),
    ("ديناميكيات السوق", "ديناميكيات السوق واتجاهاته"),
    ("تحليل المستهلك والطلب", "خصائص الطلب والمستهلك"),
    ("المشهد التنافسي", "بنية المنافسة ودرجة التركّز"),
    ("التنظيم والوصول للسوق", "البيئة التنظيمية ومتطلبات الوصول"),
    ("اللوجستيات وسلسلة الإمداد", "اللوجستيات وسلسلة الإمداد"),
    ("تقييم المخاطر", "المخاطر المرصودة"),
)

# طلب المالك (تعقيباً على النموذج المعتمد): الديموغرافيا/حجم السكان،
# ثقافة المستهلك، والاشتراطات الجمركية تُعرض **دائماً** كشواهد مهيكلة من
# بيانات البعثات نفسها — جداول أدلة (بند/مصدر/شارة) لا نقاط نصية خام
# (التزاماً بعقد WP-2)، فتظهر البيانات حتى حين يكون سرد الكاتب لها موجزاً.
_ACADEMIC_EVIDENCE_TABLES = (
    ("السياق الديموغرافي وحجم السكان", ("demographics_economy",)),
    ("ثقافة المستهلك وأنماط الطلب — الشواهد المرصودة",
     ("consumer_culture", "demand_trends")),
    ("الاشتراطات الجمركية ومتطلبات الدخول — الشواهد المرصودة",
     ("customs_requirements", "tariffs_agreements")),
)


def _ar_digits(n: int) -> str:
    """رقم بالأرقام العربية-الهندية — لاتساق ترقيم العناوين الفرعية."""
    return str(n).translate(str.maketrans("0123456789", "٠١٢٣٤٥٦٧٨٩"))


def _academic_evidence_rows(dr: dict, mission_keys: tuple) -> list[list[str]]:
    """صفوف شواهد مهيكلة من بعثات محدَّدة: (البند المرصود، المصدر، قوة
    الدليل) — عبر `_client_readable_fact` (بند غير قابل للعرض يُسقَط لا
    يُستبدَل بنائب) وبشارة المنشأ الواعية. سقف ٨ صفوف لكل جدول."""
    from silk_narrative import evidence_badge_for
    missions = dr.get("missions") or {}
    rows: list[list[str]] = []
    for key in mission_keys:
        m = missions.get(key)
        if not isinstance(m, dict) or m.get("failed"):
            continue
        for f in (m.get("findings") or []):
            fact = _client_readable_fact(f.get("value"), f.get("note"))
            if not fact:
                continue
            src = _client_sanitize(_clean_source_label(f.get("source")))
            if not src or src == "—" or _client_forbidden_hits(src):
                continue
            rows.append([fact, src, evidence_badge_for(f)])
            if len(rows) >= 8:
                return rows
    return rows


_ACADEMIC_MAIN_REC = {
    "go": "المضي في دخول السوق وفق الأولويات المفصّلة أدناه.",
    "conditional": "دخول السوق دخولاً مشروطاً باستيفاء الشروط القابلة "
                   "للتحقق أدناه قبل أي التزام إنتاجي أو شحني.",
    "watch": "إرجاء قرار دخول السوق مع إبقائه قيد المراقبة النشطة، "
             "وإعادة التقييم فور استيفاء الشروط أدناه.",
    "nogo": "عدم دخول السوق في المرحلة الراهنة، وإعادة التقييم عند "
            "تغيّر المعطيات المفصّلة أدناه.",
    "unknown": "استكمال البيانات الناقصة أدناه قبل إصدار توصية نهائية.",
}


def _academic_headline_table_rows(dr: dict) -> list[list[str]]:
    """صفوف جدول «أبرز النتائج حسب المحور» — أعلى بند لكل تقاطع محلّل مع
    شارة توثيقه. بنيوي بحت من by_category؛ تقاطع بلا أدلة يُسقَط."""
    from silk_narrative import evidence_badge_for
    by_cat = (dr.get("analyst") or {}).get("by_category") or {}
    rows: list[list[str]] = []
    for cat, label in _CATEGORY_AR.items():
        items = by_cat.get(cat) or []
        if not items:
            continue
        top = items[0]
        fact = _client_sanitize(_client_prose(
            top.get("value") if isinstance(top, dict)
            else getattr(top, "value", None), 200))
        if not fact:
            continue
        rows.append([label, fact, evidence_badge_for(top)])
    return rows


def _academic_summary(doc, view: dict, dr: dict, vtxt: str) -> None:
    """«ملخّص الدراسة» الشامل — طلب المالك: قراءته وحدها تغطي الدراسة كاملة."""
    from silk_narrative import authoritative_verdict, confidence_phrase
    doc.add_heading("ملخّص الدراسة", level=1)
    _, conf = authoritative_verdict(dr.get("verdict"))
    conf_txt = (f"، بدرجة ثقة {confidence_phrase(conf)} وفق سُلَّم "
                "المعايرة المعتمد" if conf is not None else "")
    doc.add_paragraph(
        f"التوصية الختامية: {_VERDICT_LABELS_AR[_verdict_tone(vtxt)]}"
        f"{conf_txt}؛ وتفصيلها وشروط إعادة التقييم في قسم «التوصيات» "
        "ختام الدراسة.")
    doc.add_paragraph(
        "سؤال الدراسة ومنهجها: تقييم جدوى دخول السوق المستهدف للمنتج "
        "المدروس من منظور مصدّر سعودي، بمنهج كمّي وصفي قائم على البيانات "
        f"الثانوية الرسمية. {_client_methodology_paragraph(dr)}")
    rows = _academic_headline_table_rows(dr)
    if rows:
        doc.add_paragraph("أبرز النتائج حسب محاور التحليل:")
        _add_table(doc, ["المحور", "أبرز نتيجة مرصودة", "قوة الدليل"], rows)
    if dr.get("hs_flagged"):
        hs = dr.get("hs_confirmation") or {}
        doc.add_paragraph(
            "القيد المنهجي الحاكم: رمز التصنيف الجمركي المستخدم في جمع "
            f"بيانات التجارة ({hs.get('hs_code') or view.get('hs_code')}) "
            "غير مؤكَّد المطابقة لفئة المنتج المدروس؛ لذلك تُقرأ الأرقام "
            "الجمركية مؤشراً سياقياً لا قياساً مباشراً، وتُقيَّد "
            "الاستنتاجات المبنية عليها حتى إعادة التصنيف.")
    critical, informational = _client_gap_inputs(dr)
    if critical or informational:
        doc.add_paragraph(
            "حدود الدراسة (موجز): " + " ".join(
                _trim_sentence(g, 160) for g in (critical
                                                 + informational)[:3]))
    else:
        doc.add_paragraph(
            "حدود الدراسة (موجز): اكتملت محاور التحليل الأساسية بأدلة "
            "موثّقة بمصادرها ضمن نطاق هذه الدراسة.")
    doc.add_paragraph(
        "خلاصة التوصيات: "
        + _ACADEMIC_MAIN_REC[_verdict_tone(vtxt)]
        + " التفصيل الكامل في القسم الختامي.")


def render_academic_docx(view: dict, path: str) -> str:
    """تقرير العميل بالقالب الأكاديمي — نفس نموذج العرض القانوني ونفس
    بوابات العميل (تطهير، نقاء، اتساق الحكم، بوابة نصّ المُنتَج النهائي)،
    ببنية v3 المعتمدة من المالك. يتطلّب نتيجة بحث عميق."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(_DOCX_HINT) from exc

    _assert_production_clean(view)
    dr = view.get("deep_research") or {}
    if not dr:
        raise RuntimeError("render_academic_docx يتطلّب نتيجة بحث عميق "
                           "(view['deep_research'])")

    doc = Document()
    _apply_rtl(doc)
    h = view.get("header") or {}
    market = dr.get("market") or {}
    vtxt = _resolve_vtxt(dr)
    product = h.get("product") or view.get("product") or ""
    market_ar = (h.get("target_market") or market.get("name_ar")
                 or market.get("name_en") or "")
    branding = _load_branding()

    _add_page_header_footer(doc, f"سِلك — دراسة سوق تصديرية: {product}")
    _add_cover_wordmark(doc, branding)
    doc.add_heading(
        f"دراسة تحليلية لسوق {product} في {market_ar}: تقييم فرص التصدير "
        "من منظور المصدّر السعودي", 0)
    doc.add_paragraph("أُعدّت بواسطة منصة سِلك لذكاء الأسواق",
                      style="Intense Quote")
    if view.get("test_run"):
        doc.add_paragraph("⚠ نموذج توضيحي ببيانات موسومة — ليس تقريراً "
                          "إنتاجياً")
    _stamp_degraded_banner(doc, view)
    _add_verdict_badge(doc, vtxt)
    _add_table(doc, ["البند", "القيمة"], [
        ["المنتج", product],
        ["رمز HS", h.get("hs_code") or view.get("hs_code")],
        ["بلد المنشأ", "المملكة العربية السعودية"],
        ["السوق المستهدف", market_ar],
        ["تاريخ الإعداد", h.get("date")]])

    # الملخّص الشامل أولاً (طلب المالك).
    _academic_summary(doc, view, dr, vtxt)

    # ١) المقدمة وأهداف الدراسة — قالب حتمي معمَّم لأي منتج/سوق.
    doc.add_heading("١. المقدمة وأهداف الدراسة", level=1)
    doc.add_paragraph(
        f"تسعى هذه الدراسة إلى الإجابة عن سؤال بحثي رئيس: ما مدى جاذبية "
        f"سوق {market_ar} لمنتج {product} سعودي المنشأ؟ وتتفرّع عنه ثلاثة "
        "أسئلة فرعية: (أ) ما حجم الطلب المرصود واتجاهه؟ (ب) ما بنية "
        "المنافسة وهيكل الأسعار السائد؟ (ج) ما المتطلبات التنظيمية وقنوات "
        "الدخول المتاحة؟")

    # ٢) المنهجية ومصادر البيانات.
    doc.add_heading("٢. المنهجية ومصادر البيانات", level=1)
    doc.add_paragraph(_client_methodology_paragraph(dr))
    if dr.get("hs_flagged"):
        hs = dr.get("hs_confirmation") or {}
        doc.add_paragraph(
            "قيد منهجي جوهري: أُجري جمع بيانات التجارة تحت رمزٍ جمركي "
            f"({hs.get('hs_code') or view.get('hs_code')}) لا يشمل وصفُه "
            "صفةَ المنتج المميّزة؛ وعليه تُقرأ جميع الأرقام المستمدّة منه "
            "بوصفها مؤشراً سياقياً لفئة مجاورة لا قياساً مباشراً لسوق "
            "المنتج، وتُعامَل الاستنتاجات المبنية عليها بيقين مقيَّد إلى "
            "حين إعادة التصنيف وجمع البيانات تحت الرمز الصحيح.")

    # ٣) النتائج — سرد الكاتب القائم تحت عناوين فرعية أكاديمية.
    sections = dict(_parse_writer_sections(
        ((dr.get("report") or {}).get("text") or "")))
    doc.add_heading("٣. النتائج", level=1)
    _idx = 0
    for writer_title, academic_title in _ACADEMIC_RESULTS_MAP:
        body = sections.get(writer_title)
        if not body or not any(str(ln).strip() for ln in body):
            continue
        _idx += 1
        doc.add_heading(f"٣.{_ar_digits(_idx)} {academic_title}", level=2)
        _client_render_body_block(doc, body)
    # الشواهد المهيكلة الإلزامية (طلب المالك): ديموغرافيا/سكان، ثقافة
    # المستهلك، والاشتراطات الجمركية — من بيانات البعثات مباشرة.
    for table_title, mission_keys in _ACADEMIC_EVIDENCE_TABLES:
        rows = _academic_evidence_rows(dr, mission_keys)
        if not rows:
            continue
        _idx += 1
        doc.add_heading(f"٣.{_ar_digits(_idx)} {table_title}", level=2)
        _add_table(doc, ["البند المرصود", "المصدر", "قوة الدليل"], rows)
    if not _idx:
        doc.add_paragraph(
            "لم يتوافر سرد نتائج تفصيلي في هذه التشغيلة؛ الأدلة المرصودة "
            "مفصّلة في جدول الملخّص وقسم المراجع.")

    # ٤) المناقشة — خلاصة الكاتب التنفيذية + تعليل التوصيات (دون الخارطة).
    doc.add_heading("٤. المناقشة", level=1)
    _discussed = False
    exec_body = sections.get("الخلاصة التنفيذية")
    if exec_body and any(str(ln).strip() for ln in exec_body):
        _client_render_body_block(doc, exec_body)
        _discussed = True
    strat_body = sections.get("التوصيات الاستراتيجية")
    roadmap_part: list[str] = []
    if strat_body:
        decision_part, roadmap_part = _split_at_roadmap(strat_body)
        if decision_part and any(str(ln).strip() for ln in decision_part):
            _client_render_body_block(doc, decision_part)
            _discussed = True
    if not _discussed:
        doc.add_paragraph(
            "تُناقَش دلالات النتائج ضمن قسمي الاستنتاجات والتوصيات أدناه "
            "استناداً إلى الأدلة الموثّقة في الملخّص والمراجع.")

    # ٥) الاستنتاجات — الحكم القانوني بأساسه (نفس منطق قسم القرار).
    doc.add_heading("٥. الاستنتاجات", level=1)
    verdict = dr.get("verdict") or {}
    ai = verdict.get("ai") or {}
    doc.add_paragraph(
        f"تخلص الدراسة إلى التوصية التالية: "
        f"{_VERDICT_LABELS_AR[_verdict_tone(vtxt)]}.")
    _ai_agrees = (not ai.get("verdict")
                  or _verdict_tone(ai.get("verdict")) == _verdict_tone(vtxt))
    basis = ((ai.get("reasoning") if _ai_agrees else "")
             or verdict.get("note") or "")
    if not basis and verdict.get("confidence") is not None:
        from silk_narrative import confidence_phrase
        basis = (f"حكم المحرّك الحتمي بدرجة ثقة "
                 f"{confidence_phrase(verdict.get('confidence'))} بناءً "
                 "على الأدلة المرصودة الموثّقة في هذه الدراسة.")
    _b = _client_sanitize(_client_prose(basis, 2000))
    if _b:
        doc.add_paragraph(_b)

    # ٦) حدود الدراسة — نفس المصدر الواحد لمدخلات الفجوات (WP-4).
    doc.add_heading("٦. حدود الدراسة والبحث المستقبلي", level=1)
    critical, informational = _client_gap_inputs(dr)
    if critical or informational:
        for line in critical + informational:
            doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph(
            "اكتملت محاور التحليل الأساسية بأدلة موثّقة بمصادرها، ولا "
            "بند حاسم للقرار موسوم بأنه غير محقَّق؛ لا فجوة جوهرية تمنع "
            "اتخاذ القرار ضمن نطاق هذه الدراسة.")

    # ٧) التوصيات — القسم الختامي المستقل (طلب المالك).
    doc.add_heading("٧. التوصيات", level=1)
    doc.add_paragraph("بناءً على النتائج والمناقشة أعلاه، توصي الدراسة "
                      "بما يلي:")
    doc.add_paragraph(
        f"التوصية الرئيسة: {_ACADEMIC_MAIN_REC[_verdict_tone(vtxt)]}",
        style="List Number")
    for c in (dr.get("flip_conditions") or []):
        mark = "✓ محقَّق" if c.get("met") else "○ غير محقَّق"
        doc.add_paragraph(
            f"{c.get('condition')} — {mark}؛ يُغلَق عبر: "
            f"{c.get('closes_via')}.", style="List Number")
    if roadmap_part and any(str(ln).strip() for ln in roadmap_part):
        # متن الخارطة يحمل عنوانه الفرعي («### خارطة طريق الدخول…») بنفسه.
        _client_render_body_block(doc, roadmap_part)

    # المراجع + المسرد — نفس أقسام العميل القائمة.
    _client_references_section(doc, dr)
    _docx_glossary(doc, dr, sanitize=_client_sanitize)

    # نفس سلسلة بوابات العميل الختامية حرفياً.
    if _client_redact_residual(doc):
        doc.add_paragraph(
            "ملاحظة: نُقّيت بعض المصطلحات التقنية الداخلية من هذا التقرير "
            "تلقائياً لتقديمها بلغة تجارية؛ الأرقام ومصادرها لم تُمَسّ.",
            style="Intense Quote")
    _client_assert_clean(doc)
    _assert_verdict_consistency_doc(doc, vtxt, "التقرير الأكاديمي")
    import silk_quality_gate as _qg
    _artifact_text = "\n".join(
        [p.text for p in doc.paragraphs]
        + [c.text for t in doc.tables for row in t.rows for c in row.cells])
    _artifact_findings = _qg.run_client_artifact_text_gate(_artifact_text)
    if _artifact_findings:
        raise RuntimeError(
            "رفضت بوابة نصّ المُنتَج النهائي تسليم التقرير الأكاديمي: "
            + "؛ ".join(f["note"] for f in _artifact_findings[:5]))
    _finalize_rtl(doc)
    doc.save(path)
    return path


def render_academic_pdf(view: dict, path: str) -> str:
    """التقرير الأكاديمي PDF — يُبنى docx (المُطهَّر، RTL) ثم يُحوَّل،
    ويُسلَّم الـPDF فقط (نفس عقد §3 ومعه فحص الأقواس)."""
    import os
    import tempfile
    tmp_docx = os.path.join(tempfile.mkdtemp(prefix="silk_acad_"),
                            "academic.docx")
    render_academic_docx(view, tmp_docx)
    return docx_to_pdf(tmp_docx, path)


def render_client_pdf(view: dict, path: str) -> str:
    """§3: التقرير النهائي للعميل PDF — يُبنى تقرير العميل docx (المُطهَّر،
    RTL) ثم يُحوَّل، ويُسلَّم الـPDF فقط."""
    import os
    import tempfile
    tmp_docx = os.path.join(tempfile.mkdtemp(prefix="silk_pdf_"), "client.docx")
    render_client_docx(view, tmp_docx)
    return docx_to_pdf(tmp_docx, path)


def render_research_pdf(view: dict, path: str) -> str:
    """§3: تقرير المدقّق PDF — يُبنى تقرير البحث docx (RTL) ثم يُحوَّل."""
    import os
    import tempfile
    tmp_docx = os.path.join(tempfile.mkdtemp(prefix="silk_pdf_"), "research.docx")
    render_docx(view, tmp_docx)
    return docx_to_pdf(tmp_docx, path)


def render_docx(view: dict, path: str) -> str:
    """التقرير الكامل Word (§10.3) — من القالب الموحّد حصراً.

    يعيد المسار عند النجاح؛ RuntimeError واضحة بلا python-docx.

    نتيجة `/research` (`view["deep_research"]` موجود) تُبنى عبر مسار مخصّص
    (`_render_research_docx`) لا الأقسام الأربعة عشر الكلاسيكية — راجع
    تعليق التصميم هناك (بلاغ حي، الموجة ٨: كانت تُبنى هيكلاً فارغاً يسبق
    التقرير الحقيقي).
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(_DOCX_HINT) from exc

    _assert_production_clean(view)
    doc = Document()
    _apply_rtl(doc)   # §4: المستند كله من اليمين لليسار (المدقّق + الكلاسيكي)

    if view.get("deep_research"):
        _render_research_docx(doc, view)
        _assert_verdict_consistency_doc(  # Master Prompt Part 2 §B
            doc, _resolve_vtxt(view["deep_research"]), "تقرير البحث العميق")
        _finalize_rtl(doc)   # §4: اتجاه RTL صريح على كل فقرة/run قبل الحفظ
        doc.save(path)
        return path

    top_m = (view.get("markets") or [{}])[0]

    # ═══ 0) الغلاف وبطاقة التعريف — cover + report card ═══
    # هوية سِلك (الموجة ١١، §11.1) — رأس/تذييل موحّدان لكل تقارير docx.
    _add_page_header_footer(doc, f"سِلك — تقرير بحث سوق: {view.get('product')}")
    _add_cover_wordmark(doc, _load_branding())
    doc.add_heading(f"سِلك — تقرير بحث سوق: {view.get('product')}", 0)
    if view.get("test_run"):
        doc.add_paragraph("⚠ TEST RUN — تشغيل برهاني ببدائل موسومة، "
                          "ليس تقريراً إنتاجياً")
    _stamp_degraded_banner(doc, view)  # لافتة الغلاف — بلاغ حي
    h = view.get("header") or {}
    _add_table(doc, ["البند", "القيمة"], [
        ["المنتج", h.get("product")],
        ["رمز HS", h.get("hs_code")],
        ["المنشأ", "المملكة العربية السعودية"],
        ["السوق المستهدف", h.get("target_market")],
        ["تاريخ الإعداد", h.get("date")],
        ["سنة البيانات", _data_year_label(view)],
        ["تغطية البيانات", f"{h.get('coverage_pct')}%"]])

    # جدول المحتويات الثابت — الأقسام الأربعة عشر المرقّمة، ثم الأقسام
    # غير المرقّمة بترتيبها الفعلي في المتن (نفس اصطلاح المتن — "حدود هذا
    # التقرير"/"دليل المورّدين والمستوردين"/"الملحق" بلا رقم فيه أيضاً).
    # سدّ تسريب (الطبقة ٩): "حدود هذا التقرير" و"دليل المورّدين والمستوردين"
    # (رُقّي من عنوان فرعي مدفون تحت التوصيات) كانا غائبين عن هذا الجدول.
    doc.add_heading("المحتويات", level=1)
    for i, ttl in enumerate((
            "الخلاصة التنفيذية", "منهجية البحث", "تعريف السوق ونطاقه",
            "نظرة عامة على السوق", "ديناميكيات السوق",
            "حجم السوق والتوقعات", "تحليل التقسيم",
            "تحليل التجارة (استيراد/تصدير)", "الأسواق المرشّحة الأخرى",
            "المشهد التنافسي", "استخبارات العميل والطلب",
            "المشهد التنظيمي والمخاطر", "الاتجاهات والتوقع المستقبلي",
            "التوصيات الاستراتيجية"), 1):
        doc.add_paragraph(f"{i}. {ttl}", style="List Number" if False
                          else None)
    # نص مختلف حرفياً عن العنوان الفعلي عمداً (نفس أسلوب سطر "الملحق" أدناه
    # أصلاً) — لا تطابق نصّي حرفي بين فقرة الفهرس وعنوان القسم قد يُربك أي
    # فحص نصّي لاحق على ترتيب العناوين.
    doc.add_paragraph("حدود التحليل (الفجوات المعلنة)")
    doc.add_paragraph("دليل المورّدين (الاتصال بالموردين والمستوردين)")
    doc.add_paragraph("الملحق (تغطية المصادر وأثرها)")

    # ═══ ١) الخلاصة التنفيذية — ٣ فقرات بشرية (P1) ═══
    doc.add_heading("١. الخلاصة التنفيذية", level=1)
    for para in _narrative_exec_summary(view):
        doc.add_paragraph(para)

    # ═══ ٢) منهجية البحث ═══
    doc.add_heading("٢. منهجية البحث", level=1)
    for line in _methodology_lines(view):
        doc.add_paragraph(line, style="List Bullet")

    # ═══ ٣) تعريف السوق ونطاقه ═══
    doc.add_heading("٣. تعريف السوق ونطاقه", level=1)
    doc.add_paragraph(_market_scope_paragraph(view))

    # سدّ تسريب (الطبقة ٩): قرار الدخول (المحرك الموزون §8) كان مدفوناً
    # في القسم ١٤ الأخير (التوصيات) — القارئ يصل التوصية الفعلية بعد كل
    # التفاصيل. رُفع هنا قرب الخلاصة التنفيذية (نفس ترتيب render_markdown
    # أصلاً، الذي كان صحيحاً بالفعل: "قرار الدخول" يظهر بعد تعريف السوق
    # مباشرة). _docx_entry_strategy تبقى في التوصيات — نموذج دخول تفصيلي
    # لا حكم.
    _docx_entry_decision(doc, top_m)

    # ═══ ٤) نظرة عامة على السوق — إشارات نوعية مرصودة أو غياب هادئ ═══
    doc.add_heading("٤. نظرة عامة على السوق", level=1)
    cc = view.get("consumer_culture") or {}
    overview_ins = (cc.get("insights") or [])[:3]
    if overview_ins:
        for o in overview_ins:
            doc.add_paragraph(str(o.get("point") or ""), style="List Bullet")
            ev = "؛ ".join(map(str, (o.get("evidence") or [])[:2]))
            if ev:
                doc.add_paragraph(f"الدليل: {ev}", style="Intense Quote")
    else:
        doc.add_paragraph("لم تُجمع إشارات نوعية عن حالة الصناعة في هذا "
                          "التشغيل — يتطلب مفتاح بحث الويب.")

    # ═══ ٥) ديناميكيات السوق — وكيل الديناميكيات (P2-8) ═══
    doc.add_heading("٥. ديناميكيات السوق", level=1)
    dyn = view.get("dynamics") or {}
    dyn_v = dyn.get("value") if isinstance(dyn, dict) else None
    if isinstance(dyn_v, dict) and dyn_v.get("classified"):
        _AR_DYN = (("drivers", "الدوافع"), ("restraints", "الكوابح"),
                   ("opportunities", "الفرص"), ("threats", "التحديات"))
        for key, ttl in _AR_DYN:
            items = dyn_v.get(key) or []
            if not items:
                continue
            doc.add_heading(ttl, level=2)
            for it in items[:5]:
                doc.add_paragraph(str(it.get("point")), style="List Bullet")
                ev = "؛ ".join(map(str, (it.get("evidence") or [])[:2]))
                if ev:
                    doc.add_paragraph(f"الدليل: {ev}", style="Intense Quote")
        for key, ttl, label in (("porter", "قوى المنافسة الخمس", "force"),
                                ("pestel", "تحليل PESTEL", "dimension")):
            items = dyn_v.get(key) or []
            if items:
                doc.add_heading(ttl, level=2)
                _add_table(doc, ["البُعد", "الملاحظة", "الدليل"],
                           [[it.get(label), it.get("point"),
                             "؛ ".join(map(str, (it.get("evidence") or [])[:1]))]
                            for it in items[:7]])
        if dyn_v.get("note"):
            doc.add_paragraph(str(dyn_v["note"]), style="Intense Quote")
    elif isinstance(dyn_v, dict) and dyn_v.get("raw_signals"):
        doc.add_paragraph("إشارات ويب خام (لم تُصنَّف بعد — التصنيف في "
                          "الأطر يتطلب مفتاح كلود):")
        for sig in dyn_v["raw_signals"][:6]:
            t = sig.get("title") if isinstance(sig, dict) else sig
            doc.add_paragraph(_trim_sentence(t, 180), style="List Bullet")
    else:
        doc.add_paragraph(str(dyn.get("note") or
                          "تحليل الدوافع والكوابح والفرص والتحديات يتطلب "
                          "مفتاح بحث الويب — لم يُشغَّل في هذا التحليل."))

    # ═══ ٦) حجم السوق والتوقعات — TAM/SAM/SOM + النمو ═══
    doc.add_heading("٦. حجم السوق والتوقعات", level=1)
    _docx_market_size(doc, top_m)
    tr = top_m.get("trend") or {}
    if tr.get("growth_pct") is not None or tr.get("cagr_pct") is not None:
        from silk_narrative import growth_phrase
        doc.add_paragraph(growth_phrase(tr.get("cagr_pct"),
                                        tr.get("growth_pct"),
                                        years="سنوات الدراسة"))
        doc.add_paragraph(f"المصدر: {tr.get('source') or 'UN Comtrade'}",
                          style="Intense Quote")

    # ═══ ٧) تحليل التقسيم ═══
    doc.add_heading("٧. تحليل التقسيم", level=1)
    _docx_segments(doc, top_m)

    # ═══ ٨) تحليل التجارة — قوّة كومتريد الفريدة ═══
    doc.add_heading("٨. تحليل التجارة (استيراد/تصدير)", level=1)
    from silk_narrative import fmt_money, fmt_pct
    # صفّ بلا مورّد ولا قيمة بند زائف — يُستبعد قبل الحكم على الفراغ، وإلا
    # ظهر جدول "بيانات" فارغة الخلايا فوق سطر مصدر يتيم (P2: حدود بلا محتوى).
    countries = [c for c in (top_m.get("supplier_countries") or [])
                if c.get("partner") or c.get("value_usd") is not None]
    if countries:
        _add_table(doc, ["الدولة المورّدة", "الحصة", "القيمة"],
                   [[c.get("partner"), fmt_pct(c.get("share")),
                     fmt_money(c.get("value_usd"))] for c in countries[:8]])
        doc.add_paragraph("المصدر: UN Comtrade", style="Intense Quote")
    else:
        doc.add_paragraph("لا بيانات تجارة ثنائية مرصودة لهذا التحليل — "
                          "فجوة معلنة (تتطلب with_research/UN Comtrade).")

    # ═══ ٩) الأسواق المرشّحة الأخرى — كل سوق بجمل تجارية سردية، لا تفريغ ═══
    # مكوّنات خام. سدّ تسريب (الطبقة ٩): "التحليل الإقليمي" اسم مضلِّل —
    # القسم يقارن أسواقاً مرشّحة عالمياً، لا مناطق فرعية داخل سوق واحد؛
    # نفس عنوان المرآة في markdown الآن (كانت هي وحدها سليمة التسمية).
    doc.add_heading("٩. الأسواق المرشّحة الأخرى", level=1)
    from silk_narrative import market_component_lines, confidence_phrase
    for i, m in enumerate((view.get("markets") or [])[:8], 1):
        doc.add_heading(f"٩.{i} {m.get('country')}", level=2)
        lines = market_component_lines(m)
        if lines:
            for line in lines:
                doc.add_paragraph(line, style="List Bullet")
        else:
            doc.add_paragraph("لا مكوّنات مرصودة لهذا السوق — فجوة معلنة")
        doc.add_paragraph(f"الثقة الإجمالية لهذا التقييم: "
                          f"{confidence_phrase(m.get('confidence'))}",
                          style="Intense Quote")

    # ═══ ١٠) المشهد التنافسي ═══
    doc.add_heading("١٠. المشهد التنافسي", level=1)
    _docx_competition_research(doc, top_m)
    _docx_swot(doc, top_m)
    _docx_pricing_layers(doc, top_m)
    prices = top_m.get("prices") or []
    if prices:
        doc.add_heading("أسعار المنتجات في السوق", level=2)
        for pr in prices[:8]:
            doc.add_paragraph(
                f"{pr.get('title') or 'قائمة'}: {_fmt(pr.get('price'))}"
                + (f" {pr['currency']}" if pr.get("currency") else "")
                + (f" — {pr['store']}" if pr.get("store") else ""),
                style="List Bullet")
    named = top_m.get("named_competitors") or []
    if named:
        doc.add_heading("مراجع ويب للمراجعة اليدوية (ليست أسماء منافسين)",
                        level=2)
        for n in named[:8]:
            doc.add_paragraph(str(n), style="List Bullet")
    cp = view.get("competitive_position") or {}
    doc.add_heading("موقعك التنافسي", level=2)
    if cp.get("available"):
        doc.add_paragraph(cp.get("coverage") or "")
        for f in cp.get("feasibility_threads") or []:
            p = doc.add_paragraph()
            p.add_run(f"ضد {f['competitor']}: ").bold = True
            p.add_run(f"سعر مرصود {_fmt(f['observed_price'])} — هامشك عند "
                      f"المضاهاة {f['margin_at_match_pct']}% وعند البيع "
                      f"أقل 10% {f['margin_at_10pct_below']}%")
            for gap in f.get("assumptions_and_gaps") or []:
                doc.add_paragraph(gap, style="List Bullet")
        for t in cp.get("competitor_threads") or []:
            if not t.get("observed_price"):
                doc.add_paragraph(f"مرجع ويب للمراجعة: {t['name']}",
                                  style="List Bullet")
    else:
        doc.add_paragraph(cp.get("note") or "—")

    # ═══ ١١) استخبارات العميل والطلب ═══
    doc.add_heading("١١. استخبارات العميل والطلب", level=1)
    ins = (cc.get("insights") or [])
    raw_culture = view.get("culture") or []
    if ins:
        for o in ins[:5]:
            doc.add_paragraph(str(o.get("point") or ""), style="List Bullet")
            ev = "؛ ".join(map(str, (o.get("evidence") or [])[:3]))
            if ev:
                doc.add_paragraph(f"الدليل: {ev}", style="Intense Quote")
    elif raw_culture:
        doc.add_paragraph("عناوين مرصودة من بحث الويب (لم تُحلَّل بعد):")
        for c in raw_culture[:6]:
            doc.add_paragraph(str(c.get("title"))[:200], style="List Bullet")
    else:
        doc.add_paragraph("—")
    doc.add_paragraph("صوت العميل المباشر (ما الذي يقدّره المشترون وكيف "
                      "يقيّمون المورّدين) يتطلب بحثاً أولياً — مقابلات أو "
                      "استبيانات — لم يُجرَ بعد؛ هذا القسم يعرض إشارات "
                      "ثانوية فقط.")

    # ═══ ١٢) المشهد التنظيمي والمخاطر ═══
    doc.add_heading("١٢. المشهد التنظيمي والمخاطر", level=1)
    _docx_regulatory(doc, top_m)
    ed_top = top_m.get("entry_decision") or {}
    risks = ed_top.get("risks") or []
    if risks:
        doc.add_heading("سجل المخاطر", level=2)
        _add_table(doc, ["الخطر", "الشدة", "الدليل"],
                   [[r.get("risk"), r.get("severity"), r.get("evidence")]
                    for r in risks if isinstance(r, dict)])

    # ═══ ١٣) الاتجاهات والتوقع المستقبلي ═══
    doc.add_heading("١٣. الاتجاهات والتوقع المستقبلي", level=1)
    series = [p for p in (tr.get("series") or []) if p.get("value") is not None]
    if len(series) >= 2:
        _add_table(doc, ["السنة", "قيمة الاستيراد (دولار)"],
                   [[p.get("year"), _fmt(p.get("value"))] for p in series])
        doc.add_paragraph(f"المصدر: {tr.get('source') or 'UN Comtrade'}",
                          style="Intense Quote")
        # سيناريوهات من المدى التاريخي المرصود حصراً — لا توقع مخترع:
        # التغيّرات السنوية الفعلية => أدنى/وسيط/أقصى نمو ملحوظ.
        changes = []
        for a, b in zip(series, series[1:]):
            va, vb = float(a["value"]), float(b["value"])
            if va > 0:
                changes.append((vb - va) / va * 100.0)
        if changes:
            changes.sort()
            lo, hi = changes[0], changes[-1]
            mid = changes[len(changes) // 2]
            doc.add_paragraph(
                "سيناريوهات مشتقة من المدى التاريخي المرصود (أدنى/أوسط/"
                "أعلى تغيّر سنوي فعلي خلال سنوات الدراسة — ليست تنبؤاً): "
                f"متحفّظ {lo:.1f}% | أساسي {mid:.1f}% | متفائل {hi:.1f}% "
                "سنوياً.")
    else:
        doc.add_paragraph("بيانات الاتجاه غير كافية لهذه السنوات.")

    # ═══ حدود هذا التقرير — قبل التوصيات (§10.3) ═══
    doc.add_heading("حدود هذا التقرير", level=1)
    limits = view.get("limits") or []
    if limits:
        for x in _gap_list_ar(limits[:12]):
            doc.add_paragraph(str(x), style="List Bullet")
    else:
        doc.add_paragraph("لا حدود مسجّلة لهذا التحليل.")

    # ═══ ١٤) التوصيات الاستراتيجية ═══ (قرار الدخول انتقل قرب الخلاصة
    # التنفيذية أعلاه — راجع القسم ٣)
    doc.add_heading("١٤. التوصيات الاستراتيجية", level=1)
    _docx_entry_strategy(doc, top_m)
    for line in view.get("brief") or []:
        doc.add_paragraph(line)

    # ═══ دليل المورّدين والمستوردين ═══ سدّ تسريب (الطبقة ٩): كان هذا
    # القسم مدفوناً كعنوان فرعي (level=2) داخل التوصيات — قسم رئيسي مستقل
    # الآن (level=1، بلا رقم — نفس اصطلاح "حدود هذا التقرير"/"الملحق"
    # القائم أصلاً خارج تسلسل الأقسام الأربعة عشر المرقّم، فلا يكسر قفل
    # الهيكل ١٤ الذي يختبره test_p2_report_structure.py). يظهر دوماً —
    # كما كان قبل النقل — والفجوة (لا مرشّحين) تُعلَن داخل القسم لا بإخفائه.
    doc.add_heading("دليل المورّدين والمستوردين", level=1)
    if top_m.get("suppliers"):
        doc.add_heading("الموردون والأعمال بالاسم", level=2)
        for sup in top_m["suppliers"][:10]:
            doc.add_paragraph(f"{sup.get('name')} — {sup.get('source')}",
                              style="List Bullet")
    _docx_supplier_directory(doc, top_m)

    # ═══ الملحق: تغطية المصادر وأثرها (للمحلّل) ═══
    doc.add_heading("الملحق: تغطية المصادر وأثرها", level=1)
    cov = top_m.get("section_coverage") or {}
    if cov:
        _add_table(doc, ["القسم", "المُسهم/المُحاوَل", "الدرجة"],
                   [[_SEC_AR.get(sec, sec),
                     f"{c['contributed']}/{c['attempted']}", c["score"]]
                    for sec, c in cov.items()])
    # 2B في الملحق: الأقسام دون عتبة الكفاية تُسرد هنا بجملة النقص الوحيدة
    # المسموح بها (مصادر مُحاوَلة) — شفافية المحلّل، لا صياح على وجه التقرير.
    st_all = top_m.get("section_status") or {}
    insuff = [(sec, st) for sec, st in st_all.items()
              if st.get("status") == "insufficient"]
    if insuff:
        doc.add_heading("أقسام دون عتبة الكفاية", level=2)
        from silk_render import insufficient_line
        for sec, st in insuff:
            doc.add_paragraph(insufficient_line(_SEC_AR.get(sec, sec), st))
    prov = view.get("provenance") or []
    if prov:
        doc.add_heading("أثر المصادر (المحاولات والإسهام)", level=2)
        for b in prov:
            doc.add_paragraph(f"{b['source']}: أسهم {b['contributed']} من "
                              f"{b['attempted']} محاولة")
            for f in b.get("failures") or []:
                doc.add_paragraph(f"    فشل مُسجَّل: {f}",
                                  style="Intense Quote")

    _docx_deep_research(doc, view)

    _finalize_rtl(doc)   # §4: اتجاه RTL صريح على كل فقرة/run قبل الحفظ
    doc.save(path)
    return path


def _md_cell(x: object) -> str:
    """خلية جدول Markdown آمنة — escape pipes/newlines for a table cell."""
    return str(x if x is not None else "—").replace("|", "/").replace("\n", " ")


def _md_glossary(dr: dict, L: list) -> None:
    """B1 (SPEC-v2): ذيّل «مسرد المصطلحات» في Markdown من بنية النموذج
    (`dr["glossary"]`) — المصطلحات المستعملة فعلاً فقط."""
    from silk_style_contract import GLOSSARY_HEADING
    gloss = dr.get("glossary") or []
    if not gloss:
        return
    L += [f"## {GLOSSARY_HEADING}", ""]
    L += [f"- **{g['term']}**: {g['gloss']}" for g in gloss]
    L.append("")


def _docx_glossary(doc, dr: dict, sanitize=None) -> None:
    """B1 (SPEC-v2): اعرض «مسرد المصطلحات» في مستند Word (المدقّق والعميل)
    من بنية النموذج. `sanitize` اختياري (مُطهِّر تقرير العميل)."""
    from silk_style_contract import GLOSSARY_HEADING
    gloss = dr.get("glossary") or []
    if not gloss:
        return
    doc.add_heading(GLOSSARY_HEADING, level=2)
    for g in gloss:
        line = f"{g['term']}: {g['gloss']}"
        doc.add_paragraph(sanitize(line) if sanitize else line, style="List Bullet")


_LEADS_TITLE = "قائمة مستوردين وموزعين قابلين للتواصل"
_LEADS_HEADER = ["الاسم", "العنوان", "الهاتف", "الإيميل", "الموقع", "التقييم",
                 "مستوى التوثيق"]


import functools as _functools


@_functools.lru_cache(maxsize=1)
def _country_names() -> list:
    """[(iso3, {أسماء بحروف صغيرة})] لكل دول المرجع — لفلترة جغرافيا الروابط."""
    try:
        from silk_market_resolver import _load
        out = []
        for row in _load():
            iso3 = (row.get("iso3") or "").strip().upper()
            names = {(row.get("name_en") or "").strip().lower(),
                     (row.get("name_ar") or "").strip().lower()}
            names.discard("")
            if len(iso3) == 3 and names:
                out.append((iso3, names))
        return out
    except Exception:  # noqa: BLE001 — تعذّر التحميل => لا فلترة جغرافيا (لا كسر)
        return []


def _address_wrong_geo(address, target_iso3: str, target_names: set) -> bool:
    """هل عنوانُ الرائد يُسمّي دولةً **غير** السوق المستهدفة صراحةً؟ — Wave 2
    (البند ٤). لا حقلَ دولةٍ على الرائد، فنقرأ العنوان الحرّ: إن ورد اسمُ دولةٍ
    معروفةٍ ≠ السوق => يُسقَط (فجوة معلنة). غموضٌ/لا دولة => يُبقى (متحفّظ)."""
    addr = (address or "").strip().lower()
    if not addr:
        return False
    tgt = {t for t in (target_names or set()) if t}
    for iso3, names in _country_names():
        if iso3 == (target_iso3 or "").upper():
            continue
        for nm in names:
            if nm and nm in tgt:
                continue
            if nm and re.search(r"(?<![\w])" + re.escape(nm) + r"(?![\w])", addr):
                return True
    return False


def _is_filler_lead(lead: dict) -> bool:
    """صفٌّ حشوٌ: اسمٌ (أو بلا اسم) وكلُّ حقولِ الاتصال فارغة/«—» — Wave 2 (البند ٦)."""
    for k in ("phone", "email", "website", "maps_link", "address"):
        v = lead.get(k)
        if v not in (None, "", "—"):
            return False
    return True


def _clean_leads(leads: list, dr: dict) -> list:
    """نقِّ روابط العميل عند حدّ الجدول (Wave 2، البنود ٤/٥/٦): يُسقِط جملَ
    النثر (كِيانُ اسمٍ مطلوب)، والجغرافيا الخاطئة (دولة ≠ السوق)، وصفوف الحشو
    (اسمٌ بلا أيّ اتصال). يعمل على المدوّنة المخزَّنة أيضًا (لا مسار الكشط وحده)."""
    from silk_gmaps import looks_like_name
    market = dr.get("market") or {}
    iso3 = (market.get("iso3") or "").upper()
    tnames = {(market.get("name_en") or "").strip().lower(),
              (market.get("name_ar") or "").strip().lower()}
    out = []
    for lead in leads or []:
        nm = (lead.get("name") or "").strip()
        if not nm or not looks_like_name(nm):          # البند ٥: نثر/بلا اسم
            continue
        if _is_filler_lead(lead):                        # البند ٦: حشو
            continue
        if _address_wrong_geo(lead.get("address"), iso3, tnames):  # البند ٤
            continue
        out.append(lead)
    return out


def _leads_data(dr: dict):
    il = dr.get("importer_leads") or {}
    leads = _clean_leads(il.get("leads") or [], dr)
    return leads, (il.get("note") or "")


def _lead_cells(lead: dict) -> list:
    """C5: خلايا صفّ رائد — الحقل الغائب «—» (لا اختلاق). التقييم مع عدد
    المراجعات إن توفّرا."""
    def g(k):
        v = lead.get(k)
        return str(v).strip() if v not in (None, "") else "—"
    rating, rc = lead.get("rating"), lead.get("review_count")
    rating_s = (f"{rating} ({rc})" if rating and rc
                else str(rating) if rating else "—")
    site = lead.get("website") or lead.get("maps_link") or ""
    return [g("name"), g("address"), g("phone"), g("email"),
            site.strip() or "—", rating_s, g("doc_level")]


def _md_leads(dr: dict, L: list) -> None:
    """C5: جدول «قائمة مستوردين وموزعين قابلين للتواصل» في Markdown."""
    from silk_gmaps import maps_disclaimer
    MAPS_DISCLAIMER = maps_disclaimer(dr.get("product"))   # Wave 2: بارامتري بالمنتج
    leads, note = _leads_data(dr)
    L += [f"## {_LEADS_TITLE}", ""]
    if not leads:
        L += ["لم تُرصَد جهات اتصال قابلة للتواصل في هذا التشغيل — فجوة معلنة"
              + (f" ({note})" if note else "") + ".", ""]
        return
    L += ["| " + " | ".join(_LEADS_HEADER) + " |",
          "|" + "|".join(["---"] * len(_LEADS_HEADER)) + "|"]
    for lead in leads:
        L.append("| " + " | ".join(c.replace("|", "／")
                                   for c in _lead_cells(lead)) + " |")
    L += ["", MAPS_DISCLAIMER, ""]


def _docx_leads(doc, dr: dict, sanitize=None) -> None:
    """C5: جدول الروابط في Word (المدقّق والعميل) من بنية النموذج."""
    from silk_gmaps import maps_disclaimer
    MAPS_DISCLAIMER = maps_disclaimer(dr.get("product"))   # Wave 2: بارامتري بالمنتج
    leads, note = _leads_data(dr)
    doc.add_heading(_LEADS_TITLE, level=2)
    if not leads:
        msg = ("لم تُرصَد جهات اتصال قابلة للتواصل في هذا التشغيل — فجوة معلنة"
               + (f" ({note})" if note else ""))
        doc.add_paragraph(sanitize(msg) if sanitize else msg)
        return
    rows = [[(sanitize(c) if sanitize else c) for c in _lead_cells(lead)]
            for lead in leads]
    _add_table(doc, _LEADS_HEADER, rows)
    line = MAPS_DISCLAIMER
    doc.add_paragraph(sanitize(line) if sanitize else line,
                      style="Intense Quote")


def _md_deep_research(view: dict, prefix: list[str]) -> str:
    """التقرير الكامل Markdown لنتيجة /research — يُصيَّر من `view["deep_research"]`
    (نفس مصدر اللوحة وتصدير Word عبر `_docx_deep_research`)، لا من قالب /analyze.

    بلاغ حي إنتاجي (تدقيق تمور/هولندا): `render_markdown` كان يُصيّر قالب
    /analyze حصراً — و`markets:[]` لنتيجة بحث عميق تُفرِّغ كل أقسامه («تغطية
    0.0%»، «0 أسواق»، SWOT فارغ، «with_research غير مفعّلة») بينما النصّ
    السردي الغنيّ (HHI، أسعار الرفّ، لوائح EU) موجود في `dr["report"]["text"]`
    ولا يُقرأ أبداً. هذا الفرع يُصيّره فعلاً: ترويسة صادقة + الحكم + السرد
    الكامل + ملحق الأدلة + أثر المصادر + الحدود.
    """
    dr = view.get("deep_research") or {}
    h = view.get("header") or {}
    L = list(prefix)
    verdict = dr.get("verdict") or {}
    ai = verdict.get("ai") or {}

    # ── الترويسة كجدول — نفس القالب لكن بلا حقول /analyze الفارغة ─────────────
    market = dr.get("market") or {}
    L += [f"# سِلك — تقرير بحث سوق معمّق: {view.get('product')}", "",
          "| البند | القيمة |", "| --- | --- |",
          f"| المنتج | {_md_cell(h.get('product') or view.get('product'))} |",
          f"| رمز HS | {_md_cell(h.get('hs_code') or view.get('hs_code'))} "
          f"(ثقة التصنيف {_md_cell(view.get('hs_confidence'))}) |",
          "| المنشأ | السعودية (SAU) |",
          f"| السوق المستهدف | {_md_cell(h.get('target_market') or market.get('name_ar') or market.get('name_en'))} |",
          f"| التاريخ | {_md_cell(h.get('date'))} |",
          f"| الحكم | {_md_cell(view.get('verdict_label') or dr.get('verdict_label'))} |"]
    rc = (dr.get("report") or {}).get("review_cycles") or 0
    if rc:
        L.append(f"| مراجعة التقرير | روجِع عبر {rc} دورة تنقيح |")
    L.append("")

    _bnr = _degraded_banner_text(view)
    if _bnr:
        L += [f"> {_bnr}", ""]

    # ── الحكم وأساسه — من synthesize (نقطة الحكم الوحيدة) ────────────────────
    L += ["## الحكم وأساسه", ""]
    L.append(f"- التوصية: **{view.get('verdict_label') or dr.get('verdict_label') or '—'}**")
    reasoning = ai.get("reasoning") or verdict.get("note")
    if reasoning:
        L += ["", str(reasoning)]
    L.append("")

    # Wave 6.1: شرطا قلب الحكم المهيكلان (حكم مراقبة/مشروط) — حقل مُصادَق من
    # نموذج العرض، لا نثر حظّ. كل شرط بخطوة إغلاقه (تربطها خارطة الطريق).
    _flips = dr.get("flip_conditions") or []
    if _flips:
        from silk_render import FLIP_CONDITIONS_HEADING
        L += [f"## {FLIP_CONDITIONS_HEADING}", ""]
        for i, c in enumerate(_flips, 1):
            _mark = "✓ محقَّق" if c.get("met") else "○ غير محقَّق"
            L.append(f"{i}. **{_md_cell(c.get('condition'))}** — {_mark}؛ "
                     f"يُغلَق عبر: {_md_cell(c.get('closes_via'))}")
        L.append("")

    # ── التقرير السردي الكامل (كاتب التقرير، مراجَع) — النصّ الغنيّ نفسه ──────
    report_text = (dr.get("report") or {}).get("text")
    if report_text:
        L += [str(report_text).rstrip(), ""]
        _md_glossary(dr, L)  # B1: مسرد المصطلحات المستعملة فعلاً
        _md_leads(dr, L)     # C5: قائمة المستوردين القابلين للتواصل
    else:
        fr = (dr.get("report") or {}).get("failure_reason")
        L += ["## التقرير السردي الكامل", "",
              ("تعذّر إنجاز التقرير السردي التفصيلي في هذه المحاولة"
               + (f" — {fr}" if fr else "")
               + "؛ الأدلة المرصودة والحكم أعلاه قائمة، ويمكن إعادة التوليد "
                 "دون إعادة البحث الكامل."), ""]

    # ── ملحق الأدلة الرقمية — التقاطعات (raw supporting evidence) ────────────
    by_cat = (dr.get("analyst") or {}).get("by_category") or {}
    if any(by_cat.values()):
        L += ["## ملحق: الأدلة الرقمية الداعمة", ""]
        for cat, items in by_cat.items():
            if not items:
                continue
            L.append(f"### {_CATEGORY_AR.get(cat, cat)}")
            for f in items:
                badge = f.get("confidence_badge") or ""
                L.append(f"- {_md_cell(f.get('value'))} "
                         f"[{_md_cell(_clean_source_label(f.get('source')))} {badge}] "
                         f"{_md_cell(f.get('note') or '')}".rstrip())
            L.append("")

    # ── ملحق أثر المصادر — provenance (البعثات الاثنتا عشرة وإسهامها) ────────
    L += ["## ملحق: أثر المصادر (المحاولات والإسهام)", ""]
    prov = view.get("provenance") or []
    if prov:
        for b in prov:
            L.append(f"- {b['source']}: أسهم {b['contributed']} من "
                     f"{b['attempted']} محاولة")
            for fl in b.get("failures") or []:
                L.append(f"  - فشل مُسجَّل: {fl}")
    else:
        L.append("- لا أثر مصادر مسجّلاً")
    L.append("")

    # ── حدود هذا التقرير — declared limits ──────────────────────────────────
    L += ["## حدود هذا التقرير", ""]
    limits = view.get("limits") or ["لا فجوات مرصودة"]
    for x in _gap_list_ar(limits[:12]):
        L.append(f"- {x}")
    L.append("")

    # ── التوصية / المختصر — نفس سطور المختصر (لا صياغة موازية) ───────────────
    L += ["## التوصية / المختصر", ""]
    for line in view.get("brief") or []:
        L.append(f"- {line}")
    if view.get("note"):
        L += ["", str(view["note"])]
    L.append("")
    out = "\n".join(L)
    _assert_verdict_consistency_text(  # Master Prompt Part 2 §B
        out, _resolve_vtxt(dr), "تقرير Markdown")
    return out


def render_markdown(view: dict) -> str:
    """التقرير الكامل Markdown (§7) — نفس أقسام Word وترتيبها، من القالب حصراً.

    كل رقم يليه سطر مصدره بين قوسين؛ المنمذج موسوم «مُقدَّر — نموذج بافتراضات
    معلنة» بمعادلته؛ بوابة 2B نفسها: قسم دون العتبة يطبع سطر النقص الصريح فقط.
    Full Markdown report derived from the ONE canonical view — pure display.
    نتيجة /research (`view["deep_research"]`) تُصيَّر عبر `_md_deep_research`
    (من السرد المراجَع نفسه، لا قالب /analyze الفارغ) — نفس فرع `render_docx`.
    """
    from silk_render import insufficient_line

    _assert_production_clean(view)
    h = view.get("header") or {}
    top_m = (view.get("markets") or [{}])[0]
    st_all = top_m.get("section_status") or {}
    L: list[str] = []
    if view.get("test_run"):
        # بلا اسم متغيّر البيئة على وجه التقرير — نفس صياغة render_text
        # و render_brief (اتساق المشتقات، وإزالة سباكة داخلية من نص عميل).
        L += ["> ⚠ **TEST RUN** — تشغيل برهاني ببدائل موسومة، "
              "ليس تقريراً إنتاجياً", ""]

    # بلاغ حي (تدقيق /research): نتيجة بحث عميق تُصيَّر من deep_research لا من
    # قالب /analyze (markets:[] يُفرّغه) — نفس فرع render_docx.
    if view.get("deep_research"):
        return _md_deep_research(view, L)

    # ── الترويسة كجدول — header table ────────────────────────────────────────
    L += [f"# سِلك — تقرير سوق: {view.get('product')}", "",
          "| البند | القيمة |", "| --- | --- |",
          f"| المنتج | {_md_cell(h.get('product'))} |",
          f"| رمز HS | {_md_cell(h.get('hs_code'))} "
          f"(ثقة التصنيف {_md_cell(view.get('hs_confidence'))}) |",
          "| المنشأ | السعودية (SAU) |",
          f"| السوق المستهدف | {_md_cell(h.get('target_market'))} |",
          f"| التاريخ | {_md_cell(h.get('date'))} |",
          f"| سنة البيانات | {_md_cell(_data_year_label(view))} |",
          f"| تغطية البيانات الإجمالية | {h.get('coverage_pct')}% |", ""]

    # ── ١. الخلاصة التنفيذية — narrative executive summary (حكم واحد لا حكمان) ─
    # P1: ثلاث فقرات بشرية فقط — كفاية البيانات و«قراءة أولية» في المنهجية.
    d = view.get("decision") or {}
    L += ["## الخلاصة التنفيذية", ""]
    L += [f"{para}" for para in _narrative_exec_summary(view)]
    L.append("")

    # ── ٢. منهجية البحث — قسم مستقل ظاهر (§2) ────────────────────────────────
    L += ["## منهجية البحث", ""]
    L += [f"- {line}" for line in _methodology_lines(view)]
    L.append("")

    # ── ٣. تعريف السوق ونطاقه — جملة نطاق صريحة (§3) ─────────────────────────
    L += ["## تعريف السوق ونطاقه", "", _market_scope_paragraph(view), ""]

    # ── قرار الدخول — weighted entry decision ───────────────────────────────
    L += ["## قرار الدخول", ""]
    ed, ed_absent = _entry_decision_of(top_m)
    if ed is None:
        L += [ed_absent, ""]
    else:
        from silk_decision import _WEIGHT_LABEL_AR
        from silk_narrative import confidence_phrase, verdict_ar
        sbo = ed.get("scores_by_option") or {}
        weights_label = ed.get("weights_label") or ed.get("weights_option") or ""
        other_key = "B" if ed.get("weights_option") == "A" else "A"
        other_label = _WEIGHT_LABEL_AR.get(other_key, "")
        L += [f"- الحكم: **{verdict_ar(ed.get('verdict'))}** | "
              f"النقاط: {_fmt(ed.get('score'))}"
              f" | الثقة: {confidence_phrase(ed.get('confidence'))}",
              f"- أساس الثقة: {ed.get('confidence_basis')}",
              f"- منهجية الترجيح المعتمدة: {weights_label} — النقاط "
              f"{_fmt(ed.get('score'))} (بديل مقارنة، {other_label}: "
              f"{_fmt(sbo.get(other_key))})"]
        if ed.get("weights_note"):
            L.append(f"- ملاحظة الأوزان: {ed['weights_note']}")
        L += ["", "| العمود | القيمة | الأساس | مكوّنات غائبة |",
              "| --- | --- | --- | --- |"]
        for key, p in (ed.get("pillars") or {}).items():
            missing = "، ".join(map(str, p.get("missing") or [])) or "—"
            L.append(f"| {_PILLAR_AR.get(key, key)} | {_fmt(p.get('value'))} | "
                     f"{_md_cell(p.get('basis'))} | {_md_cell(missing)} |")
        L.append("")
        if ed.get("missing_pillars"):
            L.append("- أعمدة غائبة كلياً: " + "، ".join(
                _PILLAR_AR.get(k, k) for k in ed["missing_pillars"]))
        if ed.get("critical_risk"):
            L.append("- **خطر حرج مرصود** — راجع سجل المخاطر أدناه.")
        if ed.get("conditions"):
            L += ["", "**الشروط:**",
                  *[f"- {c}" for c in ed["conditions"]]]
        if ed.get("first_steps"):
            L += ["", "**الخطوات الأولى:**",
                  *[f"{i}. {s}" for i, s in enumerate(ed["first_steps"], 1)]]
        L += ["", f"لماذا: {ed.get('why')}"]
        if ed.get("note"):
            L.append(str(ed["note"]))
        L.append("")

    # ── موقعك التنافسي — competitive position (correlation) ─────────────────
    cp = view.get("competitive_position") or {}
    L += ["## موقعك التنافسي", ""]
    if cp.get("available"):
        L.append(f"- التغطية: {cp.get('coverage')}")
        for f in cp.get("feasibility_threads") or []:
            L.append(f"- ضد {f['competitor']}: سعر مرصود "
                     f"{_fmt(f['observed_price'])} — هامشك عند المضاهاة "
                     f"{f['margin_at_match_pct']}% وعند البيع أقل 10% "
                     f"{f['margin_at_10pct_below']}%")
            for gap in f.get("assumptions_and_gaps") or []:
                L.append(f"  - {gap}")
        for t in cp.get("competitor_threads") or []:
            if not t.get("observed_price"):
                # خيوط بحث الويب مراجع لا كيانات (ثغرة ٢).
                L.append(f"- مرجع ويب للمراجعة: {t['name']} — {t['price_flag']} "
                         f"(اكتمال الخيط {t['thread_completeness']})")
    else:
        L.append(str(cp.get("note") or ""))
    L.append("")

    # ── حجم السوق TAM/SAM/SOM — من حزمة البحث ───────────────────────────────
    L += ["## حجم السوق — TAM/SAM/SOM", ""]
    r_bundle, r_absent = _research_bundle(top_m)
    if r_bundle is None:
        L += [r_absent, ""]
    else:
        ms = _ragent(top_m, "market_size")
        shown = False
        for f in ms.get("findings") or []:
            if f.get("value") is None:
                continue
            L.append(f"- {_f_text(f)} ({_f_srcline(f)})")
            shown = True
        if not shown:
            L.append("- لا اكتشافات مرصودة لحجم السوق")
        for g in ms.get("gaps") or []:
            L.append(f"- فجوة معلنة: {_gap_ar(g)}")
        L.append("")

    # ── الأسواق المرشّحة الأخرى — جمل تجارية سردية لا تفريغ مكوّنات خام ──────
    other_markets = (view.get("markets") or [])[1:8]
    if other_markets:
        from silk_narrative import market_component_lines, confidence_phrase
        L += ["## الأسواق المرشّحة الأخرى", ""]
        for m in other_markets:
            L.append(f"### {m.get('country')}")
            lines = market_component_lines(m)
            if lines:
                L += [f"- {line}" for line in lines]
            else:
                L.append("- لا مكوّنات مرصودة لهذا السوق — فجوة معلنة")
            L.append(f"- الثقة الإجمالية لهذا التقييم: "
                     f"{confidence_phrase(m.get('confidence'))}")
            L.append("")

    # ── المنافسة بطبقتيها — دولية (كومتريد) + شركات بالاسم + طبقة الإثراء ────
    L += ["## المنافسة بطبقتيها", ""]
    comp = _ragent(top_m, "competitor")
    if comp:
        from silk_narrative import internal_ar
        L.append("**الطبقة الدولية (تركّز الموردين — UN Comtrade):**")
        for metric in ("hhi", "top_supplier_share_pct", "saudi_share_pct"):
            f = _rfind(comp, metric)
            if f and f.get("value") is not None:
                L.append(f"- {_f_text(f)} ({_f_srcline(f)})")
                if f.get("note"):
                    L.append(f"  - {_gap_ar(f['note'])}")
            else:
                L.append(f"- {internal_ar(metric)}: غير مرصود")
        sc_f = _rfind(comp, "supplier_countries")
        for c in ((sc_f or {}).get("value") or [])[:6]:
            if isinstance(c, dict):
                from silk_narrative import fmt_money
                L.append(f"- {c.get('partner')}: {c.get('share')}% "
                         f"({fmt_money(c.get('value_usd'))}) ({_f_srcline(sc_f)})")
        named_rc = (_rfind(comp, "named_companies") or {}).get("value") or []
        ents_rc, refs_rc = _split_candidates(named_rc)
        L += ["", "**شركات بالاسم (كيانات Google Places، غير موثَّقة):**"]
        if ents_rc:
            for n in ents_rc[:10]:
                L.append(f"- {_entry_text(n)}")
            from silk_narrative import confidence_phrase
            L.append(f"- ملاحظة: كيانات غير موثَّقة — الثقة "
                     f"{confidence_phrase(0.4)} — أكّدها قبل أي تعاقد.")
        else:
            L.append("- لا كيانات مرصودة بالاسم — فجوة معلنة (أسماء الأعمال "
                     "تأتي من Google Places حصراً)")
        if refs_rc:
            L += ["", "**مراجع ويب للمراجعة اليدوية (ليست أسماء منافسين):**",
                  *[f"- {_entry_text(n)}" for n in refs_rc[:8]]]
        L.append("")
    else:
        L += [r_absent or "وكيل المنافسة بلا اكتشافات — فجوة معلنة", ""]
    # طبقة الإثراء المحلية خلف بوابة 2B — سطر النقص الصريح فقط دون العتبة.
    st_c = st_all.get("competitors")
    if st_c and st_c.get("status") == "insufficient":
        L += [insufficient_line("المنافسون", st_c), ""]
    else:
        from silk_narrative import fmt_money
        for c in (top_m.get("supplier_countries") or [])[:6]:
            L.append(f"- {c.get('partner')}: {c.get('share')}% "
                     f"({fmt_money(c.get('value_usd'))}) (المصدر: UN Comtrade)")
        for n in (top_m.get("named_competitors") or [])[:8]:
            # عناوين بحث (الطبقة القديمة) — مراجع لا أسماء منافسين (ثغرة ٢).
            L.append(f"- مرجع ويب للمراجعة اليدوية: {n}")
        L.append("")

    # ── التسعير بطبقتيه — border models + gated retail layer ────────────────
    L += ["## التسعير بطبقتيه", ""]
    if r_bundle is None:
        L += [r_absent, ""]
    else:
        pr = _ragent(top_m, "pricing")
        from silk_narrative import internal_ar
        L.append("**الطبقة الحدودية (قيم وحدة كومتريد):**")
        for metric in ("border_unit_value_usd_kg",
                       "saudi_border_unit_value_usd_kg", "margin_at_border_pct"):
            f = _rfind(pr, metric)
            if f and f.get("value") is not None:
                L.append(f"- {_f_text(f)} ({_f_srcline(f)})")
                if f.get("note"):
                    L.append(f"  - {_gap_ar(f['note'])}")
            else:
                L.append(f"- {internal_ar(metric)}: غير مرصود")
        L += ["", "**طبقة التجزئة:**"]
        rp = _rfind(pr, "retail_prices")
        vals = (rp or {}).get("value") or []
        st_p = st_all.get("pricing")
        if vals:
            for v in vals[:8]:
                L.append(f"- {_listing_text(v)} ({_f_srcline(rp)})")
        elif st_p and st_p.get("status") == "insufficient":
            # بوابة 2B: سطر النقص الصريح فقط + فجوة القسم المعلنة بنصّها.
            L.append(insufficient_line("الأسعار", st_p))
            gap = next((g for g in (pr.get("gaps") or [])
                        if "retail_prices" in g), None)
            if gap:
                L.append(f"- فجوة معلنة: {_gap_ar(gap)}")
        else:
            for p in top_m.get("prices") or []:
                L.append(f"- {_listing_text(p)}")
            if not top_m.get("prices"):
                gap = next((g for g in (pr.get("gaps") or [])
                            if "retail_prices" in g),
                           "retail_prices: غير مرصود — فجوة معلنة")
                L.append(f"- {_gap_ar(gap)}")
        points = (_rfind(pr, "retail_price_points") or {}).get("value") or []
        if points:
            L += ["", "**أسعار مُستخلَصة (مذكورة صراحةً في عناوين الويب — مؤشِّر لا "
                      "سعرَ رفٍّ مؤكَّد):**"]
            for p in points[:8]:
                p = p or {}
                L.append(f"- {p.get('price')} {p.get('currency')}/{p.get('unit')}"
                         .rstrip("/")
                         + (f" — {p.get('url')}" if p.get("url") else ""))
        refs = (_rfind(pr, "retail_references") or {}).get("value") or []
        if refs:
            L += ["", "**مصادر الأسعار (للاستشهاد):**"]
            for ref in refs[:3]:
                L.append(f"- {(ref or {}).get('title')} — {(ref or {}).get('url')}"
                         f" — سُحب: {(ref or {}).get('retrieved_at')}")
        elif not points:
            L.append("- مراجع الأسعار: غير مرصودة")
        L.append("")

    # ── SWOT — أربع قوائم بدليل كل بند؛ الربع الفارغ معلن ────────────────────
    sw = top_m.get("swot") or {}
    # سدّ انحراف (الطبقة ٨): عناوين الأرباع كانت ثنائية اللغة هنا
    # ("القوة Strengths") بينما docx عربية صرفة لنفس القسم — نفس الاصطلاح
    # الآن في كلا المشتقّين، لا لغة إنجليزية إضافية على وجه التقرير.
    L += ["## تحليل SWOT (قاعدي من حقائق مرصودة)", ""]
    for key, title in (("S", "القوة"), ("W", "الضعف"),
                       ("O", "الفرص"), ("T", "التهديدات")):
        L.append(f"### {title}")
        items = sw.get(key) or []
        if not items:
            L.append("- لا بند مرصوداً")
        for it in items:
            L.append(f"- {it.get('text')} — الدليل: {it.get('evidence')}")
        L.append("")
    if sw.get("note"):
        L += [f"> {sw['note']}", ""]

    # ── شرائح العملاء — segments with declared basis ─────────────────────────
    L += ["## شرائح العملاء", ""]
    segs = top_m.get("segments") or []
    if segs:
        for s in segs:
            L.append(f"- {s.get('segment')} — الأساس: {s.get('basis')}")
    else:
        L.append("بيانات غير كافية للشرائح — التقسيم السلوكي/الديموغرافي "
                "يتطلب بحثاً أولياً (مقابلات أو استبيانات) لم يُجرَ بعد؛ "
                "لا يُشتق من بيانات ثانوية")
    L.append("")

    # ── دليل المورّدين والمصنّعين — supplier directory ───────────────────────
    L += ["## دليل المورّدين والمصنّعين", ""]
    sd = top_m.get("supplier_directory") or {}
    for key, title in (("saudi", "**مورّدون ومصنّعون سعوديون:**"),
                       ("target", "**موزّعون ومستوردون في السوق المستهدف:**")):
        L.append(title)
        items = sd.get(key) or []
        if not items:
            L.append("- لا مرشّحين مرصودين — فجوة معلنة")
        for e in items[:10]:
            L.append(f"- {_entry_text(e)}")
        L.append("")
    if sd.get("note"):
        L += [f"> {sd['note']}", ""]

    # ── الاشتراطات التنظيمية — بوابة الأهلية أولاً ثم بنود L1 ────────────────
    L += ["## الاشتراطات التنظيمية", ""]
    reg = _ragent(top_m, "regulatory")
    if r_bundle is None or not reg:
        L += [r_absent or "وكيل الاشتراطات لم يعمل في هذا التحليل", ""]
    else:
        gate_f = _rfind(reg, "eligibility_gate")
        if gate_f and gate_f.get("value"):
            L += ["**تحذير — بوابة أهلية أمامية:** هذا السوق يتطلب منشأة معتمدة "
                  "(EU 2017/625) قبل أي بند لاحق؛ لا بند أدناه يُعتبر سالكاً "
                  "قبل اجتيازها.", ""]
        checklist_f = _rfind(reg, "requirements_checklist")
        checklist = (checklist_f or {}).get("value") or []
        if checklist:
            for it in checklist:
                L.append(f"- {_req_text(it)}")
            L.append(f"- ({_f_srcline(checklist_f)})")
        else:
            L.append("- لا بنود اشتراطات مرصودة في مرجع L1 لهذا السوق "
                     "— فجوة معلنة")
        tf = _rfind(reg, "tariff_applied_pct")
        if tf and tf.get("value") is not None:
            L.append(f"- {_f_text(tf)} ({_f_srcline(tf)})")
        for g in reg.get("gaps") or []:
            L.append(f"- فجوة معلنة: {_gap_ar(g)}")
        L.append("")

    # ── سجل المخاطر — decision risks + raw WGI/LPI datapoints ───────────────
    L += ["## سجل المخاطر", ""]
    ed_risks = (ed or {}).get("risks") or []
    if ed_risks:
        for rk in ed_risks:
            L.append(f"- {rk.get('risk')} (الشدة: {rk.get('severity')}) — "
                     f"الدليل: {rk.get('evidence')}")
    else:
        L.append("- لا مخاطر مسجّلة في محرك القرار (§8)" if ed else
                 f"- {ed_absent}")
    for dp in top_m.get("risk") or []:
        if dp.get("value") is not None:
            src = str(dp.get("source") or "غير مرصود")
            if dp.get("retrieved_at"):
                src += f" | سُحب: {dp['retrieved_at']}"
            L.append(f"- {dp.get('note') or 'مؤشر خطر'}: {_fmt(dp['value'])} "
                     f"(المصدر: {src})")
    L.append("")

    # ── تغطية الأقسام — لكل قسم دون العتبة سطر النقص الصريح ─────────────────
    L += ["## تغطية الأقسام", ""]
    cov = top_m.get("section_coverage") or {}
    if cov:
        for sec, c in cov.items():
            flag = " ⚠ مصدر واحد — ثقة منخفضة" if c.get("low_confidence") else ""
            L.append(f"- {_SEC_AR.get(sec, sec)}: {c['contributed']}/"
                     f"{c['attempted']} (درجة {c['score']}){flag}")
            st = st_all.get(sec)
            if st and st.get("status") == "insufficient":
                L.append(f"  - {insufficient_line(_SEC_AR.get(sec, sec), st)}")
    else:
        L.append("- لا تغطية أقسام محسوبة")
    L.append("")

    # ── ملحق أثر المصادر — provenance appendix (لا فشل صامتاً) ──────────────
    L += ["## ملحق: أثر المصادر (المحاولات والإسهام)", ""]
    prov = view.get("provenance") or []
    if prov:
        for b in prov:
            L.append(f"- {b['source']}: أسهم {b['contributed']} من "
                     f"{b['attempted']} محاولة")
            for fl in b.get("failures") or []:
                L.append(f"  - فشل مُسجَّل: {fl}")
    else:
        L.append("- لا أثر مصادر مسجّلاً")
    L.append("")

    # ── حدود هذا التقرير — declared limits before the recommendation ────────
    # سدّ انحراف (الطبقة ٨): كانت هذه القائمة تُطبع خامة بلا _gap_list_ar
    # — الآن نفس المعالجة المطبَّقة في docx (اتساق المشتقّين، لا مسارين).
    L += ["## حدود هذا التقرير", ""]
    limits = view.get("limits") or ["لا فجوات مرصودة في الأسواق العليا"]
    for x in _gap_list_ar(limits[:12]):
        L.append(f"- {x}")
    L.append("")

    # ── التوصية / المختصر — نفس سطور القالب، لا صياغة موازية ────────────────
    L += ["## التوصية / المختصر", ""]
    for line in view.get("brief") or []:
        L.append(f"- {line}")
    if view.get("note"):
        L += ["", str(view["note"])]
    L.append("")
    return "\n".join(L)
