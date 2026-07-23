#!/usr/bin/env python3
"""معايرةُ مقياس محاذاة RTL — RTL-alignment measurer calibration (Wave 2, §4).

> **أمر المُشرِف (الخيار ٣ + التعديلات الأربعة).** قبل الوثوق بأيّ رقم: ابنِ
> عيّنتين من **نفس** المحتوى العربي — A بـ`jc=start` (الوصفة الصحيحة، يجب أن
> تنحاز يمينًا)، وB بـ`jc=right` (البناء المقلوب المعروف، ينحاز يسارًا).
>
> **العقد (التعديل ١): صفرُ محاذاةٍ يسارًا للأسطر العربية.** fixture-A أثبت أنّ
> الوصفة تُصيَّر **100%** يمينًا — فأيّ سطرٍ **عربيّ الأغلبية** مُحاذًى يسارًا
> عيبٌ حقيقيّ. الأسطرُ لاتينيّةُ الأغلبية (أسماء مصادر، تواريخ، أرقام، رموز HS)
> **تُستثنى** من النسبة كليًّا. التساهلُ الهندسيّ 10 نقاط فقط.
>
> الاحتياطيّ القديم `pdftotext -bbox` كان يعدّ **كلّ كلمةٍ سطرًا** بلا تجميع y
> (لا يميّز A من B) => حُذِف نهائيًّا. القياسُ الآن `pymupdf/fitz` بتجميعٍ صحيحٍ
> للأسطر على y — تبعيةٌ صلبةٌ لوظيفة e2e.

يُشغَّل في وظيفة e2e-live-shape (soffice + fitz حاضران هناك):
    python3 tools/rtl_calibration.py
يطبع نسخة LibreOffice + خُلاصة تصنيفٍ لكلّ عيّنة + حكم A/B، وينهي بـ0 دائمًا
(تشخيصيّ). الحارسُ الدائم لنفس العقد يعيش في اختبار (test_report_output_overhaul).
"""
from __future__ import annotations

import collections
import os
import shutil
import subprocess
import sys
import tempfile

_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if _ROOT not in sys.path:
    sys.path.insert(0, _ROOT)

# أسطرٌ عربيةٌ قصيرةٌ متعرّجة (غير مشكّلة كي يكون تجريدُ الحركات لا-أثر) — كلُّ
# فقرةٍ سطرٌ قصيرٌ مفردٌ فتكثر الأسطرُ الراجعةُ القصيرةُ القابلةُ للقياس.
_LINES = [
    "هذا سطر اختبار قصير", "السوق ينمو سنويا", "الاسعار مستقرة",
    "المنافسة معتدلة", "الطلب مرتفع", "التصدير ممكن", "الجودة عالية",
    "المستوردون كثر", "الشحن بحري", "التقرير جاهز", "الخلاصة واضحة",
    "التوصية بالدخول", "المخاطر محدودة", "الفرصة قائمة", "النتيجة ايجابية",
]

# نطاقاتُ الحروف العربية (بلا حركات — الحركاتُ فئةُ Mn تُتجاهَل في عدّ الأغلبية).
_AR_RANGES = ((0x0600, 0x06FF), (0x0750, 0x077F), (0x08A0, 0x08FF),
              (0xFB50, 0xFDFF), (0xFE70, 0xFEFF))


def _is_arabic_letter(ch: str) -> bool:
    o = ord(ch)
    return any(lo <= o <= hi for lo, hi in _AR_RANGES)


def is_arabic_majority(text: str) -> bool:
    """سطرٌ عربيُّ الأغلبية إن كان عددُ حروفه العربية > اللاتينية و> صفر.
    الأرقامُ والترقيمُ لا تُحسَب — فسطرُ `080410` أو `UN Comtrade` ليس عربيًّا."""
    ar = sum(1 for c in text if _is_arabic_letter(c))
    la = sum(1 for c in text if c.isascii() and c.isalpha())
    return ar > 0 and ar > la


def build_fixture(jc: str) -> str:
    """ابنِ docx بكلّ فقرةٍ bidi + قيمة jc معطاة، من نفس المحتوى."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    doc = Document()
    for ln in _LINES:
        p = doc.add_paragraph()
        r = p.add_run(ln)
        r.font.name = "Amiri"
        ppr = p._p.get_or_add_pPr()
        if ppr.find(qn("w:bidi")) is None:
            ppr.append(OxmlElement("w:bidi"))
        el = ppr.find(qn("w:jc"))
        if el is None:
            el = OxmlElement("w:jc")
            ppr.append(el)
        el.set(qn("w:val"), jc)
        rpr = r._r.get_or_add_rPr()
        if rpr.find(qn("w:rtl")) is None:
            rpr.append(OxmlElement("w:rtl"))
    path = os.path.join(tempfile.mkdtemp(prefix=f"rtlcal_{jc}_"), "f.docx")
    doc.save(path)
    return path


# WP-5 — معايرة فحص انعكاس الأقواس: أسطر عربية بمقاطع لاتينية/رقمية بين
# قوسين — C (بعزل RLM عبر silk_reports._bidi_isolate_brackets، الوصفة
# الصحيحة) مقابل D (بلا عزل — البناء الخام الذي أنتج «) ... (» المُسلَّم).
_BRACKET_LINES = [
    "الواردات (UN Comtrade) في نمو مستمر",
    "متوسط السعر (6 USD/kg) مؤشر سياقي",
    "مؤشر التركز (HHI 2500) مرتفع نسبيا",
    "النمو السنوي (CAGR 5%) خلال ثلاث سنوات",
    "التقييم وفق (World Bank LPI) مستقر",
]


def build_bracket_fixture(isolate: bool) -> str:
    """ابنِ docx أسطره العربية تحوي أقواساً لاتينية — بعزل RLM أو بدونه."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import silk_reports
    doc = Document()
    for ln in _BRACKET_LINES:
        text = silk_reports._bidi_isolate_brackets(ln) if isolate else ln
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Amiri"
        ppr = p._p.get_or_add_pPr()
        if ppr.find(qn("w:bidi")) is None:
            ppr.append(OxmlElement("w:bidi"))
        el = OxmlElement("w:jc")
        el.set(qn("w:val"), "start")
        ppr.append(el)
        rpr = r._r.get_or_add_rPr()
        if rpr.find(qn("w:rtl")) is None:
            rpr.append(OxmlElement("w:rtl"))
    tag = "iso" if isolate else "raw"
    path = os.path.join(tempfile.mkdtemp(prefix=f"rtlcal_br_{tag}_"), "f.docx")
    doc.save(path)
    return path


def bracket_suspicious_count(pdf_path: str) -> "int | None":
    """عدد الأقواس الافتتاحية المعلّقة في نص الـPDF المستخرَج — None بلا fitz."""
    try:
        import fitz  # pymupdf
    except Exception:  # noqa: BLE001
        return None
    import silk_reports
    try:
        with fitz.open(pdf_path) as pdf:
            text = "\n".join(page.get_text() for page in pdf)
    except Exception:  # noqa: BLE001
        return None
    return silk_reports.count_suspicious_brackets(text)


def to_pdf(docx_path: str) -> str:
    import silk_reports
    return silk_reports.docx_to_pdf(
        docx_path, os.path.join(os.path.dirname(docx_path), "f.pdf"))


def measure_lines(pdf_path: str):
    """أرجِع (عرض_الصفحة، [(x0,x1,text) لكلّ **سطرٍ مُجمَّع** على y]) عبر fitz،
    أو None إن غاب pymupdf. النقاطُ بوحدة النقطة."""
    try:
        import fitz  # pymupdf
    except Exception:  # noqa: BLE001
        return None
    try:
        doc = fitz.open(pdf_path)
    except Exception:  # noqa: BLE001
        return None
    pw = doc[0].rect.width if doc.page_count else None
    if not pw:
        return None
    lines = []
    for page in doc:
        rows = collections.defaultdict(list)
        for b in page.get_text("dict")["blocks"]:
            for line in b.get("lines", []):
                for s in line["spans"]:
                    if s["text"].strip():
                        rows[round(s["bbox"][3] / 2)].append(s)
        for sp in rows.values():
            sp.sort(key=lambda s: s["bbox"][0])
            lines.append((min(s["bbox"][0] for s in sp),
                          max(s["bbox"][2] for s in sp),
                          " ".join(s["text"] for s in sp)))
    return pw, lines


def classify(pw: float, lines, tol: float = 10.0) -> dict:
    """صنِّف الأسطرَ العربية وفق عقد «صفرُ **تثبيتٍ يساريّ**» — الإمضاءُ الدقيق
    للانقلاب المنطقيّ لـjc، مُعايَرٌ على fixture-B (المقلوب المعروف).

    **الدرسُ من المصفوفة (لِمَ لا نكتفي بـ«لم يبلغ حدّ اليمين»):** تقريرٌ حقيقيّ
    يحوي أسطرًا عربيةً قصيرةً **مشروعةً** لا تبلغ حدّ اليمين دون أيّ انقلاب —
    عناوينُ انفصل رقمُها في الاستخراج، فقراتٌ بنمطٍ مُزاحٍ (Intense Quote:
    `silk_reports.py:1427`)، أسطرٌ موسَّطة، خلايا أضيقُ من الصفحة. «لم يبلغ
    اليمين» يخلطها بالانقلاب. الإمضاءُ الحقيقيّ للانقلاب (كما في fixture-B:
    `x0≈90` = هامشُ اليسار) هو **بدءُ السطر من هامش اليسار وهو قصير**.

    مراجعُ هندسية (متينةٌ ضدّ الشواذّ):
    - `left_margin`: أدنى `x0` — حدُّ اليسار للنصّ في الصفحة.
    - `block_right`: المئينُ الـ90 لـ`x1` — حدُّ اليمين للكتلة.

    دِلاءُ الأسطر العربية:
    - **`arabic_right`**: تبلغ حدّ اليمين (`block_right - x1 <= tol`) — سليمة.
    - **`arabic_centered`**: مركزُها ≈ مركزُ الصفحة — عنوان/تذييل موسَّط مشروع.
    - **`arabic_left`** (عيبٌ حقيقيّ): تبدأ من هامش اليسار (`x0 - left_margin
      <= 0.10*pw`) ولا تبلغ اليمين — إمضاءُ انقلاب jc.
    - **`arabic_short`**: قصيرةٌ لكنّها ليست يساريّةً ولا موسَّطة (مُزاحة/خليّة)
      — ليست عيبًا.
    اللاتينيُّ (أسماءُ مصادر، تواريخ، رموزُ HS، أرقام) يُستثنى كليًّا."""
    if not lines:
        return {"block_right": 0.0, "left_margin": 0.0, "arabic_right": 0,
                "arabic_left": 0, "arabic_centered": 0, "arabic_short": 0,
                "latin_excluded": 0, "arabic_left_texts": [],
                "arabic_short_texts": [], "arabic_centered_texts": [],
                "right_short_texts": [], "latin_texts": []}
    left_margin = min(x0 for x0, _, _ in lines)
    xs = sorted(x1 for _, x1, _ in lines)
    block_right = xs[min(int(len(xs) * 0.90), len(xs) - 1)]
    # مراجعُ **مطلقةٌ** بكسور الصفحة (لا مشتقّةٌ من البيانات كي لا تنهار على
    # مستندٍ مقلوبٍ بالكامل بأطوالٍ متماثلة): يمينُ الصفحة = آخِرُ 40%، وهامشُ
    # اليسار = أوّلُ 22%. إمضاءُ الانقلاب = بدءٌ من هامش اليسار دون بلوغِ يمين.
    RIGHT_SIDE = 0.60 * pw
    LEFT_ZONE = 0.22 * pw
    CENTER_TOL = 0.06 * pw
    ar_right = 0
    ar_left_texts, ar_short_texts = [], []
    ar_centered_texts, right_short_texts, latin_texts = [], [], []
    for x0, x1, t in lines:
        if not is_arabic_majority(t):
            latin_texts.append((round(x0), round(x1), t))
            continue
        reaches_right = x1 >= RIGHT_SIDE
        at_left = x0 <= LEFT_ZONE
        centered = abs(((x0 + x1) / 2) - pw / 2) <= CENTER_TOL
        if at_left and not reaches_right and not centered:
            ar_left_texts.append((round(x0), round(x1), t))   # عيبٌ: انقلاب
        elif reaches_right:
            ar_right += 1
            if block_right - x1 > tol:      # على الجانب الأيمن لكن دون الحافّة
                right_short_texts.append((round(x0), round(x1), t))
        elif centered:
            ar_centered_texts.append((round(x0), round(x1), t))
        else:
            ar_short_texts.append((round(x0), round(x1), t))
    return {"block_right": round(block_right, 1),
            "left_margin": round(left_margin, 1), "arabic_right": ar_right,
            "arabic_left": len(ar_left_texts),
            "arabic_centered": len(ar_centered_texts),
            "arabic_short": len(ar_short_texts),
            "latin_excluded": len(latin_texts),
            "arabic_left_texts": ar_left_texts,
            "arabic_short_texts": ar_short_texts,
            "arabic_centered_texts": ar_centered_texts,
            "right_short_texts": right_short_texts,
            "latin_texts": latin_texts}


def format_digest(title: str, d: dict) -> str:
    """خُلاصةٌ قابلةٌ للفحص تُطبَع **دائمًا** (لا فقط عند الفشل) — التعديل ٢.
    تُظهِر كلَّ دلوٍ بنصوصه كي يبقى الأخضرُ مفحوصًا: يرى المُشرِفُ **لماذا** لم
    يُوسَم كلُّ سطرٍ قصيرٍ عيبًا (هامشُ يساره/مركزه بعيد عن إمضاء الانقلاب)."""
    out = [f"----- {title} -----",
           f"left_margin(min x0)={d['left_margin']}  block_right(90pct x1)="
           f"{d['block_right']}  arabic_right={d['arabic_right']}  "
           f"arabic_left={d['arabic_left']}  arabic_centered={d['arabic_centered']}"
           f"  arabic_short={d['arabic_short']}  latin_excluded={d['latin_excluded']}"]

    def _dump(label, key):
        if d[key]:
            out.append(f"  {label}:")
            for x0, x1, t in d[key]:
                out.append(f"    [x0={x0} x1={x1}] {t!r}")
    _dump("ARABIC-LEFT offenders (real defects — left-margin anchored)",
          "arabic_left_texts")
    _dump("right-side but short of exact margin (indent/heading-split — ok)",
          "right_short_texts")
    _dump("arabic short, mid-page not left-anchored (indent/cell — ok)",
          "arabic_short_texts")
    _dump("arabic centered (legit heading/footer — ok)", "arabic_centered_texts")
    _dump("latin (excluded from ratio)", "latin_texts")
    return "\n".join(out)


def _lo_version() -> str:
    for b in ("soffice", "libreoffice"):
        exe = shutil.which(b)
        if exe:
            try:
                return subprocess.run([exe, "--version"], capture_output=True,
                                      text=True, timeout=30).stdout.strip()
            except Exception:  # noqa: BLE001
                pass
    return "unknown"


def main() -> int:
    print("LibreOffice:", _lo_version())
    digests = {}
    for label, jc in (("A(jc=start,correct)", "start"),
                      ("B(jc=right,inverted)", "right")):
        try:
            measured = measure_lines(to_pdf(build_fixture(jc)))
            if measured is None:
                print(f"{label}: no fitz measurer available")
                continue
            pw, lines = measured
            d = classify(pw, lines)
            digests[label] = d
            print(format_digest(label, d))
        except Exception as e:  # noqa: BLE001
            print(f"{label}: ERROR {e}")
    # WP-5: معايرة فحص انعكاس الأقواس — C (بعزل RLM) مقابل D (خام).
    # تشخيصية هنا؛ الفحص الصلب يعيش في silk_reports.docx_to_pdf نفسه —
    # نرفع عتبته مؤقتاً كي تُقاس عيّنة D بلا رفع استثناء.
    os.environ["SILK_PDF_BRACKET_FAIL_MAX"] = "9999"
    try:
        for label, iso in (("C(bracket,RLM-isolated)", True),
                           ("D(bracket,raw)", False)):
            try:
                n = bracket_suspicious_count(to_pdf(build_bracket_fixture(iso)))
                print(f"{label}: suspicious_open_brackets={n}")
            except Exception as e:  # noqa: BLE001
                print(f"{label}: ERROR {e}")
    finally:
        os.environ.pop("SILK_PDF_BRACKET_FAIL_MAX", None)

    # عقدُ الصلاحية الدائم: A بلا يسارٍ عربيّ (تمرّ)؛ B بيسارٍ عربيّ (تفشل).
    a = digests.get("A(jc=start,correct)")
    b = digests.get("B(jc=right,inverted)")
    a_ok = bool(a) and a["arabic_left"] == 0 and a["arabic_right"] >= 3
    b_catches = bool(b) and b["arabic_left"] > 0
    print(f"\nCALIBRATION (zero-left Arabic contract): "
          f"A passes={a_ok} | B is caught={b_catches} | "
          f"VALID MEASURER={a_ok and b_catches}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
